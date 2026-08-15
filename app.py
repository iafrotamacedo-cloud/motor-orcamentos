# -*- coding: utf-8 -*-
# =====================================================================
#  CONTADOR DE REVISÕES DESTE app.py: 116
#  (some +1 sempre que uma versão nova for gerada)
# =====================================================================
"""
MOTOR DE ORÇAMENTOS — Frota Macedo  (100% online: FastAPI no Render)

O operador abre a página, sobe as NOTAS/DAVs (PDF ou imagem) e clica em
"Gerar e enviar ao Dropbox". O motor faz TUDO no Dropbox online (não devolve ZIP):
  - gera 1 orçamento (PDF) por ticket em ORCAMENTOS MONTADOS/<mês>/<loja>/
    e uma cópia em "1 - ORÇAMENTOS NÃO LANÇADOS" (p/ o robô do Trílogo);
  - roteia cada nota (página a página) em NOTAS INCLUIDAS ORCAMENTO/<aba>,
    SEM TICKET ou TICKET NAO ASSOCIADO;
  - mantém a planilha de controle do mês atualizada (e oferece p/ baixar no app);
  - atualiza as planilhas de correção (SEM TICKET / TICKET NAO ASSOCIADO).

Leitura das notas (visão): Groq (Llama/Qwen) se GROQ_API_KEY existir; senão Gemini.
Ticket -> loja/aba: tabela `chamados` no Supabase (alimentada pelo robô do Trílogo).
Regras (dos scripts): +20% no unitário (sem mencionar), rateio de entrega,
1 orçamento por ticket, dedup pelo NÚMERO DA NOTA (permite +1 orçamento por ticket).

Interface: FastAPI + uvicorn (sem Gradio). Login por HTTP Basic (APP_USER/APP_PASS).

Segredos no Render (Environment):
  Leitura:  GROQ_API_KEY   (recomendado)  ou  GEMINI_API_KEY
  Supabase: SUPABASE_URL, SUPABASE_SERVICE_KEY
  Login:    APP_USER, APP_PASS
  Dropbox:  DROPBOX_APP_KEY, DROPBOX_APP_SECRET, DROPBOX_REFRESH_TOKEN, DROPBOX_BASE
  Opcionais: GROQ_MODEL, GEMINI_MODEL, FAT_NOME, FAT_CNPJ
"""
import os, re, io, json, subprocess, tempfile, datetime, unicodedata, urllib.parse, urllib.request, urllib.error
import google.generativeai as genai
import dropbox_rateio

# Fuso de Fortaleza/Brasília (UTC-3, sem horário de verão). O Render roda em UTC;
# sem isto, o sistema datava/classificava com a hora do servidor (até 3h à frente).
TZ_OFFSET_H = int(os.environ.get("TZ_OFFSET_H", "-3"))
def _agora(): return datetime.datetime.utcnow() + datetime.timedelta(hours=TZ_OFFSET_H)
def _hoje():  return _agora().date()

BASE = os.path.dirname(os.path.abspath(__file__))
SCRIPTS = os.path.join(BASE, "scripts")
ASSETS = os.path.join(BASE, "assets")
LOGO = os.path.join(ASSETS, "logo_frota.jpg")
CAD = json.load(open(os.path.join(ASSETS, "cadastro_lojas.json"), encoding="utf-8"))

GEMINI_KEY  = os.environ.get("GEMINI_API_KEY", "")
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash-lite")   # ajustável; o motor ainda tem fallback abaixo
_GEMINI_OK = None   # modelo que funcionou (cache entre chamadas)
GROQ_KEY   = os.environ.get("GROQ_API_KEY", "")
GROQ_MODEL = os.environ.get("GROQ_MODEL", "qwen/qwen3.6-27b")
SB_URL = os.environ.get("SUPABASE_URL", "").rstrip("/")
SB_KEY = os.environ.get("SUPABASE_SERVICE_KEY", "")
# anon key (pública) — usada para validar o token do usuário do FrotaHub
SB_ANON = os.environ.get("SUPABASE_ANON_KEY",
    "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6ImZhYWxnZmJ1Z3Zla2J1aGh0YXR0Iiwicm9sZSI6ImFub24iLCJpYXQiOjE3ODU4NzA0OTIsImV4cCI6MjEwMTQ0NjQ5Mn0.Vwba3hsm43bwOXMXf2iQMkCWXqrGZaHKojHvV6mhFSI")
# --- GitHub (disparo do robô do Trílogo via workflow_dispatch) ---
GH_TOKEN    = os.environ.get("GITHUB_TOKEN", "") or os.environ.get("GH_TOKEN", "")
GH_REPO     = os.environ.get("GH_REPO", "")                     # ex.: "usuario/trilogo_robo"
GH_WORKFLOW = os.environ.get("GH_WORKFLOW", "trilogo-chamados.yml")
GH_REF      = os.environ.get("GH_REF", "main")
# segredo compartilhado com o Supabase (pg_cron/pg_net) para o agendador
AGENDADOR_SECRET = os.environ.get("AGENDADOR_SECRET", "")
# base do PCO no Dropbox (novo layout FROTAHUB)
PCO_BASE = os.environ.get("PCO_BASE", "/FROTAHUB/1 - ADMINISTRATIVO/1 - PCO")
# --- ENVIAR_PCO: SMTP + destinatários (padrão = TESTE p/ igor@; trocar por env p/ ir ao ar) ---
SMTP_HOST = os.environ.get("SMTP_HOST", "mail.frotamacedo.com.br")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "465"))
SMTP_USER = os.environ.get("SMTP_USER", "")
SMTP_PASS = os.environ.get("SMTP_PASS", "")
# Render bloqueia portas de SMTP -> preferimos a API HTTP do Brevo quando houver chave
BREVO_API_KEY = os.environ.get("BREVO_API_KEY", "")
PCO_FROM  = os.environ.get("PCO_FROM", SMTP_USER or "pco@frotamacedo.com.br")
def _lista_env(k, padrao):
    return [x.strip() for x in os.environ.get(k, padrao).split(",") if x.strip()]
PCO_TO      = _lista_env("PCO_TO", "igor@frotamacedo.com.br")
PCO_CC      = _lista_env("PCO_CC", "")
PCO_BLOQ_TO = _lista_env("PCO_BLOQ_TO", "igor@frotamacedo.com.br")
ASSINATURA = ("<b>Isabele Melissa</b><br><b>Administrativo</b><br>"
              "<b>(85) 98795-5735</b><br><b>Frota Macedo Engenharia LTDA</b>")
FAT_NOME = os.environ.get("FAT_NOME", "FROTA MACEDO ENGENHARIA LTDA")
FAT_CNPJ = os.environ.get("FAT_CNPJ", "27.363.223/0001-70")
if GEMINI_KEY:
    genai.configure(api_key=GEMINI_KEY)

MESES = ["JANEIRO","FEVEREIRO","MARÇO","ABRIL","MAIO","JUNHO","JULHO","AGOSTO",
         "SETEMBRO","OUTUBRO","NOVEMBRO","DEZEMBRO"]

def slug(s):
    return re.sub(r"\s+", "_", (s or "").strip())

def norm(s):
    s = unicodedata.normalize("NFKD", str(s or "")).encode("ascii","ignore").decode().upper()
    return re.sub(r"[^A-Z0-9 ]", " ", s).strip()

def fmt_cnpj(s):
    d = re.sub(r"\D", "", str(s or ""))
    if len(d) == 14:
        return f"{d[:2]}.{d[2:5]}.{d[5:8]}/{d[8:12]}-{d[12:]}"
    return str(s).strip() if s else "—"

def cidade_de(endereco):
    """Extrai 'Cidade - UF' de um endereço do cadastro (ex.: '..., Fortaleza - CE, 60175-395')."""
    m = re.search(r",\s*([^,]+?\s*-\s*[A-Z]{2})\b", str(endereco or ""))
    return m.group(1).strip() if m else ""

RE_TICKET = re.compile(r"(?:ticket|ticker|tickte|tikett|tiket|tcket|tck|tk|chamado|cham|\bos\b|or[çc]amento|#)\s*[:/#\.\-]?\s*(\d{5,6})\b", re.I)
def ticket_do_texto(t):
    """Extrai o ticket (5-6 dígitos) de um texto de observação, tolerando rótulos errados."""
    t = str(t or "")
    m = RE_TICKET.search(t)
    if m: return m.group(1)
    m = re.search(r"(?<!\d)(\d{5,6})(?!\d)", t)   # fallback: número isolado de 5-6 dígitos
    return m.group(1) if m else ""

# ---------- leitura das notas (render por página) ----------
def paginas_imagens(path, workdir):
    """Retorna lista de (indice, caminho_png). PDF: render por página; imagem: usa direto."""
    ext = path.lower().rsplit(".", 1)[-1]
    out = []
    if ext in ("jpg", "jpeg", "png"):
        return [(1, path)]
    if ext == "pdf":
        try:
            info = subprocess.run(["pdfinfo", path], capture_output=True, text=True).stdout
            m = re.search(r"Pages:\s+(\d+)", info); n = int(m.group(1)) if m else 1
        except Exception:
            n = 1
        for p in range(1, n+1):
            pref = os.path.join(workdir, f"pg_{os.path.basename(path)}_{p}")
            subprocess.run(["pdftoppm", "-r", "170", "-png", "-f", str(p), "-l", str(p), path, pref],
                           capture_output=True)
            # pdftoppm gera pref-<p>.png (com zero à esquerda dependendo do total)
            cand = [f for f in os.listdir(workdir) if f.startswith(os.path.basename(pref)) and f.endswith(".png")]
            if cand: out.append((p, os.path.join(workdir, sorted(cand)[0])))
        return out
    return out

PROMPT = """Você lê UMA nota fiscal / DAV / orçamento de fornecedor (imagem).
O EMITENTE (Razão Social de quem emitiu) é o "fornecedor". Ignore o destinatário.
Extraia SOMENTE o que estiver na nota e responda em JSON puro (sem texto fora do JSON):
{
 "ticket": "<número do ticket/chamado, 5-6 dígitos. Quase sempre aparece no campo Observação / Dados Complementares / Dados adicionais, precedido por um rótulo que PODE estar com erro de digitação ou abreviado: TICKET, TICKER, TICKTE, TIKET, TCKET, TK, CHAMADO, OS, ORÇAMENTO, ou apenas #. Retorne SÓ os dígitos. Exemplos: 'TICKER:126486 AMAURI COCO' -> 126486 ; 'Obs: tk 125411' -> 125411 ; '#126486' -> 126486. Só use null se não houver mesmo nenhum número de 5-6 dígitos com esse sentido>",
 "num_documento": "<Nº do Documento/Nota, só os dígitos e SEM zeros à esquerda (ex.: '0000018747' -> '18747'); se não houver, null>",
 "fornecedor": "<razão social do EMITENTE>",
 "forma_pagamento": "<ex.: Boleto 30 dias / Plano de Pagamento, se aparecer; senão null>",
 "obs": "<TRANSCREVA VERBATIM, exatamente como está na nota, todo o conteúdo dos campos Observação / Dados Complementares / Dados Adicionais / Informações Complementares (é onde costuma estar o ticket). Se não houver, ''>",
 "itens": [
   {"desc":"<descrição do item, LIMPA: remova o código/SKU do início (ex.: '00000000001633 - CIMENTO TODAS AS OBRAS' -> 'CIMENTO TODAS AS OBRAS') e não inclua NCM/CFOP>", "quant":<número>, "unid":"<UN/PC/SV/KG/…>", "unit":<valor UNITÁRIO BRUTO, sem imposto/desconto/acréscimo>}
 ]
}
Regras: NÃO aplique nenhum acréscimo/desconto — use o valor unitário BRUTO. Inclua item de ENTREGA/FRETE se houver (desc contendo 'ENTREGA' ou 'FRETE'). quant e unit são números (ponto decimal)."""

def _parse_json(txt):
    txt = (txt or "").strip()
    try:
        return json.loads(txt)
    except Exception:
        m = re.search(r"\{.*\}", txt, re.S)
        return json.loads(m.group(0)) if m else {"ticket": None, "itens": []}

def gemini_le(png_path):
    """Tenta uma lista de modelos e usa o 1º que funcionar (memoriza no _GEMINI_OK)."""
    global _GEMINI_OK
    from PIL import Image
    img = Image.open(png_path)
    candidatos = [_GEMINI_OK] if _GEMINI_OK else list(dict.fromkeys([
        GEMINI_MODEL, "gemini-2.5-flash-lite", "gemini-flash-latest",
        "gemini-2.0-flash-lite", "gemini-2.5-flash", "gemini-2.0-flash",
    ]))
    ultimo = None
    for nome in candidatos:
        if not nome: continue
        try:
            model = genai.GenerativeModel(nome)
            r = model.generate_content([PROMPT, img],
                generation_config={"temperature": 0, "response_mime_type": "application/json"})
            _GEMINI_OK = nome
            return _parse_json(r.text)
        except Exception as e:
            ultimo = e
            m = str(e).lower()
            if any(s in m for s in ("not available", "not found", "404", "no longer",
                                    "unsupported", "permission", "does not exist")):
                continue                 # modelo indisponível p/ essa conta -> tenta o próximo
            raise                        # erro real (ex.: cota) -> propaga
    raise ultimo or RuntimeError("nenhum modelo Gemini disponível")

def groq_le(png_path):
    import base64, io
    from PIL import Image
    # converte SEMPRE p/ PNG de verdade (evita descasar bytes JPEG com rótulo image/png)
    im = Image.open(png_path).convert("RGB")
    if im.width > 1600:                      # reduz p/ caber no limite da API
        im = im.resize((1600, int(im.height * 1600 / im.width)))
    buf = io.BytesIO(); im.save(buf, "PNG")
    b64 = base64.b64encode(buf.getvalue()).decode()
    payload = {
        "model": GROQ_MODEL, "temperature": 0,
        "messages": [{"role": "user", "content": [
            {"type": "text", "text": PROMPT},
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}},
        ]}],
    }
    req = urllib.request.Request("https://api.groq.com/openai/v1/chat/completions",
        data=json.dumps(payload).encode(),
        headers={"Authorization": f"Bearer {GROQ_KEY}", "Content-Type": "application/json",
                 "Accept": "application/json",
                 "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                               "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"})
    try:
        raw = urllib.request.urlopen(req, timeout=90).read().decode()
    except urllib.error.HTTPError as e:
        body = ""
        try: body = e.read().decode("utf-8", "ignore")
        except Exception: pass
        raise RuntimeError(f"Groq {e.code}: {body[:300]}") from None
    r = json.loads(raw)
    return _parse_json(r["choices"][0]["message"]["content"])

def ler_nota(png_path):
    """Usa o Groq (visão) se a chave estiver setada; senão o Gemini."""
    return groq_le(png_path) if GROQ_KEY else gemini_le(png_path)

# ---------- ticket -> loja (Supabase chamados) ----------
def busca_chamado(ticket):
    if not (SB_URL and SB_KEY and ticket): return None
    q = urllib.parse.urlencode({"numero": f"eq.{ticket}", "select": "loja,aba,descricao", "limit": "1"})
    req = urllib.request.Request(f"{SB_URL}/rest/v1/chamados?{q}",
        headers={"apikey": SB_KEY, "authorization": f"Bearer {SB_KEY}"})
    try:
        data = json.loads(urllib.request.urlopen(req, timeout=20).read().decode())
        return data[0] if data else None
    except Exception:
        return None

def loja_cadastro(loja_str):
    """Acha a loja no cadastro_lojas.json a partir da 'Unidade' do chamado."""
    m = re.search(r"LOJA\s*0*(\d{1,2})", norm(loja_str))
    if m:
        e = CAD.get(m.group(1).zfill(2))
        if e: return e
    alvo = norm(loja_str)
    for e in CAD.values():
        nomes = [norm(e.get("nome"))] + [norm(a) for a in e.get("apelidos", [])]
        if any(nm and nm in alvo for nm in nomes):
            return e
    return None

# ---------- geração de 1 orçamento ----------
def gera_orcamento(ticket, chamado, itens_brutos, workdir, outdir, data_str, forma_pag):
    prep = json.loads(subprocess.run(["python3", os.path.join(SCRIPTS, "preparar_itens.py")],
        input=json.dumps({"itens": itens_brutos}), capture_output=True, text=True).stdout)
    e = loja_cadastro(chamado.get("loja")) or {}
    num = e.get("numero") or "00"
    nome_loja = e.get("nome") or re.sub(r"^LOJA\s*\d+\s*-\s*", "", chamado.get("loja") or "").strip() or "LOJA"
    cliente = {"nome": f"Mercadinhos São Luiz — {nome_loja.title()}",
               "cnpj": fmt_cnpj(e.get("cnpj")), "endereco": e.get("endereco") or "—",
               "cidade": cidade_de(e.get("endereco"))}
    obra = f"{norm(chamado.get('descricao'))}  #{ticket}"
    dados = {"numero_ticket": str(ticket), "revisao": 1, "data": data_str, "obra": obra,
             "faturamento": {"nome": FAT_NOME, "cnpj": FAT_CNPJ,
                             "forma_pagamento": forma_pag or "", "data_faturamento": data_str},
             "cliente": cliente, "itens": prep["itens"],
             "total_geral": prep["total_geral"], "extenso": prep["extenso"]}
    dj = os.path.join(workdir, f"dados_{ticket}.json")
    json.dump(dados, open(dj, "w"), ensure_ascii=False)
    base = f"{num}_{slug(nome_loja)}_{ticket}"
    pdf = os.path.join(workdir, base + ".pdf")
    subprocess.run(["python3", os.path.join(SCRIPTS, "gerar_pdf.py"), dj, pdf, LOGO],
                   capture_output=True)
    return dict(ticket=ticket, num=num, nome_loja=nome_loja, loja=f"{num} - {nome_loja}",
                total=prep["total_geral"], pdf=pdf, base=base, pdf_ok=os.path.exists(pdf))

# ---------- orquestração ----------
def _nota_nome_pdf(info):
    """Nome da nota roteada no Dropbox: 'TICKET <n> - NOTA <doc>.pdf' (sempre .pdf)."""
    num = re.sub(r'[\\/:*?"<>|]', "", str(info.get("num") or "")).strip()
    tk = info.get("ticket") or ""
    if info["cat"] == "SEM TICKET":
        base = f"TICKET SEMTICKET - NOTA {num}" if num else f"NOTA p{info.get('page', 1)}"
    elif tk:
        base = f"TICKET {tk} - NOTA {num}" if num else f"TICKET {tk} - p{info.get('page', 1)}"
    else:
        base = f"NOTA p{info.get('page', 1)}"
    return base + ".pdf"

def _pagina_pdf(src, page, is_img, workdir, tag):
    """Produz um PDF de 1 página a partir de uma página do arquivo (imagem -> PDF; PDF -> só a página)."""
    out = os.path.join(workdir, f"nota_{tag}.pdf")
    try:
        if is_img:
            from PIL import Image
            Image.open(src).convert("RGB").save(out, "PDF")
        else:
            subprocess.run(["pdfseparate", "-f", str(page), "-l", str(page), src, out], capture_output=True)
        return out if os.path.exists(out) else None
    except Exception:
        return None

def processar(arquivos, planilha_controle=None):
    if not (GEMINI_KEY or GROQ_KEY):
        return None, "⚠️ Configure o segredo GROQ_API_KEY (ou GEMINI_API_KEY)."
    if not (SB_URL and SB_KEY):
        return None, "⚠️ Faltam os segredos SUPABASE_URL / SUPABASE_SERVICE_KEY."
    if not dropbox_rateio.ativo():
        return None, "⚠️ Configure os segredos do Dropbox (DROPBOX_APP_KEY/SECRET/REFRESH_TOKEN/BASE)."
    try:
        access = dropbox_rateio.obter_token()
        if not access: raise RuntimeError("token vazio")
    except Exception as e:
        return None, f"⚠️ Falha ao autenticar no Dropbox: {e}"

    work = tempfile.mkdtemp(prefix="orc_")
    hoje = _hoje()
    mes = f"{MESES[hoje.month-1]} {hoje.year}"
    data_str = hoje.strftime("%d/%m/%Y")
    B = dropbox_rateio.BASE
    dbx_mes = f"{B}/{dropbox_rateio.ORCAMENTOS}/{mes}"
    ATUALIZAR = os.path.join(SCRIPTS, "atualizar_planilha_mensal.py")
    PENDENTES = os.path.join(SCRIPTS, "pendentes.py")
    erros_dbx = []

    # 1) baixa a planilha de controle e o registro de notas já feitas neste mês
    ctrl = os.path.join(work, f"ORÇAMENTOS MONTADOS - {mes}.xlsx")
    ctrl_dbx = f"{dbx_mes}/ORÇAMENTOS MONTADOS - {mes}.xlsx"
    try:
        atual = dropbox_rateio.baixar(access, ctrl_dbx)
        if atual: open(ctrl, "wb").write(atual)
    except Exception as e:
        erros_dbx.append(f"baixar controle: {e}")
    reg_dbx = f"{dbx_mes}/.notas_processadas.json"
    notas_feitas = set()
    try:
        rb = dropbox_rateio.baixar(access, reg_dbx)
        if rb: notas_feitas = set(str(x) for x in json.loads(rb.decode("utf-8")))
    except Exception:
        pass

    # 2) lê as notas (por página); dedup pelo NÚMERO DA NOTA; agrupa por ticket
    por_ticket = {}
    sem_ticket, nao_assoc = [], []
    paginas = []          # cada página vira 1 arquivo roteado no Dropbox
    novas_notas = set()   # notas que VIRARAM orçamento (vão pro registro do mês)
    vistos = set()        # guarda contra repetir a MESMA nota dentro deste lote
    duplicadas = 0
    diag = []             # diagnóstico do que o Gemini leu (p/ notas SEM TICKET)
    for f in (arquivos or []):
        path = f.name if hasattr(f, "name") else f
        is_img = path.lower().rsplit(".", 1)[-1] in ("jpg", "jpeg", "png")
        pgs = paginas_imagens(path, work)
        if not pgs:
            paginas.append({"src": path, "page": 1, "is_img": is_img, "cat": "RESIDUAL", "ticket": "", "num": ""})
            continue
        for idx, png in pgs:
            try:
                nota = ler_nota(png)
            except Exception as e:
                nota = {"ticket": None, "itens": [], "fornecedor": None, "_erro": str(e)}
            try:
                print("LEITURA:", json.dumps({k: nota.get(k) for k in ("ticket","num_documento","fornecedor","obs","_erro")}, ensure_ascii=False)[:500], flush=True)
            except Exception:
                pass
            ticket = re.sub(r"\D", "", str(nota.get("ticket") or ""))
            if not ticket:                                   # rede de segurança: extrai do texto da observação
                ticket = ticket_do_texto(nota.get("obs"))
            itens = nota.get("itens") or []
            numdoc = str(nota.get("num_documento") or "").strip()
            # dedup por número da nota: nota já ORÇADA no mês, ou repetida neste mesmo lote -> ignora
            if numdoc and (numdoc in notas_feitas or numdoc in vistos):
                duplicadas += 1; continue
            if numdoc: vistos.add(numdoc)
            reg = {"nota": numdoc, "fornecedor": nota.get("fornecedor") or "", "ticket": ticket, "loja": ""}
            pg = {"src": path, "page": idx, "is_img": is_img, "ticket": ticket, "num": numdoc}
            if not ticket:
                diag.append({"forn": nota.get("fornecedor") or "?", "obs": str(nota.get("obs") or "")[:140],
                             "nitens": len(itens), "err": nota.get("_erro")})
                sem_ticket.append({**reg, "status": "SEM TICKET"}); pg["cat"] = "SEM TICKET"; paginas.append(pg); continue
            ch = busca_chamado(ticket)
            if not ch:
                nao_assoc.append({**reg, "status": "TICKET NÃO ASSOCIADO"}); pg["cat"] = "NAO ASSOCIADO"; paginas.append(pg); continue
            pg["cat"] = (ch.get("aba") or "CIVIL").upper(); paginas.append(pg)
            g = por_ticket.setdefault(ticket, {"chamado": ch, "itens": [], "forma": nota.get("forma_pagamento"), "notas": set()})
            g["itens"].extend(itens)
            if numdoc: g["notas"].add(numdoc)
            if nota.get("forma_pagamento") and not g["forma"]: g["forma"] = nota.get("forma_pagamento")

    # 3) gera 1 orçamento por ticket -> sobe PDF (mês/loja + não lançados) + acrescenta no controle
    feitos = []
    for ticket, g in por_ticket.items():
        if not g["itens"]: continue
        try:
            d = gera_orcamento(ticket, g["chamado"], g["itens"], work, work, data_str, g["forma"])
        except Exception as e:
            nao_assoc.append({"nota": "", "fornecedor": "", "status": f"ERRO: {e}", "ticket": ticket, "loja": ""}); continue
        feitos.append(d)
        novas_notas |= g.get("notas", set())      # só REGISTRA as notas que viraram orçamento
        if d["pdf_ok"]:
            subprocess.run(["python3", ATUALIZAR, "--xlsx", ctrl, "--ticket", str(ticket),
                            "--loja", d["nome_loja"], "--pdf", d["pdf"], "--data", data_str, "--append"], capture_output=True)
            nome_pdf = d["base"] + ".pdf"
            try: dropbox_rateio.subir(access, d["pdf"], f"{dbx_mes}/{d['num']}_{slug(d['nome_loja'])}/{nome_pdf}")
            except Exception as e: erros_dbx.append(f"PDF {ticket}: {e}")
            try: dropbox_rateio.subir(access, d["pdf"], f"{B}/{dropbox_rateio.NAO_LANCADOS}/{nome_pdf}")
            except Exception as e: erros_dbx.append(f"não lançados {ticket}: {e}")

    # 4) sobe controle atualizado + registro de notas processadas
    ctrl_ok = os.path.exists(ctrl)
    if ctrl_ok:
        try: dropbox_rateio.subir(access, ctrl, ctrl_dbx, overwrite=True)
        except Exception as e: erros_dbx.append(f"subir controle: {e}")
    if novas_notas:
        try:
            dropbox_rateio.subir_bytes(access, json.dumps(sorted(notas_feitas | novas_notas)).encode("utf-8"),
                                       reg_dbx, overwrite=True)
        except Exception as e:
            erros_dbx.append(f"registro notas: {e}")

    # 5) planilhas de correção por categoria (baixa -> adiciona -> sobe)
    def atualiza_correcao(lista, subpasta, arq):
        if not lista: return
        dbxp = f"{B}/{subpasta}/{arq}"
        local = os.path.join(work, arq)
        try:
            atual = dropbox_rateio.baixar(access, dbxp)
            if atual: open(local, "wb").write(atual)
        except Exception as e:
            erros_dbx.append(f"baixar {arq}: {e}")
        inp = os.path.join(work, "p_" + re.sub(r"\W+", "_", subpasta) + ".json")
        json.dump(lista, open(inp, "w"), ensure_ascii=False)
        subprocess.run(["python3", PENDENTES, "add", "--xlsx", local, "--input", inp], capture_output=True)
        if os.path.exists(local):
            try: dropbox_rateio.subir(access, local, dbxp, overwrite=True)
            except Exception as e: erros_dbx.append(f"subir {arq}: {e}")
    atualiza_correcao(sem_ticket, "SEM TICKET", "NOTAS - SEM TICKET.xlsx")
    atualiza_correcao(nao_assoc, "TICKET NAO ASSOCIADO", "NOTAS - TICKET NAO ASSOCIADO.xlsx")

    # 6) roteia as notas página a página (multipágina: cada página vai pro seu destino)
    itens_dbx = []
    for i, pg in enumerate(paginas):
        local = _pagina_pdf(pg["src"], pg["page"], pg["is_img"], work, str(i))
        if not local:
            erros_dbx.append(f"extrair pág {pg['page']} de {os.path.basename(pg['src'])}"); continue
        itens_dbx.append({"local": local, "categoria": pg["cat"], "nome_destino": _nota_nome_pdf(pg)})
    ok_r, erros_r = dropbox_rateio.ratear(access, itens_dbx)

    # mensagem
    msg = [f"✅ {len(feitos)} orçamento(s) gerado(s) e enviado(s) ao Dropbox."]
    for d in feitos: msg.append(f"• Ticket {d['ticket']} — {d['loja']} — {d['total']}" + ("" if d['pdf_ok'] else "  (⚠️ PDF falhou)"))
    msg.append(f"📁 {ok_r} nota(s) roteada(s) no Dropbox.")
    if duplicadas: msg.append(f"↩️ {duplicadas} nota(s) já processada(s) neste mês (ignoradas pelo número da nota).")
    if sem_ticket: msg.append(f"⚠️ {len(sem_ticket)} SEM TICKET (planilha de correção atualizada).")
    if diag:
        msg.append(f"🔎 Leitor: {'GROQ ' + GROQ_MODEL if GROQ_KEY else 'GEMINI ' + (_GEMINI_OK or GEMINI_MODEL)}")
    for x in diag[:5]:
        det = x.get("err") and f" | ERRO: {x['err']}" or ""
        msg.append(f"   🔎 forn='{x['forn']}' | itens={x.get('nitens')} | obs=«{x['obs']}»{det}")
    if nao_assoc:  msg.append(f"⚠️ {len(nao_assoc)} TICKET NÃO ASSOCIADO (planilha de correção atualizada).")
    if ctrl_ok:    msg.append("🧾 Planilha de controle do mês atualizada no Dropbox — baixe abaixo p/ enviar ao cliente.")
    errs = erros_dbx + erros_r
    if errs: msg.append("⚠️ Erros no Dropbox: " + "; ".join(errs[:4]))
    return (ctrl if ctrl_ok else None), "\n".join(msg)

# ---------- 2ª passagem: ORÇAMENTOS CORRIGIDOS (lê a planilha de pendentes no Dropbox) ----------
def _pendentes_list(xlsx_local):
    out = subprocess.run(["python3", os.path.join(SCRIPTS, "pendentes.py"), "list", "--xlsx", xlsx_local],
                         capture_output=True, text=True)
    try: return json.loads(out.stdout.strip() or "[]")
    except Exception: return []

def _acha_nota_pdf(arqs, numero):
    """Casa o nº da nota com um arquivo da pasta (ex.: 9107 -> 'TICKET 126486 - NOTA 9107.pdf')."""
    n = re.sub(r"\D", "", str(numero or ""))
    if not n: return None
    cands = [a for a in arqs if n in re.sub(r"\D", "", a)]
    # prefere o que casa o número exato entre separadores
    exatos = [a for a in cands if re.search(rf"(?<!\d){n}(?!\d)", a)]
    return (exatos or cands or [None])[0]

def processar_corrigidos(subpasta="TICKET NAO ASSOCIADO", arq_xlsx="NOTAS - TICKET NAO ASSOCIADO.xlsx"):
    """Gera os orçamentos das linhas RODAR=SIM da planilha de pendentes, roteia as
    notas p/ NOTAS INCLUIDAS ORCAMENTO/<classe> e remove as linhas processadas."""
    if not (GEMINI_KEY or GROQ_KEY):
        return None, "⚠️ Configure GEMINI_API_KEY (ou GROQ_API_KEY) no Render."
    if not (SB_URL and SB_KEY):
        return None, "⚠️ Faltam SUPABASE_URL / SUPABASE_SERVICE_KEY."
    if not dropbox_rateio.ativo():
        return None, "⚠️ Faltam os segredos do Dropbox."
    access = dropbox_rateio.obter_token()
    B = dropbox_rateio.BASE
    work = tempfile.mkdtemp(prefix="corr_")
    hoje = _hoje()
    mes = f"{MESES[hoje.month-1]} {hoje.year}"
    data_str = hoje.strftime("%d/%m/%Y")
    dbx_mes = f"{B}/{dropbox_rateio.ORCAMENTOS}/{mes}"
    ATUALIZAR = os.path.join(SCRIPTS, "atualizar_planilha_mensal.py")
    PENDENTES = os.path.join(SCRIPTS, "pendentes.py")
    erros = []

    pasta_notas = f"{B}/{subpasta}"
    xlsx_dbx = f"{pasta_notas}/{arq_xlsx}"
    xlsx_local = os.path.join(work, arq_xlsx)
    raw = dropbox_rateio.baixar(access, xlsx_dbx)
    if not raw:
        return None, f"⚠️ Planilha não encontrada: {xlsx_dbx}"
    open(xlsx_local, "wb").write(raw)
    linhas = _pendentes_list(xlsx_local)
    if not linhas:
        return None, "Nenhuma linha com RODAR = SIM na planilha de pendentes."
    arqs = dropbox_rateio.listar(access, pasta_notas)

    # planilha de controle do mês (append com dedup por ticket)
    ctrl = os.path.join(work, f"ORÇAMENTOS MONTADOS - {mes}.xlsx")
    ctrl_dbx = f"{dbx_mes}/ORÇAMENTOS MONTADOS - {mes}.xlsx"
    try:
        a = dropbox_rateio.baixar(access, ctrl_dbx)
        if a: open(ctrl, "wb").write(a)
    except Exception as e: erros.append(f"baixar controle: {e}")

    # 1) agrupa por ticket (várias notas do mesmo ticket -> 1 orçamento)
    por_ticket, notas_por_ticket, achadas = {}, {}, []
    faltando = []
    for row in linhas:
        ticket = re.sub(r"\D", "", str(row.get("ticket") or ""))
        if not ticket:
            faltando.append(f"nota {row.get('nota')}: sem ticket na planilha"); continue
        nome_arq = _acha_nota_pdf(arqs, row.get("nota"))
        if not nome_arq:
            faltando.append(f"nota {row.get('nota')}: PDF não encontrado na pasta"); continue
        pdf_bytes = dropbox_rateio.baixar(access, f"{pasta_notas}/{nome_arq}")
        if not pdf_bytes:
            faltando.append(f"nota {row.get('nota')}: falha ao baixar"); continue
        local_pdf = os.path.join(work, nome_arq)
        open(local_pdf, "wb").write(pdf_bytes)
        # lê os itens da nota (todas as páginas)
        itens, forma = [], None
        for idx, png in paginas_imagens(local_pdf, work):
            try:
                nota = ler_nota(png)
            except Exception as e:
                nota = {"itens": [], "_erro": str(e)}
            itens.extend(nota.get("itens") or [])
            if nota.get("forma_pagamento") and not forma: forma = nota.get("forma_pagamento")
        if not itens:
            faltando.append(f"nota {row.get('nota')} (ticket {ticket}): não consegui ler itens"); continue
        # chamado: Supabase; senão usa a Loja preenchida na planilha
        ch = busca_chamado(ticket) or {"loja": row.get("loja") or "", "aba": None, "descricao": ""}
        g = por_ticket.setdefault(ticket, {"chamado": ch, "itens": [], "forma": forma})
        g["itens"].extend(itens)
        if forma and not g["forma"]: g["forma"] = forma
        notas_por_ticket.setdefault(ticket, []).append({"nota": row.get("nota"), "arq": nome_arq, "aba": ch.get("aba")})
        achadas.append(row.get("nota"))

    # 2) gera 1 orçamento por ticket + sobe PDFs + atualiza controle
    feitos = []
    for ticket, g in por_ticket.items():
        try:
            d = gera_orcamento(ticket, g["chamado"], g["itens"], work, work, data_str, g["forma"])
        except Exception as e:
            erros.append(f"orçamento ticket {ticket}: {e}"); continue
        feitos.append(d)
        if d["pdf_ok"]:
            subprocess.run(["python3", ATUALIZAR, "--xlsx", ctrl, "--ticket", str(ticket),
                            "--loja", d["nome_loja"], "--pdf", d["pdf"], "--data", data_str, "--append"],
                           capture_output=True)
            nome_pdf = d["base"] + ".pdf"
            try: dropbox_rateio.subir(access, d["pdf"], f"{dbx_mes}/{d['num']}_{slug(d['nome_loja'])}/{nome_pdf}")
            except Exception as e: erros.append(f"PDF {ticket}: {e}")
            try: dropbox_rateio.subir(access, d["pdf"], f"{B}/{dropbox_rateio.NAO_LANCADOS}/{nome_pdf}")
            except Exception as e: erros.append(f"não lançados {ticket}: {e}")

    # 3) sobe controle atualizado
    if os.path.exists(ctrl):
        try: dropbox_rateio.subir(access, ctrl, ctrl_dbx, overwrite=True)
        except Exception as e: erros.append(f"subir controle: {e}")

    # 4) MOVE cada nota p/ NOTAS INCLUIDAS ORCAMENTO/<classe> (só dos tickets que viraram orçamento)
    tickets_ok = {d["ticket"] for d in feitos}
    movidas = 0
    incl = "NOTAS INCLUIDAS ORCAMENTO"
    for ticket, itens_n in notas_por_ticket.items():
        if ticket not in tickets_ok: continue
        for it in itens_n:
            classe = (it.get("aba") or "SEM CLASSIFICACAO").upper()
            destino_pasta = f"{B}/{incl}/{classe}"
            dropbox_rateio.criar_pasta(access, destino_pasta)
            try:
                dropbox_rateio.mover(access, f"{pasta_notas}/{it['arq']}", f"{destino_pasta}/{it['arq']}")
                movidas += 1
            except Exception as e:
                erros.append(f"mover nota {it['nota']}: {e}")

    # 5) remove da planilha de pendentes só as notas que viraram orçamento; sobe a planilha
    notas_feitas = [it["nota"] for t in tickets_ok for it in notas_por_ticket.get(t, [])]
    for n in notas_feitas:
        subprocess.run(["python3", PENDENTES, "remove", "--xlsx", xlsx_local, "--nota", str(n)],
                       capture_output=True)
    if notas_feitas and os.path.exists(xlsx_local):
        try: dropbox_rateio.subir(access, xlsx_local, xlsx_dbx, overwrite=True)
        except Exception as e: erros.append(f"subir pendentes: {e}")

    # mensagem
    msg = [f"✅ {len(feitos)} orçamento(s) corrigido(s) gerado(s) e enviado(s) ao Dropbox."]
    for d in feitos: msg.append(f"• Ticket {d['ticket']} — {d['loja']} — {d['total']}" + ("" if d['pdf_ok'] else "  (⚠️ PDF falhou)"))
    msg.append(f"📁 {movidas} nota(s) movida(s) p/ NOTAS INCLUIDAS ORCAMENTO.")
    msg.append(f"🧾 {len(notas_feitas)} linha(s) removida(s) da planilha de pendentes.")
    if faltando:
        msg.append("⚠️ Ficaram de fora: " + "; ".join(faltando[:8]))
    if erros:
        msg.append("⚠️ Erros: " + "; ".join(erros[:6]))
    return (ctrl if os.path.exists(ctrl) else None), "\n".join(msg)

# ---------- FrotaHub: auth por token do Supabase + permissão por rotina ----------
def _oc_pdf_texto(pdf_bytes):
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(pdf_bytes); p = f.name
    try:
        out = subprocess.run(["pdftotext", "-layout", p, "-"], capture_output=True, text=True, timeout=60)
        return out.stdout or ""
    except Exception:
        return ""
    finally:
        try: os.unlink(p)
        except Exception: pass

def _valor_br(s):
    t = re.sub(r"[^\d,.\-]", "", str(s or ""))
    if not t: return None
    if "," in t: t = t.replace(".", "").replace(",", ".")
    try: return round(float(t), 2)
    except Exception: return None

FART_BASE = "03720882"
# rótulos que aparecem no bloco de dados (usados p/ não colar o endereço no nome)
_OC_LABELS = ("Nome:", "CNPJ:", "Telefone:", "Vendedor:", "E-mail:", "Endereço:", "Endereco:", "I.E.:")

def extrai_oc(txt):
    """Extrai os campos de uma O.C. do Obra Prima (mesma lógica do PCO manual):
    lê os BLOCOS 'DADOS DO FATURAMENTO' (tomador Fartura) e 'DADOS DO FORNECEDOR'
    (o fornecedor de verdade — Nome + CNPJ), não um CNPJ solto qualquer."""
    d = {"numero":None,"data_oc":None,"centro_custo":None,"cnpj_centro_custo":None,
         "fornecedor":None,"cnpj_fornecedor":None,"valor":None}
    # número da O.C.
    m = re.search(r"ORDEM DE COMPRA\s+(\d+)", txt, re.I)
    if not m: m = re.search(r"(?:ordem de compra|o\.?c\.?)\D{0,12}(\d{4,7})", txt, re.I)
    if m: d["numero"] = m.group(1).zfill(6)
    # data (preferir Criação/Data; senão a 1ª data do documento)
    m = re.search(r"(?:cria[çc][ãa]o|data)\D{0,12}(\d{2}/\d{2}/\d{4})", txt, re.I) or re.search(r"(\d{2}/\d{2}/\d{4})", txt)
    if m:
        p=m.group(1).split("/"); d["data_oc"]=f"{p[2]}-{p[1]}-{p[0]}"
    # centro de custo (delimitado por CNO:)
    m = re.search(r"OBRA\s*/?\s*CENTRO DE CUSTO\s*:?\s*(.+?)\s{2,}CNO:", txt, re.I)
    if not m: m = re.search(r"OBRA\s*/?\s*CENTRO DE CUSTO\s*:?\s*(.+)", txt, re.I)
    if m: d["centro_custo"] = re.sub(r"\s+"," ",m.group(1)).strip()[:120]
    # bloco FATURAMENTO (tomador Fartura) -> CNPJ do centro de custo
    ftb = re.search(r"DADOS DO FATURAMENTO(.*?)DADOS DO FORNECEDOR", txt, re.DOTALL|re.I)
    if ftb:
        mfc = re.search(r"CNPJ:\s*([\d./\-]+)", ftb.group(1))
        if mfc and mfc.group(1).strip(): d["cnpj_centro_custo"] = fmt_cnpj(mfc.group(1))
    # bloco FORNECEDOR -> Nome + CNPJ do fornecedor de verdade
    fb = re.search(r"DADOS DO FORNECEDOR(.*?)(?:OBRA\s*/?\s*CENTRO DE CUSTO|DADOS DO|ITENS|PRODUTOS)", txt, re.DOTALL|re.I) \
         or re.search(r"DADOS DO FORNECEDOR(.*)", txt, re.DOTALL|re.I)
    if fb:
        block = fb.group(1); lines = block.splitlines()
        for i, line in enumerate(lines):
            mn = re.search(r"Nome:\s+(.+?)(?:\s{2,}Endere[çc]o:.*)?$", line)
            if mn:
                nome = mn.group(1).strip()
                if i+1 < len(lines):
                    nxt = lines[i+1].strip()
                    if nxt and not any(lbl in lines[i+1] for lbl in _OC_LABELS):
                        nome = (nome + " " + re.sub(r"\s{2,}.*$","",nxt)).strip()
                d["fornecedor"] = re.sub(r"\s+"," ",nome).strip()[:160]
                break
        mc = re.search(r"CNPJ:\s*([\d./\-]*)", block)
        if mc and mc.group(1).strip(): d["cnpj_fornecedor"] = fmt_cnpj(mc.group(1))
    # total
    totais = re.findall(r"(?mi)^\s*Total\s+([\d.,]+)\s*$", txt)
    if totais: d["valor"] = _valor_br(totais[-1])
    else:
        m = re.search(r"(?:total geral|valor total|total)\s*[:\-]?\s*R?\$?\s*([\d\.\,]+)", txt, re.I)
        if m: d["valor"] = _valor_br(m.group(1))
    # validação PCO
    fat_ok = bool(d["cnpj_centro_custo"]) and re.sub(r"\D","",d["cnpj_centro_custo"] or "").startswith(FART_BASE)
    d["valido"] = bool(d["cnpj_fornecedor"]) and fat_ok
    d["motivo"] = "" if d["valido"] else ("sem CNPJ do fornecedor" if not d["cnpj_fornecedor"] else "faturamento não é Fartura")
    return d

def _bearer(request):
    h = request.headers.get("authorization","") or request.headers.get("Authorization","")
    return h[7:].strip() if h.lower().startswith("bearer ") else ""

def _sb_json(url, key, tok=None, data=None, method="GET"):
    hdrs = {"apikey": key, "authorization": f"Bearer {tok or key}"}
    if data is not None: hdrs["content-type"]="application/json"
    req = urllib.request.Request(url, data=(json.dumps(data).encode() if data is not None else None),
                                 headers=hdrs, method=method)
    with urllib.request.urlopen(req, timeout=20) as r:
        b=r.read().decode()
        return json.loads(b) if b else None

def _sb_delete(url):
    req = urllib.request.Request(url, method="DELETE",
        headers={"apikey": SB_KEY, "authorization": f"Bearer {SB_KEY}", "prefer": "return=minimal"})
    try: urllib.request.urlopen(req, timeout=30)
    except urllib.error.HTTPError as e: print("sb delete erro:", e.read().decode()[:200], flush=True)

def auth_user(tok):
    if not tok: return None
    try: return _sb_json(f"{SB_URL}/auth/v1/user", SB_ANON, tok)
    except Exception: return None

def perfil_de(uid):
    q=urllib.parse.urlencode({"id":f"eq.{uid}",
        "select":"papel,nome,ativo,categoria_id,primeiro_acesso,must_change_pw,usuario,nome_completo,categorias(nivel)",
        "limit":"1"})
    try:
        d=_sb_json(f"{SB_URL}/rest/v1/perfis?{q}", SB_KEY)
        if not d: return None
        r=d[0]; cat=r.get("categorias")
        r["nivel"]=(cat.get("nivel") if isinstance(cat,dict) else None) or ("builder" if r.get("papel")=="builder" else ("gerente" if r.get("papel")=="gerente" else "comum"))
        return r
    except Exception: return None

def pode_rotina(perfil, rotina):
    nivel=(perfil or {}).get("nivel") or ("builder" if (perfil or {}).get("papel")=="builder" else "comum")
    if nivel in ("builder","gerente"): return True   # gerente também tem acesso pleno às rotinas
    cat=(perfil or {}).get("categoria_id") or (perfil or {}).get("papel")
    q=urllib.parse.urlencode({"categoria_id":f"eq.{cat}","rotina":f"eq.{rotina}","pode":"is.true","select":"rotina","limit":"1"})
    try:
        return bool(_sb_json(f"{SB_URL}/rest/v1/categoria_permissoes?{q}", SB_KEY))
    except Exception: return False

def exige(request, rotina):
    """Valida o token do FrotaHub e a permissão. Devolve (user, perfil) ou levanta HTTPException."""
    from fastapi import HTTPException
    u=auth_user(_bearer(request))
    if not u or not u.get("id"): raise HTTPException(401, "não autenticado")
    p=perfil_de(u["id"])
    if not p or p.get("ativo") is False: raise HTTPException(403, "usuário sem perfil ativo")
    if not pode_rotina(p, rotina): raise HTTPException(403, f"sem permissão para {rotina}")
    return u, p

def log_frotahub(uid, papel, rotina, acao, alvo="", detalhe=None):
    try:
        _sb_json(f"{SB_URL}/rest/v1/log_atividades", SB_KEY, data={
            "user_id":uid,"papel":papel,"rotina":rotina,"acao":acao,"alvo":alvo,
            "detalhe":detalhe or {}}, method="POST")
    except Exception: pass

# ---------- interface (FastAPI puro — sem Gradio) ----------
from fastapi import FastAPI, UploadFile, File, Request
from fastapi.responses import HTMLResponse, FileResponse, PlainTextResponse, JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title="Motor de Orçamentos — Frota Macedo")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
RESULTS = {}   # token -> caminho da planilha de controle p/ download

PAGINA = """<!doctype html><html lang="pt-br"><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Motor de Orçamentos — Frota Macedo</title>
<style>
 body{font-family:system-ui,Segoe UI,Arial,sans-serif;max-width:640px;margin:28px auto;padding:0 18px;color:#1e2733}
 h1{font-size:21px;color:#1F3864} .card{border:1px solid #C9CFDA;border-radius:12px;padding:20px;background:#fff}
 label{font-weight:bold;display:block;margin:14px 0 6px} input[type=file]{width:100%}
 button{margin-top:18px;background:#1F3864;color:#fff;border:0;border-radius:8px;padding:12px 20px;font-size:16px;cursor:pointer}
 button:disabled{background:#9aa4b4} small{color:#5A6472} #msg{white-space:pre-wrap;margin-top:16px}
 .ok{color:#1a7f37} .err{color:#b3261e} a.dl{display:inline-block;margin-top:14px;background:#1a7f37;color:#fff;padding:11px 18px;border-radius:8px;text-decoration:none}
</style></head><body>
<h1>Motor de Orçamentos — Frota Macedo</h1>
<div class="card">
 <p><small>Suba as notas/DAVs. O motor gera os orçamentos, envia tudo para o seu Dropbox
 (orçamentos, notas roteadas, planilhas) e deixa a <b>planilha de controle do mês</b> pronta para baixar e enviar ao cliente.</small></p>
 <form id="f">
  <label>Notas / DAVs (PDF ou imagem — pode selecionar várias)</label>
  <input type="file" name="notas" multiple accept=".pdf,.jpg,.jpeg,.png" required>
  <button id="b" type="submit">Gerar e enviar ao Dropbox</button>
 </form>
 <hr style="margin:22px 0;border:0;border-top:1px solid #E1E5EC">
 <p><small><b>Orçamentos corrigidos</b> — depois de corrigir o Ticket/Loja e marcar
  <b>RODAR = SIM</b> na planilha <i>NOTAS - TICKET NAO ASSOCIADO</i> (no Dropbox),
  clique abaixo. O motor gera os orçamentos dessas notas, roteia e remove as linhas.</small></p>
 <button id="bc" type="button" style="background:#1a7f37">Rodar corrigidos (TICKET NÃO ASSOCIADO)</button>
 <div id="msg"></div>
</div>
<script>
const f=document.getElementById('f'), b=document.getElementById('b'), msg=document.getElementById('msg');
const bc=document.getElementById('bc');
bc.addEventListener('click', async ()=>{
 bc.disabled=true; bc.textContent='Processando… (pode levar 1–2 min)'; msg.textContent='';
 try{
  const r=await fetch('corrigidos',{method:'POST'});
  const j=await r.json();
  if(j.erro){ msg.innerHTML='<span class="err">'+j.erro+'</span>'; }
  else{
   const warn=(j.status||'').trim().indexOf('⚠️')===0;
   let h='<span class="'+(warn?'err':'ok')+'">'+(j.status||'').replace(/</g,'&lt;')+'</span>';
   if(j.token){ h+='<br><a class="dl" href="baixar?t='+j.token+'">⬇ Baixar planilha de controle</a>'; }
   msg.innerHTML=h;
  }
 }catch(err){ msg.innerHTML='<span class="err">Falhou: '+err+'</span>'; }
 bc.disabled=false; bc.textContent='Rodar corrigidos (TICKET NÃO ASSOCIADO)';
});
f.addEventListener('submit', async (e)=>{
 e.preventDefault(); b.disabled=true; b.textContent='Processando… (pode levar 1–2 min)'; msg.textContent='';
 try{
  const r=await fetch('processar',{method:'POST',body:new FormData(f)});
  const j=await r.json();
  if(j.erro){ msg.innerHTML='<span class="err">'+j.erro+'</span>'; }
  else{
   const warn=(j.status||'').trim().indexOf('⚠️')===0;
   let h='<span class="'+(warn?'err':'ok')+'">'+(j.status||'').replace(/</g,'&lt;')+'</span>';
   if(j.token){ h+='<br><a class="dl" href="baixar?t='+j.token+'">⬇ Baixar planilha de controle</a>'; }
   msg.innerHTML=h;
  }
 }catch(err){ msg.innerHTML='<span class="err">Falhou: '+err+'</span>'; }
 b.disabled=false; b.textContent='Gerar e enviar ao Dropbox';
});
</script></body></html>"""

@app.get("/", response_class=HTMLResponse)
def home():
    return PAGINA

@app.post("/processar")
async def processar_endpoint(request: Request):
    form = await request.form()
    tmp = tempfile.mkdtemp(prefix="up_")
    caminhos = []
    for uf in form.getlist("notas"):
        if not getattr(uf, "filename", ""): continue
        dest = os.path.join(tmp, os.path.basename(uf.filename))
        with open(dest, "wb") as w: w.write(await uf.read())
        caminhos.append(dest)
    if not caminhos:
        return {"erro": "Selecione ao menos uma nota."}
    try:
        path, status = processar(caminhos)
    except Exception as e:
        return {"erro": f"Erro ao processar: {e}"}
    token = None
    if path and os.path.exists(path):
        import secrets as _s
        token = _s.token_urlsafe(10)
        RESULTS[token] = path
    return {"status": status, "token": token}

@app.post("/corrigidos")
def corrigidos_endpoint():
    try:
        path, status = processar_corrigidos()
    except Exception as e:
        return {"erro": f"Erro ao processar corrigidos: {e}"}
    token = None
    if path and os.path.exists(path):
        import secrets as _s
        token = _s.token_urlsafe(10)
        RESULTS[token] = path
    return {"status": status, "token": token}

@app.get("/baixar")
def baixar(t: str):
    path = RESULTS.get(t)
    if not path or not os.path.exists(path):
        return PlainTextResponse("Arquivo expirado — gere novamente.", status_code=404)
    return FileResponse(path, filename=os.path.basename(path),
                        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# ============ API do FrotaHub (protegida por token do Supabase) ============
@app.get("/api/ping")
def api_ping(): return {"ok": True, "motor": "frotahub", "rev": 116}

@app.get("/api/me")
def api_me(request: Request):
    from fastapi import HTTPException
    u=auth_user(_bearer(request))
    if not u: raise HTTPException(401,"não autenticado")
    p=perfil_de(u["id"]) or {}
    return {"email":u.get("email"),"nome":p.get("nome"),"papel":p.get("papel")}

PCO_ENVIAR = PCO_BASE + "/0 - ENVIAR (ADICIONAR AQUI)"
PCO_RESIDUAL = PCO_BASE + "/100 - ARQUIVOS RESIDUAIS (NÃO MEXER)"

@app.post("/pco/listar")
def pco_listar(request: Request):
    """CONFERIR_LISTA_PCO — lê a pasta 0-ENVIAR, extrai os dados de cada O.C. e
    devolve a lista (filtra as que já têm PCO enviado no banco)."""
    from fastapi import HTTPException
    u,p = exige(request, "CONFERIR_LISTA_PCO")
    if not dropbox_rateio.ativo(): raise HTTPException(500,"Dropbox não configurado")
    access = dropbox_rateio.obter_token()
    arqs = [a for a in dropbox_rateio.listar(access, PCO_ENVIAR) if a.lower().endswith(".pdf")]
    # números de O.C. já enviadas (não mostrar de novo)
    try:
        env = _sb_json(f"{SB_URL}/rest/v1/ocs?pco_status=eq.enviado&select=numero", SB_KEY) or []
        ja = {str(r["numero"]) for r in env}
    except Exception: ja=set()
    itens=[]
    for nome in sorted(arqs):
        try:
            pdf = dropbox_rateio.baixar(access, f"{PCO_ENVIAR}/{nome}")
            d = extrai_oc(_oc_pdf_texto(pdf)) if pdf else {}
        except Exception as e:
            d = {"_erro": str(e)[:120]}
        if d.get("numero") and d["numero"] in ja:   # já enviada
            continue
        d["arquivo"]=nome
        itens.append(d)
    return {"pasta": PCO_ENVIAR, "total": len(itens), "itens": itens}

@app.post("/pco/arquivos")
def pco_arquivos(request: Request):
    """Lista rápida só com os NOMES dos PDFs da pasta 0-ENVIAR (p/ barra de progresso)."""
    from fastapi import HTTPException
    exige(request, "CONFERIR_LISTA_PCO")
    if not dropbox_rateio.ativo(): raise HTTPException(500,"Dropbox não configurado")
    access = dropbox_rateio.obter_token()
    arqs = [a for a in dropbox_rateio.listar(access, PCO_ENVIAR) if a.lower().endswith(".pdf")]
    return {"arquivos": sorted(arqs)}

@app.get("/pco/oc")
def pco_oc(request: Request, arquivo: str):
    """Lê UMA O.C. (extrai os campos) — usado no carregamento com progresso."""
    from fastapi import HTTPException
    exige(request, "CONFERIR_LISTA_PCO")
    access = dropbox_rateio.obter_token()
    nome = os.path.basename(arquivo)
    pdf = dropbox_rateio.baixar(access, f"{PCO_ENVIAR}/{nome}")
    d = extrai_oc(_oc_pdf_texto(pdf)) if pdf else {}
    d["arquivo"] = nome
    d["ja_enviada"] = False
    if d.get("numero"):
        try:
            q = urllib.parse.urlencode({"numero": f"eq.{d['numero']}", "pco_status":"eq.enviado","select":"numero","limit":"1"})
            d["ja_enviada"] = bool(_sb_json(f"{SB_URL}/rest/v1/ocs?{q}", SB_KEY))
        except Exception: pass
    return d

@app.get("/pco/visualizar")
def pco_visualizar(request: Request, arquivo: str):
    from fastapi import HTTPException
    exige(request, "CONFERIR_LISTA_PCO")
    access = dropbox_rateio.obter_token()
    pdf = dropbox_rateio.baixar(access, f"{PCO_ENVIAR}/{os.path.basename(arquivo)}")
    if not pdf: raise HTTPException(404,"arquivo não encontrado")
    return Response(content=pdf, media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{os.path.basename(arquivo)}"'})

@app.post("/pco/excluir")
async def pco_excluir(request: Request):
    from fastapi import HTTPException
    u,p = exige(request, "CONFERIR_LISTA_PCO")
    body = await request.json()
    arq = os.path.basename(body.get("arquivo",""))
    if not arq: raise HTTPException(400,"arquivo?")
    access = dropbox_rateio.obter_token()
    try:
        dropbox_rateio.apagar(access, f"{PCO_ENVIAR}/{arq}")   # vai p/ lixeira do Dropbox (recuperável)
    except Exception as e:
        raise HTTPException(500, f"falha ao excluir: {e}")
    log_frotahub(u["id"], p["papel"], "CONFERIR_LISTA_PCO", "EXCLUIU_OC", arq)
    return {"ok": True}

@app.post("/pco/reenviar")
async def pco_reenviar(request: Request):
    """Libera uma O.C. já enviada para ser enviada de novo: APAGA o registro de envio
    (numero) no banco. O PDF continua em 0-ENVIAR e volta a aparecer na lista normal."""
    from fastapi import HTTPException
    u,p = exige(request, "ENVIAR_PCO")           # reabrir um envio exige permissão de envio
    body = await request.json()
    numero = str(body.get("numero","")).strip()
    if not numero: raise HTTPException(400,"numero?")
    req = urllib.request.Request(
        f'{SB_URL}/rest/v1/ocs?numero=eq.{urllib.parse.quote(numero)}&pco_status=eq.enviado',
        method="DELETE",
        headers={"apikey": SB_KEY, "authorization": f"Bearer {SB_KEY}", "prefer": "return=minimal"})
    try: urllib.request.urlopen(req, timeout=30)
    except urllib.error.HTTPError as e: raise HTTPException(500, "apagar registro: " + e.read().decode()[:150])
    log_frotahub(u["id"], p["papel"], "ENVIAR_PCO", "REENVIAR_LIBEROU", numero)
    return {"ok": True, "numero": numero}

# ---------------- OC_INVALIDA (O.C. bloqueadas) ----------------
def _sugerir_cnpj_cc(centro):
    """Sugere o CNPJ de faturamento certo (Fartura) para um centro de custo,
    olhando outra O.C. do mesmo centro que já tenha CNPJ válido."""
    if not centro: return None
    try:
        q = urllib.parse.urlencode({"centro_custo": f"eq.{centro}",
            "cnpj_centro_custo": f"like.{FART_BASE}*", "select": "cnpj_centro_custo", "limit": "1"})
        r = _sb_json(f"{SB_URL}/rest/v1/ocs?{q}", SB_KEY) or []
        return r[0]["cnpj_centro_custo"] if r else None
    except Exception: return None

def _bloq_row(numero):
    q = urllib.parse.urlencode({"numero": f"eq.{numero}", "pco_status": "eq.bloqueado", "limit": "1",
        "select": "numero,data_oc,centro_custo,arquivo_oc_path"})
    r = _sb_json(f"{SB_URL}/rest/v1/ocs?{q}", SB_KEY) or []
    return r[0] if r else None

def _pdf_carimbo_correcao(pdf_bytes, linhas):
    """Carimba um banner de correção no topo da 1ª página do PDF."""
    import io as _io
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas
    from reportlab.lib.colors import HexColor
    reader = PdfReader(_io.BytesIO(pdf_bytes))
    page0 = reader.pages[0]
    w = float(page0.mediabox.width); h = float(page0.mediabox.height)
    bh = 22 + 12 * len(linhas)
    buf = _io.BytesIO(); c = canvas.Canvas(buf, pagesize=(w, h))
    c.setFillColor(HexColor("#7A1517")); c.rect(0, h - bh, w, bh, fill=1, stroke=0)
    c.setFillColor(HexColor("#FFFFFF"))
    c.setFont("Helvetica-Bold", 9); c.drawString(14, h - 16, "CORREÇÃO — FROTA MACEDO ENGENHARIA")
    c.setFont("Helvetica", 8); y = h - 30
    for ln in linhas:
        c.drawString(14, y, (ln or "")[:150]); y -= 12
    c.save(); buf.seek(0)
    overlay = PdfReader(buf).pages[0]
    page0.merge_page(overlay)
    writer = PdfWriter(); writer.add_page(page0)
    for pg in reader.pages[1:]: writer.add_page(pg)
    out = _io.BytesIO(); writer.write(out); return out.getvalue()

def _oc_words(pdf_bytes):
    """Palavras da 1ª página com bounding box (via pdftotext -bbox)."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f: f.write(pdf_bytes); pth = f.name
    outp = pth + ".html"
    try:
        subprocess.run(["pdftotext", "-bbox", "-f", "1", "-l", "1", pth, outp], capture_output=True, timeout=30)
        h = open(outp, encoding="utf-8").read()
    except Exception: return [], 595.0, 842.0
    finally:
        for x in (pth, outp):
            try: os.unlink(x)
            except Exception: pass
    mp = re.search(r'<page width="([\d.]+)" height="([\d.]+)"', h)
    pageW, pageH = (float(mp.group(1)), float(mp.group(2))) if mp else (595.0, 842.0)
    words = [(float(a), float(b), float(c), float(d), w) for a, b, c, d, w in
             re.findall(r'<word xMin="([\d.]+)" yMin="([\d.]+)" xMax="([\d.]+)" yMax="([\d.]+)">([^<]*)</word>', h)]
    return words, pageW, pageH

def corrige_oc_inplace(pdf_bytes, campos):
    """Cobre o valor ERRADO e escreve o CERTO no mesmo lugar (sem banner) — deixa a O.C. com
    aparência de original. campos: cnpj_centro_custo, fornecedor, cnpj_fornecedor, endereco_fornecedor.
    Retorna bytes do PDF corrigido, ou None se o layout não bater (aí o chamador usa o original)."""
    try:
        return _corrige_oc_inplace(pdf_bytes, campos)
    except Exception as e:
        print("corrige_oc_inplace falhou (usa original):", str(e)[:160], flush=True)
        return None   # qualquer erro (PdfReader/pypdf/reportlab) -> cai no PDF original, nunca quebra

def _corrige_oc_inplace(pdf_bytes, campos):
    import io as _io
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas as _cv
    from reportlab.lib.colors import white as _wh, black as _bk
    try:
        words, pageW, pageH = _oc_words(pdf_bytes)
    except Exception:
        return None
    def yhdr(nm):
        for x0, y0, x1, y1, w in words:
            if w == nm: return y0
        return None
    y_fat = yhdr("FATURAMENTO"); y_forn = yhdr("FORNECEDOR"); y_obra = yhdr("OBRA/CENTRO") or 9999
    if y_fat is None or y_forn is None: return None
    def find_lab(label, ylo, yhi, xmin=0, xmax=250):
        c = [(x0, y0, x1, y1) for x0, y0, x1, y1, w in words if w == label and ylo < y0 < yhi and xmin <= x0 < xmax]
        return c[0] if c else None
    ovl = []
    def esq(lab, txt): ovl.append((lab[2] + 2, 350.0, lab, txt))
    def dire(lab, txt): ovl.append((lab[2] + 3, pageW - 28, lab, txt))
    if campos.get("cnpj_centro_custo"):
        lab = find_lab("CNPJ:", y_fat, y_forn, xmax=160)
        if lab: esq(lab, campos["cnpj_centro_custo"])
    if campos.get("fornecedor"):
        lab = find_lab("Nome:", y_forn, y_obra, xmax=160)
        if lab: esq(lab, campos["fornecedor"])
    if campos.get("cnpj_fornecedor"):
        lab = find_lab("CNPJ:", y_forn, y_obra, xmax=160)
        if lab: esq(lab, campos["cnpj_fornecedor"])
    if campos.get("endereco_fornecedor"):
        lab = find_lab("Endereço:", y_forn, y_obra, xmin=340, xmax=430)
        if lab: dire(lab, campos["endereco_fornecedor"])
    if not ovl: return None
    buf = _io.BytesIO(); c = _cv.Canvas(buf, pagesize=(pageW, pageH))
    for cx0, cx1, lab, txt in ovl:
        lx0, ly0, lx1, ly1 = lab
        c.setFillColor(_wh); c.rect(cx0, pageH - ly1 - 1.5, cx1 - cx0, (ly1 - ly0) + 4.5, stroke=0, fill=1)
        c.setFillColor(_bk); c.setFont("Helvetica", 8.5)
        c.drawString(cx0 + 1, pageH - ly1 + 1.2, str(txt))
    c.save(); buf.seek(0)
    ov = PdfReader(buf).pages[0]
    reader = PdfReader(_io.BytesIO(pdf_bytes)); p0 = reader.pages[0]; p0.merge_page(ov)
    wr = PdfWriter(); wr.add_page(p0)
    for pg in reader.pages[1:]: wr.add_page(pg)
    out = _io.BytesIO(); wr.write(out); return out.getvalue()

@app.get("/pco/bloqueadas")
def pco_bloqueadas(request: Request):
    from fastapi import HTTPException
    exige(request, "OC_INVALIDA")
    q = urllib.parse.urlencode({"pco_status": "eq.bloqueado", "order": "numero",
        "select": "numero,data_oc,centro_custo,cnpj_centro_custo,fornecedor,cnpj_fornecedor,valor,endereco_fornecedor,bloqueio_motivo,bloqueio_em,arquivo_oc_path"})
    try: rows = _sb_json(f"{SB_URL}/rest/v1/ocs?{q}", SB_KEY) or []
    except Exception as e: raise HTTPException(500, f"ler bloqueadas: {e}")
    for r in rows:
        cc = re.sub(r"\D", "", r.get("cnpj_centro_custo") or "")
        if not cc.startswith(FART_BASE):
            r["cnpj_cc_sugerido"] = _sugerir_cnpj_cc(r.get("centro_custo"))
    return {"itens": rows, "total": len(rows)}

@app.get("/pco/bloqueada_pdf")
def pco_bloqueada_pdf(request: Request, numero: str):
    from fastapi import HTTPException
    exige(request, "OC_INVALIDA")
    row = _bloq_row(numero)
    if not row or not row.get("arquivo_oc_path"): raise HTTPException(404, "O.C. não encontrada")
    access = dropbox_rateio.obter_token()
    pdf = dropbox_rateio.baixar(access, f"{PCO_BASE}/{row['arquivo_oc_path']}")
    if not pdf: raise HTTPException(404, "PDF não encontrado")
    return Response(content=pdf, media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{numero}.pdf"'})

def _oc_corrigir_aplicar(access, numero, campos, orig_pdf, destino_nome):
    """Valida os campos, carimba o PDF, sobe em 0-ENVIAR/destino_nome (overwrite) e grava o
    registro pendente+corrigida (que o envio passa a usar no lugar do PDF). Retorna (ok, motivo)."""
    cc_cnpj  = fmt_cnpj(campos.get("cnpj_centro_custo"))
    forn     = (campos.get("fornecedor") or "").strip()
    forn_cnpj= fmt_cnpj(campos.get("cnpj_fornecedor"))
    endereco = (campos.get("endereco_fornecedor") or "").strip()
    centro   = (campos.get("centro_custo") or "").strip()
    valor    = _valor_br(campos.get("valor")) if str(campos.get("valor") or "").strip() else None
    fat_ok  = bool(cc_cnpj) and re.sub(r"\D","",cc_cnpj).startswith(FART_BASE)
    forn_ok = len(re.sub(r"\D","",forn_cnpj or "")) == 14
    if not (fat_ok and forn_ok):
        m = []
        if not fat_ok:  m.append("CNPJ de faturamento precisa começar com 03.720.882")
        if not forn_ok: m.append("CNPJ do fornecedor inválido/ausente")
        return False, "; ".join(m)
    if not orig_pdf: return False, "PDF da O.C. não encontrado"
    # correção in-place: cobre o valor errado e escreve o certo no mesmo lugar (O.C. fica com cara de original)
    campos_ov = {"cnpj_centro_custo": cc_cnpj, "fornecedor": forn,
                 "cnpj_fornecedor": forn_cnpj, "endereco_fornecedor": endereco or None}
    corrig = corrige_oc_inplace(orig_pdf, campos_ov) or orig_pdf
    dropbox_rateio.criar_pasta(access, PCO_ENVIAR)
    dropbox_rateio.subir_bytes(access, corrig, f"{PCO_ENVIAR}/{destino_nome}", overwrite=True)
    sb_upsert_ocs([{ "numero": numero, "data_oc": campos.get("data_oc"),
        "centro_custo": centro or None, "cnpj_centro_custo": cc_cnpj,
        "fornecedor": forn, "cnpj_fornecedor": forn_cnpj, "endereco_fornecedor": endereco or None,
        "valor": valor, "pco_status": "pendente", "corrigida": True, "bloqueio_motivo": None,
        "arquivo_oc_path": f"0 - ENVIAR (ADICIONAR AQUI)/{destino_nome}" }])
    return True, ""

@app.post("/pco/corrigir")
async def pco_corrigir(request: Request):
    """Correção DEPOIS do envio: O.C. já bloqueadas no banco (pasta 2 - OCS BLOQUEADAS)."""
    from fastapi import HTTPException
    u, p = exige(request, "OC_INVALIDA")
    itens = (await request.json()).get("itens") or []
    access = dropbox_rateio.obter_token()
    res = []
    for it in itens:
        numero = str(it.get("numero", "")).strip()
        if not numero: continue
        row = _bloq_row(numero)
        if not row or not row.get("arquivo_oc_path"):
            res.append({"numero": numero, "ok": False, "motivo": "PDF da O.C. não encontrado"}); continue
        it2 = dict(it); it2.setdefault("data_oc", row.get("data_oc")); it2.setdefault("centro_custo", row.get("centro_custo"))
        try:
            orig = dropbox_rateio.baixar(access, f"{PCO_BASE}/{row['arquivo_oc_path']}")
            # apaga o registro bloqueado; o corrigido entra como pendente+corrigida
            _sb_delete(f"{SB_URL}/rest/v1/ocs?numero=eq.{urllib.parse.quote(numero)}&pco_status=eq.bloqueado")
            ok, motivo = _oc_corrigir_aplicar(access, numero, it2, orig, f"{numero}.pdf")
            if ok:
                try: dropbox_rateio.apagar(access, f"{PCO_BASE}/{row['arquivo_oc_path']}")   # tira da pasta de bloqueadas
                except Exception: pass
            res.append({"numero": numero, "ok": ok, "motivo": motivo})
        except Exception as e:
            res.append({"numero": numero, "ok": False, "motivo": str(e)[:150]})
    ok = sum(1 for r in res if r.get("ok"))
    log_frotahub(u["id"], p["papel"], "OC_INVALIDA", "CORRIGIU_OC", f"{ok}/{len(res)} corrigidas")
    return {"resultados": res, "corrigidas": ok, "total": len(res)}

@app.post("/pco/corrigir_previa")
async def pco_corrigir_previa(request: Request):
    """Correção ANTES do envio: O.C. ainda em 0-ENVIAR (não bloqueadas no banco)."""
    from fastapi import HTTPException
    u, p = exige(request, "ENVIAR_PCO")
    itens = (await request.json()).get("itens") or []
    access = dropbox_rateio.obter_token()
    res = []
    for it in itens:
        numero  = str(it.get("numero", "")).strip()
        arquivo = os.path.basename(it.get("arquivo", "") or "")
        if not (numero and arquivo):
            res.append({"numero": numero, "ok": False, "motivo": "arquivo/número ausente"}); continue
        try:
            orig = dropbox_rateio.baixar(access, f"{PCO_ENVIAR}/{arquivo}")
            ok, motivo = _oc_corrigir_aplicar(access, numero, it, orig, arquivo)   # sobrescreve o próprio arquivo
            res.append({"numero": numero, "ok": ok, "motivo": motivo})
        except Exception as e:
            res.append({"numero": numero, "ok": False, "motivo": str(e)[:150]})
    ok = sum(1 for r in res if r.get("ok"))
    log_frotahub(u["id"], p["papel"], "ENVIAR_PCO", "CORRIGIU_PREVIA", f"{ok}/{len(res)} corrigidas")
    return {"resultados": res, "corrigidas": ok, "total": len(res)}

@app.post("/pco/bloqueada_excluir")
async def pco_bloqueada_excluir(request: Request):
    from fastapi import HTTPException
    u, p = exige(request, "OC_INVALIDA")
    numero = str((await request.json()).get("numero", "")).strip()
    if not numero: raise HTTPException(400, "numero?")
    row = _bloq_row(numero)
    access = dropbox_rateio.obter_token()
    if row and row.get("arquivo_oc_path"):
        try: dropbox_rateio.apagar(access, f"{PCO_BASE}/{row['arquivo_oc_path']}")
        except Exception: pass
    _sb_delete(f"{SB_URL}/rest/v1/ocs?numero=eq.{urllib.parse.quote(numero)}&pco_status=eq.bloqueado")
    log_frotahub(u["id"], p["papel"], "OC_INVALIDA", "EXCLUIU_BLOQUEADA", numero)
    return {"ok": True, "numero": numero}

# ================= NOTAS FISCAIS — CONFERIR_NOTA (obra) =================
@app.get("/notas/buscar_oc")
def notas_buscar_oc(request: Request, q: str = ""):
    """Busca O.C. ENVIADAS por número, fornecedor ou centro de custo (a partir de 3 letras)."""
    from fastapi import HTTPException
    exige(request, "CONFERIR_NOTA")
    q = (q or "").strip()
    if len(q) < 3: return {"itens": []}
    val = f"(numero.ilike.*{q}*,fornecedor.ilike.*{q}*,centro_custo.ilike.*{q}*)"
    url = (f"{SB_URL}/rest/v1/v_ocs?pco_status=eq.enviado&or={urllib.parse.quote(val)}"
           f"&select=id,numero,data_oc,centro_custo,fornecedor,valor,qtd_notas,valor_notas&order=numero&limit=60")
    try: rows = _sb_json(url, SB_KEY) or []
    except Exception as e: raise HTTPException(500, f"buscar: {e}")
    for r in rows:
        try: r["saldo"] = round(float(r.get("valor") or 0) - float(r.get("valor_notas") or 0), 2)
        except Exception: r["saldo"] = None
    return {"itens": rows}

@app.get("/notas/da_oc")
def notas_da_oc(request: Request, oc_id: str):
    from fastapi import HTTPException
    exige(request, "CONFERIR_NOTA")
    url = (f"{SB_URL}/rest/v1/notas?oc_id=eq.{urllib.parse.quote(oc_id)}"
           f"&select=numero_nota,emissao,valor,tem_boleto,recebimento,divergencia,entregue&order=criado_em")
    try: return {"itens": _sb_json(url, SB_KEY) or []}
    except Exception as e: raise HTTPException(500, f"notas: {e}")

@app.post("/notas/conferir")
async def notas_conferir(request: Request):
    """Gera uma N.F. ligada à O.C. (parcial ou total). Total com valor divergente exige observação."""
    from fastapi import HTTPException
    u, p = exige(request, "CONFERIR_NOTA")
    b = await request.json()
    oc_id       = str(b.get("oc_id", "")).strip()
    numero_nota = (b.get("numero_nota") or "").strip()
    emissao     = (b.get("emissao") or "").strip() or None
    valor       = _valor_br(b.get("valor")) if str(b.get("valor") or "").strip() else None
    tem_boleto  = bool(b.get("tem_boleto"))
    recebimento = b.get("recebimento") if b.get("recebimento") in ("parcial", "total") else "total"
    divergencia = (b.get("divergencia") or "").strip() or None
    if not oc_id: raise HTTPException(400, "oc_id?")
    if not (numero_nota and emissao and valor is not None):
        raise HTTPException(400, "Preencha número da nota, data de emissão e valor.")
    oc = _sb_json(f"{SB_URL}/rest/v1/ocs?id=eq.{urllib.parse.quote(oc_id)}&select=numero,valor&limit=1", SB_KEY) or []
    if not oc: raise HTTPException(404, "O.C. não encontrada")
    oc_valor, oc_num = oc[0].get("valor"), oc[0].get("numero")
    if recebimento == "total" and oc_valor is not None and abs(float(valor) - float(oc_valor)) > 0.005 and not divergencia:
        raise HTTPException(400, "Recebimento TOTAL com valor diferente da O.C.: informe a observação.")
    nota = {"oc_id": oc_id, "numero_nota": numero_nota, "emissao": emissao, "valor": valor,
            "tem_boleto": tem_boleto, "recebimento": recebimento, "divergencia": divergencia,
            "conferida": True, "conferida_por": u["id"]}
    req = urllib.request.Request(f"{SB_URL}/rest/v1/notas",
        data=json.dumps(nota, ensure_ascii=False).encode(), method="POST",
        headers={"apikey": SB_KEY, "authorization": f"Bearer {SB_KEY}", "content-type": "application/json",
                 "prefer": "return=minimal"})
    try: urllib.request.urlopen(req, timeout=30)
    except urllib.error.HTTPError as e: raise HTTPException(500, "gravar nota: " + e.read().decode()[:200])
    log_frotahub(u["id"], p["papel"], "CONFERIR_NOTA", "CONFERIU_NOTA", f"{oc_num}/NF {numero_nota} ({recebimento})")
    vo = _sb_json(f"{SB_URL}/rest/v1/v_ocs?id=eq.{urllib.parse.quote(oc_id)}&select=valor,qtd_notas,valor_notas&limit=1", SB_KEY) or []
    resumo = vo[0] if vo else {}
    if resumo:
        try: resumo["saldo"] = round(float(resumo.get("valor") or 0) - float(resumo.get("valor_notas") or 0), 2)
        except Exception: pass
    return {"ok": True, "oc_numero": oc_num, "resumo": resumo}

# ================= NOTAS FISCAIS — ENTREGUES (visualizar) =================
@app.get("/notas/entregues")
def notas_entregues(request: Request):
    """Todas as notas já entregues pela obra (via física recebida no adm)."""
    u,p=exige(request,"PROCURAR_NOTA")
    rows=_sb_json(f"{SB_URL}/rest/v1/notas?entregue=eq.true&order=entregue_em.desc"
                  "&select=id,oc_id,numero_nota,emissao,valor,forma_pagamento,vencimento,entregue_em,tem_boleto,recebimento,divergencia&limit=5000",SB_KEY) or []
    ocids=list({str(r.get("oc_id")) for r in rows if r.get("oc_id")})
    ocmap={}
    if ocids:
        inlist=",".join(urllib.parse.quote(x) for x in ocids[:1500])
        for o in (_sb_json(f"{SB_URL}/rest/v1/ocs?id=in.({inlist})&select=id,numero,fornecedor,centro_custo",SB_KEY) or []):
            ocmap[str(o.get("id"))]=o
    itens=[]
    for r in rows:
        o=ocmap.get(str(r.get("oc_id")),{})
        itens.append({"id":r.get("id"),"oc":o.get("numero"),"fornecedor":o.get("fornecedor"),
            "centro_custo":o.get("centro_custo"),"numero_nota":r.get("numero_nota"),"emissao":r.get("emissao"),
            "valor":_numf(r.get("valor")),"forma_pagamento":r.get("forma_pagamento"),"vencimento":r.get("vencimento"),
            "entregue_em":(r.get("entregue_em") or "")[:10],"tem_boleto":bool(r.get("tem_boleto")),
            "recebimento":r.get("recebimento"),"divergencia":r.get("divergencia")})
    return {"itens":itens,"total":len(itens),"pode_excluir":(p.get("nivel") in ("builder","gerente"))}

@app.post("/notas/excluir")
async def notas_excluir(request: Request):
    """Exclui uma nota entregue — só gerente/builder."""
    from fastapi import HTTPException
    u,p=exige(request,"PROCURAR_NOTA")
    if p.get("nivel") not in ("builder","gerente"): raise HTTPException(403,"apenas gerente e builder")
    b=await request.json(); nid=str(b.get("id") or "").strip()
    if not nid: raise HTTPException(400,"id?")
    rq=urllib.request.Request(f"{SB_URL}/rest/v1/notas?id=eq.{urllib.parse.quote(nid)}",method="DELETE",
        headers={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}","prefer":"return=minimal"})
    try: urllib.request.urlopen(rq,timeout=20)
    except urllib.error.HTTPError as e: raise HTTPException(500,"excluir nota: "+e.read().decode()[:150])
    log_frotahub(u["id"],p.get("papel"),"PROCURAR_NOTA","EXCLUIU_NOTA",nid)
    return {"ok":True}

# ================= NOTAS FISCAIS — RECEBER_NOTA (adm) =================
@app.get("/notas/receber_lista")
def notas_receber_lista(request: Request, q: str = ""):
    """Notas já conferidas na obra e ainda NÃO entregues (via física pendente no adm)."""
    from fastapi import HTTPException
    exige(request, "RECEBER_NOTA")
    url = (f"{SB_URL}/rest/v1/notas?entregue=eq.false&conferida=eq.true&order=criado_em"
           f"&select=id,numero_nota,emissao,valor,tem_boleto,recebimento,divergencia,ocs(numero,fornecedor,centro_custo,valor)")
    try: rows = _sb_json(url, SB_KEY) or []
    except Exception as e: raise HTTPException(500, f"listar: {e}")
    ql = (q or "").strip().lower()
    out = []
    for r in rows:
        oc = r.get("ocs") or {}
        item = {"id": r.get("id"), "numero_nota": r.get("numero_nota"), "emissao": r.get("emissao"),
                "valor": r.get("valor"), "tem_boleto": r.get("tem_boleto"), "recebimento": r.get("recebimento"),
                "divergencia": r.get("divergencia"), "oc_numero": oc.get("numero"), "fornecedor": oc.get("fornecedor"),
                "centro_custo": oc.get("centro_custo"), "oc_valor": oc.get("valor")}
        if len(ql) >= 3:
            blob = " ".join(str(x or "") for x in (item["oc_numero"], item["fornecedor"], item["centro_custo"], item["numero_nota"])).lower()
            if ql not in blob: continue
        out.append(item)
    return {"itens": out[:300], "total": len(out)}

@app.post("/notas/receber")
async def notas_receber(request: Request):
    """Marca a nota como entregue (via física recebida) e grava forma de pagamento + vencimento."""
    from fastapi import HTTPException
    u, p = exige(request, "RECEBER_NOTA")
    b = await request.json()
    nota_id = str(b.get("nota_id", "")).strip()
    forma   = (b.get("forma_pagamento") or "").strip() or None
    venc    = (b.get("vencimento") or "").strip() or None
    if not nota_id: raise HTTPException(400, "nota_id?")
    if not forma:   raise HTTPException(400, "Informe a forma de pagamento.")
    patch = {"entregue": True, "entregue_em": _agora().isoformat(),
             "entregue_por": u["id"], "forma_pagamento": forma, "vencimento": venc}
    req = urllib.request.Request(f"{SB_URL}/rest/v1/notas?id=eq.{urllib.parse.quote(nota_id)}",
        data=json.dumps(patch, ensure_ascii=False).encode(), method="PATCH",
        headers={"apikey": SB_KEY, "authorization": f"Bearer {SB_KEY}", "content-type": "application/json",
                 "prefer": "return=minimal"})
    try: urllib.request.urlopen(req, timeout=30)
    except urllib.error.HTTPError as e: raise HTTPException(500, "receber: " + e.read().decode()[:200])
    log_frotahub(u["id"], p["papel"], "RECEBER_NOTA", "RECEBEU_NOTA", nota_id)
    return {"ok": True}

# ================= NOTAS FISCAIS — PROCURAR / PENDÊNCIAS / VENCIMENTOS =================
def _data_br(s):
    m = re.match(r"(\d{4})-(\d{2})-(\d{2})", str(s or ""))
    return f"{m.group(3)}/{m.group(2)}/{m.group(1)}" if m else (str(s or ""))

def _lista_pdf(titulo, colunas, linhas, subtitulo="", aligns=None):
    """PDF paisagem, formatado com a identidade visual da Frota (cabeçalho maroon)."""
    import io as _io
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=landscape(A4), leftMargin=1.1*cm, rightMargin=1.1*cm,
                            topMargin=1.1*cm, bottomMargin=1.1*cm, title=titulo)
    ss = getSampleStyleSheet()
    hh = ParagraphStyle('hh', parent=ss['Title'], textColor=colors.HexColor('#7A1517'), fontSize=16, spaceAfter=1, alignment=0)
    sb = ParagraphStyle('sb', parent=ss['Normal'], textColor=colors.HexColor('#666666'), fontSize=9, spaceAfter=10)
    el = [Paragraph("FROTA MACEDO ENGENHARIA", hh), Paragraph(titulo + (f" — {subtitulo}" if subtitulo else ""), sb)]
    t = Table([colunas] + linhas, repeatRows=1)
    stl = [('BACKGROUND',(0,0),(-1,0),colors.HexColor('#7A1517')),
           ('TEXTCOLOR',(0,0),(-1,0),colors.white),
           ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
           ('FONTSIZE',(0,0),(-1,-1),8.5),
           ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white, colors.HexColor('#f6f2f2')]),
           ('GRID',(0,0),(-1,-1),0.4,colors.HexColor('#dddddd')),
           ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
           ('TOPPADDING',(0,0),(-1,-1),4),('BOTTOMPADDING',(0,0),(-1,-1),4),
           ('LEFTPADDING',(0,0),(-1,-1),5),('RIGHTPADDING',(0,0),(-1,-1),5)]
    for c, a in (aligns or {}).items():
        stl.append(('ALIGN',(c,1),(c,-1),a))
    t.setStyle(TableStyle(stl)); el.append(t)
    doc.build(el); return buf.getvalue()

@app.get("/notas/procurar")
def notas_procurar(request: Request, q: str = ""):
    """Panorama dos 4 marcos (PCO solicitado / conferida / entregue / protocolada) por O.C."""
    from fastapi import HTTPException
    exige(request, "PROCURAR_NOTA")
    q = (q or "").strip()
    if len(q) < 3: return {"itens": []}
    SEL = "id,numero,data_oc,centro_custo,fornecedor,valor,pco_enviado_em,qtd_notas,valor_notas,fisica_ok,protocolo_ok,ciclo_fechado"
    val = f"(numero.ilike.*{q}*,fornecedor.ilike.*{q}*,centro_custo.ilike.*{q}*)"
    url = (f"{SB_URL}/rest/v1/v_ocs?pco_status=eq.enviado&or={urllib.parse.quote(val)}"
           f"&select={SEL}&order=numero&limit=40")
    try: rows = _sb_json(url, SB_KEY) or []
    except Exception as e: raise HTTPException(500, f"procurar: {e}")
    # também busca pelo NÚMERO DA NOTA -> traz as O.C. dessas notas
    have = {r["id"] for r in rows}
    try:
        nrows = _sb_json(f"{SB_URL}/rest/v1/notas?numero_nota=ilike.{urllib.parse.quote('*'+q+'*')}&select=oc_id", SB_KEY) or []
        extra = [n["oc_id"] for n in nrows if n.get("oc_id") and n["oc_id"] not in have]
        if extra:
            idlist = ",".join(f'"{i}"' for i in dict.fromkeys(extra))
            e = _sb_json(f"{SB_URL}/rest/v1/v_ocs?id=in.({idlist})&pco_status=eq.enviado&select={SEL}", SB_KEY) or []
            rows.extend(e)
    except Exception: pass
    for r in rows:
        nurl = (f"{SB_URL}/rest/v1/notas?oc_id=eq.{urllib.parse.quote(r['id'])}"
                f"&select=numero_nota,emissao,valor,tem_boleto,recebimento,divergencia,entregue,vencimento,forma_pagamento&order=criado_em")
        try: r["notas"] = _sb_json(nurl, SB_KEY) or []
        except Exception: r["notas"] = []
        try: r["saldo"] = round(float(r.get("valor") or 0) - float(r.get("valor_notas") or 0), 2)
        except Exception: r["saldo"] = None
    return {"itens": rows}

_MOT_LABEL = {"sem_nota": "Sem nota (aguardando obra)", "a_receber": "A receber (via física)",
              "a_protocolar": "A protocolar", "divergencia": "Divergência"}

def _pendencias_calc(q="", cats=None):
    rows = _sb_json(f"{SB_URL}/rest/v1/v_ocs?pco_status=eq.enviado&ciclo_fechado=eq.false"
                    f"&select=id,numero,data_oc,centro_custo,fornecedor,valor,qtd_notas,valor_notas,fisica_ok,protocolo_ok&order=data_oc", SB_KEY) or []
    dv = _sb_json(f"{SB_URL}/rest/v1/notas?divergencia=not.is.null&select=oc_id", SB_KEY) or []
    dvset = {r["oc_id"] for r in dv}
    ql = (q or "").strip().lower(); out = []
    for r in rows:
        mot = []
        if (r.get("qtd_notas") or 0) == 0: mot.append("sem_nota")
        elif not r.get("fisica_ok"): mot.append("a_receber")
        elif not r.get("protocolo_ok"): mot.append("a_protocolar")
        if r["id"] in dvset: mot.append("divergencia")
        if not mot: continue
        if cats and not (set(mot) & set(cats)): continue
        if len(ql) >= 3:
            blob = " ".join(str(x or "") for x in (r.get("numero"), r.get("fornecedor"), r.get("centro_custo"))).lower()
            if ql not in blob: continue
        r["motivos"] = mot; out.append(r)
    return out

@app.get("/notas/pendencias")
def notas_pendencias(request: Request, q: str = ""):
    from fastapi import HTTPException
    exige(request, "PENDENCIAS_NOTA")
    try: return {"itens": _pendencias_calc(q)}
    except Exception as e: raise HTTPException(500, f"pendencias: {e}")

@app.get("/notas/pendencias_pdf")
def notas_pendencias_pdf(request: Request, q: str = "", cats: str = ""):
    from fastapi import HTTPException
    exige(request, "PENDENCIAS_NOTA")
    catl = [c for c in (cats or "").split(",") if c]
    itens = _pendencias_calc(q, catl or None)
    linhas = [[r.get("numero") or "", _data_br(r.get("data_oc")), (r.get("centro_custo") or "")[:28],
               (r.get("fornecedor") or "")[:30], _fmt_reais(r.get("valor")),
               ", ".join(_MOT_LABEL.get(m, m) for m in r.get("motivos", []))] for r in itens]
    pdf = _lista_pdf("Pendências de notas", ["O.C.", "Data", "Centro de custo", "Fornecedor", "Valor", "Pendência"],
                     linhas, subtitulo=f"{len(itens)} pendência(s)", aligns={4: 'RIGHT'})
    return Response(content=pdf, media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="pendencias_notas.pdf"'})

def _vencimentos_calc(q="", filtro="todas"):
    rows = _sb_json(f"{SB_URL}/rest/v1/v_notas_app?vencimento=not.is.null"
                    f"&select=numero_nota,valor,vencimento,forma_pagamento,tem_boleto,oc_numero,fornecedor,centro_custo&order=vencimento", SB_KEY) or []
    hoje = _hoje(); ql = (q or "").strip().lower(); out = []
    for r in rows:
        try: dias = (datetime.date.fromisoformat(str(r["vencimento"])[:10]) - hoje).days
        except Exception: dias = None
        r["dias"] = dias
        if filtro == "vencidas" and not (dias is not None and dias < 0): continue
        if filtro == "a_vencer" and not (dias is not None and dias >= 0): continue
        if len(ql) >= 3:
            blob = " ".join(str(x or "") for x in (r.get("oc_numero"), r.get("fornecedor"), r.get("centro_custo"), r.get("numero_nota"))).lower()
            if ql not in blob: continue
        out.append(r)
    return out

@app.get("/notas/vencimentos")
def notas_vencimentos(request: Request, q: str = "", filtro: str = "todas"):
    from fastapi import HTTPException
    exige(request, "VENCIMENTOS_NOTA")
    try: return {"itens": _vencimentos_calc(q, filtro)}
    except Exception as e: raise HTTPException(500, f"vencimentos: {e}")

@app.get("/notas/vencimentos_pdf")
def notas_vencimentos_pdf(request: Request, q: str = "", filtro: str = "todas"):
    from fastapi import HTTPException
    exige(request, "VENCIMENTOS_NOTA")
    itens = _vencimentos_calc(q, filtro)
    def sit(d):
        if d is None: return ""
        if d < 0: return f"vencida há {-d} d"
        if d == 0: return "vence hoje"
        return f"em {d} d"
    linhas = [[r.get("oc_numero") or "", (r.get("fornecedor") or "")[:28], r.get("numero_nota") or "",
               _fmt_reais(r.get("valor")), r.get("forma_pagamento") or "", _data_br(r.get("vencimento")), sit(r.get("dias"))]
              for r in itens]
    titf = {"vencidas": "vencidas", "a_vencer": "a vencer"}.get(filtro, "todas")
    pdf = _lista_pdf("Vencimentos de notas", ["O.C.", "Fornecedor", "NF", "Valor", "Pagamento", "Vencimento", "Situação"],
                     linhas, subtitulo=f"{len(itens)} nota(s) · {titf}", aligns={3: 'RIGHT'})
    return Response(content=pdf, media_type="application/pdf",
        headers={"Content-Disposition": 'attachment; filename="vencimentos_notas.pdf"'})

# ================= NOTAS FISCAIS — PROTOCOLO (4º marco) =================
@app.get("/notas/protocolo_lista")
def notas_protocolo_lista(request: Request, q: str = ""):
    """O.C. totalmente entregues (via física recebida) e ainda NÃO protocoladas."""
    from fastapi import HTTPException
    exige(request, "GERAR_PROTOCOLO")
    url = (f"{SB_URL}/rest/v1/v_ocs?pco_status=eq.enviado&fisica_ok=eq.true&protocolo_ok=eq.false&qtd_notas=gt.0"
           f"&select=id,numero,data_oc,centro_custo,fornecedor,valor,qtd_notas,valor_notas&order=centro_custo,numero")
    try: rows = _sb_json(url, SB_KEY) or []
    except Exception as e: raise HTTPException(500, f"listar: {e}")
    ql = (q or "").strip().lower()
    if len(ql) >= 3:
        rows = [r for r in rows if ql in " ".join(str(x or "") for x in (r.get("numero"), r.get("fornecedor"), r.get("centro_custo"))).lower()]
    return {"itens": rows, "total": len(rows)}

@app.post("/notas/protocolar")
async def notas_protocolar(request: Request):
    """Marca a O.C. como protocolada (4º marco): cria o protocolo e fecha o ciclo."""
    from fastapi import HTTPException
    u, p = exige(request, "GERAR_PROTOCOLO")
    b = await request.json()
    oc_id = str(b.get("oc_id", "")).strip()
    numero_prot = (b.get("numero") or "").strip() or None
    if not oc_id: raise HTTPException(400, "oc_id?")
    oc = _sb_json(f"{SB_URL}/rest/v1/ocs?id=eq.{urllib.parse.quote(oc_id)}&select=numero,centro_custo&limit=1", SB_KEY) or []
    if not oc: raise HTTPException(404, "O.C. não encontrada")
    prot = {"data": _hoje().isoformat(), "centro_custo": oc[0].get("centro_custo"),
            "numero": numero_prot or oc[0].get("numero"), "criado_por": u["id"]}
    req = urllib.request.Request(f"{SB_URL}/rest/v1/protocolos",
        data=json.dumps(prot, ensure_ascii=False).encode(), method="POST",
        headers={"apikey": SB_KEY, "authorization": f"Bearer {SB_KEY}", "content-type": "application/json",
                 "prefer": "return=representation"})
    try: pid = (json.loads(urllib.request.urlopen(req, timeout=30).read().decode()) or [{}])[0].get("id")
    except urllib.error.HTTPError as e: raise HTTPException(500, "criar protocolo: " + e.read().decode()[:200])
    patch = {"protocolo": True, "protocolado_em": _agora().isoformat(), "protocolo_id": pid}
    req2 = urllib.request.Request(f"{SB_URL}/rest/v1/ocs?id=eq.{urllib.parse.quote(oc_id)}",
        data=json.dumps(patch, ensure_ascii=False).encode(), method="PATCH",
        headers={"apikey": SB_KEY, "authorization": f"Bearer {SB_KEY}", "content-type": "application/json",
                 "prefer": "return=minimal"})
    try: urllib.request.urlopen(req2, timeout=30)
    except urllib.error.HTTPError as e: raise HTTPException(500, "marcar protocolo: " + e.read().decode()[:200])
    log_frotahub(u["id"], p["papel"], "GERAR_PROTOCOLO", "PROTOCOLOU", f"{oc[0].get('numero')}")
    return {"ok": True}

# ---------------- ENVIAR_PCO ----------------
def _fmt_reais(v):
    try: return "R$ " + f"{float(v):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception: return "—"

def _pco_email_html(aprovadas, data_str):
    th = "style='background:#000;color:#fff;padding:7px 9px;text-align:left;font-size:13px'"
    td = "style='padding:7px 9px;border-bottom:1px solid #ddd;font-size:13px'"
    linhas = ""; total = 0.0
    for o in aprovadas:
        try: total += float(o.get("valor") or 0)
        except Exception: pass
        linhas += (f"<tr><td {td}>{o.get('centro_custo') or '—'}</td>"
                   f"<td {td}>{o.get('numero') or '—'}</td>"
                   f"<td {td}>{o.get('fornecedor') or '—'}</td>"
                   f"<td {td}>{o.get('cnpj_fornecedor') or '—'}</td>"
                   f"<td {td} align='right'>{_fmt_reais(o.get('valor'))}</td></tr>")
    total_row = (f"<tr><td {td} colspan='4'><b>Total ({len(aprovadas)} ordens)</b></td>"
                 f"<td {td} align='right'><b>{_fmt_reais(total)}</b></td></tr>")
    return f"""<div style="font-family:Arial,sans-serif;color:#1e2733;font-size:14px">
<p>Prezados, equipe Distribuidora de Alimentos Fartura S/A,</p>
<p style="margin-bottom:18px">Seguem as Ordens de Compra (PCO) para solicitação. Os PDFs correspondentes
seguem no arquivo em anexo.</p>
<table style="border-collapse:collapse;width:100%;max-width:720px">
<thead><tr><th {th}>CENTRO DE CUSTO</th><th {th}>OC</th><th {th}>FORNECEDOR</th><th {th}>CNPJ</th><th {th}>VALOR (R$)</th></tr></thead>
<tbody>{linhas}{total_row}</tbody></table>
<p style="margin-top:16px">Arquivo em anexo: <b>Pedidos_PCO_{data_str.replace('/','-')}.zip</b></p>
<p style="margin-top:24px">Atenciosamente,</p>
<p>{ASSINATURA}</p></div>"""

def _bloq_email_html(bloqueadas, data_str):
    itens = "".join(f"<li>O.C. <b>{o.get('numero') or o.get('arquivo')}</b> — {o.get('motivo') or 'bloqueada'}</li>" for o in bloqueadas)
    return f"""<div style="font-family:Arial,sans-serif;color:#1e2733;font-size:14px">
<p>As seguintes Ordens de Compra <b>NÃO</b> foram enviadas no PCO de {data_str} (precisam de correção):</p>
<ul>{itens}</ul>
<p style="margin-top:20px">Atenciosamente,</p><p>{ASSINATURA}</p></div>"""

def _enviar_brevo(assunto, html, para, cc, anexos):
    import base64
    payload = {"sender": {"email": PCO_FROM, "name": "Frota Macedo Engenharia"},
               "to": [{"email": e} for e in para], "subject": assunto, "htmlContent": html}
    if cc: payload["cc"] = [{"email": e} for e in cc]
    if anexos:
        payload["attachment"] = [{"name": n, "content": base64.b64encode(c).decode()} for n, c, _ in anexos]
    req = urllib.request.Request("https://api.brevo.com/v3/smtp/email",
        data=json.dumps(payload).encode(),
        headers={"api-key": BREVO_API_KEY, "content-type": "application/json", "accept": "application/json"})
    try:
        urllib.request.urlopen(req, timeout=60).read()
    except urllib.error.HTTPError as e:
        raise RuntimeError(f"Brevo {e.code}: {e.read().decode()[:300]}") from None

def enviar_email(assunto, html, para, cc=None, anexos=None):
    # Preferir API HTTP (Brevo) — o Render bloqueia SMTP.
    if BREVO_API_KEY:
        return _enviar_brevo(assunto, html, para, cc, anexos)
    import smtplib, ssl
    from email.message import EmailMessage
    msg = EmailMessage()
    msg["Subject"] = assunto; msg["From"] = PCO_FROM
    msg["To"] = ", ".join(para)
    if cc: msg["Cc"] = ", ".join(cc)
    msg.set_content("Seu cliente de e-mail não suporta HTML.")
    msg.add_alternative(html, subtype="html")
    for nome, conteudo, mime in (anexos or []):
        mt, st = mime.split("/", 1)
        msg.add_attachment(conteudo, maintype=mt, subtype=st, filename=nome)
    destinos = list(para) + list(cc or [])
    ctx = ssl.create_default_context()
    if SMTP_PORT == 465:
        with smtplib.SMTP_SSL(SMTP_HOST, SMTP_PORT, context=ctx, timeout=60) as s:
            s.login(SMTP_USER, SMTP_PASS); s.send_message(msg, to_addrs=destinos)
    else:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=60) as s:
            s.ehlo(); s.starttls(context=ctx); s.login(SMTP_USER, SMTP_PASS); s.send_message(msg, to_addrs=destinos)

def _oc_row(d, status, path, motivo=None):
    row = {"numero": d.get("numero"), "data_oc": d.get("data_oc"),
           "centro_custo": d.get("centro_custo"), "cnpj_centro_custo": d.get("cnpj_centro_custo"),
           "fornecedor": d.get("fornecedor"), "cnpj_fornecedor": d.get("cnpj_fornecedor"),
           "valor": d.get("valor"), "pco_status": status, "arquivo_oc_path": path}
    if status == "enviado": row["pco_enviado_em"] = _agora().isoformat()
    if motivo: row["bloqueio_motivo"] = motivo
    return row

def sb_upsert_ocs(rows):
    rows = [r for r in rows if r.get("numero")]
    if not rows: return
    req = urllib.request.Request(f"{SB_URL}/rest/v1/ocs?on_conflict=numero",
        data=json.dumps(rows, ensure_ascii=False).encode(), method="POST",
        headers={"apikey": SB_KEY, "authorization": f"Bearer {SB_KEY}", "content-type": "application/json",
                 "prefer": "resolution=merge-duplicates,return=minimal"})
    try: urllib.request.urlopen(req, timeout=30)
    except urllib.error.HTTPError as e: print("ocs upsert erro:", e.read().decode()[:300], flush=True)

def _oc_corrigida(numero):
    """Se existe registro de O.C. corrigida (pendente + corrigida=true), devolve os campos
    corretos do banco — usados no lugar do que foi relido do PDF (que ainda traz o erro)."""
    try:
        q = urllib.parse.urlencode({"numero": f"eq.{numero}", "pco_status": "eq.pendente",
            "corrigida": "eq.true", "limit": "1",
            "select": "centro_custo,cnpj_centro_custo,fornecedor,cnpj_fornecedor,valor,data_oc,endereco_fornecedor"})
        r = _sb_json(f"{SB_URL}/rest/v1/ocs?{q}", SB_KEY) or []
        return r[0] if r else None
    except Exception: return None

def _pco_coleta(access):
    """Baixa os PDFs de 0-ENVIAR, extrai e separa aprovadas/bloqueadas (ignora já enviadas)."""
    arqs = [a for a in dropbox_rateio.listar(access, PCO_ENVIAR) if a.lower().endswith(".pdf")]
    try:
        env = _sb_json(f"{SB_URL}/rest/v1/ocs?pco_status=eq.enviado&select=numero", SB_KEY) or []
        ja = {str(r["numero"]) for r in env}
    except Exception: ja = set()
    aprov, bloq = [], []
    for nome in sorted(arqs):
        pdf = dropbox_rateio.baixar(access, f"{PCO_ENVIAR}/{nome}")
        if not pdf: continue
        d = extrai_oc(_oc_pdf_texto(pdf)); d["arquivo"] = nome; d["_pdf"] = pdf
        num = d.get("numero")
        if num and num in ja: continue
        corr = _oc_corrigida(num) if num else None
        if corr:                       # foi corrigida na OC_INVALIDA -> usa os dados do banco
            for k in ("centro_custo","cnpj_centro_custo","fornecedor","cnpj_fornecedor","valor","data_oc"):
                if corr.get(k) not in (None, ""): d[k] = corr[k]
            d["valido"] = True; d["motivo"] = ""
        (aprov if d.get("valido") else bloq).append(d)
    return aprov, bloq

# ---- config chave/valor (lembra escolhas do usuário) ----
def _cfg_get(chave, default=None):
    try:
        r=_sb_json(f"{SB_URL}/rest/v1/app_config?chave=eq.{urllib.parse.quote(chave)}&select=valor&limit=1",SB_KEY) or []
        return r[0]["valor"] if r else default
    except Exception: return default
def _cfg_set(chave, valor):
    body=json.dumps([{"chave":chave,"valor":valor,"atualizado_em":_agora().isoformat()}],ensure_ascii=False).encode()
    rq=urllib.request.Request(f"{SB_URL}/rest/v1/app_config?on_conflict=chave",data=body,method="POST",
        headers={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}","content-type":"application/json","prefer":"resolution=merge-duplicates,return=minimal"})
    urllib.request.urlopen(rq,timeout=20)
def _emails_norm(lst):
    out=[]
    for e in (lst or []):
        e=str(e).strip()
        if e and re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", e) and e not in out: out.append(e)
    return out
def _pco_dest():
    """Destinatários salvos (última escolha) ou o padrão do ambiente."""
    d=_cfg_get("pco_destinatarios")
    if not d: return {"to":PCO_TO,"cc":PCO_CC}
    return {"to":(_emails_norm(d.get("to")) or PCO_TO),"cc":_emails_norm(d.get("cc"))}

@app.post("/pco/destinatarios_set")
async def pco_dest_set(request: Request):
    from fastapi import HTTPException
    u,p=exige(request,"ENVIAR_PCO")
    b=await request.json()
    to=_emails_norm(b.get("to")); cc=_emails_norm(b.get("cc"))
    if not to: raise HTTPException(400,"informe ao menos um e-mail em 'Para'")
    _cfg_set("pco_destinatarios",{"to":to,"cc":cc})
    log_frotahub(u["id"],p.get("papel"),"ENVIAR_PCO","DESTINATARIOS",f"{len(to)} para / {len(cc)} cc")
    return {"ok":True,"to":to,"cc":cc}

@app.post("/pco/enviar_previa")
def pco_enviar_previa(request: Request):
    from fastapi import HTTPException
    exige(request, "ENVIAR_PCO")
    if not dropbox_rateio.ativo(): raise HTTPException(500, "Dropbox não configurado")
    access = dropbox_rateio.obter_token()
    aprov, bloq = _pco_coleta(access)
    limpo = lambda d: {k: d.get(k) for k in ("arquivo","numero","data_oc","centro_custo","fornecedor","cnpj_fornecedor","valor","motivo")}
    def limpo_bl(d):
        r = {k: d.get(k) for k in ("arquivo","numero","data_oc","centro_custo","cnpj_centro_custo","fornecedor","cnpj_fornecedor","valor","motivo")}
        cc = re.sub(r"\D", "", r.get("cnpj_centro_custo") or "")
        if not cc.startswith(FART_BASE): r["cnpj_cc_sugerido"] = _sugerir_cnpj_cc(r.get("centro_custo"))
        return r
    dest=_pco_dest()
    return {"aprovadas": [limpo(d) for d in aprov], "bloqueadas": [limpo_bl(d) for d in bloq],
            "destinatarios": {"to": dest["to"], "cc": dest["cc"], "bloqueadas": PCO_BLOQ_TO},
            "smtp_ok": bool(BREVO_API_KEY or (SMTP_USER and SMTP_PASS))}

@app.post("/pco/enviar")
def pco_enviar(request: Request):
    import io, zipfile
    from fastapi import HTTPException
    u, p = exige(request, "ENVIAR_PCO")
    if not (BREVO_API_KEY or (SMTP_USER and SMTP_PASS)):
        raise HTTPException(500, "Envio não configurado — defina BREVO_API_KEY (recomendado) no Render.")
    access = dropbox_rateio.obter_token()
    hoje = _hoje().strftime("%d/%m/%Y")
    dest = _pco_dest()   # última escolha salva (ou padrão do ambiente)
    aprov, bloq = _pco_coleta(access)
    if not aprov and not bloq:
        return {"ok": True, "enviados": 0, "bloqueados": 0, "msg": "Nada novo para enviar."}
    erros = []
    # ---- APROVADAS: e-mail com zip + move + banco ----
    if aprov:
        buf = io.BytesIO(); z = zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED)
        for d in aprov: z.writestr(d["arquivo"], d["_pdf"])
        z.close(); zipbytes = buf.getvalue()
        # nome do zip com data+hora -> nunca sobrescreve um envio anterior do mesmo dia
        zipnome = f"Pedidos_PCO_{_agora().strftime('%d-%m-%Y_%H%M%S')}.zip"
        try:
            enviar_email(f"Ordens de Compra (PCO) - {hoje} - Frota Macedo Engenharia",
                         _pco_email_html(aprov, hoje), dest["to"], dest["cc"],
                         [(zipnome, zipbytes, "application/zip")])
        except Exception as e:
            raise HTTPException(500, f"falha ao enviar e-mail: {e}")
        # SÓ o ZIP fica arquivado em "1 - PEDIDOS FEITOS" (sem subpasta APAGAR).
        feitos = f"{PCO_BASE}/1 - PEDIDOS FEITOS"
        dropbox_rateio.criar_pasta(access, feitos)
        zip_rel = f"1 - PEDIDOS FEITOS/{zipnome}"
        zip_ok = True
        try: dropbox_rateio.subir_bytes(access, zipbytes, f"{feitos}/{zipnome}", overwrite=False)
        except Exception as e: erros.append(f"zip: {e}"); zip_ok = False
        rows = []
        for d in aprov:
            if zip_ok:
                # a via solta sai da pasta 0-ENVIAR (fica preservada dentro do ZIP);
                # vai para a lixeira do Dropbox (recuperável ~30 dias)
                try: dropbox_rateio.apagar(access, f"{PCO_ENVIAR}/{d['arquivo']}")
                except Exception as e: erros.append(f"remover {d['arquivo']}: {e}")
                path = f"{zip_rel}|{d['arquivo']}"        # <zip>|<arquivo interno>
            else:
                path = f"0 - ENVIAR (ADICIONAR AQUI)/{d['arquivo']}"
            rows.append(_oc_row(d, "enviado", path))
        sb_upsert_ocs(rows)
    # ---- BLOQUEADAS: move + banco + aviso ----
    if bloq:
        pasta_bloq = f"{PCO_BASE}/2 - OCS BLOQUEADAS"; dropbox_rateio.criar_pasta(access, pasta_bloq)
        rows = []
        for d in bloq:
            novo = f"OC_BLOC_{d.get('numero') or os.path.splitext(d['arquivo'])[0]}.pdf"
            try: dropbox_rateio.mover(access, f"{PCO_ENVIAR}/{d['arquivo']}", f"{pasta_bloq}/{novo}")
            except Exception as e: erros.append(f"mover bloq {d['arquivo']}: {e}")
            rows.append(_oc_row(d, "bloqueado", f"2 - OCS BLOQUEADAS/{novo}", motivo=d.get("motivo")))
        sb_upsert_ocs(rows)
        try: enviar_email(f"O.C. bloqueadas no PCO - {hoje}", _bloq_email_html(bloq, hoje), PCO_BLOQ_TO)
        except Exception as e: erros.append(f"e-mail bloqueadas: {e}")
    log_frotahub(u["id"], p["papel"], "ENVIAR_PCO", "ENVIOU_PCO",
                 f"{len(aprov)} enviadas / {len(bloq)} bloqueadas", {"to": dest["to"]})
    return {"ok": True, "enviados": len(aprov), "bloqueados": len(bloq),
            "to": dest["to"], "cc": dest["cc"], "erros": erros}

# ---------------- PCOS_ENVIADOS (planilha do banco) ----------------
def _pco_enviados_query(desde, ate):
    parts = ["pco_status=eq.enviado",
             "select=numero,data_oc,centro_custo,fornecedor,cnpj_fornecedor,valor,pco_enviado_em",
             "order=pco_enviado_em.desc"]
    if desde: parts.append(f"pco_enviado_em=gte.{desde}")
    if ate:   parts.append(f"pco_enviado_em=lte.{ate}T23:59:59")
    try: return _sb_json(f"{SB_URL}/rest/v1/ocs?" + "&".join(parts), SB_KEY) or []
    except Exception: return []

@app.get("/pco/enviados")
def pco_enviados(request: Request, desde: str = "", ate: str = ""):
    exige(request, "PCOS_ENVIADOS")
    itens = _pco_enviados_query(desde, ate)
    total = sum((x.get("valor") or 0) for x in itens)
    return {"itens": itens, "total": total}

@app.get("/pco/enviados_xlsx")
def pco_enviados_xlsx(request: Request, desde: str = "", ate: str = ""):
    exige(request, "PCOS_ENVIADOS")
    import io, openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    itens = _pco_enviados_query(desde, ate)
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "PCOs enviados"
    heads = ["O.C.", "Data O.C.", "Centro de custo", "Fornecedor", "CNPJ", "Valor (R$)", "Enviado em"]
    ws.append(heads)
    for c in ws[1]:
        c.font = Font(name="Arial", bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="A11F22"); c.alignment = Alignment(horizontal="center")
    def br(s): return "/".join(reversed(str(s)[:10].split("-"))) if s else ""
    for it in itens:
        ws.append([it.get("numero"), br(it.get("data_oc")), it.get("centro_custo"),
                   it.get("fornecedor"), it.get("cnpj_fornecedor"), it.get("valor"), br(it.get("pco_enviado_em"))])
    for i, w in enumerate([12, 12, 28, 34, 20, 14, 12], 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    buf = io.BytesIO(); wb.save(buf)
    return Response(content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": 'attachment; filename="PCOs_enviados.xlsx"'})

# ================= FERRAMENTAS DO BUILDER (Config) =================
# Migração antigo -> novo (1:1) e Desfazer por data. SÓ builder.
OLD_BASES = {
    "ADM":   ("/AUTOMACAO ADMINISTRATIVO", "/FROTAHUB/1 - ADMINISTRATIVO"),
    "MANUT": ("/AUTOMACAO MANUTENCAO",     "/FROTAHUB/2 - MANUTENCAO"),
}
_MIG_JUNK = ("ARQUIVOS RESIDUAIS", "SKILLS", "APP (", "APP(", "_RODAR LOCAL")
def _mig_pular(nome):
    u = (nome or "").upper()
    return nome.startswith("_") or any(j in u for j in _MIG_JUNK)

def _exige_builder(request):
    from fastapi import HTTPException
    u, p = exige(request, "CONFIG_BUILDER")   # sem linha em permissoes -> só builder passa
    if p.get("papel") != "builder":
        raise HTTPException(403, "apenas builder")
    return u, p

@app.post("/migrar/plano")
async def migrar_plano(request: Request):
    from fastapi import HTTPException
    _exige_builder(request)
    dep = (await request.json()).get("depto")
    if dep not in OLD_BASES: raise HTTPException(400, "depto inválido")
    de_base, para_base = OLD_BASES[dep]
    access = dropbox_rateio.obter_token()
    itens = []
    for e in dropbox_rateio.listar_entradas(access, de_base):
        if _mig_pular(e["name"]): continue
        if e["dir"] and e["name"].upper() == "ORCAMENTOS MONTADOS":
            for m in dropbox_rateio.listar_entradas(access, f"{de_base}/{e['name']}"):
                if _mig_pular(m["name"]): continue
                itens.append({"nome": f"{e['name']}/{m['name']}",
                              "de": f"{de_base}/{e['name']}/{m['name']}",
                              "para": f"{para_base}/{e['name']}/{m['name']}"})
        else:
            itens.append({"nome": e["name"], "de": f"{de_base}/{e['name']}", "para": f"{para_base}/{e['name']}"})
    return {"depto": dep, "de_base": de_base, "para_base": para_base, "itens": itens}

@app.post("/migrar/limpar")
async def migrar_limpar(request: Request):
    from fastapi import HTTPException
    u, p = _exige_builder(request)
    dep = (await request.json()).get("depto")
    if dep not in OLD_BASES: raise HTTPException(400, "depto inválido")
    _, para_base = OLD_BASES[dep]
    access = dropbox_rateio.obter_token()
    apagados = 0
    for e in dropbox_rateio.listar_entradas(access, para_base):
        try: dropbox_rateio.apagar(access, f"{para_base}/{e['name']}"); apagados += 1
        except Exception: pass
    log_frotahub(u["id"], p["papel"], "CONFIG_BUILDER", "MIGRAR_LIMPAR", dep, {"apagados": apagados})
    return {"ok": True, "apagados": apagados}

@app.post("/migrar/copiar")
async def migrar_copiar(request: Request):
    from fastapi import HTTPException
    _exige_builder(request)
    b = await request.json(); de = b.get("de"); para = b.get("para")
    if not de or not para: raise HTTPException(400, "de/para?")
    access = dropbox_rateio.obter_token()
    dropbox_rateio.criar_pasta(access, para.rsplit("/", 1)[0])
    try:
        dropbox_rateio.copiar(access, de, para)
    except Exception as e:
        return {"ok": False, "erro": str(e)[:200]}
    return {"ok": True}

@app.post("/desfazer")
async def desfazer(request: Request):
    from fastapi import HTTPException
    u, p = _exige_builder(request)
    b = await request.json(); dep = b.get("depto"); desde = (b.get("desde") or "").strip()
    if not desde: raise HTTPException(400, "informe a data")
    access = dropbox_rateio.obter_token()
    if dep != "ADM":
        return {"ok": False, "depto": dep,
                "msg": "Desfazer da Manutenção virá com o journal. Para resetar a Manutenção nos testes, use a Migração (recopiar do antigo)."}
    # ADM/PCO: reverte ocs criadas a partir da data -> volta o arquivo p/ 0-ENVIAR e apaga a linha
    q = urllib.parse.urlencode({"criado_em": f"gte.{desde}", "select": "numero,arquivo_oc_path", "order": "criado_em.desc"})
    try: rows = _sb_json(f"{SB_URL}/rest/v1/ocs?{q}", SB_KEY) or []
    except Exception as e: raise HTTPException(500, f"ler ocs: {e}")
    import io, zipfile
    voltas, erros, zip_cache = 0, [], {}
    dropbox_rateio.criar_pasta(access, PCO_ENVIAR)
    for r in rows:
        path = r.get("arquivo_oc_path")
        if not path: continue
        try:
            if "|" in path:                         # enviada: extrai a via de dentro do ZIP
                zrel, inner = path.split("|", 1)
                zbytes = zip_cache.get(zrel)
                if zbytes is None:
                    zbytes = dropbox_rateio.baixar(access, f"{PCO_BASE}/{zrel}")
                    zip_cache[zrel] = zbytes
                if not zbytes: erros.append(f"{r.get('numero')}: ZIP não encontrado"); continue
                with zipfile.ZipFile(io.BytesIO(zbytes)) as zf: pdf = zf.read(inner)
                dropbox_rateio.subir_bytes(access, pdf, f"{PCO_ENVIAR}/{inner}", overwrite=True)
                voltas += 1
            else:                                   # bloqueada/solta: apenas move de volta
                nome = path.rsplit("/", 1)[-1]
                dropbox_rateio.mover(access, f"{PCO_BASE}/{path}", f"{PCO_ENVIAR}/{nome}")
                voltas += 1
        except Exception as e: erros.append(f"{r.get('numero')}: {str(e)[:70]}")
    nums = [r["numero"] for r in rows if r.get("numero")]
    if nums:
        inlist = ",".join(f'"{n}"' for n in nums)
        req = urllib.request.Request(f'{SB_URL}/rest/v1/ocs?numero=in.({inlist})', method="DELETE",
            headers={"apikey": SB_KEY, "authorization": f"Bearer {SB_KEY}", "prefer": "return=minimal"})
        try: urllib.request.urlopen(req, timeout=30)
        except urllib.error.HTTPError as e: erros.append("apagar ocs: " + e.read().decode()[:150])
    log_frotahub(u["id"], p["papel"], "CONFIG_BUILDER", "DESFAZER_ADM", desde, {"ocs": len(rows)})
    return {"ok": True, "depto": "ADM", "ocs_revertidas": len(rows), "arquivos_voltaram": voltas, "erros": erros}

# ================= MANUTENÇÃO — CONFERIR_LISTA_ORCAMENTOS / GERAR_ORCAMENTOS =================
import base64 as _b64
from copy import deepcopy as _dcopy
from reportlab.lib.pagesizes import A4 as _A4
from reportlab.lib.units import mm as _mm
from reportlab.lib import colors as _col
from reportlab.lib.styles import ParagraphStyle as _PS
from reportlab.platypus import (SimpleDocTemplate as _SDT, Table as _T, TableStyle as _TS,
                                Paragraph as _P, Spacer as _SP, Image as _IMG)
from reportlab.lib.enums import TA_RIGHT as _TR, TA_CENTER as _TC

MANUT_BASE   = os.environ.get("MANUT_BASE", "/FROTAHUB/2 - MANUTENCAO").rstrip("/")
ORC_NOTAS    = MANUT_BASE + "/0 - NOTAS PARA ORCAMENTO (COLOCAR AQUI)"
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
GEMINI_MODEL   = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
GEMINI_RPD     = int(os.environ.get("GEMINI_RPD", "500"))   # cota diária do plano grátis (requisições/dia)
GROQ_TEXT_MODEL= os.environ.get("GROQ_TEXT_MODEL", "llama-3.3-70b-versatile")   # leitor de PDF (texto) — gratuito
GROQ_VIS_MODEL = os.environ.get("GROQ_VIS_MODEL", "")  # visão do Groq (vazio = conta sem modelo de visão -> usa OCR)
_ORC_GEMINI_OK = None   # memoriza o 1º modelo que funcionar no leitor de notas
ORC_EXTRAPOLA  = float(os.environ.get("ORC_EXTRAPOLA", "500"))   # nota (antes do +20%) acima disso -> extrapolado
RATEIO_TETO    = float(os.environ.get("RATEIO_TETO", "600"))     # NENHUM orçamento de rateio (JÁ com +20%) pode passar disto
_LOGO_ORC_B64  = "/9j/4AAQSkZJRgABAQEAuAC4AAD/2wBDAAMCAgMCAgMDAwMEAwMEBQgFBQQEBQoHBwYIDAoMDAsKCwsNDhIQDQ4RDgsLEBYQERMUFRUVDA8XGBYUGBIUFRT/2wBDAQMEBAUEBQkFBQkUDQsNFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBQUFBT/wgARCABkAMgDASIAAhEBAxEB/8QAHAABAAIDAQEBAAAAAAAAAAAAAAEHBQYIBAID/8QAGgEBAAIDAQAAAAAAAAAAAAAAAAMFAgQGAf/aAAwDAQACEAMQAAAB6oJISISIB8TUlYw2HUjkPCQ2HZev8heeLf6g2HjjNe4dxfWMydhyZIhIhIhIiYkAA+cJm/EcO4zeLCqu1pDaNsxKb3bDXevJegM7pt/bvKeHLROxVAAAARMSAAImDX6J6PriLd5V8/tvev62ntj8NpZY630f6JsOP+pic9cAAACJiQABEiEwaZx/3hzRBb1DnLO6D1rLGZ74+7Dl5DwAAACJiSMTluNzq70cjYg7Q/Hlv5Oo/rmjNHQf3xD1Gbn6Kl0Y6T/bmbOFtZXWaROk/ZzbsReevUvYpZ01HbgBExJFCX5BTOhdRjmXyXflDUtTsrHFO75vuVKNz2w5Yr/CXX4TRvJuWxlS/pve8FSY+7hqe2xIBEgAB8fQPgJgPsH5fqAAEgAA/8QAJhAAAQQCAgEEAwEBAAAAAAAABAIDBQYAAQcREhMVMEAUFiAhQf/aAAgBAQABBQL4/LWOlstaetEUKl/kOGYyAsolhR9G1XTUCQ7ybJuY7dpgjHZk97Fb9TNp6zesgpdyEkRSmzB/nlDm4sB89w8p5PjtgAwnY1SmH8HoEq5tvjN7eM8aCaxnj+GbwIBiNG+d5lD7dvqqq4bSBGwIMnks/Sl3yYdx6zypGPlEv746sClZ/wA+jJxrEqJfjm41jfTyW23FKHgZMjGqRLvJB48lGiG+/D6PeW+sIsQQEeQVKWOf3X8et8uXqjWdw36z8OK3IEnuHGua32CHIKdiiXjAPq8g1j0FcaQ3qI8U6+d2QGYX7kL4slsE4+U0I0HIjHJclREbbkhXl6MZ28Tpr0I8McEP1keq4822tS9ITq1RCnXJIVlWpAbaGiG3tOOJZQxZ4op7+uQUibvk+HGM0evstptLzDl7u9srOqS7bawDqBp9Wjg4cCTLDmbytDlN4+m9vQ9ckXJvkO4OvTEzY5L9pN/SoL8WaTFi3i0AgJg56mjQkeccu5HyFGhjAOOpN6SgP5lascXfr/Xip2GPpMxHmytOmR5ByqWS3GXKDelKyZAS6aC/xaZ7J7PKk0CaqMkQ0qvEA2GBoIftjFIO9v8AyrntouKng7NYI6wzsecDZ7O3N031hiU26VHgIVmvxfwTxC2Nyq1sxYBBCzTTNNyE06tiPDV5DRphy3TiCmjGSdsxEIQSpgYoj3pxS9SkqeQwQnfevj3rvOs6zaNbzrPHSddZ1m9f5nWdZ1/f/8QAJREAAQQBAgUFAAAAAAAAAAAAAQACAwQFEjEREyEiMCAyQlBR/9oACAEDAQE/AfRDRnsDi0KPCTv3KZgf1ysYZkcfZutvA1C6yOu0xps16T2gBcnIP+YCtzWo36JH+Kq8RyAu2T7Z1gRdSVZmuwM19FPO+y7U/wAeNsNjk7layugmIDij9b//xAAkEQABAwMBCQAAAAAAAAAAAAACAAEDESEwIhITFCAjMVBRYf/aAAgBAgEBPwHkKYWRTiuK+IZixbvUumtoPSjEHa2IxqNk0dqkhYHKiEGHtjmGoqOGurx3/8QAQxAAAQIDBQMGCggFBQAAAAAAAQIDAAQRBRITITEiQVEUI2GRobEGFTAyQFJxcoHRECAzNEJic8EkU2PC4UNUgpLw/9oACAEBAAY/AvJaRrFXHEJ94xzk+z/3B7oydL36aDDpl7yS0aKS4KH0IS7MuJl2l41NAnhHNssNdaoznLg/poEc7Ovq/wCZipJUek1jIU+huZRUp0cTxG+G3mlBTaxeSRv9Afmncm2gVGHZh37RxRV/iKjQ7uEc1KvOe62TGzIOgf1KJ742wyyPzLr3RV6eQn3UE95gFyafe92gjOXL36iyYSxLthllOiRoPQFNrSFJUKFJi+0CqQdOwr1D6p/aJm1ZgC6oGlR+EfMxRqUZQnco1OUZPob91sfvGc+97EGndFXJh5z3lkx4reXmnaZqdRvT8Iy9CclphF9pYzES1iy4utoQCsflGQH7xQ67jwiiG1uH8qSY2ZB/2qQR3x91CP1FiG3xMsS7iDeFKmhhIXS9TP0OqKJnGhzTn9p6DDdn3C3MFdxaTqnj1CGLPs5CLzaBUrFacBH3wt1/lpAhUhNuX3xtNuLOaxvHtHozlqol706GikU3/wCYfffrjLWag7jw+EX0CtdQIbflZZ/EQahQbORhl2YYVLPEZtK3ejKtWVRsn7wkbvzfOHLRdTrzbde0/tGWXl7rsw02rgtYBi/ylm5Wl7EFKwcF5DtNcNQNIK3nENNjVThAEEy801MU1w1hVOqClc0yhW+rgqICG5lpajolCwaxhB1BdH+neF4QoPXcMjO8aCG5eWASwhNEAHdAbK04lKhNc+qEpUpKVKyAJpWKrNANSYwhaUtieriCLjky0hXBSwIvB9ot+bevileEc24lfuKrF9aglI1JMYLVoyzjvqhwfX/jgvktxGJhedSh0hl2ycYyz07e/iaVqEkRZAsVbylG5ym9oP5g6U0iZlH3VtyUqVAITuSDTLpJ3xKWnZMw4jbu7R36/GvCFeESC7ymbKHS2ogpBXrlTpiS8InVupeabU6rMXcqjhwiX8I3U0aenFBauPEdR7ItIggpLQIPRUQqTm+bfkQK397ZFQeqEzyk0adad5PX1BkP7on5yXdoLDQnDFfOXWquzuiwLObdWzI2gjGcunNQ9XsMYHixjDprc2uvWJsWjJqm5YSyEobbQVUIApp0RYaJOXMvJTM6FFpQKdRQ14Q9aVjKcs+al039lZKVgaggx4PSD5LUrMy/Kn0pNL5FcuzthUuiUal6J2HWxRSTuNY59eM4w4pnEP4gNPrS1qBpCpFFy8SoVyB3RLy1ntIK0vYhTUJFKH5xZ9o2QhCZlLaA82lQTRYFD0EHfDds2YppNoL2phls5BZ86ldQeBhnx0Uyso3+FNPjQcT0wZCQbTfqi6kqoKA8YlLHYaBmvNd5wABNSdeqAlE+4t9KQoSZ+zvbxr2w5Zr7Q8YBvBG2KGhyz9kWcZDmZhUsJKd2gNmg698SszIMJVLSsgWEbQBKs6dfGErteVS/aLilOOKqdSeiMAuCWmpB9S5CZrWqSa0Px74wuQyCV/7nE7aRMWnIy8tNYrCWSp1y7mAKmntESpck2BNy82HA205kUgcT0xySbblrKklfaltd9ShwizzZbvJZ2z0BLDh0I4H/ANvjki2pGQSdlcy2sk030EMybFVJTmVHUnj5GUCXMJtayFm/d3eww8tlZxAjZWMzDzUwKYaE0UPNXWufdD/KJh1ltoJU220aX/nnlSHFtm6oFOYNPxCsIIXfy1v3u3fEoh8lxK0KVigUr0EbiO2EoYQFJwFqoo0zFKbobmXaqUGQtWWZNIcbmwoTLas60zBzGnV8IcQ4tQZK1BPBWmQ4EZ+2GED7NTayR0gpp3mGAwlWG1zj+nm6fM/D0AZafRl5b//EACgQAAIBAgUDBAMBAAAAAAAAAAABESExECBAQXFRYaEwgZHwscHR8f/aAAgBAQABPyH03akCJ20Uo+AJCWF8TZAD0M4aJ/D8OIeCJHk0yzQlOaN5aQW1NAyvm0N7lxyIQg6GoErCCYNBf7EAX1rQpf7ZtLiG1iSRgfIEh2yDkWeR0aXih7qdGAaV6v2LvkQQpBIHfsvmKPDihS0zwCq5baddFSXcz6mALb5A7+pqAZ7SYWefS2QBn26Ev2FeJWODNBto0dcGzNpvLdFop9IVXOx96CFkxAjjgIIUKIYWiJBeI0YHG9zwgnbRLQKxkahGim3AQvKTQ4DWopNFBeIoF3iPwruoYT7h0PtcUIkc9abOh4H8O16kQFJ7DAXHSnJgC5iVSuj4MD9xiSC9N8UDUhiAdZo8AP3JAI6ahDBLIjEj7Dg7uINnRgIzrvgEB8sCxiBP3UgCTqJAe6gwgpp+Y0jv2gkDFxEigYK5BAC7umaAHnul/AfIGhyU/AG0+YIHoihvOAqJ68yagGVLnRBgJ8plDbpXUBteftwACnrBOaChbSPgAI/lMgkxnAOQ8JAbUxOgHzdRBX9E6jtJ3cC778nByzFeofoWuGVO7AYx/wATEoIweqVzxYA3+G4BMkwYRED6P0oCoQksrVtQADnWhwAYj2Lq0fw9UyLiBAq3sECwIIEBWCBAgLN//9oADAMBAAIAAwAAABAQwwBiUOiAwwwxTzy1ZrOrTzzzxTzygMOi1TzzzxTzywgECnzzzzxQAQRQwBijSBTxQjBAwRhTjTBTxzzwACACAADzzz//xAAdEQACAQQDAAAAAAAAAAAAAAABETEAECAhMEBQ/9oACAEDAQE/EM8bK3eoiPBPdySIU7cbU0nC7vHECZfm/wD/xAAbEQACAQUAAAAAAAAAAAAAAAAAATEgMEBQYf/aAAgBAgEBPxCt2POyxzneUqIJsXF6/f/EACIQAAEDAwQDAQAAAAAAAAAAAAEQEUAAMDEgIUFxUFGBof/aAAgBAQABPxCw59I1CHKs5tk9/uenxQdnP6SK7zr7JmM/awHX8RrrjJ1+oax1z/5ado70kxFXP7psT4gYNcX9jJWOt72fTY/xTzxwbnXlnAlTL60ND7h7/s3Wvp+am4pd/gaiDz+eJ70fOYWUpwIUXIfe8AtQDPx1cN/kPjtCJAr5l/on/wCNuchwyJnJxR1/4GCtzPXTxFu8c06P9T1zyX+jTz1/Zugveb/Va7e2z8//ALT7ObLp4Cl6YF55WGV334nzFvgjyMs259/PbSWtv/bkH9x3+f8A1o1f8RiD+D2/78/+yGWOnizmI/d7kT6emr9/2F/0Deu/nUc7+zfT/wBEPGfUAi+nzv8A0dFvjaut/wD0QCu/ylf77d2m801ffNrbb8Fdv/svkxW/PO1r+KxbztHvDX+s9+8RJf8A/Wlu+dwWTt6rW/BIgf1Im9htWH0Ro8bxv+DNJx+Redu62xI8RAb/AIBuX/eF3ea7X9nFlh/6iT/cm47/ACM2LWzg/gk9tdFrhbdX/9k="
_MODELO_DOCX_B64 = "UEsDBAoAAAAAAC1g/VwAAAAAAAAAAAAAAAAFAAAAd29yZC9QSwMECgAAAAAALWD9XAAAAAAAAAAAAAAAAAsAAAB3b3JkL19yZWxzL1BLAwQKAAAACAAtYP1cgwyTCjUBAABHBQAAHAAAAHdvcmQvX3JlbHMvZG9jdW1lbnQueG1sLnJlbHOtlF9vgyAUxb+K4X2CTqsutX1ZlvR1cR8A4aJuAgbosn77sax/7NKZPvB4D3DOLyc3rLdfcow+wdhBqxolMUERKKb5oLoavTUvDyXabtavMFLnb9h+mGzknyhbo9656Qljy3qQ1MZ6AuVPhDaSOj+aDk+UfdAOcErICpu5B7r2jHa8RmbHExQ1hwnu8dZCDAyeNdtLUO5GBLbuMIL1jtR04Gr0O8feB+Hb8WnIeLWXLRjf44XgLC1BPIaEEFo7pd28hrO0BJGFhADF/zCclCWEPOgugHO+9/k2HJUlhFVIBKblz9EM4aQsIRShtwHM9SqASZbyy5D5g/S/wSVeAh8oJqQogFesTMqSizbNsqLgOW+rFamEyFie5UAKkcTvU/cfZRW2JeUa2o4wL+oonarCVz/i5htQSwMECgAAAAgALWD9XIoRD9eNDAAAhrQAABEAAAB3b3JkL2RvY3VtZW50LnhtbO1d23LiOBp+FRW7FzNVSYwNGEhNZtYBk85WgCyQ7N4qtgBt25ZHNiTpq6nai3mB3QfomYt+h73cvMk8yUo+cXI4J+GgVDfItvRL+v/vP+i3sX746cm2wBBRDxPnIiOfZTMAOQYxsdO7yNx1aqelDPB86JjQIg66yDwjL/PTjz88npvEGNjI8YFtnF/3HELhg8WuP8p58CgXwKMr5zOAEXe880fXuMj0fd89lyTP6CMbemc2NijxSNc/M4gtkW4XG0h6JNSUlKycDUouJQbyPDaSCnSG0IvJ2bPUiIscdrFLqA19dkh7kg3p54F7yqi70McP2ML+M6OdVWMy5CIzoM55ROI0GRBvch4OKPqKW9Bl+g2bVCPuBD1KFFlsDMTx+tgdTWNdauxiPyYynDeJoW2NRCDnN5NBlcJH9jUiuMzwzbCRbYUjn09Rzi4hEU4iabHMECb7jEdiQ+yMOl6LNWPMlQurEVCmCbi9zYRzRcnAHVHDm1G7dj4ntLjOr0ArEvL41LzNBtPuQzfRQONpOWIR7ji9vGT0IfXR04iGvDKRglSWSrOElDUIsQkq8iyp3MqkVImPaobQklieIsRGNUNpSVBPU0qZnLoeJWWWUnE9SrlZSqX1KM3AiRmSz2uQwiMdg3bOXJlCUbKJiazcyBjKqoGWVI9Y10qRskrGaD6cDl5yPDEdNaGDx8ez3mDGCHimb/ZXoqLEtlnibaEP+9Drj1NczZwxfY3JPduMRzzweSDmM//2H6zo65ZGhb8D9vXssj7MJ5hhB8xDlYvZbEaKKlwyYizWCo6IyyoMoXWRcVhkxWsbxCIsxqgFf/yE9+UiEzW2UNdfpf4D8X1ir9KC4l5/pS6w42ETfVq9yf3yTaRJtkmT/L6i2OTFHvuuECtkuFyIGT5xulBOPa3ksmMdxQT9sAcj/Iz6Mw5IetLUfIaahXtO3N5gsSaiSb1w/m7wEZY9FxrMELH6sMtqcqKs/IBY4IVGPUSVg4/IdLGiy1BgYQcBE3t+J2jKS5dJ6SYptXgpaIKefL7W4P5flvNMZqyG8cykWkzEzSp1u8jw9bCqH1AJh/YQfFoxs1y+eLmlgJs5OQMcaLNBsx6RZ1Be8LHPlzJRXaMxvKLQ7WOjRllNPifIMDQ6c0OMz15sy9eIisNY1CGVPnR6SPNcNgk+sICJ8/vftNcxUlVmLMGAzoaOi0m52PAHFDFqrHTuJsNipY2pOcNbHACQHzBWRGLLzopNGtUJW0A+gFA4s8wdnaKUPPYRNL2Y55NUpJlRPFjYrWHL4j3wMqDnyH5AbFT02mTLZYOtl30UaaHEa3nUaLF+w7JPkW/0ebHLiETnpbEL0mQn/MhjqgQeHuvM5V9k4MAngeSeutTm38xngaeAJ88RxCFXmHnaIo1au9TzrxCzNrzA5sAGFFCHwxsvGlpcJRqb5yaMYf+DGmMgGj8OERzqe2ASEjMgBXZBCsyKNDK1wuBuy+DSqFwjju/xdp6BmXJXoIUfKA4G600cIuj5mofhxMm+xjz26EzIq/Cz4gXfwZTjeci1XEnNh9W8L/FZJTlT8SbPSck4fW4zgmmy+bgUeYgOmQ2otZodDdS1il5tAr1xpTc+aa1rDdx0qlrAq0kk7S7TJthU0NR8UZlmk1yYZVN4biGbdKd3Bj4h7LM+TASaFh4iTCHQrIfBzwNE2f8ToJQL4I9f/g0q2ISskkk8UBs4BibOy1eK2RG/WCPUhxb6AqWKLhicMPi7UuF7oMgl+VRmCAeMVb8B0KXEhzaraRLEPInDFoYY/qXHPKvF1xJRrUrj9q9AKZ7l1NyZouQkZo3l02I2hbnCDm7DDj6e/9OI6QWD2zkYqikwVJeCoVKWskW+zFW3oZuHzKkmffkGec6cAOd//wWyIucLc3ROCtd8UrKsH2ele2nSWeXgyUlrHO2R84vQLifmL2o9IxBZnRCJOmEvY5GKLMNuZhlihq+fN1iAH2VWIsu1WAmlSopUlmkybca9vsku8wXNePWofbBgSUy8hSCNoGrUIR1xZRrTsjKJypnr6jQKX6UQz+8VElIylv0OvyNoT1nRXHbWiuayy1nR1suvWl1vdJqgMc+K7i63JvhTVasVXd2il2np99ftl381gbyma5l1CsoHMmuMJ6UUnpSW4kkqPIQbOxo3ppf1Sk1exY2ltljgxlLbzHdjKU3mubFR9c3cWHlTL1Y+LieWnkPawCA1L1vaOQATdul9/c7Wp1TXGncdvVHRmqB619DaAPxpoxXOorWJcEPCDb2lG8oXU2/O5lLP5jf2WXspvElfNWbjlwR1fjGoZ1xmpVypVbVVfHlqiwW+PLXNfF+e2mQK3DONYgfOm+RTwb24yQbgFqmCN0sVFBbEWMqiGKuwKMZSDirGSk8UyMUUb1tcyttWtWqzDapNcNvS2x120FoYiWyoAvtkmqZVYAGaS5uiufRWaM5/JIK3AtQGsdH5O68FtjLwGr/ZCOrB3UagJ7cb17wTvv+C5LdV91KQS90KPnjx1fgDYPxZBRf2wvt0eynMS2Ihn4BcFpgYescoyOC5QSbHLvQHdI8lOffW+qs5FCm5dpxrz12Nfd8qhXTUwhaJBpFoWCgXkWiYuzQTiYYtJRo6zbpIM4g0w45GxXubZqgjakATO33igfbL7wTcDPCX4Mnw6sA5zhXO3mYasrmzopI9K5UUnmlQTnPlYxSf7jB7iF6+7eeyVBueRdk+hCnhP+hA1CdMNb+eAJn9gVNwC11sDI5RtOEvWSTd86G5n+JNfm7D5Jj6g5slMg+rPdmhZA9wWb7oaZaPnPOc6PiV54LKKawpLxcdX7crrev6dUN7+ZU/lspj5euO3min8mx7T8MsswKeGzSn1F8QMqe0mB8wv7oo37l1fPrDMHk59bScfjr9bQf8EbE5p1MTBzQZ7CfErC2N6sWn11tWVQtVTVdWWValtliAkdQ281GS0mQX8grZjfMK2am12OMb//Rugt6HW9qt5yGYVa1vnh8WenEUenGQGlDVeawRxBlCEYQiCAcxqR5/u9ManTOhGUIzhGZMasZd47oqFEMoxuqK8f4vqXhPvbjXbpotwLRDuA2hHUI7prSj0+xoN299t3+fleTo8P4RYVLcdxdaHsrE8E85uxXMz3mrhID4oUL8UMHcubtsAr19BW7vKyCfrddB5/qqNeeep4D4oUL8yKz4SXYLb3sUKBco32WUi7thRw7xD1h+vi/CW38G5ZN8SeBc4Py4cS5yLeMJyVquVqiVRELyGIMeRTgDoRdHndW5vKt80pK0TiH7jyCzw49vrkSCR6iIcB2vuQ6RFRKqIVQj/QEhoRhCMcTCPGVhnj9RRQJKKIdQjvSsVW7xy7uOOGt1dKA//FApJ5zB8UH8UMF8c3cfppLa1/XbG70NCslTQpJW1xoV8UOzI0T7cRl0kRYSKD98lIsMz7Fj/BjWo7kTVdhygfODx3nxRFm88Vx63oW/hqjtwoT9+Tg9dwS6MJ64rFZ0Wdc2T1yqmyqQuqoC7a2axIoxrQrrv5gu+IkiuNJb836oeJRmfw2o70F6ZEc2GTh2pAij+PZGcf1XUrL4QFFPsvmFAQL/WuL9paUDVPglXtm6EyjY4N0WkL8bwCc+tIDLSujJR45H3n1bZhx+bnNqQ8xWDAABD2EPUATZJwI/D6BPCeCrCTgkXtoLPR7f8v2r+7Rtwc6+h7Us9lP56LewiP1U3vi1Ks3Ltt6656+L/o+e9pborT48sU9WacXM1aLNVNRFQF64mYq6JpDzuxEwrQ/RP3757b3DhO2oFn35Fm6ICIYvXy1skiD0KYLvPOSj74MdI4FBKPMrxAMQuJD6mAITsn/hnorIxp738jtJDx+E1Hdx3E0PDHmwi5hIDWxDgB3DGiAb2JAJBrP4l4eKdIhfvpFAxo5PUQ/OeeXbauujwmHu77C9WHkv92SciZGXb3K/fJM1YuR8+mYFufQtDLYfUJe0ck7LxrNQd1R8KwYT67r6nXusYCJ2LWhqvqjMGAo1xVCoSxmKFbdEXzGtupd2YhJoe7DW2dUcurArx2tXNANhHwFmVCoW5pML9j/kWzoCCfD9zleI1Txk+GFnXUIYm1qoiyhyDDQSD+rCgeVnAD3H5kWGXptR8On22l9AtMdPOasGvOBzKOWCX7gQysfGxk2oTyH240ZMsiAAAJN3uAoIQMUgXQjlGIAyucgxnVzrBzv88B2Igi7CISeHvYEfwSHqqjGwO2wOwZFJDO7QOUXsoFvsG2ywuQRxMR8kPgDzOSiwJgO+Pvrx/1BLAwQKAAAACAAtYP1cmNPGwg0DAAATEQAADwAAAHdvcmQvc3R5bGVzLnhtbOVWW0/bMBT+K1HeIU2aFqgoiBUqkKYNcdGeXcdpLBw7sx1K+fWzEzu9pKGFBiZtT825+PP5zqU+p+cvKXGeEReY0aHrH3ZcB1HIIkynQ/fxYXxw7DpCAhoBwigaunMk3POz09lAyDlBwknh4GZKGQcToqwzP3Rmfs91FCoVgxQO3UTKbOB5AiYoBeKQZYgqY8x4CqQS+dRLAX/KswPI0gxIPMEEy7kXdDp9C8N3QWFxjCG6ZDBPEZXFeY8johAZFQnOhEWb7YI2YzzKOINICJWJlJR4KcC0gvHDGlCKIWeCxfJQkTERFVDquN8pvlKyAOi9DyCwADr9EYOXKAY5kUKL/JYb0UjFz5hRKZzZAAiI8dAdAYInHLtKA8WKiICQFwKDFWVyQcXSKe/s1DPQ3vqFWSWVXmvRFb2iEOU8U02SAQ6mHGSJvqQw3URD9wFLggpqFKTa+RkQq/W0egIEin5Sa/mhq0VKE0UvcpP+97goqbeUE/FqHXv90km8jsSqbolmEd6uFK4R0HPj11gYg+O3yQQywrj1Da6Owm89S8hqu0GdYqnbk2LQSDH4YorBhioGbVSx20ix+2kU/XF4eXRcoxhuoBi2QDFspBi2SREXAh4J742a7kml10il9wUNuWfw/cbg+1/Qah8N/l5yRqe10I26xbgnJVbRPx8N9jsW8rayrMesrc7CvC32RYzNYcBEwUGJ+GrBlY0TTJ/qFa8sm243j2kVon7YS8cc33LMuFqZrO/JibHQBEfoV4Loo8JqbIROr98dmYcpt0q99JTv7vaEb2Y6ZkxSJtEdihFXG2X9aY+Nh8Mrl7aoC5TiaxxFiG7JhFp85QXB0+o2kasyCMhxJveZDcv+QXV5M3GprduaTfeE1S/DjlTa989DZraiDED9f6NWxVhVUnWFpqOuRvqpqYS7XC/5IJfMJMccr+1WQWfDk9Vpo58q6utZtQ6O9nAW2dm5nZoS3VqzfWZ6rmj09rSh0uFfHDbDfeOsWdrvHrUl0P9s0taZr6fU2FuZs+XS/d0xs1/i7A9QSwMECgAAAAAALWD9XAAAAAAAAAAAAAAAAAkAAABkb2NQcm9wcy9QSwMECgAAAAgALWD9XPgrwD83AQAAgwIAABEAAABkb2NQcm9wcy9jb3JlLnhtbKWS0W7CIBSGX6XhvqXUpG6kxWRbvJrJkmm27I7AUckKJcCsvv1o1aqZd7uE/+PLf05bzfa6SXbgvGpNjUiWowSMaKUymxqtlvP0ASU+cCN50xqo0QE8mrFKWCpaB2+uteCCAp9Ej/FU2BptQ7AUYy+2oLnPImFiuG6d5iEe3QZbLr75BnCR5yXWELjkgeNemNrRiE5KKUal/XHNIJACQwMaTPCYZARf2ABO+7sPhuSK1CocLNxFz+FI770awa7rsm4yoLE/wZ+L1/dh1FSZflMCEKukoMIBD61jK5MarkFW+OqyX2DDfVjETa8VyKfDFfc363EHO9V/JUYGYjxWp6GPbpBJLEuPo52Tj8nzy3KOWJEXZZpP0+JxSQqaE1qU2bScfvXVbhwXqT6V+Jf1LGFD89sfh/0CUEsDBAoAAAAIAC1g/VweKelacAIAAGQMAAASAAAAd29yZC9udW1iZXJpbmcueG1szZdLbtswEIavInDvUHLkB4QoQdsghYu+gKYHoCXaJsIXSEqKz9BFd+22Z+tJOpQs+VEgsGUE8Ma0ODPf/BQ5Q+jm7lnwoKTGMiVTFF2FKKAyUzmTyxR9f3wYTFFgHZE54UrSFK2pRXe3N1UiCzGnBtwCkSWzpVSGzDk4VFEcVNEoqHQUowDo0iaVzlK0ck4nGNtsRQWxV4JlRlm1cFeZElgtFiyjuFImx8MwCut/2qiMWgs53hFZEtvixP80pakE40IZQRw8miUWxDwVegB0TRybM87cGtjhuMWoFBVGJhvEoBPkQ5JG0GZoI8wxeZuQe5UVgkpXZ8SGctCgpF0xvV1GXxoYVy2kfGkRpeDbLYji8/bg3pAKhi3wGPl5EyR4o/xlYhQesSMe0UUcI2E/Z6tEECa3iXu9mp2XG41OAwwPAXp53ua8N6rQWxo7jzaTTx3LF/0JrM0m7y7Nnifm24poinzLIXPrDMnc50IEe0+zHFoX8m0nMRS6lfGTTXd6s3DUvDWUPKUorCmi4I59pCXlj2tNAVQSDgrXc8PyT97GvQ1h78tLDg4MBh9dJ3BQhlDLJfUpvU+dr8VETRw0xwfRTc4LzqnriI/0uTP9/f2zm/+QtbOcLjbu+qvxA5M52Px0iiZDryRZEbmsm/T1OPS+eOOMa9ah+Oh1xP84VXwUxz3UD19F/a8/p6ofRuMe6q8v5OAMp9Me6uMLOTkgtof60YWcnPi6T9WOL+TkjMI+VTu5FPWTPlU7vRD14/i4qsV7N+JGVVD/NtfjwQ06yw8WAZQv8CEAtyDdufO6Je/YtlF4L6x+lj453vk+uP0HUEsDBAoAAAAAAC1g/VwAAAAAAAAAAAAAAAAGAAAAX3JlbHMvUEsDBAoAAAAIAC1g/Vwfo5KW5gAAAM4CAAALAAAAX3JlbHMvLnJlbHOtks9KAzEQh18lzL0721ZEpGkvUuhNpD5ASGZ3g80fJlOtb28oilbq2kOPmfzmyzdDFqtD2KlX4uJT1DBtWlAUbXI+9hqet+vJHayWiyfaGamJMvhcVG2JRcMgku8Rix0omNKkTLHedImDkXrkHrOxL6YnnLXtLfJPBpwy1cZp4I2bgtq+Z7qEnbrOW3pIdh8oypknfiUq2XBPouEtsUP3WW4qFvC8zexym78nxUBinBGDNjFNMtduFk/lW6i6PNZyOSbGhObXXA8dhKIjN65kch4zurmmkd0XSeGfFR0zX0p48jGXH1BLAwQKAAAACAAtYP1c0nf8t20AAAB7AAAAGwAAAHdvcmQvX3JlbHMvZm9vdGVyMS54bWwucmVsc02MQQ4CIQxFr0K6d4oujDHDzG4OYPQADVYgDoVQYjy+LF3+vPf+vH7zbj7cNBVxcJwsGBZfnkmCg8d9O1xgXeYb79SHoTFVNSMRdRB7r1dE9ZEz6VQqyyCv0jL1MVvASv5NgfFk7Rnb/wfg8gNQSwMECgAAAAgALWD9XDHFkUDWAQAAtQUAABAAAAB3b3JkL2Zvb3RlcjEueG1spZRZbtswEIavQvDdIiVnq2A5EOy6aNEWAZoegKEpi4m4YEhLbZ96lh6tJwllLbZbIHDiF1Kc4XzzzwzE2e0PVaFagJNGZziOKEZCc7OWepPh7/eryQ2+nc+atPCAwlXt0sbyDJfe25QQx0uhmIuU5GCcKXzEjSKmKCQXpDGwJgmN6e7LguHCucBdMF0zh3uc+p9mrNDBWRhQzIcjbIhi8LS1k0C3zMsHWUn/M7Dp1YAxGd6CTnvEZBTUhqSdoH4bIuCUvF3I0vCtEtrvMhIQVdBgtCul3ZfxVlpwlgOkfqmIWlV4HEF8cd4MlsCasO2Bp8hfd0Gq6pS/TIzpCRNpEWPEKRKOcw5KFJN6n/hNrTlobnz5OkDyL8BuzhvOBzBbu6fJ82gf9dPI0uJVrH7Ih6W588R8K5kVuH1Q7G65g3Z75KhJa1ZlmIf/QgAm8xkZvd3Sf6+M9i7cZo7L0JgFq+QDSBws3B0dBXM+d5IdGctcu4Mo0iK5qQwM+W/yd9Ocdg73a7DGF4Nl4Y5tZFTm2xalzjIeemxBOAF1qHQFxjP0JVjXBr3XG6FLBpKhz/fLHKG/v/8gtPh69wkl19H0aholyZRQSuPJNW3ZvsvQdWO3hnd4/gxQSwMECgAAAAgALWD9XMCyc5ujAQAAuAgAABMAAABbQ29udGVudF9UeXBlc10ueG1stVbLTsMwEPyVKFfUuHBACLXlwOMIHOADXHuTGmKvZW8K/D3r9CEFmlKguWU9MzsT70bK5Ord1tkSQjTopvlpMc4zcAq1cdU0f366G13kV7PJ04eHmDHVxWm+IPKXQkS1ACtjgR4cIyUGK4nLUAkv1ausQJyNx+dCoSNwNKLUI59NbqCUTU3Z9eo8tZ7mxia+d1We3b7z8SpOqsVexYuHrqQ9+LXmJ8nc+o4i1fsVlSk7ilTvV8RldcL32FHxWa9Kel8bJYmJYun0lzmM1jMoAtQtJy6Mj98MGI0HOXwVpvqPybAsjQKNqrEsKXBeNpHZoO+4SccENVF7bQ+8ocFo+I/PGwbtAyqIkZfb1sUWsdK41c08ykD30nJvkehiS1m/7iA5In3UEHcHWGH/st8sgsIAIzb2EMjs8OOAj4xGkYjHfGHVREJ7mHVLPaY5pG3SoA+y59aDTto1dg6Bn3cPewsPGqJEJIfUt3FbeNAQPJM9GTbosJ8dEPFT34e3RgeNoNAmoCfCBh14G7iRnNfQtw1rePCVhNC/jxBON/6i/RWZfQJQSwMECgAAAAgALWD9XFh52yKSAAAA5AAAABMAAABkb2NQcm9wcy9jdXN0b20ueG1snc5BCsIwEIXhq5TZ21QXIqVpN+LaRXUf0mkbaGZCJi329kYED+Dy8cPHa7qXX4oNozgmDceyggLJ8uBo0vDob4cLFJIMDWZhQg07CnRtc48cMCaHUmSARMOcUqiVEjujN1LmTLmMHL1JecZJ8Tg6i1e2q0dK6lRVZ2VXSewP4cfB16u39C85sP28k2e/h+yp9g1QSwMECgAAAAgALWD9XOL8ndqTAAAA5gAAABAAAABkb2NQcm9wcy9hcHAueG1snc5BCsIwEIXhq4TsbaoLkdK0G3HtoroPybQNNDMhE0t7eyOCB3D5+OHjtf0WFrFCYk+o5bGqpQC05DxOWj6G2+EiBWeDziyEoOUOLPuuvSeKkLIHFgVA1nLOOTZKsZ0hGK5KxlJGSsHkMtOkaBy9hSvZVwDM6lTXZwVbBnTgDvEHyq/YrPlf1JH9/OPnsMfiqe4NUEsDBAoAAAAIAC1g/VycicmRzgEAAK0GAAASAAAAd29yZC9mb290bm90ZXMueG1s1ZTNTuMwEMdfJfK9dVIBWkVNOYBA3BDdfQDjOI2F7bFsJ6Fvv5PETbosqgo9cYm/Zn7zn5nY69t3rZJWOC/BFCRbpiQRhkMpza4gf34/LH6RxAdmSqbAiILshSe3m3WXVwDBQBA+QYLxeWd5QeoQbE6p57XQzC+15A48VGHJQVOoKskF7cCVdJVm6TCzDrjwHsPdMdMyTyJO/08DKwweVuA0C7h0O6qZe2vsAumWBfkqlQx7ZKc3BwwUpHEmj4jFJKh3yUdBcTh4uHPiji73wBstTBgiUicUagDja2nnNL5Lw8P6AGlPJdFqRaYWZFeX9eDesQ6HGXiO/HJ00mpUfpqYpWd0pEdMHudI+DfmQYlm0syBv1Wao+Jm118DrD4C7O6y5jw6aOxMk5fRnszbxOov9hdYscnHqfnLxGxrZvEGap4/7Qw49qpQEbYswaon/W9Njp+cpMvD3qKFF5Y5FsAR3JJlQRbZYGiHz7PrB28ZxwhowKog8HanvbGSfc6rq2nx0vQhWROA0M2aTu7jJ863Ya/66C1TBXmIal5EJRy+mSI6RuNqPo77E26SPR3QQTOdvT5Nl4MJ0jTDK7P9mHr6EzL/NINTVTha+M1fUEsDBAoAAAAIAC1g/VzSd/y3bQAAAHsAAAAdAAAAd29yZC9fcmVscy9mb290bm90ZXMueG1sLnJlbHNNjEEOAiEMRa9CuneKLowxw8xuDmD0AA1WIA6FUGI8vixd/rz3/rx+824+3DQVcXCcLBgWX55JgoPHfTtcYF3mG+/Uh6ExVTUjEXUQe69XRPWRM+lUKssgr9Iy9TFbwEr+TYHxZO0Z2/8H4PIDUEsDBAoAAAAIAC1g/Vw/So6NwQEAAJIGAAARAAAAd29yZC9lbmRub3Rlcy54bWzNlNtu4yAQhl/F4j7BjrrVyorTix5Wvaua3QegGMeowCDA9ubtd3wIzrZVlDY3vTGnmW/+mTGsb/5qlbTCeQmmINkyJYkwHEppdgX58/th8ZPcbNZdLkxpIAifoL3xeWd5QeoQbE6p57XQzC+15A48VGHJQVOoKskF7cCVdJVm6TCzDrjwHuG3zLTMkwmn39PACoOHFTjNAi7djmrmXhu7QLplQb5IJcMe2en1AQMFaZzJJ8QiCupd8lHQNBw83DlxR5c74I0WJgwRqRMKNYDxtbRzGl+l4WF9gLSnkmi1IrEF2dVlPbhzrMNhBp4jvxydtBqVnyZm6Rkd6RHR4xwJ/8c8KNFMmjnwl0pzVNzsx+cAq7cAu7usOb8cNHamyctoj+Y1soz4FGtq8nFq/jIx25pZvIGa5487A469KFSELUuw6kn/W5OjFyfp8rC3aOCFZY4FcAS3ZFmQRTbY2eHz5PrBW8YxABqwKgi83GlvrGSf8uoqLp6bPiJrAhC6WdPoPn6m+TbsVR+9Zaog96OYZ1EJh++jmPwmWxFPp+0Ii6LjAR0U0+j0UaocTJCmGR6Y7du00++f9Yf6T1RgnvvNP1BLAwQKAAAACAAtYP1c0nf8t20AAAB7AAAAHAAAAHdvcmQvX3JlbHMvZW5kbm90ZXMueG1sLnJlbHNNjEEOAiEMRa9CuneKLowxw8xuDmD0AA1WIA6FUGI8vixd/rz3/rx+824+3DQVcXCcLBgWX55JgoPHfTtcYF3mG+/Uh6ExVTUjEXUQe69XRPWRM+lUKssgr9Iy9TFbwEr+TYHxZO0Z2/8H4PIDUEsDBAoAAAAIAC1g/VxNn8rKoQEAAHMFAAARAAAAd29yZC9zZXR0aW5ncy54bWyllN1u2zAMhV/F0H0iu1iLwahbdCvW9WLYRbcHYCXZFiJRgiTby9uPjuO4P0CRNFeSQfE7R6TF69t/1mS9ClE7rFixzlmmUDipsanY3z8/Vl9ZFhOgBONQVWyrIru9uR7KqFKiQzEjAMZy8KJibUq+5DyKVlmIa6tFcNHVaS2c5a6utVB8cEHyi7zIdzsfnFAxEug7YA+R7XH2Pc15hRSsXbCQ6DM03ELYdH5FdA9JP2uj05bY+dWMcRXrApZ7xOpgaEwpJ0P7Zc4Ix+hOKfdOdFZh2inyoAx5cBhb7ZdrfJZGwXaG9B9doreGHVpQfDmvB/cBBloW4DH25ZRkzeT8Y2KRH9GREXHIOMbCa83ZiQWNi/CnSvOiuMXlaYCLtwDfnNech+A6v9D0ebRH3BxY47s+gbVv8surxfPMPLXg6QVaUT426AI8G3JELcuo6tn4W7Nx4kgdvYHtNxCbhmqBcpfGx5DqFd6h/C3lTwWSplk2lD2YitVgomK7M9OUWHZP0wCbTxaXjLYIlqRfDZRfTqox1IUTSj5K8kWTL/Py5j9QSwMECgAAAAgALWD9XIuGOcTFAQAAxggAABEAAAB3b3JkL2NvbW1lbnRzLnhtbKXU3XLiIBgG4FtxOFeSWFM307Qnne30eNsLoIDCNPwMoNG7X1IlSZedToJH6iTfk5fXwMPTSTSLIzWWK1mDfJWBBZVYES73NXh/+73cgoV1SBLUKElrcKYWPD0+tBVWQlDp7MID0lb4VAPmnK4gtJhRgexKcGyUVTu38vdCtdtxTCExqPU2LLL8DmKGjKMn0Bv5bGQDf8FtDBUJUJ7BIo+p9WyqhF2qCLpLgnyqSNqkSf9ZXJkmFbF0nyatY2mbJkWvk8ARpDSV/uJOGYGc/2n2UCDzedBLD2vk+AdvuDt7MysDg7j8TEjkp3pBrMls4R4KRWizJkFRNTgYWV3nl/18F726zF8/woSZsv7LyLPCh247f60cGtr4LpS0jGvb15mq+YssIMefFnEUTbiv1fnE7dIqQ7q+sq9v2ihMrfUdPl+qHMAp8a/9i+aS/Gcxzyb8Ix3RT0yJ8P2ZIYnwb+Hw4KRqRuXmEw+QABQRUGI68cAPxvZqQDzs0M7hE7dGcMre4WTkpIUZAZY4wmYpRegVdrPIIYYsG4t0XqhNz53FqCO9v20jvBh10IPGb9Neh2OtlfMWmJX/tq7tbWH+MKQpgI9/AVBLAwQKAAAACAAtYP1c0nf8t20AAAB7AAAAHAAAAHdvcmQvX3JlbHMvY29tbWVudHMueG1sLnJlbHNNjEEOAiEMRa9CuneKLowxw8xuDmD0AA1WIA6FUGI8vixd/rz3/rx+824+3DQVcXCcLBgWX55JgoPHfTtcYF3mG+/Uh6ExVTUjEXUQe69XRPWRM+lUKssgr9Iy9TFbwEr+TYHxZO0Z2/8H4PIDUEsDBAoAAAAIAC1g/Vxj7V7WHQEAAEMDAAASAAAAd29yZC9mb250VGFibGUueG1sndHdbsIgFAfwVyHcK7WZjWms3ixLdr89AAK1RA6n4eDUtx+ttmvijd0VEPL/5Xxs91dw7McEsugrvlpmnBmvUFt/rPj318diwxlF6bV06E3Fb4b4fre9lDX6SCylPZWgKt7E2JZCkGoMSFpia3z6rDGAjOkZjgJkOJ3bhUJoZbQH62y8iTzLCv5gwisK1rVV5h3VGYyPfV4E45KInhrb0qBdXtEuGHQbUBmi1DG4uwfS+pFZvT1BYFVAwjouUzOPinoqxVdZfwP3B6znAfkTUChznWdsHoZIyalj9TynGB2rJ87/ipkApKNuZin5MFfRZWWUjaRmKpp5Ra1H7gbdjECVn0ePQR5cktLWWVoc62F2n1x3sPsy2NACF7tfUEsDBAoAAAAIAC1g/VzSd/y3bQAAAHsAAAAdAAAAd29yZC9fcmVscy9mb250VGFibGUueG1sLnJlbHNNjEEOAiEMRa9CuneKLowxw8xuDmD0AA1WIA6FUGI8vixd/rz3/rx+824+3DQVcXCcLBgWX55JgoPHfTtcYF3mG+/Uh6ExVTUjEXUQe69XRPWRM+lUKssgr9Iy9TFbwEr+TYHxZO0Z2/8H4PIDUEsDBAoAAAAAAC1g/VwAAAAAAAAAAAAAAAALAAAAd29yZC9tZWRpYS9QSwMECgAAAAgALWD9XCT9uJzVDQAA0Q4AADcAAAB3b3JkL21lZGlhLzAwNzdlZDljODE4OGRmYjI0NDc3ZDVkYjk2MDlmZjRjNTQ1ZTA3ZjEuanBnnVZ7OJT5238GhbCl6DDIpshZzlthOlEpISEzDm9ZYiakHENTbc4rK2GxmhCTw7AYMzHMrIYUSQ5jZDCZHNJgjMEMY+b5jb323d0/3j/e972f676u7/Pc9/e+ns/nvu/v9waHwc/ADkeH8w4ABAIBmiQPADKA04Dc1q2yW7fIycrKysvLbVNUUVJUUFDct3PXdhUNqOZ+Dai6+veHjA9/r2V4UF1d11rP8IiphYWF5uGjtj+Y2RibW5htBoHIy8srKijuVVLaa3ZA/YDZ/1nAPwBlOeBH4K00RAuQUoZIK0PAdkATACBbIH8K8JdApKRltmyVlZPfpiBxaNwBSEGkpaVkpLdskZGRWBMkdkBGecvOA6Ynt+5yvSarFa5i9iCrWO7gqbo21ct9i4fMr99+KL9t9569+6DaOod19fQtLK2sfzh67PQZe4ez5847ul1x9/C86gX3/zEg8EZQMPJORGRUdEzs3Z8eJSYlp6SmPcl+mpOb92t+QUnpi7Jy7MuKyvoGfCOB+Kqp+TW1veNN59t3Xf0Dg7Qh+vCnkQnWl8mp6Zmvs9+4S7zllVW+YG19ExcEkIb8t/yPuJQluKRkZKRlZDdxQaSiNx2UZbYcMN2686Sr7LXwXVpmD+RUTmUV17XJHzS/vKh6/Xbftt2HLCa0uZvQ/kT2vwP28P+F7G9g/+AaARSlIZLkSSsDMID/TbfsuUxJbhrgqBckWaSZn1Aqqz3O7W0PWGz8NjxyFTf/pteUb0dzXqDLw95kKjuJS/q7mIs3qqfguuzShzgvX10+RaCkRhe+Jt/GsJzTFsSrVVMmItNE+4BellHojNguMp2+/p72W1U7G5vtkTu3fj2OjYgcs23Xhc2WPrjKt7NcpKXdAwFEg+Z5d8oydK2D8WGtKJ8hKh9zqQ8RVhFK2XZnr5hdWrKdZrPoiUYNhNmF7ZL4Vf6vyaYDNggl6nBsytA84Ux+FLGutiAf5UyFijqTGt9ayhEt0FEsGAfLdYX265Bd5dz41naVZZVW1h+RVi5jjHO1il2HMvIEib41DW3DumUl93emwXUhfy/0gvWCgJI03U2FGBggBueXnaCcy83HrxEcslrZHpWjjI5ImVyeQ61azVBGs/KJM2sebPdXtEi1dGFp/kzWC9iGk2uTU0aTHRu9nY+q5ILAg6nFpKoZhAE56xYINKHXJpmjYaJckklSQQI5Z0KNs5NZORcZUpgn9qoFARUadU70ZdauSc4l/cFUVLNvVQzz60jIJHMB0eAhwtxv8YwU155pcgeBcoYbCJyqSeJ/MSbUqxIGDFmZAh+idbdX/4Yr16dzrOWh5vpMKEVoqUmyyPD+UaySFTow4ZwaeG3GqKwZBIhWHNFWwjWBacQ9VGugHAiMJhZK4iFZC/EoFoJCwY2jSbtCQOAjCvpN6NIcQt9Dz0bTz+e/nE72d/Z/nMc4t6oW9hP8ciIqw3H1i6in0sCDMeH8aELTNn1Ucz3E3PKOn1yLx8L2ZBPttYEZvCEzFp0K+2zJZlKXPJ1+aMNtl1VF7JJfuhTgY32hvq/6YP3vxAtxV8VzKL4HXk8P5fZItyQN+Lc6wh3hf7/o6QXrSmvC++Yae8wSjolNrMYDqCYr+1cmJQyunLOIO64pf9X47fVE302iUEuVIDBHm4+r8/bM3FDdno7HfZxzcvxkHmJdK1QSd1+5ZP04+RIUqbort7N3nsCyxfEpiRSSC1WJ+q5eLbY8bP2GmiW8f7mQzdlBLY41UmL3THZl8LOFTH4XHRdD7XCeMA8uDjrSNn1yGyfYM4q5UDvUpGprcpD5TI9X5uzRa9WwSEwJBQFaUB5WhE3Mk/CMe2kHjSmfgk1C69D04KfuCEzKft3BPeJTtd8tfHxT4cOhmkBHkYhlp8ufu6MwLQ+DEdM56FT0ODzc5XbdD2nDzUdmPrYP2O0UPB5fVaAN3sQdTnyU4nknBLn66nSKa6D+jwye/1B5Wc2rtrP+M2x6Ueydi3AJgaU7/qrtv2v83wUPvNhsg58unD+h3kNPXh6ObGq5oPpY5GpUYMnYQIvykGtkvvurwRUEidHDROYGGuM2jERMgVJCsK8G16/NLyDWxqm8Q0gAgcxJvwkYbwAE0qqrHLi2is0t19zGbVgBIeMRBqTocXecj/kluu+0Qo+1cyuPFfJOi0Hsoo1tQBc5RYrrK+/94vi7hRTuNAh8Qg0TtETM3Dqsg5j7KfX9mlp5Ttcx1fnx4lKGN/PT+YCIbh6D4m6Jc8J6erFMg0rSZFRc7YF/q7aRthGgfaO3Kni1IdbTeuj3d6JPge7Ij7UF9wzKf2XuPbv6ukeEmaSuZjaAQDWjkTwPJMiKGsrXfMNWM7A5JEni8MZj4zmWjaKiaU69D9WO2Tr5FkcgR3A9iiYY6Ee5RLu9QsTVsW9CEBggLIzW1Q9FJxp9dk7xVXeWd4paOvKIyZm4AgLdASCAtMVDlqdoK7hEklN0uOKO2vHYInekEcxzoNEi8PfBkNBVEUcnIDtMxPOVv4MK58JYPav5cXcsjvn+tqTfjqFCrZhP0bGFggqb6M/tgpBwbEwydc9Ah5BpmBxfNDwGAlJKohGAXyHpi8NBU4M3v/xyq61GxWygeMiCnZq711LN61BYc9bFaniLdSd5W19eW3/Mwp5mYdSLj6N44299VU58sg9l3gNL5Uh/PRadZEHRGO2JyPMYWvJ+NU72625CkyThV7UHjeBzI7hBpVut13ied8UqbBpBqJ4e5b+075aQgWSRvfpXcM6XuPaYzFB/g7DIymNdIDAuXgIBphEIPKa3xjs90SPvQYVW49QYBgZ3a6aef1d4fmAuosdYQcNW1ehA5k19HCa3XHhRayJkJKi9I1o/dC0FBKJaeBOwVN8Z6s7gY8VLocuPfMbDuCe6mxDitRq319DWzlQ7q7PVLfYhC0Y1Wx5z1lgUfUnnL0PmB7xPKtZ1Hqp78ST+ZTXjQoK+Z871n1uqVa8E5L/qiLnRrpBO2SHAetJuGp+qfVSSOccmcXpIERgBR/wTnHxkyDljxU8uyIp/mTELH/RDvlpiKLJBwIYHAhAksGzP7Uwcn5nrrp0djwpvmFfwnJaQHWtTsaJw9YV4Llu/gAr7zgYaAQJ16NfHAybn33AM+Bg+fbzWkAuNCe1kjZcUrGYbF18viDelKpqa50S6M/joT8+3h/kgnIPiaa2EoDsN73l5oZwF2q/G8wueR7ztTUS3f7GMyczAtyrbaE+cwxduxfdzVgKycJMep3nGQ2OikRGuM7UHRoVr01NsDhuuXyr21yxautijXkmJ9YLft4UHlxIf9VikZdUavh+WqsPFoyaq2JR6OHsknq45ikjsiX2MrXTibvdaDE2/y9mKQfuLimOyCMyjedHdkgyCwBfaCJKp2KL38mv/VWo5fEQ1k2+9o4aeQUkkushGXvCNXEz3Ig3F790y94aFOp7cyCsUw/VFI8vwfw5zxyXb6D/We0DgzQVN5TK8SZn95NrP/cacHATM6O68NhCL6fCTR/oeXmx9hilmnOovY/72rtXfh4SQHCs1Z68IQli4w+59KxwF6ja/s2Hbxk480Q/J90Sh67uYHb4r9kBMIGmR15FbMLjfcHE5t6L6ok03O8hvyquQMz/Gz6E003zDpATCzpGyCi6q0x97EgQMvxhIOmikVFBIqxrccIu2bJjt1mgjJnnlB4Ydv6GWENXoRw8oYFLpbvCLrwOS+bdr3zPhwMy9A4JADmJxV1CBdkugy+JaOL7wij6RX/Vs0MaA9VRgjY2Bd9HUXGCkJtF6ZXVcPpkv4vm8yYtG61RFouJJ8SmlDC0QqCoMD/sV/aXF3eiCy7nIKgrRz5txMrSutcVqnvhBo1foZ91NuG8zxP78uDJep64D12s04lCvUzC8ZNPL8l4Agc4q3H5NrtFMW8UGLA5JYRiErV7ZfzvnFm6piLlA92EdnWrs08mV1dPCo80IzLUgEvk3j6RnE8uCIm3ePb/W6TWMgM91kX0qyPzD7nsSoc/PXmtkZXmlfyVsXH+j0BOzSklrZL72HLmRgRedPVOfMWchUsUdELGqLDBWTCFbSGsoRULf5x27af0z76zFJRAIusIgiXL9JgI2/+20G2nwOe1gUpdqhNWdc+6MIMWyKAzphDo6kaw6CxtW7mZlWB3HIVnndV4lzFPN0Q1eURSlSMycc50NO+BlEL3+Qqlf2Oi+fcSQMOfNQaXWhdCKQiudfum9ELubFpg9tZZPoPkpjX3b6dUbVnkoPUEGYzVPMSO+p0SziWLpE+pYBtmwkfgpKbPkUKvdaKgZbJzhzlmGQWfFdq3c7zKuYZkJJR+OB8alZxqGuvv0EEVTNbkf1oXinRWsZ8G7aWVZ/jLx9f9lrIpu7uahWxFuNlUstXO/ZCnm8sxNhI7viJvDE8Dr3R28/yjXLhXb+ib7enNGbnlDtRVf08ChSiS5zJCYilmMsg1KdjUQG4YNuuudLNIuzvvwxF48H1kwXZGAigsiuDpeI4pE01aSevgwiv7q1CkZIh1G8ZCZe9qC2NV0JY2No5GrGm72qf4N+WS2CR+xAQKqbi3OadCp1AGdxj48PnyqJ0HxfQ9RQMDdzVGB9MZ4UqOeqSyKarIDvfph+UVkXt03Wt/lyx8h9iouUpf7/rkpLZ/ux5MkHyF4t4vnGLdLp2btVQCSvU6H/ZqDbsl9DbwrXuJl+jQEK9mkZq/2p6uLxLI7Tfe51F+6Hfz0H1BLAQIUAAoAAAAAAC1g/VwAAAAAAAAAAAAAAAAFAAAAAAAAAAAAEAAAAAAAAAB3b3JkL1BLAQIUAAoAAAAAAC1g/VwAAAAAAAAAAAAAAAALAAAAAAAAAAAAEAAAACMAAAB3b3JkL19yZWxzL1BLAQIUAAoAAAAIAC1g/VyDDJMKNQEAAEcFAAAcAAAAAAAAAAAAAAAAAEwAAAB3b3JkL19yZWxzL2RvY3VtZW50LnhtbC5yZWxzUEsBAhQACgAAAAgALWD9XIoRD9eNDAAAhrQAABEAAAAAAAAAAAAAAAAAuwEAAHdvcmQvZG9jdW1lbnQueG1sUEsBAhQACgAAAAgALWD9XJjTxsINAwAAExEAAA8AAAAAAAAAAAAAAAAAdw4AAHdvcmQvc3R5bGVzLnhtbFBLAQIUAAoAAAAAAC1g/VwAAAAAAAAAAAAAAAAJAAAAAAAAAAAAEAAAALERAABkb2NQcm9wcy9QSwECFAAKAAAACAAtYP1c+CvAPzcBAACDAgAAEQAAAAAAAAAAAAAAAADYEQAAZG9jUHJvcHMvY29yZS54bWxQSwECFAAKAAAACAAtYP1cHinpWnACAABkDAAAEgAAAAAAAAAAAAAAAAA+EwAAd29yZC9udW1iZXJpbmcueG1sUEsBAhQACgAAAAAALWD9XAAAAAAAAAAAAAAAAAYAAAAAAAAAAAAQAAAA3hUAAF9yZWxzL1BLAQIUAAoAAAAIAC1g/Vwfo5KW5gAAAM4CAAALAAAAAAAAAAAAAAAAAAIWAABfcmVscy8ucmVsc1BLAQIUAAoAAAAIAC1g/VzSd/y3bQAAAHsAAAAbAAAAAAAAAAAAAAAAABEXAAB3b3JkL19yZWxzL2Zvb3RlcjEueG1sLnJlbHNQSwECFAAKAAAACAAtYP1cMcWRQNYBAAC1BQAAEAAAAAAAAAAAAAAAAAC3FwAAd29yZC9mb290ZXIxLnhtbFBLAQIUAAoAAAAIAC1g/VzAsnObowEAALgIAAATAAAAAAAAAAAAAAAAALsZAABbQ29udGVudF9UeXBlc10ueG1sUEsBAhQACgAAAAgALWD9XFh52yKSAAAA5AAAABMAAAAAAAAAAAAAAAAAjxsAAGRvY1Byb3BzL2N1c3RvbS54bWxQSwECFAAKAAAACAAtYP1c4vyd2pMAAADmAAAAEAAAAAAAAAAAAAAAAABSHAAAZG9jUHJvcHMvYXBwLnhtbFBLAQIUAAoAAAAIAC1g/VycicmRzgEAAK0GAAASAAAAAAAAAAAAAAAAABMdAAB3b3JkL2Zvb3Rub3Rlcy54bWxQSwECFAAKAAAACAAtYP1c0nf8t20AAAB7AAAAHQAAAAAAAAAAAAAAAAARHwAAd29yZC9fcmVscy9mb290bm90ZXMueG1sLnJlbHNQSwECFAAKAAAACAAtYP1cP0qOjcEBAACSBgAAEQAAAAAAAAAAAAAAAAC5HwAAd29yZC9lbmRub3Rlcy54bWxQSwECFAAKAAAACAAtYP1c0nf8t20AAAB7AAAAHAAAAAAAAAAAAAAAAACpIQAAd29yZC9fcmVscy9lbmRub3Rlcy54bWwucmVsc1BLAQIUAAoAAAAIAC1g/VxNn8rKoQEAAHMFAAARAAAAAAAAAAAAAAAAAFAiAAB3b3JkL3NldHRpbmdzLnhtbFBLAQIUAAoAAAAIAC1g/VyLhjnExQEAAMYIAAARAAAAAAAAAAAAAAAAACAkAAB3b3JkL2NvbW1lbnRzLnhtbFBLAQIUAAoAAAAIAC1g/VzSd/y3bQAAAHsAAAAcAAAAAAAAAAAAAAAAABQmAAB3b3JkL19yZWxzL2NvbW1lbnRzLnhtbC5yZWxzUEsBAhQACgAAAAgALWD9XGPtXtYdAQAAQwMAABIAAAAAAAAAAAAAAAAAuyYAAHdvcmQvZm9udFRhYmxlLnhtbFBLAQIUAAoAAAAIAC1g/VzSd/y3bQAAAHsAAAAdAAAAAAAAAAAAAAAAAAgoAAB3b3JkL19yZWxzL2ZvbnRUYWJsZS54bWwucmVsc1BLAQIUAAoAAAAAAC1g/VwAAAAAAAAAAAAAAAALAAAAAAAAAAAAEAAAALAoAAB3b3JkL21lZGlhL1BLAQIUAAoAAAAIAC1g/Vwk/bic1Q0AANEOAAA3AAAAAAAAAAAAAAAAANkoAAB3b3JkL21lZGlhLzAwNzdlZDljODE4OGRmYjI0NDc3ZDVkYjk2MDlmZjRjNTQ1ZTA3ZjEuanBnUEsFBgAAAAAaABoAoQYAAAM3AAAAAA=="

_NAVY=_col.HexColor("#24344D"); _GRAY=_col.HexColor("#666666"); _LGRAY=_col.HexColor("#EDEDED"); _LN=_col.HexColor("#C9CDD3")
_MESES=["JANEIRO","FEVEREIRO","MARÇO","ABRIL","MAIO","JUNHO","JULHO","AGOSTO","SETEMBRO","OUTUBRO","NOVEMBRO","DEZEMBRO"]
def _rs(v): return "R$ " + f"{float(v):,.2f}".replace(",","X").replace(".",",").replace("X",".")
def _num_limpo(s):
    d=re.sub(r"\D","",str(s or "")); d=d.lstrip("0"); return d or ("0" if re.search(r"\d",str(s or "")) else "")
def _data_iso(s):
    m=re.search(r"(\d{2})/(\d{2})/(\d{4})",str(s or ""))
    if m: return f"{m.group(3)}-{m.group(2)}-{m.group(1)}"
    m=re.search(r"(\d{4})-(\d{2})-(\d{2})",str(s or ""))
    return m.group(0) if m else None
def _mes_da_data(s):
    iso=_data_iso(s)
    if iso:
        y,mo,_=iso.split("-"); return f"{_MESES[int(mo)-1]} {y}"
    h=_hoje(); return f"{_MESES[h.month-1]} {h.year}"
def _mes_atual():
    h=_hoje(); return f"{_MESES[h.month-1]} {h.year}"

_EU=["zero","um","dois","três","quatro","cinco","seis","sete","oito","nove","dez","onze","doze","treze","quatorze","quinze","dezesseis","dezessete","dezoito","dezenove"]
_ED=["","","vinte","trinta","quarenta","cinquenta","sessenta","setenta","oitenta","noventa"]
_EC=["","cento","duzentos","trezentos","quatrocentos","quinhentos","seiscentos","setecentos","oitocentos","novecentos"]
def _ext999(n):
    if n==0: return ""
    if n==100: return "cem"
    p=[];c=n//100;r=n%100
    if c:p.append(_EC[c])
    if 0<r<20:p.append(_EU[r])
    else:
        d=r//10;u=r%10
        if d:p.append(_ED[d])
        if u:p.append(_EU[u])
    return " e ".join(p)
def _extenso_reais(v):
    v=round(float(v),2);reais=int(v);cent=int(round((v-reais)*100))
    def g(n):
        mil=n//1000;resto=n%1000;o=[]
        if mil:o.append("mil" if mil==1 else _ext999(mil)+" mil")
        if resto:o.append(_ext999(resto))
        return " e ".join([x for x in o if x])
    rp="zero reais" if reais==0 else ("um real" if reais==1 else g(reais)+" reais")
    if cent==0:return rp+"."
    cp="um centavo" if cent==1 else _ext999(cent)+" centavos"
    return f"{rp} e {cp}."

def _oc_hdr(t): return _P(t,_PS("h",fontName="Helvetica-Bold",fontSize=8.5,textColor=_col.white,leading=11))
def _oc_kv(l,v): return _P(f"<b>{l}:</b> {v}",_PS("kv",fontName="Helvetica",fontSize=8.5,textColor=_col.HexColor('#1e2733'),leading=13))

def gera_orcamento_pdf(d):
    buf=io.BytesIO()
    doc=_SDT(buf,pagesize=_A4,leftMargin=16*_mm,rightMargin=16*_mm,topMargin=12*_mm,bottomMargin=14*_mm,title=f"Orçamento {d['num']}")
    W=doc.width; el=[]
    emp=_P("<b>FROTA MACEDO ENGENHARIA LTDA</b><br/><font size=7 color='#666666'>"
        "Eng. Heitor de Oliveira Albuquerque, 295 — Cidade dos Funcionários — Fortaleza/CE<br/>"
        "(85) 2181-1386 • frotamacedoengenharia@gmail.com • CNPJ 27.363.223/0001-70</font>",
        _PS("emp",fontName="Helvetica",fontSize=11,textColor=_NAVY,leading=13))
    dire=_P(f"<font size=8 color='#666666'>{d['data']}<br/>Orçamento nº {d['num']}</font>",_PS("dir",alignment=_TR,fontName="Helvetica",fontSize=8,leading=12))
    try: logo=_IMG(io.BytesIO(_b64.b64decode(_LOGO_ORC_B64)),width=22*_mm,height=11*_mm)
    except Exception: logo=_P("",_PS("x"))
    h=_T([[logo,emp,dire]],colWidths=[24*_mm,W-24*_mm-40*_mm,40*_mm])
    h.setStyle(_TS([("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LINEBELOW",(0,0),(-1,-1),1.2,_NAVY),("BOTTOMPADDING",(0,0),(-1,-1),8)]))
    el+=[h,_SP(1,8)]
    tit=_T([[_P(f"<font size=17 color='white'><b>ORÇAMENTO N° {d['num']}</b></font><br/><font size=8 color='white'>REVISÃO {d.get('revisao',1)}</font>",_PS("t",leading=20))]],colWidths=[W])
    tit.setStyle(_TS([("BACKGROUND",(0,0),(-1,-1),_NAVY),("LEFTPADDING",(0,0),(-1,-1),12),("TOPPADDING",(0,0),(-1,-1),10),("BOTTOMPADDING",(0,0),(-1,-1),10)]))
    el+=[tit,_SP(1,8)]
    obra=_T([[_P(f"<b>OBRA:</b>  MANUTENCAO {d['loja_nome']}  #{d['num']}",_PS("o",fontName="Helvetica",fontSize=9,leading=12))]],colWidths=[W])
    obra.setStyle(_TS([("BACKGROUND",(0,0),(-1,-1),_LGRAY),("LEFTPADDING",(0,0),(-1,-1),10),("TOPPADDING",(0,0),(-1,-1),7),("BOTTOMPADDING",(0,0),(-1,-1),7)]))
    el+=[obra,_SP(1,10)]
    pv=d["prestador"]; tv=d["tomador"]; cw=(W-6)/2
    pb=[_oc_kv("Nome",pv["nome"]),_oc_kv("CNPJ",pv["cnpj"]),_oc_kv("Forma de pagamento",pv["forma"]),_oc_kv("Data de faturamento",d["data"])]
    tb=[_oc_kv("Nome",tv["nome"]),_oc_kv("CNPJ",tv.get("cnpj") or "—"),_oc_kv("Endereço",tv.get("endereco") or "—"),_oc_kv("Cidade/Estado",tv.get("cidade") or "—")]
    def bloco(tt,ls):
        inner=_T([[_oc_hdr(tt)]]+[[x] for x in ls],colWidths=[cw-2])
        inner.setStyle(_TS([("BACKGROUND",(0,0),(0,0),_NAVY),("LEFTPADDING",(0,0),(-1,-1),8),("RIGHTPADDING",(0,0),(-1,-1),8),
            ("TOPPADDING",(0,0),(0,0),5),("BOTTOMPADDING",(0,0),(0,0),5),("TOPPADDING",(0,1),(-1,-1),2),("BOTTOMPADDING",(0,1),(-1,-1),2),
            ("BOX",(0,1),(0,-1),0.6,_LN),("TOPPADDING",(0,1),(0,1),6),("BOTTOMPADDING",(0,-1),(0,-1),6)]))
        return inner
    pt=_T([[bloco("DADOS DO PRESTADOR",pb),bloco("DADOS DO TOMADOR",tb)]],colWidths=[cw,cw],hAlign="LEFT")
    pt.setStyle(_TS([("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(0,0),0),("RIGHTPADDING",(1,0),(1,0),0)]))
    el+=[pt,_SP(1,12)]
    el.append(_P("<b>DISCRIMINAÇÃO DOS ITENS</b>",_PS("di",fontName="Helvetica-Bold",fontSize=9,textColor=_NAVY,spaceAfter=4)))
    data=[["ITEM","DESCRIÇÃO","QUANT.","UNID.","VALOR UNIT.","TOTAL"]]; tg=0.0
    ds=_PS("ds",fontName="Helvetica",fontSize=8.5,leading=11)
    for i,it in enumerate(d["itens"],1):
        vu=round(float(it["valor_unit"])*1.20,2); tot=round(vu*float(it["quant"]),2); tg+=tot
        data.append([str(i),_P(it["descricao"],ds),f"{float(it['quant']):.2f}".replace(".",","),it.get("unid","UN"),_rs(vu),_rs(tot)])
    data.append(["TOTAL GERAL","","","","",_rs(tg)]); n=len(data)-1
    itb=_T(data,colWidths=[13*_mm,W-13*_mm-22*_mm-16*_mm-26*_mm-26*_mm,22*_mm,16*_mm,26*_mm,26*_mm])
    itb.setStyle(_TS([("BACKGROUND",(0,0),(-1,0),_NAVY),("TEXTCOLOR",(0,0),(-1,0),_col.white),("FONTNAME",(0,0),(-1,0),"Helvetica-Bold"),
        ("FONTSIZE",(0,0),(-1,-1),8.5),("ALIGN",(2,0),(5,-1),"CENTER"),("ALIGN",(4,1),(5,-1),"RIGHT"),("ALIGN",(0,0),(0,-1),"CENTER"),
        ("VALIGN",(0,0),(-1,-1),"MIDDLE"),("GRID",(0,0),(-1,-1),0.5,_LN),("ROWBACKGROUNDS",(0,1),(-1,n-1),[_col.white,_col.HexColor("#F6F7F9")]),
        ("SPAN",(0,n),(4,n)),("BACKGROUND",(0,n),(-1,n),_LGRAY),("FONTNAME",(0,n),(-1,n),"Helvetica-Bold"),("ALIGN",(0,n),(4,n),"RIGHT"),
        ("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5),("LEFTPADDING",(1,0),(1,-1),6)]))
    el+=[itb,_SP(1,8)]
    el.append(_P(f"<b>Valor total por extenso:</b>  <i>{_extenso_reais(tg)}</i>",_PS("ex",fontName="Helvetica",fontSize=9,leading=12)))
    el.append(_SP(1,12))
    obs=["Orçamento válido por 7 (sete) dias corridos a partir da data de emissão.","Os valores acima incluem material e serviço de entrega."]
    obt=_T([[_oc_hdr("OBSERVAÇÕES")]]+[[_P("• "+o,_PS("ob",fontName="Helvetica",fontSize=8.5,leading=12))] for o in obs],colWidths=[W-2])
    obt.setStyle(_TS([("BACKGROUND",(0,0),(0,0),_NAVY),("LEFTPADDING",(0,0),(-1,-1),8),("TOPPADDING",(0,0),(0,0),5),("BOTTOMPADDING",(0,0),(0,0),5),
        ("BOX",(0,1),(0,-1),0.6,_LN),("TOPPADDING",(0,1),(0,1),6),("BOTTOMPADDING",(0,-1),(0,-1),6),("TOPPADDING",(0,1),(-1,-1),2),("BOTTOMPADDING",(0,1),(-1,-1),2)]))
    el+=[obt,_SP(1,26)]
    ass=_T([[_P("Frota Macedo Engenharia LTDA",_PS("a1",alignment=_TC,fontSize=8.5,textColor=_GRAY)),
             _P("Aceite do Cliente — Nome / Data",_PS("a2",alignment=_TC,fontSize=8.5,textColor=_GRAY))]],colWidths=[(W-10)/2,(W-10)/2])
    ass.setStyle(_TS([("LINEABOVE",(0,0),(0,0),0.6,_GRAY),("LINEABOVE",(1,0),(1,0),0.6,_GRAY),("TOPPADDING",(0,0),(-1,-1),4),("LEFTPADDING",(0,0),(0,0),20),("RIGHTPADDING",(1,0),(1,0),20)]))
    el+=[ass]
    def _rod(c,dd):
        c.saveState(); c.setFont("Helvetica",7.5); c.setFillColor(_GRAY)
        c.drawCentredString(_A4[0]/2,9*_mm,"Frota Macedo Engenharia LTDA  •  CNPJ 27.363.223/0001-70"); c.restoreState()
    doc.build(el,onFirstPage=_rod,onLaterPages=_rod)
    return buf.getvalue()

# ---- gerador DOCX (usa o modelo Word do usuário como base) ----
def _dx_set(cell, lines):
    p=cell.paragraphs[0]; base=p.runs[0] if p.runs else None
    bold=base.bold if base else None; size=base.font.size if base else None; name=base.font.name if base else None
    for ex in cell.paragraphs[1:]: ex._element.getparent().remove(ex._element)
    for r in list(p.runs): r._element.getparent().remove(r._element)
    for i,ln in enumerate(lines):
        if i>0: p.add_run().add_break()
        rr=p.add_run(ln)
        if bold is not None: rr.bold=bold
        if size is not None: rr.font.size=size
        if name is not None: rr.font.name=name
def gera_orcamento_docx(d):
    from docx import Document
    doc=Document(io.BytesIO(_b64.b64decode(_MODELO_DOCX_B64))); T=doc.tables
    _dx_set(T[0].rows[0].cells[2],[d["data"],f"Orçamento nº {d['num']}"])
    _dx_set(T[1].rows[0].cells[0],[f"ORÇAMENTO Nº {d['num']}",f"REVISÃO {d.get('revisao',1)}"])
    _dx_set(T[2].rows[0].cells[0],[f"OBRA:  MANUTENCAO {d['loja_nome']}  #{d['num']}"])
    itbl=T[4]; total_tr=itbl.rows[-1]._tr; modelo_tr=_dcopy(itbl.rows[1]._tr)
    for row in list(itbl.rows[1:-1]): itbl._tbl.remove(row._tr)
    tg=0.0
    for it in d["itens"]:
        total_tr.addprevious(_dcopy(modelo_tr))
    for i,(row,it) in enumerate(zip(itbl.rows[1:-1],d["itens"]),1):
        vu=round(float(it["valor_unit"])*1.20,2); tot=round(vu*float(it["quant"]),2); tg+=tot
        vals=[str(i),it["descricao"],f"{float(it['quant']):.2f}".replace(".",","),it.get("unid","UN"),_rs(vu),_rs(tot)]
        for c,v in zip(row.cells,vals): _dx_set(c,[v])
    _dx_set(itbl.rows[-1].cells[-1],[_rs(tg)])
    for p in doc.paragraphs:
        if p.text.strip().lower().startswith("valor total por extenso"):
            base=p.runs[0] if p.runs else None
            for r in list(p.runs): r._element.getparent().remove(r._element)
            rr=p.add_run("Valor total por extenso:  "+_extenso_reais(tg))
            if base is not None and base.font.size: rr.font.size=base.font.size
            break
    out=io.BytesIO(); doc.save(out); return out.getvalue()

_ORC_PROMPT=(
"Você recebe um 'DOCUMENTO AUXILIAR DE VENDA - PEDIDO' (nota de material). Pode haver MAIS DE UMA nota (uma por página). "
"Para CADA nota extraia em JSON:\n"
"- ticket: número do chamado, no campo 'Observação' (aparece como 'TICKET', 'TICKER' ou '#' seguido de dígitos). Só os dígitos; se não houver, null.\n"
"- nota_numero: o 'Nº do Documento' da nota (só dígitos).\n"
"- data_nota: a data de emissão ('Dt. Emis') no formato DD/MM/AAAA.\n"
"- itens: linhas da tabela de produtos. Para cada item: descricao (nome do produto SEM o código numérico inicial e sem ' - '), "
"quant (número), unid (coluna 'Embalagem': KG, UN, M...), valor_unit (o 'Preço Unitário', número). Inclua 'SERVICO DE ENTREGA' se existir.\n"
"Responda só JSON: {\"notas\":[{\"ticket\":\"126486\",\"nota_numero\":\"18747\",\"data_nota\":\"03/08/2026\",\"itens\":[{\"descricao\":\"...\",\"quant\":1.0,\"unid\":\"UN\",\"valor_unit\":1.70}]}]}. Use ponto decimal."
)
def _ler_notas_gemini(file_bytes, mime="application/pdf"):
    """Tenta uma lista de modelos e usa o 1º que funcionar (memoriza em _ORC_GEMINI_OK)."""
    global _ORC_GEMINI_OK
    import google.generativeai as genai
    genai.configure(api_key=GEMINI_API_KEY)
    candidatos=[_ORC_GEMINI_OK] if _ORC_GEMINI_OK else list(dict.fromkeys([
        GEMINI_MODEL, "gemini-2.5-flash", "gemini-2.0-flash", "gemini-flash-latest",
        "gemini-2.5-flash-lite", "gemini-2.0-flash-lite",
    ]))
    ultimo=None
    for nome in candidatos:                  # tenta cada modelo 1x (cotas separadas), sem esperas longas
        if not nome: continue
        try:
            model=genai.GenerativeModel(nome)
            r=model.generate_content([{"mime_type":mime,"data":file_bytes},_ORC_PROMPT],
                generation_config={"response_mime_type":"application/json","temperature":0})
            _ORC_GEMINI_OK=nome
            _gemini_conta_inc()   # contabiliza 1 requisição na cota do dia
            try: return (json.loads(r.text) or {}).get("notas") or []
            except Exception: return []
        except Exception as e:
            ultimo=e; m=str(e).lower()
            if any(s in m for s in ("not available","not found","404","no longer","unsupported","is not supported")):
                continue           # modelo inexistente → próximo
            if any(s in m for s in ("429","quota","exceeded","rate limit","resource_exhausted")):
                continue           # cota nesse modelo → tenta o próximo (cotas separadas)
            raise                  # chave/rede/outro → propaga
    raise RuntimeError(f"Gemini indisponível: {ultimo}")

# ---------- CONTADOR DE COTA DO GEMINI (por dia) ----------
def _hoje_iso():
    return _hoje().isoformat()

def _gemini_conta_inc():
    """Soma +1 na cota do dia (atômico via RPC gemini_inc). Best-effort."""
    try:
        _sb_json(f"{SB_URL}/rest/v1/rpc/gemini_inc", SB_KEY, data={"p_dia":_hoje_iso()}, method="POST")
    except Exception: pass

def _gemini_cota():
    """Devolve {'dia','limite','usados','restantes'} da cota do dia."""
    usados=0
    try:
        r=_sb_json(f"{SB_URL}/rest/v1/gemini_uso?dia=eq.{_hoje_iso()}&select=usados&limit=1", SB_KEY) or []
        if r: usados=int(r[0].get("usados") or 0)
    except Exception: pass
    return {"dia":_hoje_iso(), "limite":GEMINI_RPD, "usados":usados,
            "restantes":max(0, GEMINI_RPD-usados)}

# ---------- LEITURA DE NOTAS: Gemini (imagem) + Groq (PDF) ----------
def _pdf_texto(fb):
    """Extrai o texto de um PDF (camada de texto). Vazio => PDF escaneado/imagem."""
    import tempfile
    try:
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(fb); pth=f.name
        out=subprocess.run(["pdftotext","-layout",pth,"-"],capture_output=True,text=True,timeout=60)
        try: os.unlink(pth)
        except Exception: pass
        return (out.stdout or "").strip()
    except Exception:
        return ""

GROQ_TEXT_FALLBACK = [GROQ_TEXT_MODEL, "llama-3.1-8b-instant", "openai/gpt-oss-20b"]
GROQ_VIS_FALLBACK  = [GROQ_VIS_MODEL, "meta-llama/llama-4-maverick-17b-128e-instruct"]

def _groq_post(messages, model, json_mode=True):
    """1 chamada ao Groq. Devolve o texto da resposta ou levanta RuntimeError('Groq <code>: <body>')."""
    payload={"model":model,"temperature":0,"messages":messages,"max_tokens":4096}
    if json_mode: payload["response_format"]={"type":"json_object"}
    req=urllib.request.Request("https://api.groq.com/openai/v1/chat/completions",
        data=json.dumps(payload).encode(),
        headers={"Authorization":f"Bearer {GROQ_KEY}","Content-Type":"application/json","Accept":"application/json",
                 # Cloudflare barra sem UA de navegador (erro 1010 browser_signature_banned)
                 "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                              "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"})
    try:
        raw=urllib.request.urlopen(req,timeout=90).read().decode()
    except urllib.error.HTTPError as e:
        body=""
        try: body=e.read().decode("utf-8","ignore")
        except Exception: pass
        raise RuntimeError(f"Groq {e.code}: {body[:400]}") from None
    return json.loads(raw)["choices"][0]["message"]["content"]

def _groq_models():
    """Lista os modelos que a conta Groq tem acesso (GET /models)."""
    req=urllib.request.Request("https://api.groq.com/openai/v1/models",
        headers={"Authorization":f"Bearer {GROQ_KEY}","Accept":"application/json",
                 "User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                              "(KHTML, like Gecko) Chrome/124.0 Safari/537.36"})
    raw=urllib.request.urlopen(req,timeout=30).read().decode()
    return [m.get("id") for m in (json.loads(raw).get("data") or []) if m.get("id")]

def _groq_chat(messages, modelos):
    """Tenta uma lista de modelos; no 1º que responder, devolve o texto. Guarda o último erro."""
    ultimo=None
    for m in modelos:
        if not m: continue
        try:
            return _groq_post(messages, m, json_mode=True)
        except RuntimeError as e:
            s=str(e); ultimo=e
            # modelo não aceita response_format json -> tenta sem
            if "400" in s and ("json" in s.lower() or "response_format" in s.lower()):
                try: return _groq_post(messages, m, json_mode=False)
                except RuntimeError as e2: ultimo=e2
            # modelo inexistente/desativado -> próximo
            if any(x in s.lower() for x in ("not found","does not exist","decommission","not supported","400","404")):
                continue
            raise
    raise (ultimo or RuntimeError("Groq: nenhum modelo respondeu"))

def _notas_de_conteudo(c):
    d=_parse_json(c)
    if isinstance(d,list): return d
    if isinstance(d,dict):
        if isinstance(d.get("notas"),list): return d["notas"]
        if d.get("itens"): return [d]              # formato de nota única
    return []

def _groq_notas_texto(texto):
    """Lê as notas a partir do TEXTO do PDF (grátis, sem gastar cota do Gemini)."""
    msgs=[{"role":"system","content":_ORC_PROMPT},
          {"role":"user","content":"TEXTO EXTRAÍDO DO PDF (pode conter mais de uma nota):\n\n"+texto[:24000]}]
    return _notas_de_conteudo(_groq_chat(msgs, GROQ_TEXT_FALLBACK))

def _pdf_png_bytes(fb):
    """Rasteriza a 1ª página do PDF em PNG (o Groq visão só aceita imagem, não PDF)."""
    import tempfile, glob
    d=tempfile.mkdtemp()
    pth=os.path.join(d,"n.pdf"); open(pth,"wb").write(fb)
    pref=os.path.join(d,"pg")
    subprocess.run(["pdftoppm","-r","170","-png","-f","1","-l","1",pth,pref],capture_output=True,timeout=60)
    pngs=sorted(glob.glob(pref+"*.png"))
    return open(pngs[0],"rb").read() if pngs else None

def _img_png_b64(fb, mime, nome):
    """Normaliza qualquer nota (PDF escaneado ou imagem) para PNG base64 que o Groq aceita."""
    import base64, io
    from PIL import Image
    is_pdf=(nome or "").lower().endswith(".pdf") or mime=="application/pdf"
    png = _pdf_png_bytes(fb) if is_pdf else fb
    if not png: return None
    im=Image.open(io.BytesIO(png)).convert("RGB")
    if im.width>1600: im=im.resize((1600,int(im.height*1600/im.width)))
    buf=io.BytesIO(); im.save(buf,"PNG")
    return base64.b64encode(buf.getvalue()).decode()

def _groq_notas_img_bytes(fb, mime, nome=""):
    """PDF escaneado/imagem -> converte para PNG e manda ao Groq (visão)."""
    b64=_img_png_b64(fb, mime, nome)
    if not b64: raise RuntimeError("Groq visão: não consegui rasterizar a página")
    msgs=[{"role":"user","content":[
        {"type":"text","text":_ORC_PROMPT},
        {"type":"image_url","image_url":{"url":f"data:image/png;base64,{b64}"}}]}]
    return _notas_de_conteudo(_groq_chat(msgs, GROQ_VIS_FALLBACK))

OCR_MAX_PAG = int(os.environ.get("OCR_MAX_PAG", "3"))   # nº de páginas por nota que passam pelo OCR
OCR_DPI     = os.environ.get("OCR_DPI", "200")          # 200 dpi (limitado a poucas páginas -> continua rápido)
def _ocr_texto(fb, mime, nome, max_pag=None):
    """OCR local (tesseract) da nota escaneada/imagem -> texto. Sem IA e sem cota.
       Leve e limitado (poucas páginas, dpi baixo) para não travar o lote."""
    import tempfile, glob, io as _io
    max_pag = max_pag or OCR_MAX_PAG
    d=tempfile.mkdtemp(); pngs=[]
    is_pdf=(nome or "").lower().endswith(".pdf") or mime=="application/pdf"
    if is_pdf:
        pth=os.path.join(d,"n.pdf"); open(pth,"wb").write(fb)
        pref=os.path.join(d,"pg")
        try:
            subprocess.run(["pdftoppm","-r",OCR_DPI,"-png","-f","1","-l",str(max_pag),pth,pref],
                           capture_output=True,timeout=90)
        except Exception: pass
        pngs=sorted(glob.glob(pref+"*.png"))
    else:
        try:
            from PIL import Image
            im=Image.open(_io.BytesIO(fb)).convert("RGB")
            if im.width>2000: im=im.resize((2000,int(im.height*2000/im.width)))
            p=os.path.join(d,"img.png"); im.save(p,"PNG"); pngs=[p]
        except Exception: pngs=[]
    env=dict(os.environ, OMP_THREAD_LIMIT="1")      # tesseract 1 thread -> previsível, não engasga a CPU
    textos=[]
    for pg in pngs:
        _prep_ocr(pg)                               # pré-processa a imagem (melhora o OCR do scan)
        txt=""
        for lang in ("por","eng"):                  # tenta português; cai p/ inglês se faltar o pacote
            try:
                r=subprocess.run(["tesseract",pg,"stdout","-l",lang,"--psm","6"],
                                 capture_output=True,text=True,timeout=45,env=env)
                if r.returncode==0 and (r.stdout or "").strip(): txt=r.stdout; break
            except Exception: pass
        if txt: textos.append(txt)
    return "\f".join(textos).strip()

def _prep_ocr(path):
    """Limpa a imagem antes do OCR: tons de cinza, contraste e upscale. Melhora scans/fotos."""
    try:
        from PIL import Image, ImageOps, ImageFilter
        im=Image.open(path).convert("L")            # tons de cinza
        im=ImageOps.autocontrast(im, cutoff=1)      # estica o contraste
        if im.width<1800:                           # amplia imagens pequenas (mais nitidez p/ o OCR)
            f=1800/im.width; im=im.resize((1800,int(im.height*f)))
        im=im.filter(ImageFilter.SHARPEN)
        im.save(path)
    except Exception: pass

class _CotaExcedida(Exception):
    """Gemini indisponível por limite diário — pula a nota e avisa."""
    pass

def _num_br(s):
    s=re.sub(r"[^\d,.\-]","",str(s or "").strip())
    if not s: return 0.0
    if "," in s: s=s.replace(".","").replace(",",".")   # vírgula = decimal
    try: return float(s)
    except Exception: return 0.0

def _notas_seguras(notas):
    """Leitura 'segura' = pelo menos 1 nota, com itens, e todo item com quant>0 e valor_unit>0."""
    if not notas: return False
    for nt in notas:
        itens=nt.get("itens") or []
        if not itens: return False
        for it in itens:
            if _num_br(it.get("quant"))<=0 or _num_br(it.get("valor_unit"))<=0: return False
    return True

def _reconcilia(q, vu, vt):
    """Se quant*unit não bate com o total impresso da linha (erro de OCR),
       confia no TOTAL e recalcula o unitário. Mantém o dinheiro correto."""
    if vt>0 and q>0 and abs(q*vu - vt) > max(0.02, 0.02*vt):
        return round(vt/q, 4)
    return vu

def _parse_danfe_texto(txt):
    """LEITOR (sem IA) da DANFE / Nota Fiscal Eletrônica (ex.: DISMONZA).
       Extrai nº da nota (N°. com pontos de milhar), ticket, itens (código, desc,
       NCM, CST, CFOP, UNID, QUANT, VALOR UNIT, VALOR TOTAL)."""
    T=txt or ""; up=T.upper()
    if "DADOS DO PRODUTO" not in up and "DOCUMENTO AUXILIAR DA" not in up:
        return []
    mn=re.search(r"N[ºo°]\.?\s*([\d][\d.]*\d)", T);  nota=_num_limpo(mn.group(1)) if mn else None
    mt=re.search(r"(?:TICK\w*|#)\s*[:\-]?\s*(\d{4,})", T, re.I); ticket=mt.group(1) if mt else None
    md=re.search(r"DATA DA EMISS[ÃA]O.*?(\d{2}/\d{2}/\d{4})", T, re.S|re.I); data=md.group(1) if md else None
    m0=re.search(r"DADOS DO PRODUTO", up); seg=T
    if m0:
        m1=re.search(r"C[ÁA]LCULO DO ISSQN|DADOS ADICIONAIS", up[m0.end():])
        seg=T[m0.end(): m0.end()+(m1.start() if m1 else len(T))]
    IT=re.compile(r"^\s*(\d{3,})\s+(.+?)\s+(\d{8})\s+\d{1,3}\s+\d{3,4}\s+([A-Za-zÇç][A-Za-z0-9²³ºÇç]{0,4})\s+([\d.]+,\d+)\s+([\d.]+,\d+)\s+([\d.]+,\d+)")
    itens=[]
    for ln in seg.splitlines():
        m=IT.search(ln)
        if m:
            q=_num_br(m.group(5)); vu=_num_br(m.group(6)); vt=_num_br(m.group(7))
            itens.append({"descricao":m.group(2).strip()[:120],"quant":q,"unid":m.group(4).upper(),
                          "valor_unit":_reconcilia(q,vu,vt)})
        else:
            t=ln.strip()
            if itens and t and not re.search(r"\d{8}",t) and len(t)>=3 and not re.match(r"^[\d,.\s%]+$",t) and not re.match(r"^\s*C[ÓO]DIGO",t,re.I):
                itens[-1]["descricao"]=(itens[-1]["descricao"]+" "+t)[:160]
    if not itens: return []
    o={"ticket":ticket,"nota_numero":nota,"data_nota":data,"itens":itens}
    return [o] if _notas_seguras([o]) else []

def _parse_dav_texto(txt):
    """LEITOR (sem IA) do 'DOCUMENTO AUXILIAR DE VENDA - PEDIDO' (ex.: RODRIGUES).
       Item = 'código - descrição UNID QUANT PREÇO ... VALOR TOTAL'."""
    T=txt or ""
    mn=re.search(r"N[ºo°]?\s*d[eo]?\s*Documento\s*[:\-]?\s*(\d+)", T, re.I); nota=_num_limpo(mn.group(1)) if mn else None
    mt=re.search(r"(?:Observa\w*|TICK\w*|#)\s*[:\-]?\s*(\d{4,})", T, re.I); ticket=mt.group(1) if mt else None
    md=re.search(r"Dt\.?\s*Emis\S*\s*[:\-]?\s*(\d{2}/\d{2}/\d{4})", T, re.I); data=md.group(1) if md else None
    up=T.upper(); seg=T
    mh=re.search(r"QUANTIDADE|PRE[ÇC]O UNIT|PREGO UNIT|EMBALAGEM|MBALAGEM", up)
    me=re.search(r"TOTAL BRUTO|TOTAL A PAGAR|PLANO DE PAGAMENTO", up)
    if mh: seg=T[mh.end(): (me.start() if me else len(T))]
    IT=re.compile(r"^\s*(\d{5,})\s*-\s*(.+?)\s+([A-Za-zÇç]{1,4})\s+([\d.]+,\d+)\s+([\d.]+,\d+)")
    itens=[]
    for ln in seg.splitlines():
        m=IT.search(ln)
        if m:
            q=_num_br(m.group(4)); vu=_num_br(m.group(5))
            dec=re.findall(r"[\d.]+,\d+", ln); vt=_num_br(dec[-1]) if dec else 0
            itens.append({"descricao":m.group(2).strip()[:120],"quant":q,"unid":m.group(3).upper(),
                          "valor_unit":_reconcilia(q,vu,vt)})
        else:
            t=ln.strip()
            if itens and t and not re.match(r"^\s*\d{5,}\s*-",t) and len(t)>=3 and not re.match(r"^[\d,.\s%]+$",t):
                itens[-1]["descricao"]=(itens[-1]["descricao"]+" "+t)[:160]
    if not itens: return []
    o={"ticket":ticket,"nota_numero":nota,"data_nota":data,"itens":itens}
    return [o] if _notas_seguras([o]) else []

def _parse_nota_local(txt):
    """Leitor local sem IA: tenta DANFE (NF-e) e depois DAV (pedido)."""
    return _parse_danfe_texto(txt) or _parse_dav_texto(txt)

def _data_br_any(v):
    """Aceita datetime ou texto e devolve DD/MM/AAAA."""
    if v is None: return None
    try:
        if hasattr(v,"strftime"): return v.strftime("%d/%m/%Y")
    except Exception: pass
    m=re.search(r"(\d{2})/(\d{2})/(\d{4})", str(v))
    if m: return m.group(0)
    m=re.search(r"(\d{4})-(\d{2})-(\d{2})", str(v))
    return f"{m.group(3)}/{m.group(2)}/{m.group(1)}" if m else None

def _parse_excel_nota(fb):
    """Leitor da nota digitada em EXCEL (quando o PDF está ilegível).
       Cabeçalho: FORNECEDOR, CNPJ, NUMERO DA NOTA, DATA DE EMISSAO, COD. PRODUTO,
       DESCRICAO DO ITEM, UNIDADE, QUANTIDADE, PRECO UNITARIO, PRECO TOTAL.
       1 linha por item; a linha 'TOTAL DA NOTA' é ignorada."""
    import openpyxl, io as _io
    try:
        wb=openpyxl.load_workbook(_io.BytesIO(fb), data_only=True)
    except Exception:
        return []
    ws=wb.worksheets[0]
    rows=[list(r) for r in ws.iter_rows(values_only=True)]
    if not rows: return []
    hdr=[str(c or "").strip().upper() for c in rows[0]]
    def col(*names):
        for n in names:
            for i,h in enumerate(hdr):
                if n in h: return i
        return None
    ci_forn=col("FORNECEDOR"); ci_cnpj=col("CNPJ"); ci_num=col("NUMERO DA NOTA","NUMERO","NOTA")
    ci_data=col("DATA"); ci_desc=col("DESCRICAO","DESCRIÇÃO"); ci_un=col("UNIDADE","UNID")
    ci_q=col("QUANTIDADE","QUANT"); ci_vu=col("PRECO UNITARIO","PREÇO UNIT","VALOR UNIT")
    ci_vt=col("PRECO TOTAL","PREÇO TOTAL","VALOR TOTAL")
    if ci_desc is None or ci_q is None: return []
    g=lambda r,i: (r[i] if (i is not None and i<len(r)) else None)
    fornecedor=cnpj=nota=data=None; itens=[]
    for r in rows[1:]:
        if "TOTAL" in str(g(r,ci_q) or "").upper(): continue     # linha de total
        desc=g(r,ci_desc)
        if not desc or not str(desc).strip(): continue
        q=_num_br(g(r,ci_q)); vu=_num_br(g(r,ci_vu)); vt=_num_br(g(r,ci_vt))
        if vu<=0 and vt>0 and q>0: vu=round(vt/q,4)
        if q<=0 or vu<=0: continue
        itens.append({"descricao":str(desc).strip()[:120],"unid":str(g(r,ci_un) or "UN").strip()[:6],
                      "quant":q,"valor_unit":_reconcilia(q,vu,vt)})
        fornecedor=fornecedor or (str(g(r,ci_forn)).strip() if g(r,ci_forn) else None)
        cnpj=cnpj or (re.sub(r"\D","",str(g(r,ci_cnpj))) if g(r,ci_cnpj) else None)
        nota=nota or (_num_limpo(g(r,ci_num)) if g(r,ci_num) else None)
        data=data or _data_br_any(g(r,ci_data))
    if not itens: return []
    return [{"fornecedor":fornecedor,"cnpj":cnpj,"nota_numero":nota,"data_nota":data,"itens":itens}]

def _via_gemini(fb, mime, it, st):
    """Gemini como ÚLTIMO recurso, respeitando a cota diária."""
    if st.get("gemini_bloqueado"): raise _CotaExcedida()
    if _gemini_cota()["restantes"]<=0:
        st["gemini_bloqueado"]=True; raise _CotaExcedida()
    it["etapa"]="gemini"; it["status"]="lendo com Gemini"
    try:
        return _ler_notas_gemini(fb, mime)
    except Exception as e:
        if any(s in str(e).lower() for s in ("429","quota","exceeded","resource_exhausted","rate limit")):
            st["gemini_bloqueado"]=True; raise _CotaExcedida()
        raise

# ---------- CONFERÊNCIA POR TOTAL (garante que a leitura fechou) ----------
def _soma_itens(notas):
    s=0.0
    for nt in (notas or []):
        for it in (nt.get("itens") or []):
            s += _num_br(it.get("valor_unit"))*_num_br(it.get("quant"))
    return round(s,2)

def _bate_total(notas, texto):
    """Confere a soma dos itens contra os valores 'total' impressos na nota.
       True = bate (confiável) · False = há total mas NÃO bate (suspeito) · None = sem total p/ conferir."""
    soma=_soma_itens(notas)
    if soma<=0: return False
    toks=[_num_br(x) for x in re.findall(r"\d{1,3}(?:\.\d{3})*,\d{2}\b", texto or "")]
    toks=[t for t in toks if t>0]
    if not toks: return None
    tol=max(0.02, 0.01*soma)
    return any(abs(t-soma)<=tol for t in toks)

def _leitura_ok(notas, texto):
    """Aceita a leitura se ela é segura E a soma bate com um total impresso (ou não há total p/ conferir)."""
    if not _notas_seguras(notas): return False
    return _bate_total(notas, texto) is not False

def _ler_notas_rota(fb, mime, nome, it, st):
    """CASCATA DE LEITURA (barata/confiável -> cara/último caso):
       Camada 0: vira TEXTO — PDF digital (pdftotext) ou OCR (scan/imagem, com pré-processamento).
       Camada 1: parsers locais DANFE/DAV (sem IA), conferidos pelo total da nota.
       Camada 2: Groq no texto (grátis, sem cota diária), também conferido pelo total.
       Camada 3: Gemini na imagem (último recurso; respeita a cota diária).
       Devolve (notas, reader)."""
    is_pdf = (nome or "").lower().endswith(".pdf") or mime=="application/pdf"

    # ---- Camada 0: obter o texto ----
    texto=""; origem="digital"
    if is_pdf:
        it["etapa"]="pdf"; it["status"]="lendo com pdf reader"
        texto=_pdf_texto(fb)
    if len(texto) < 40:                               # imagem, ou PDF sem camada de texto -> OCR
        it["etapa"]="ocr"; it["status"]="extraindo texto (OCR)"
        try: texto=_ocr_texto(fb, mime, nome); origem="ocr"
        except Exception as e: it["diag"]=f"OCR: {str(e)[:150]}"; texto=""

    # ---- Camada 1: parsers locais (sem IA) ----
    if texto and len(texto)>=30:
        notas=_parse_nota_local(texto)
        if notas and _leitura_ok(notas, texto):
            _formato_registra(texto, _qual_formato(texto), notas[0].get("nota_numero"))   # aprende o layout
            return notas, ("pdf" if origem=="digital" else "ocr-pdf")
        if notas and not _notas_seguras(notas):
            it["diag"]="leitor local: itens incompletos"
        elif notas:
            it["diag"]="leitor local: soma não bateu com o total"

        # ---- Camada 2: Groq no texto (grátis) ----
        if GROQ_KEY:
            it["etapa"]="groq"; it["status"]=("lendo com Groq (OCR)" if origem=="ocr" else "lendo com Groq")
            try:
                g=_groq_notas_texto(texto)
                if g and _leitura_ok(g, texto):
                    _formato_registra(texto, "groq", g[0].get("nota_numero"))
                    return g, ("groq-ocr" if origem=="ocr" else "groq")
                it["diag"]="Groq leu incompleto / total não bateu"
            except Exception as e: it["diag"]=f"Groq: {str(e)[:160]}"; st["groq_erro"]=it["diag"]

    # ---- Camada 3: Gemini na imagem (último recurso) ----
    return _via_gemini(fb, mime, it, st), "gemini"

# ---------- PADRONIZAÇÃO DO NOME DA LOJA (item: "LOJA NN - NOME"; CD -> Centro de Distribuição) ----------
_LOJA_NOME_FULL = {20:"RUI BARBOSA", 23:"JÚLIO VENTURA", 101:"CENTRO DE DISTRIBUIÇÃO"}   # nomes por extenso/oficiais
def _cad_por_numero(n):
    return CAD.get(str(int(n)).zfill(2)) or CAD.get(str(int(n)))
def _loja_padrao(raw):
    """Normaliza o nome da loja para 'LOJA NN - NOME' (consolida grafias diferentes da mesma loja).
       CD vira 'Centro de Distribuição'."""
    s=(raw or "").strip()
    if not s: return "—"
    num=None
    m=re.search(r"LOJA\s*0*(\d{1,3})", s, re.I) or re.search(r"^\s*0*(\d{1,3})\b", s)
    if m: num=int(m.group(1))
    if num is None:                                  # sem número: casa por nome/apelido/nome-completo
        alvo=norm(s)
        for e in CAD.values():
            n_e=int(e.get("numero"))
            cands=[e.get("nome"), _LOJA_NOME_FULL.get(n_e)]+list(e.get("apelidos") or [])
            for c in cands:
                cn=norm(c)
                if cn and re.search(rf"\b{re.escape(cn)}\b", alvo):
                    num=n_e; break
            if num is not None: break
    if num is None: return s                          # desconhecida: devolve o original
    if num==101: return "CENTRO DE DISTRIBUIÇÃO"       # CD não é loja
    e=_cad_por_numero(num) or {}
    nome=_LOJA_NOME_FULL.get(num) or e.get("nome") or ""
    rot=f"LOJA {num:02d}" if num<100 else f"LOJA {num}"
    return f"{rot} - {nome}".strip() if nome else rot

FROTAHUB_ROOT = os.environ.get("FROTAHUB_ROOT", "/FROTAHUB")
def _find_dir(access, base, pattern):
    """Acha, em 'base', a 1ª subpasta cujo nome bate 'pattern'. base='' = raiz real do token."""
    try:
        for e in dropbox_rateio.listar_entradas(access, base):
            if e["dir"] and re.search(pattern, e["name"], re.I):
                return f"{base}/{e['name']}" if base else f"/{e['name']}"
    except Exception: pass
    return None
def _manut_base(access):
    """Acha o depto Manutenção ('2 -') — a raiz do token pode ser a conta (tem /FROTAHUB)
       OU já ser a própria FROTAHUB (App folder: os deptos ficam direto na raiz '')."""
    for root in (FROTAHUB_ROOT, ""):            # tenta /FROTAHUB e a raiz real do token
        p=_find_dir(access, root, r"^\s*2\s*-")
        if p: return p
    fh=_find_dir(access, "", r"^\s*FROTAHUB\s*$") # ou uma pasta 'FROTAHUB' na raiz real
    if fh:
        p=_find_dir(access, fh, r"^\s*2\s*-")
        if p: return p
    return MANUT_BASE
def _pasta_manut(access, n, base=None):
    base = base or _manut_base(access)
    try:
        for e in dropbox_rateio.listar_entradas(access, base):
            if e["dir"] and re.match(rf"^\s*{n}\s*-", e["name"]):
                return f"{base}/{e['name']}"
    except Exception: pass
    return None

def _orc_base(access, mb):
    """Acha a pasta-módulo dos orçamentos: a que CONTÉM a subpasta '0 - ...'.
       As numeradas 0..10 podem estar direto no depto OU dentro de um módulo
       (ex.: '2 - MANUTENÇÃO/<algo> - ORÇAMENTOS/0..10'), como o PCO fica em
       '1 - ADMINISTRATIVO/1 - PCO/...'. Procura no depto e um nível abaixo."""
    if _pasta_manut(access, 0, mb):        # numeradas direto no depto
        return mb
    try:
        for e in dropbox_rateio.listar_entradas(access, mb):
            if not e["dir"]: continue
            sub = f"{mb}/{e['name']}"
            if _pasta_manut(access, 0, sub):
                return sub
    except Exception: pass
    return mb

def _orc_diag(access):
    """Radiografia da estrutura para depurar a leitura da pasta 0."""
    def _ls(p):
        try:
            es=dropbox_rateio.listar_entradas(access,p)
            return {"dirs":[e["name"] for e in es if e["dir"]],"arqs":len([e for e in es if not e["dir"]])}
        except Exception as e:
            return {"erro":str(e)[:180]}
    d={"root":FROTAHUB_ROOT}
    d["api_root"]=_ls("")                    # raiz REAL do token Dropbox
    d["frotahub"]=_ls(FROTAHUB_ROOT)         # /FROTAHUB
    d["pco_base"]=_ls(PCO_BASE)              # âncora conhecida que funciona
    mb=_manut_base(access); d["manut_base"]=mb; d["manut"]=_ls(mb)
    ob=_orc_base(access,mb); d["orc_base"]=ob; d["orc"]=_ls(ob)
    p0=_pasta_manut(access,0,ob); d["pasta0"]=p0
    d["pasta0_arqs"]=(_ls(p0).get("arqs") if p0 else 0)
    return d

def _loja_do_ticket(ticket):
    q=urllib.parse.urlencode({"numero":f"eq.{ticket}","select":"loja,aba","limit":"1"})
    ch=_sb_json(f"{SB_URL}/rest/v1/chamados?{q}",SB_KEY) or []
    if not ch: return None
    unidade=ch[0].get("loja") or ""
    m=re.search(r"LOJA\s*0*(\d{1,3})",unidade,re.I) or re.search(r"\b0*(\d{1,3})\b",unidade)
    lj=None
    if m: lj=(_sb_json(f"{SB_URL}/rest/v1/lojas?numero=eq.{int(m.group(1))}&limit=1",SB_KEY) or [None])[0]
    if not lj and unidade:
        alvo=re.sub(r"^LOJA\s*\d*\s*-?\s*","",unidade,flags=re.I).strip()
        if alvo: lj=(_sb_json(f"{SB_URL}/rest/v1/lojas?nome=ilike.*{urllib.parse.quote(alvo[:14])}*&limit=1",SB_KEY) or [None])[0]
    return {"aba":ch[0].get("aba"),"unidade":unidade,"loja":lj}

def _slug_loja(loja, unidade):
    num=loja.get("numero") if loja else None
    nome=(loja.get("nome") if loja else None) or re.sub(r"^LOJA\s*\d*\s*-?\s*","",unidade or "",flags=re.I).strip() or "SEM_LOJA"
    slug=re.sub(r"\s+","_",nome.strip())
    if num is None: return slug
    return (f"{num:02d}" if num<100 else str(num))+"_"+slug

def _split_pdf(pdf_bytes):
    from pypdf import PdfReader, PdfWriter
    out=[]
    try:
        r=PdfReader(io.BytesIO(pdf_bytes))
        for pg in r.pages:
            w=PdfWriter(); w.add_page(pg); b=io.BytesIO(); w.write(b); out.append(b.getvalue())
    except Exception: return [pdf_bytes]
    return out or [pdf_bytes]

def _orc_registra(row):
    req=urllib.request.Request(f"{SB_URL}/rest/v1/notas_orcamento",
        data=json.dumps([row],ensure_ascii=False).encode(),method="POST",
        headers={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}","content-type":"application/json","prefer":"return=minimal"})
    try: urllib.request.urlopen(req,timeout=30)
    except urllib.error.HTTPError as e: print("notas_orcamento erro:",e.read().decode()[:200],flush=True)

def _nota_ja_gerada(ticket, nota_num, base_nome=None):
    """True se já existe orçamento gerado para este ticket/nota — evita reprocessar a MESMA nota.
    Cobre também nota SEM número (SN): nesse caso usa o nome-base do arquivo (determinístico por
    ticket+loja+nota), fechando o furo que deixava notas sem número serem rateadas várias vezes."""
    # 1) por ticket + número da nota (quando há número)
    if nota_num and str(nota_num)!="SN":
        q=urllib.parse.urlencode({"ticket":f"eq.{ticket}","nota_numero":f"eq.{nota_num}","status":"eq.gerado","select":"ticket","limit":"1"})
        try:
            if _sb_json(f"{SB_URL}/rest/v1/notas_orcamento?{q}",SB_KEY): return True
        except Exception: pass
    # 2) por nome-base do arquivo gerado (cobre SN e reprocessamento do mesmo arquivo)
    if base_nome:
        like=urllib.parse.quote(f"*{base_nome}*")
        q2=f"status=eq.gerado&or=(arquivo_pdf.ilike.{like},arquivo_doc.ilike.{like})&select=ticket&limit=1"
        try:
            if _sb_json(f"{SB_URL}/rest/v1/notas_orcamento?{q2}",SB_KEY): return True
        except Exception: pass
    return False

@app.post("/orc/listar")
def orc_listar(request: Request):
    from fastapi import HTTPException
    exige(request,"CONFERIR_LISTA_ORCAMENTOS")
    if not dropbox_rateio.ativo(): raise HTTPException(500,"Dropbox não configurado")
    access=dropbox_rateio.obter_token()
    _mb=_manut_base(access); _ob=_orc_base(access,_mb)
    ORC_NOTAS=_pasta_manut(access,0,_ob) or (_ob + "/0 - NOTAS PARA ORCAMENTO (COLOCAR AQUI)")
    arqs=[a for a in dropbox_rateio.listar(access,ORC_NOTAS) if a.lower().endswith((".pdf",".jpg",".jpeg",".png"))]
    out={"pasta":ORC_NOTAS,"total":len(arqs),"arquivos":sorted(arqs)}
    if not arqs: out["diag"]=_orc_diag(access)   # ajuda a achar a estrutura quando vem vazio
    return out

_ST_MAP={"1":"Aberto","3":"Arquivado","5":"Vistoriado","6":"Em execução","7":"Executado"}
def _chamados_query(q="", aba="", status="", desde="", ate=""):
    parts=["select=numero,aba,loja,status,tipo_predial,prioridade,solicitante,responsavel,data_criacao,prazo",
           "order=data_criacao.desc","limit=3000"]
    if aba:    parts.append(f"aba=eq.{urllib.parse.quote(aba)}")
    if desde:  parts.append(f"data_criacao=gte.{desde}")
    if ate:    parts.append(f"data_criacao=lte.{ate}")
    if q and len(q)>=3:
        parts.append("or="+urllib.parse.quote(f"(numero.ilike.*{q}*,loja.ilike.*{q}*,solicitante.ilike.*{q}*)"))
    rows=_sb_json(f"{SB_URL}/rest/v1/chamados?"+"&".join(parts),SB_KEY) or []
    for r in rows:                       # normaliza status numérico legado -> texto
        s=str(r.get("status") or "")
        if s in _ST_MAP: r["status"]=_ST_MAP[s]
        r["loja"]=_loja_padrao(r.get("loja"))   # padroniza 'LOJA NN - NOME' / CD -> Centro de Distribuição
    if status:                           # filtra por status já normalizado (pega legado e novo)
        alvo=status.strip().lower()
        rows=[r for r in rows if (r.get("status") or "").strip().lower()==alvo]
    return rows

@app.get("/orc/chamados")
def orc_chamados(request: Request, q: str="", aba: str="", status: str="", desde: str="", ate: str=""):
    from fastapi import HTTPException
    exige(request,"CHAMADOS_TRILOGO")
    try: rows=_chamados_query(q,aba,status,desde,ate)
    except Exception as e: raise HTTPException(500,f"chamados: {e}")
    return {"itens":rows,"total":len(rows)}

def _abacurta(a): return "Instalações" if (a or "").upper().startswith("INST") else ("Civil" if (a or "").upper().startswith("CIV") else (a or ""))

@app.get("/orc/chamado_det")
def orc_chamado_det(request: Request, numero: str="", aba: str=""):
    from fastapi import HTTPException
    exige(request,"CHAMADOS_TRILOGO")
    parts=["select=*","limit=1",f"numero=eq.{urllib.parse.quote(numero)}"]
    if aba: parts.append(f"aba=eq.{urllib.parse.quote(aba)}")
    rows=_sb_json(f"{SB_URL}/rest/v1/chamados?"+"&".join(parts),SB_KEY) or []
    if not rows: raise HTTPException(404,"chamado não encontrado")
    r=rows[0]; s=str(r.get("status") or "")
    if s in _ST_MAP: r["status"]=_ST_MAP[s]
    r["loja"]=_loja_padrao(r.get("loja"))
    return r

@app.get("/orc/chamados_pdf")
def orc_chamados_pdf(request: Request, q: str="", aba: str="", status: str="", desde: str="", ate: str=""):
    from fastapi import HTTPException
    exige(request,"CHAMADOS_TRILOGO")
    rows=_chamados_query(q,aba,status,desde,ate)
    linhas=[[r.get("numero") or "", _abacurta(r.get("aba")), (r.get("loja") or "")[:30], r.get("status") or "",
             (r.get("tipo_predial") or "")[:22], r.get("prioridade") or "", _data_br(r.get("data_criacao")), _data_br(r.get("prazo"))] for r in rows]
    sub=" · ".join([x for x in [_abacurta(aba) if aba else "", status if status else "", (f"{desde}→{ate}" if (desde or ate) else "")] if x])
    pdf=_lista_pdf("Chamados do Trílogo", ["Nº","Conta","Loja","Status","Tipo","Prior.","Criado","Prazo"], linhas,
                   subtitulo=f"{len(rows)} chamado(s)"+(f" · {sub}" if sub else ""))
    return Response(content=pdf, media_type="application/pdf", headers={"Content-Disposition":'attachment; filename="chamados_trilogo.pdf"'})

@app.get("/orc/chamados_xlsx")
def orc_chamados_xlsx(request: Request, q: str="", aba: str="", status: str="", desde: str="", ate: str=""):
    from fastapi import HTTPException
    exige(request,"CHAMADOS_TRILOGO")
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    rows=_chamados_query(q,aba,status,desde,ate)
    wb=openpyxl.Workbook(); ws=wb.active; ws.title="Chamados"
    heads=["Nº","Conta","Loja","Status","Tipo predial","Prioridade","Solicitante","Responsável","Criado em","Prazo"]
    ws.append(heads)
    for c in ws[1]:
        c.font=Font(name="Arial",bold=True,color="FFFFFF"); c.fill=PatternFill("solid",fgColor="7A1517"); c.alignment=Alignment(horizontal="center")
    for r in rows:
        ws.append([r.get("numero"),_abacurta(r.get("aba")),r.get("loja"),r.get("status"),r.get("tipo_predial"),
                   r.get("prioridade"),r.get("solicitante"),r.get("responsavel"),_data_br(r.get("data_criacao")),_data_br(r.get("prazo"))])
    for i,w in enumerate([10,12,30,14,20,12,26,22,12,12],1): ws.column_dimensions[get_column_letter(i)].width=w
    ws.freeze_panes="A2"
    buf=io.BytesIO(); wb.save(buf)
    return Response(content=buf.getvalue(), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition":'attachment; filename="chamados_trilogo.xlsx"'})

# ================= ATUALIZAÇÃO REAL DO TRÍLOGO (dispara o robô no GitHub) =================
def _github_dispatch(modo="rotina"):
    """Aciona o workflow do robô (login real no Trílogo -> Supabase) via workflow_dispatch."""
    if not (GH_TOKEN and GH_REPO):
        raise RuntimeError("Configure GITHUB_TOKEN e GH_REPO no Render.")
    url=f"https://api.github.com/repos/{GH_REPO}/actions/workflows/{urllib.parse.quote(GH_WORKFLOW)}/dispatches"
    data=json.dumps({"ref":GH_REF,"inputs":{"modo":modo}}).encode()
    req=urllib.request.Request(url,data=data,method="POST",headers={
        "Authorization":f"Bearer {GH_TOKEN}","Accept":"application/vnd.github+json",
        "X-GitHub-Api-Version":"2022-11-28","User-Agent":"frotahub-motor","Content-Type":"application/json"})
    urllib.request.urlopen(req,timeout=30)   # 204 No Content em caso de sucesso
    return True

# --- Lançar orçamentos no Trílogo (robô próprio via GitHub Actions) ---
GH_WF_LANCAR = os.environ.get("GH_WORKFLOW_LANCAR", "trilogo-lancar.yml")
ROBOT_KEY    = os.environ.get("ROBOT_KEY", "")   # segredo compartilhado motor <-> robô

def _gh_dispatch_wf(workflow, inputs=None):
    if not (GH_TOKEN and GH_REPO): raise RuntimeError("Configure GITHUB_TOKEN e GH_REPO no Render.")
    url=f"https://api.github.com/repos/{GH_REPO}/actions/workflows/{urllib.parse.quote(workflow)}/dispatches"
    data=json.dumps({"ref":GH_REF,"inputs":(inputs or {})}).encode()
    req=urllib.request.Request(url,data=data,method="POST",headers={
        "Authorization":f"Bearer {GH_TOKEN}","Accept":"application/vnd.github+json",
        "X-GitHub-Api-Version":"2022-11-28","User-Agent":"frotahub-motor","Content-Type":"application/json"})
    urllib.request.urlopen(req,timeout=30); return True

def _lancar_itens(access):
    """Orçamentos a lançar: pastas 1 (normais) e 4 (rateio). Casa cada arquivo com o
    notas_orcamento pra pegar ticket/valor/aba. Só entra o que ainda está em 1/4."""
    ob=_orc_base(access,_manut_base(access))
    P1=_pasta_manut(access,1,ob); P4=_pasta_manut(access,4,ob)
    orcs=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?status=eq.gerado&select=ticket,aba,valor_orcamento,arquivo_pdf,rateio,loja_nome,lancado&limit=8000",SB_KEY) or []
    def _match(nome, origem):
        alvo=f"{origem}/{nome}"
        for o in orcs:
            if str(o.get("arquivo_pdf") or "")==alvo: return o
        for o in orcs:                                   # fallback: pelo ticket no nome
            tk=str(o.get("ticket") or "")
            if tk and (f"_{tk}_NOTA_" in nome or nome.endswith(f"_{tk}.pdf")): return o
        return None
    itens=[]
    for origem,base in (("1",P1),("4",P4)):
        if not base: continue
        for e in dropbox_rateio.listar_entradas(access,base):
            if e["dir"] or not e["name"].lower().endswith(".pdf"): continue
            o=_match(e["name"],origem) or {}
            itens.append({"arquivo":e["name"],"origem":origem,
                "ticket":str(o.get("ticket") or ""),"aba":o.get("aba") or "",
                "valor":round(_numf(o.get("valor_orcamento")),2) if o.get("valor_orcamento") is not None else None,
                "loja":_loja_padrao(o.get("loja_nome")),"rateio":bool(o.get("rateio")),
                "lancado":bool(o.get("lancado"))})
    itens.sort(key=lambda x:(x["origem"],x["arquivo"].lower()))
    return itens

@app.get("/orc/lancar_worklist")
def orc_lancar_worklist(request: Request):
    from fastapi import HTTPException
    exige(request,"GERAR_ORCAMENTOS")
    if not dropbox_rateio.ativo(): raise HTTPException(500,"Dropbox não configurado")
    itens=_lancar_itens(dropbox_rateio.obter_token())
    return {"total":len(itens),"itens":itens}

_LANCAR_PROG={"itens":{},"atualizado":None}   # {arquivo: {status,pct}} — progresso ao vivo do robô

@app.get("/orc/lancar_ver")
def orc_lancar_ver(request: Request, origem: str="", nome: str=""):
    """Visualiza o PDF do orçamento (pasta 1 ou 4) pelo site."""
    from fastapi import HTTPException
    exige(request,"GERAR_ORCAMENTOS")
    if origem not in ("1","4") or not _seg_ok(nome): raise HTTPException(400,"parâmetros inválidos")
    access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
    base=_pasta_manut(access,int(origem),ob)
    data=dropbox_rateio.baixar(access,f"{base}/{nome}") if base else None
    if data is None: raise HTTPException(404,"pdf não encontrado")
    return Response(content=data,media_type="application/pdf",
                    headers={"Content-Disposition":f'inline; filename="{urllib.parse.quote(nome)}"'})

@app.post("/orc/lancar_disparar")
async def orc_lancar_disparar(request: Request):
    """alvo vazio = LANÇAR TODOS (exige PIN); alvo 'origem/arquivo' = lançar só um (sem PIN)."""
    from fastapi import HTTPException
    u,p=exige(request,"GERAR_ORCAMENTOS")
    b={}
    try: b=await request.json()
    except Exception: pass
    alvo=(b.get("alvo") or "").strip()
    if not alvo:                                   # todos -> exige PIN do usuário
        if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")):
            raise HTTPException(403,"PIN incorreto")
    _LANCAR_PROG["itens"]={}; _LANCAR_PROG["atualizado"]=_agora().isoformat()   # zera o progresso
    try: _gh_dispatch_wf(GH_WF_LANCAR,{"modo":"lancar","alvo":alvo})
    except urllib.error.HTTPError as e: raise HTTPException(500,f"GitHub {e.code}: {e.read().decode()[:200]}")
    except Exception as e: raise HTTPException(500,f"disparo falhou: {e}")
    log_frotahub(u["id"],p.get("papel"),"GERAR_ORCAMENTOS","DISPAROU_LANCAR", alvo or "TODOS")
    return {"ok":True,"alvo":alvo}

@app.get("/orc/lancar_status")
def orc_lancar_status(request: Request):
    exige(request,"GERAR_ORCAMENTOS")
    out={"progresso":_LANCAR_PROG.get("itens",{}),"atualizado":_LANCAR_PROG.get("atualizado"),"gh":None}
    if GH_TOKEN and GH_REPO:
        url=f"https://api.github.com/repos/{GH_REPO}/actions/workflows/{urllib.parse.quote(GH_WF_LANCAR)}/runs?per_page=1"
        req=urllib.request.Request(url,headers={"Authorization":f"Bearer {GH_TOKEN}","Accept":"application/vnd.github+json",
            "X-GitHub-Api-Version":"2022-11-28","User-Agent":"frotahub-motor"})
        try:
            r=json.loads(urllib.request.urlopen(req,timeout=20).read().decode())
            run=(r.get("workflow_runs") or [None])[0]
            if run: out["gh"]={"estado":run.get("status"),"conclusao":run.get("conclusion"),"url":run.get("html_url")}
        except Exception: pass
    return out

# ---- endpoints do ROBÔ (autenticados por ROBOT_KEY, não por login) ----
def _robot_ok(request):
    from fastapi import HTTPException
    k=request.headers.get("x-robot-key") or request.query_params.get("key")
    if not ROBOT_KEY or k!=ROBOT_KEY: raise HTTPException(403,"robot key inválida")

@app.get("/robot/lancar_worklist")
def robot_lancar_worklist(request: Request):
    _robot_ok(request)
    access=dropbox_rateio.obter_token()
    return {"itens":_lancar_itens(access)}

@app.post("/robot/lancar_progresso")
async def robot_lancar_progresso(request: Request):
    """O robô reporta o andamento de cada orçamento: {arquivo,status,pct}."""
    _robot_ok(request)
    b=await request.json()
    lista=b.get("itens") if isinstance(b.get("itens"),list) else [b]
    for it in lista:
        a=it.get("arquivo")
        if a: _LANCAR_PROG["itens"][a]={"status":it.get("status") or "","pct":int(it.get("pct") or 0)}
    _LANCAR_PROG["atualizado"]=_agora().isoformat()
    return {"ok":True}

@app.get("/robot/lancar_pdf")
def robot_lancar_pdf(request: Request, origem: str="", nome: str=""):
    from fastapi import HTTPException
    _robot_ok(request)
    if origem not in ("1","4") or not _seg_ok(nome): raise HTTPException(400,"parâmetros inválidos")
    access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
    base=_pasta_manut(access,int(origem),ob)
    data=dropbox_rateio.baixar(access,f"{base}/{nome}") if base else None
    if data is None: raise HTTPException(404,"pdf não encontrado")
    return Response(content=data,media_type="application/pdf")

@app.post("/robot/lancar_ok")
async def robot_lancar_ok(request: Request):
    """O robô chama após o custo entrar no Trílogo: move 1->2 / 4->5 e marca lançado.
    Idempotente — se o arquivo já saiu da origem, não faz nada."""
    from fastapi import HTTPException
    _robot_ok(request)
    b=await request.json(); origem=str(b.get("origem")); nome=b.get("nome")
    if origem not in ("1","4") or not _seg_ok(nome): raise HTTPException(400,"parâmetros inválidos")
    _LANCAR_PROG["itens"][nome]={"status":"lançado","pct":100}   # marca pronto no progresso
    return _mover_marcar_lancado(origem, nome)

def _mover_marcar_lancado(origem, nome):
    """Move o PDF (1->2 / 4->5) e marca 'lancado' no banco. Idempotente."""
    from fastapi import HTTPException
    access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
    dst_n=2 if origem=="1" else 5
    src=_pasta_manut(access,int(origem),ob); dst=_pasta_manut(access,dst_n,ob)
    if not src or not dst: raise HTTPException(500,"pasta de origem/destino não encontrada")
    try:
        body=json.dumps({"lancado":True,"lancado_em":_agora().isoformat()}).encode()
        rq=urllib.request.Request(f"{SB_URL}/rest/v1/notas_orcamento?status=eq.gerado&or=(arquivo_pdf.eq.{urllib.parse.quote(origem+'/'+nome)},arquivo_doc.ilike.*{urllib.parse.quote(os.path.splitext(nome)[0])}*)",
            data=body,method="PATCH",headers={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}","content-type":"application/json","prefer":"return=minimal"})
        urllib.request.urlopen(rq,timeout=20)
    except Exception: pass
    moved=False
    try:
        dropbox_rateio.mover(access,f"{src}/{nome}",f"{dst}/{nome}"); moved=True
    except Exception as e:
        if dropbox_rateio.baixar(access,f"{src}/{nome}") is None: moved=True
        else: raise HTTPException(500,f"mover: {str(e)[:120]}")
    return {"ok":True,"movido":moved}

def _mover_volta(origem, nome):
    """Desfaz o move: volta o PDF de 2->1 / 5->4 (NÃO mexe no banco aqui)."""
    from fastapi import HTTPException
    access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
    dst_n=2 if origem=="1" else 5
    p_orig=_pasta_manut(access,int(origem),ob); p_lanc=_pasta_manut(access,dst_n,ob)
    if not p_orig or not p_lanc: raise HTTPException(500,"pasta não encontrada")
    try:
        dropbox_rateio.mover(access,f"{p_lanc}/{nome}",f"{p_orig}/{nome}"); return True
    except Exception as e:
        if dropbox_rateio.baixar(access,f"{p_lanc}/{nome}") is None: return True   # já voltou
        raise HTTPException(500,f"mover volta: {str(e)[:120]}")

@app.get("/orc/conferencia_movidas")
def orc_conferencia_movidas(request: Request, desde: str=""):
    """Lista os orçamentos marcados como lançados (candidatos a terem sido movidos
    pela conferência). Se 'desde' for informado, filtra por lancado_em >= desde."""
    from fastapi import HTTPException
    u,p=exige(request,"GERAR_ORCAMENTOS")
    if p.get("nivel") not in ("builder","gerente"): raise HTTPException(403,"apenas builder e gerente")
    q=("lancado=eq.true&select=id,ticket,arquivo_pdf,valor_orcamento,lancado_em,loja_nome,aba,rateio"
       "&order=lancado_em.desc&limit=5000")
    if desde: q+=f"&lancado_em=gte.{urllib.parse.quote(desde)}"
    rows=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?{q}",SB_KEY) or []
    for r in rows:
        ap=str(r.get("arquivo_pdf") or ""); r["origem"]=ap.split("/")[0] if "/" in ap else ""
        r["pasta_atual"]=("2 - ORÇAMENTOS LANÇADOS" if r["origem"]=="1"
                          else "5 - ORÇAMENTOS DE RATEIO LANÇADOS" if r["origem"]=="4" else "?")
    return {"total":len(rows),"itens":rows}

@app.post("/orc/conferencia_desfazer")
async def orc_conferencia_desfazer(request: Request):
    """Desfaz uma rodada de conferência: volta os PDFs de 2->1 / 5->4 e zera 'lancado'.
    Escopo obrigatório por 'desde' (timestamp). Builder+PIN OU robot key (para a Action)."""
    from fastapi import HTTPException
    b={}
    try: b=await request.json()
    except Exception: pass
    rk=request.headers.get("x-robot-key")
    if rk and ROBOT_KEY and rk==ROBOT_KEY:
        actor="robot"
    else:
        u,p=exige(request,"GERAR_ORCAMENTOS")
        if p.get("nivel")!="builder": raise HTTPException(403,"apenas builder")
        if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
        actor=u["id"]
    desde=(b.get("desde") or "").strip()
    if not desde: raise HTTPException(400,"informe 'desde' (timestamp da conferência) por segurança")
    q=f"lancado=eq.true&lancado_em=gte.{urllib.parse.quote(desde)}&status=eq.gerado&select=id,arquivo_pdf&limit=5000"
    rows=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?{q}",SB_KEY) or []
    voltou=[]; erros=[]
    for r in rows:
        ap=str(r.get("arquivo_pdf") or ""); origem=ap.split("/")[0] if "/" in ap else ""
        nome=ap.split("/",1)[1] if "/" in ap else ""
        if origem not in ("1","4") or not nome: erros.append(ap or str(r.get("id"))); continue
        try:
            _mover_volta(origem, nome)
            _sb_write(f"notas_orcamento?id=eq.{r['id']}", {"lancado":False,"lancado_em":None}, "PATCH")
            voltou.append(f"{origem}/{nome}")
        except Exception as e:
            erros.append(f"{ap}: {str(e)[:80]}")
    try: log_frotahub(actor,"builder","GERAR_ORCAMENTOS","DESFEZ_CONFERENCIA",f"{len(voltou)} desde {desde}")
    except Exception: pass
    return {"ok":True,"voltaram":voltou,"erros":erros,"total":len(voltou)}

@app.post("/orc/aplicar_conferencia")
async def orc_aplicar_conferencia(request: Request):
    """Aplica a conferência do Trílogo (feita fora): marca 'lancado' no banco e ROTEIA as
    pastas do Dropbox (não lançado -> 1/4, lançado -> 2/5), casando por TICKET + VALOR.
    dry_run=true (padrão) só devolve o PLANO, não mexe em nada. Builder+PIN OU robot key."""
    from fastapi import HTTPException
    b={}
    try: b=await request.json()
    except Exception: pass
    rk=request.headers.get("x-robot-key"); actor=None
    if rk and ROBOT_KEY and rk==ROBOT_KEY:
        actor="robot"
    else:
        u,p=exige(request,"GERAR_ORCAMENTOS")
        if p.get("nivel")!="builder": raise HTTPException(403,"apenas builder")
        if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
        actor=u["id"]
    dry=bool(b.get("dry_run", True))
    conf={}
    for it in (b.get("itens") or []):
        try: conf[(str(it["ticket"]), round(float(it["valor"]),2))]=bool(it["lancado"])
        except Exception: pass
    if not conf: raise HTTPException(400,"itens vazios")
    access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
    orcs=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?status=eq.gerado&select=id,ticket,valor_orcamento,arquivo_pdf,rateio,lancado&limit=20000",SB_KEY) or []
    planos=[]; fora=0; achados=set()
    for o in orcs:
        key=(str(o.get("ticket") or ""), round(_numf(o.get("valor_orcamento")),2))
        if key not in conf: fora+=1; continue
        achados.add(key)
        deseja=conf[key]; rateio=bool(o.get("rateio"))
        desired=(5 if rateio else 2) if deseja else (4 if rateio else 1)
        ap=str(o.get("arquivo_pdf") or ""); cur=ap.split("/")[0] if "/" in ap else ""
        nome=ap.split("/",1)[1] if "/" in ap else ""
        flip=bool(o.get("lancado"))!=deseja
        move=(cur in ("1","2","4","5")) and cur!=str(desired) and bool(nome)
        if flip or move:
            planos.append({"id":o["id"],"ticket":key[0],"valor":key[1],"nome":nome,
                "para_lancado":deseja,"de_pasta":cur,"para_pasta":str(desired),
                "mover":bool(move),"marcar":bool(flip)})
    nao_no_db=[{"ticket":k[0],"valor":k[1]} for k in conf if k not in achados]
    resumo={"db_total":len(orcs),"na_conferencia":len(achados),"fora_da_conferencia":fora,
        "vao_mudar":len(planos),
        "marcar_lancado":sum(1 for x in planos if x["marcar"] and x["para_lancado"]),
        "marcar_nao_lancado":sum(1 for x in planos if x["marcar"] and not x["para_lancado"]),
        "mover_p_lancados_2_5":sum(1 for x in planos if x["mover"] and x["para_pasta"] in ("2","5")),
        "mover_p_naolancados_1_4":sum(1 for x in planos if x["mover"] and x["para_pasta"] in ("1","4")),
        "conf_sem_par_no_db":len(nao_no_db)}
    if dry:
        return {"dry_run":True,"resumo":resumo,"amostra":planos[:20],"sem_par_no_db":nao_no_db[:20]}
    feitos=0; erros=[]
    for x in planos:
        try:
            if x["mover"]:
                src=_pasta_manut(access,int(x["de_pasta"]),ob); dst=_pasta_manut(access,int(x["para_pasta"]),ob)
                if not src or not dst: raise HTTPException(500,"pasta não encontrada")
                try: dropbox_rateio.mover(access,f"{src}/{x['nome']}",f"{dst}/{x['nome']}")
                except Exception:
                    if dropbox_rateio.baixar(access,f"{src}/{x['nome']}") is not None: raise  # ainda no src -> erro real
            patch={"lancado":x["para_lancado"]}
            if x["mover"]: patch["arquivo_pdf"]=f"{x['para_pasta']}/{x['nome']}"
            patch["lancado_em"]=_agora().isoformat() if x["para_lancado"] else None
            _sb_write(f"notas_orcamento?id=eq.{x['id']}", patch, "PATCH")
            feitos+=1
        except Exception as e:
            erros.append(f"{x['ticket']}/{x['nome']}: {str(e)[:80]}")
    try: log_frotahub(actor,"builder","GERAR_ORCAMENTOS","APLICOU_CONFERENCIA",f"{feitos} de {len(planos)}")
    except Exception: pass
    return {"dry_run":False,"resumo":resumo,"aplicados":feitos,"erros":erros[:40]}

# ---- CONFERÊNCIA de duplicidade (custos já lançados no Trílogo) ----
_CONFERIR={"itens":{},"atualizado":None}

@app.post("/robot/conferir_resultado")
async def robot_conferir_resultado(request: Request):
    """O robô reporta, por orçamento pendente, os custos que já existem no ticket."""
    _robot_ok(request)
    b=await request.json()
    lista=b.get("itens") if isinstance(b.get("itens"),list) else [b]
    for it in lista:
        a=it.get("arquivo")
        if a: _CONFERIR["itens"][a]=it
    _CONFERIR["atualizado"]=_agora().isoformat()
    return {"ok":True}

@app.post("/orc/conferir_disparar")
async def orc_conferir_disparar(request: Request):
    """Dispara o robô em modo CONFERÊNCIA (só lê custos, não lança). Builder/gerente."""
    from fastapi import HTTPException
    u,p=exige(request,"GERAR_ORCAMENTOS")
    if p.get("nivel") not in ("builder","gerente"): raise HTTPException(403,"apenas builder e gerente")
    _CONFERIR["itens"]={}; _CONFERIR["atualizado"]=_agora().isoformat()
    try: _gh_dispatch_wf(GH_WF_LANCAR,{"modo":"conferir","alvo":""})
    except urllib.error.HTTPError as e: raise HTTPException(500,f"GitHub {e.code}: {e.read().decode()[:200]}")
    except Exception as e: raise HTTPException(500,f"disparo falhou: {e}")
    log_frotahub(u["id"],p.get("papel"),"GERAR_ORCAMENTOS","CONFERIR_DUP","conferência de duplicidade")
    return {"ok":True}

@app.get("/orc/conferir_status")
def orc_conferir_status(request: Request):
    from fastapi import HTTPException
    u,p=exige(request,"GERAR_ORCAMENTOS")
    if p.get("nivel") not in ("builder","gerente"): raise HTTPException(403,"apenas builder e gerente")
    itens=_CONFERIR.get("itens",{})
    vals=list(itens.values())
    resumo={
        "total": len(vals),
        "duplicatas": sum(1 for x in vals if x.get("duplicata")),
        "reconciliados": sum(1 for x in vals if x.get("reconciliado")),
        "a_lancar": sum(1 for x in vals if x.get("aberto") and not x.get("duplicata")),
        "fora_conta": sum(1 for x in vals if x.get("outra_conta")),
    }
    out={"itens":itens,"atualizado":_CONFERIR.get("atualizado"),"resumo":resumo,"gh":None}
    if GH_TOKEN and GH_REPO:
        url=f"https://api.github.com/repos/{GH_REPO}/actions/workflows/{urllib.parse.quote(GH_WF_LANCAR)}/runs?per_page=1"
        req=urllib.request.Request(url,headers={"Authorization":f"Bearer {GH_TOKEN}","Accept":"application/vnd.github+json",
            "X-GitHub-Api-Version":"2022-11-28","User-Agent":"frotahub-motor"})
        try:
            r=json.loads(urllib.request.urlopen(req,timeout=20).read().decode()); run=(r.get("workflow_runs") or [None])[0]
            if run: out["gh"]={"estado":run.get("status"),"conclusao":run.get("conclusion"),"url":run.get("html_url")}
        except Exception: pass
    return out

@app.post("/orc/lancar_reconciliar")
async def orc_lancar_reconciliar(request: Request):
    """Marca um orçamento como JÁ lançado (custo já existe no Trílogo): move 1->2/4->5 e
    marca lancado, SEM relançar. Builder/gerente."""
    from fastapi import HTTPException
    u,p=exige(request,"GERAR_ORCAMENTOS")
    if p.get("nivel") not in ("builder","gerente"): raise HTTPException(403,"apenas builder e gerente")
    b=await request.json(); origem=str(b.get("origem")); nome=b.get("nome")
    if origem not in ("1","4") or not _seg_ok(nome): raise HTTPException(400,"parâmetros inválidos")
    res=_mover_marcar_lancado(origem, nome)
    _CONFERIR["itens"].pop(nome,None)
    log_frotahub(u["id"],p.get("papel"),"GERAR_ORCAMENTOS","RECONCILIOU",f"{origem}/{nome}")
    return res

@app.post("/orc/trilogo_run")
async def orc_trilogo_run(request: Request):
    """Atualização REAL Trílogo->Supabase, sob demanda. SOMENTE builder."""
    from fastapi import HTTPException
    u,p=exige(request,"CHAMADOS_TRILOGO")
    if p.get("papel")!="builder": raise HTTPException(403,"Somente o builder pode disparar a atualização manual.")
    body={}
    try: body=await request.json()
    except Exception: pass
    modo=(body.get("modo") or "rotina").strip()
    if modo not in ("rotina","inicial"): modo="rotina"
    try: _github_dispatch(modo)
    except urllib.error.HTTPError as e:
        raise HTTPException(500,f"GitHub {e.code}: {e.read().decode()[:200]}")
    except Exception as e:
        raise HTTPException(500,f"disparo falhou: {e}")
    log_frotahub(u["id"],p["papel"],"CHAMADOS_TRILOGO","DISPAROU_ROBO",modo)
    return {"ok":True,"modo":modo,"msg":"Robô disparado no GitHub. A atualização leva ~2 min — recarregue a lista depois."}

@app.get("/orc/trilogo_status")
def orc_trilogo_status(request: Request):
    """Status da última execução do robô no GitHub (para a tela mostrar 'rodando/concluído')."""
    from fastapi import HTTPException
    exige(request,"CHAMADOS_TRILOGO")
    if not (GH_TOKEN and GH_REPO): return {"ok":False,"msg":"GitHub não configurado no Render"}
    url=f"https://api.github.com/repos/{GH_REPO}/actions/workflows/{urllib.parse.quote(GH_WORKFLOW)}/runs?per_page=1"
    req=urllib.request.Request(url,headers={"Authorization":f"Bearer {GH_TOKEN}","Accept":"application/vnd.github+json",
        "X-GitHub-Api-Version":"2022-11-28","User-Agent":"frotahub-motor"})
    try:
        r=json.loads(urllib.request.urlopen(req,timeout=20).read().decode())
        run=(r.get("workflow_runs") or [None])[0]
        if not run: return {"ok":True,"estado":"nenhuma execução"}
        return {"ok":True,"estado":run.get("status"),"conclusao":run.get("conclusion"),
                "em":run.get("run_started_at") or run.get("created_at"),"url":run.get("html_url")}
    except Exception as e:
        return {"ok":False,"msg":str(e)[:160]}

# ================= AGENDADOR (executor chamado pelo Supabase pg_cron) =================
def _ag_log(ag_id, rotina, status, detalhe=""):
    try:
        _sb_json(f"{SB_URL}/rest/v1/agendamento_exec", SB_KEY,
                 data={"agendamento_id":ag_id,"rotina":rotina,"status":status,"detalhe":str(detalhe)[:400]},
                 method="POST")
    except Exception: pass

def _run_rotina_auto(rotina):
    """Executa uma rotina de forma automática (sem UI). Só as automatizáveis; as demais levantam erro."""
    if rotina=="CHAMADOS_TRILOGO":
        _github_dispatch("rotina"); return "robô do Trílogo disparado"
    if rotina=="GERAR_ORCAMENTOS":
        import threading
        _ORC_JOB_SEQ[0]+=1; job=str(_ORC_JOB_SEQ[0])
        _ORC_JOBS[job]={"estado":"rodando","previa":False,"total":0,"feitas":0,"gerados":0,
                        "res":[],"pausa":False,"retoma_em":0,"lote":ORC_LOTE,"descanso":ORC_DESCANSO}
        threading.Thread(target=_orc_job_run,args=(job,False,None,"agendador"),daemon=True).start()
        return f"geração de orçamentos iniciada (job {job})"
    raise RuntimeError(f"rotina '{rotina}' ainda não é automatizável")

@app.post("/orc/agendador_exec")
async def orc_agendador_exec(request: Request):
    """Chamado pelo Supabase (pg_cron/pg_net). Executa a fila de rotinas do agendamento."""
    from fastapi import HTTPException
    if not AGENDADOR_SECRET or request.headers.get("x-agendador-secret","")!=AGENDADOR_SECRET:
        raise HTTPException(401,"não autorizado")
    body={}
    try: body=await request.json()
    except Exception: pass
    ag_id=body.get("agendamento_id")
    if not ag_id: raise HTTPException(400,"agendamento_id ausente")
    itens=_sb_json(f"{SB_URL}/rest/v1/agendamento_itens?agendamento_id=eq.{urllib.parse.quote(str(ag_id))}&order=ordem&select=rotina,ordem",SB_KEY) or []
    res=[]
    for it in itens:
        rot=it.get("rotina")
        try:
            det=_run_rotina_auto(rot); _ag_log(ag_id,rot,"ok",det); res.append({"rotina":rot,"ok":True,"detalhe":det})
        except Exception as e:
            _ag_log(ag_id,rot,"erro",str(e)[:200]); res.append({"rotina":rot,"ok":False,"erro":str(e)[:200]})
    return {"ok":True,"itens":res}

# ================= ESTATÍSTICAS (Manutenção) =================
# "atendidos" = apenas Executado + Vistoriado (aberto/em execução NÃO contam)
_ATENDIDOS=("Executado","Vistoriado")
def _numf(x):
    try: return float(x or 0)
    except Exception: return 0.0
def _reais(v): return "R$ "+f"{_numf(v):,.2f}".replace(",","·").replace(".",",").replace("·",".")

_MESES_CURTO=["jan","fev","mar","abr","mai","jun","jul","ago","set","out","nov","dez"]
def _mes_label(ym):   # 'YYYY-MM' -> 'jul/2026'
    try: y,m=ym.split("-"); return f"{_MESES_CURTO[int(m)-1]}/{y}"
    except Exception: return ym
def _stat_chamados(q="",aba="",desde="",ate=""):
    # "atendidos" = chamados que atingiram o MARCO do vistoriado (atendido=true),
    # no período pela DATA DA VISTORIA (atendido_em) — mesmo que hoje estejam fechados/arquivados.
    parts=["select=numero,aba,loja,tipo_predial,atendido_em","atendido=is.true","limit=10000"]
    if aba:   parts.append(f"aba=eq.{urllib.parse.quote(aba)}")
    if desde: parts.append(f"atendido_em=gte.{desde}")
    if ate:   parts.append(f"atendido_em=lte.{ate}")
    if q and len(q)>=3: parts.append("or="+urllib.parse.quote(f"(numero.ilike.*{q}*,loja.ilike.*{q}*)"))
    rows=_sb_json(f"{SB_URL}/rest/v1/chamados?"+"&".join(parts),SB_KEY) or []
    por_aba={"Civil":0,"Instalações":0}; por_loja={}; por_tipo={}; por_mes={}
    for r in rows:
        ab=_abacurta(r.get("aba"))
        if ab in por_aba: por_aba[ab]+=1
        lj=_loja_padrao(r.get("loja")); por_loja[lj]=por_loja.get(lj,0)+1
        tp=r.get("tipo_predial") or "—"; por_tipo[tp]=por_tipo.get(tp,0)+1
        ym=(r.get("atendido_em") or "")[:7]
        if ym: por_mes[ym]=por_mes.get(ym,0)+1
    lojas=sorted(({"loja":k,"total":v} for k,v in por_loja.items()),key=lambda x:-x["total"])
    tipos=sorted(({"tipo":k,"total":v} for k,v in por_tipo.items()),key=lambda x:-x["total"])
    meses=[{"mes":_mes_label(k),"total":v} for k,v in sorted(por_mes.items())]
    return {"total":len(rows),"atendidos":len(rows),"por_aba":por_aba,"por_loja":lojas,"por_tipo":tipos,"por_mes":meses}

@app.get("/orc/stat_chamados")
def orc_stat_chamados(request: Request, q: str="", aba: str="", desde: str="", ate: str=""):
    from fastapi import HTTPException
    exige(request,"CHAMADOS_ATENDIDOS")
    try: return _stat_chamados(q,aba,desde,ate)
    except Exception as e: raise HTTPException(500,f"stat: {e}")

@app.get("/orc/stat_chamados_pdf")
def orc_stat_chamados_pdf(request: Request, q: str="", aba: str="", desde: str="", ate: str=""):
    from fastapi import HTTPException
    exige(request,"CHAMADOS_ATENDIDOS")
    d=_stat_chamados(q,aba,desde,ate)
    linhas=[[l["loja"][:34], str(l["total"])] for l in d["por_loja"]]
    sub=(f"{_abacurta(aba)} · " if aba else "")+(f"{desde}→{ate}" if (desde or ate) else "")
    civ=d["por_aba"].get("Civil",0); ins=d["por_aba"].get("Instalações",0)
    resumo=f"Atendidos (vistoriados) {d['atendidos']} · Civil {civ} · Instalações {ins}"
    pdf=_lista_pdf("Estatística — Chamados atendidos", ["Loja","Atendidos"], linhas,
                   subtitulo=resumo+(f"  |  {sub}" if sub else ""), aligns={1:'RIGHT'})
    return Response(content=pdf, media_type="application/pdf",
        headers={"Content-Disposition":'attachment; filename="estatistica_chamados.pdf"'})

def _stat_financeiro(aba="",desde="",ate="",mes="",loja=""):
    parts=["select=ticket,loja_nome,aba,valor_nota,valor_orcamento,mes_ref,extrapolado",
           "status=eq.gerado","limit=5000"]
    if aba:  parts.append(f"aba=eq.{urllib.parse.quote(aba)}")
    if mes:  parts.append(f"mes_ref=eq.{urllib.parse.quote(mes)}")
    if desde:parts.append(f"criado_em=gte.{desde}")
    if ate:  parts.append(f"criado_em=lte.{ate}T23:59:59")
    if loja and len(loja)>=3: parts.append("loja_nome=ilike.*"+urllib.parse.quote(loja)+"*")
    notas=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?"+"&".join(parts),SB_KEY) or []
    tickets=set(); v_orc=0.0; v_nota=0.0; por_loja={}; por_aba={"Civil":{"n":0,"orc":0.0},"Instalações":{"n":0,"orc":0.0}}; por_mes={}
    for r in notas:
        t=r.get("ticket")
        if t: tickets.add(str(t))
        vo=_numf(r.get("valor_orcamento")); vn=_numf(r.get("valor_nota")); v_orc+=vo; v_nota+=vn
        lj=_loja_padrao(r.get("loja_nome")); dd=por_loja.setdefault(lj,{"n":0,"orc":0.0,"nota":0.0}); dd["n"]+=1; dd["orc"]+=vo; dd["nota"]+=vn
        ab=_abacurta(r.get("aba"))
        if ab in por_aba: por_aba[ab]["orc"]+=vo; por_aba[ab]["n"]+=1
        mr=r.get("mes_ref") or "—"; mm=por_mes.setdefault(mr,{"n":0,"orc":0.0}); mm["n"]+=1; mm["orc"]+=vo
    # denominador: chamados ATENDIDOS (marco vistoriado) no período, pela data da vistoria
    ap=["select=numero","atendido=is.true","limit=10000"]
    if aba:   ap.append(f"aba=eq.{urllib.parse.quote(aba)}")
    if desde: ap.append(f"atendido_em=gte.{desde}")
    if ate:   ap.append(f"atendido_em=lte.{ate}")
    try: n_atend=len(_sb_json(f"{SB_URL}/rest/v1/chamados?"+"&".join(ap),SB_KEY) or [])
    except Exception: n_atend=0
    com_mat=len(tickets)
    lojas=sorted(({"loja":k,**v} for k,v in por_loja.items()),key=lambda x:-x["orc"])
    meses=[{"mes":k,**v} for k,v in por_mes.items()]
    return {"orcamentos":len(notas),"com_material":com_mat,"valor_orcamento":round(v_orc,2),
            "valor_nota":round(v_nota,2),"margem":round(v_orc-v_nota,2),"chamados_atendidos":n_atend,
            "pct_material":(round(com_mat*100.0/n_atend,1) if n_atend else None),
            "por_aba":por_aba,"por_loja":lojas,"por_mes":meses}

@app.get("/orc/stat_financeiro")
def orc_stat_financeiro(request: Request, aba: str="", desde: str="", ate: str="", mes: str="", loja: str=""):
    from fastapi import HTTPException
    exige(request,"FINANCEIRO_MATERIAIS")
    try: return _stat_financeiro(aba,desde,ate,mes,loja)
    except Exception as e: raise HTTPException(500,f"stat: {e}")

@app.get("/orc/stat_financeiro_pdf")
def orc_stat_financeiro_pdf(request: Request, aba: str="", desde: str="", ate: str="", mes: str="", loja: str=""):
    from fastapi import HTTPException
    exige(request,"FINANCEIRO_MATERIAIS")
    d=_stat_financeiro(aba,desde,ate,mes,loja)
    linhas=[[l["loja"][:30], str(l["n"]), _reais(l["orc"]), _reais(l.get("nota",0))] for l in d["por_loja"]]
    sub=(f"{_abacurta(aba)} · " if aba else "")+(mes+" · " if mes else "")+(f"{desde}→{ate}" if (desde or ate) else "")
    resumo=(f"Orçamentos {d['orcamentos']} · Chamados c/ material {d['com_material']} · "
            f"Valor orçamentos {_reais(d['valor_orcamento'])} · Notas {_reais(d['valor_nota'])} · Margem {_reais(d['margem'])}")
    pdf=_lista_pdf("Estatística — Financeiro de materiais", ["Loja","Qtd","Orçamentos","Notas"], linhas,
                   subtitulo=resumo+(f"  |  {sub}" if sub else ""), aligns={1:'RIGHT',2:'RIGHT',3:'RIGHT'})
    return Response(content=pdf, media_type="application/pdf",
        headers={"Content-Disposition":'attachment; filename="estatistica_financeiro.pdf"'})

# ================= PLANILHA GERAL DE ORÇAMENTOS (controle) =================
_MAROM="7A1517"
def _conta_da_aba(aba):
    a=_abacurta(aba)
    if a=="Civil": return "PREDIAL CIVIL - MANUENÇÃO CORRETIVA"
    if a=="Instalações": return "PREDIAL INSTALAÇÕES - MANUENÇÃO CORRETIVA"
    return ""
def _planilha_orc_rows(desde="",ate="",faixa=""):
    parts=["select=id,ticket,nota_numero,loja_nome,aba,valor_orcamento,criado_em,rateio","status=eq.gerado","order=criado_em.asc","limit=8000"]
    if desde: parts.append(f"criado_em=gte.{desde}")
    if ate:   parts.append(f"criado_em=lte.{ate}T23:59:59")
    notas=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?"+"&".join(parts),SB_KEY) or []
    val=lambda n: _numf(n.get("valor_orcamento"))
    if faixa=="ate600":   notas=[n for n in notas if val(n)<=600]
    elif faixa=="acima600": notas=[n for n in notas if val(n)>600]
    tickets=list({str(n.get("ticket")) for n in notas if n.get("ticket")})
    lojamap={}
    if tickets:
        inlist=",".join(urllib.parse.quote(t) for t in tickets[:1500])
        ch=_sb_json(f"{SB_URL}/rest/v1/chamados?numero=in.({inlist})&select=numero,loja,aba",SB_KEY) or []
        for c in ch: lojamap[str(c.get("numero"))]={"loja":c.get("loja"),"aba":c.get("aba")}
    rows=[]
    for i,n in enumerate(notas,1):
        tk=str(n.get("ticket") or ""); cm=lojamap.get(tk,{})
        aba=cm.get("aba") or n.get("aba") or ""
        rows.append({"n":i,"id":n.get("id"),"ticket":tk,"nota":n.get("nota_numero"),"loja":_loja_padrao(cm.get("loja") or n.get("loja_nome")),
            "valor":round(val(n),2),"data":(n.get("criado_em") or "")[:10],"conta":_conta_da_aba(aba),"rateio":bool(n.get("rateio"))})
    return rows

def _planilha_header(desde,ate,faixa,resp,total,qtd):
    per = (f"{_data_br(desde)} a {_data_br(ate)}" if (desde or ate) else "Todos os períodos")
    fx  = {"ate600":"Orçamentos ≤ R$ 600","acima600":"Orçamentos > R$ 600"}.get(faixa,"Todas as faixas")
    hoje=_hoje().strftime("%d/%m/%Y")
    return {"periodo":per,"faixa":fx,"gerado":hoje,"responsavel":resp or "—",
            "total":_reais(total),"qtd":qtd,
            "titulo":"CUSTOS DE MATERIAIS DOS CHAMADOS DE MANUTENÇÃO"}

@app.get("/orc/planilha_orcamentos")
def orc_planilha(request: Request, desde: str="", ate: str="", faixa: str=""):
    from fastapi import HTTPException
    u,p=exige(request,"ORCAMENTOS_GERADOS")
    rows=_planilha_orc_rows(desde,ate,faixa); total=round(sum(r["valor"] for r in rows),2)
    h=_planilha_header(desde,ate,faixa,(p.get("nome_completo") or p.get("nome")),total,len(rows))
    # quem pode remover orçamentos (builder/gerente) — o front usa isso para mostrar a opção
    return {"itens":rows,"total":total,"header":h,"pode_remover":(p.get("nivel") in ("builder","gerente"))}

# ================= GERENCIADOR DE ARQUIVOS (fluxo pelas pastas, sem Dropbox direto) =================
_ARQ_EXTS_DOC={".pdf",".jpg",".jpeg",".png"}
_ARQ_TREE_NUMS=[1,2,4,5,9]
def _seg_ok(nome): return bool(nome) and ("/" not in nome) and ("\\" not in nome) and (".." not in nome)
def _ctype(nome):
    e=os.path.splitext(nome)[1].lower()
    return {".pdf":"application/pdf",".jpg":"image/jpeg",".jpeg":"image/jpeg",".png":"image/png"}.get(e,"application/octet-stream")
def _arq_area(access, area):
    if area=="pco_oc":
        return {"base":PCO_ENVIAR,"perm":"ENVIAR_PCO","exts":{".pdf"},"rotulo":"Ordens de compra (PCO)"}
    if area=="notas_orc":
        ob=_orc_base(access,_manut_base(access))
        base=_pasta_manut(access,0,ob) or (ob+"/0 - NOTAS PARA ORCAMENTO (COLOCAR AQUI)")
        return {"base":base,"perm":"GERAR_ORCAMENTOS","exts":_ARQ_EXTS_DOC,"rotulo":"Notas e DAVs"}
    return None

@app.get("/arq/listar")
def arq_listar(request: Request, area: str=""):
    from fastapi import HTTPException
    if not dropbox_rateio.ativo(): raise HTTPException(500,"Dropbox não configurado")
    access=dropbox_rateio.obter_token(); a=_arq_area(access,area)
    if not a: raise HTTPException(400,"área inválida")
    exige(request,a["perm"])
    ents=dropbox_rateio.listar_entradas(access,a["base"])
    itens=[{"nome":e["name"],"tamanho":e.get("size")} for e in ents
           if not e["dir"] and os.path.splitext(e["name"])[1].lower() in a["exts"]]
    itens.sort(key=lambda x:x["nome"].lower())
    return {"area":area,"rotulo":a["rotulo"],"total":len(itens),"itens":itens}

@app.post("/arq/upar")
async def arq_upar(request: Request):
    from fastapi import HTTPException
    if not dropbox_rateio.ativo(): raise HTTPException(500,"Dropbox não configurado")
    access=dropbox_rateio.obter_token()
    form=await request.form(); area=form.get("area"); a=_arq_area(access,area)
    if not a: raise HTTPException(400,"área inválida")
    u,p=exige(request,a["perm"])
    files=form.getlist("arquivos")
    if not files: raise HTTPException(400,"nenhum arquivo enviado")
    salvos=[]; erros=[]
    for f in files:
        nome=os.path.basename(getattr(f,"filename","") or "")
        if not nome: continue
        ext=os.path.splitext(nome)[1].lower()
        if ext not in a["exts"]: erros.append(f"{nome}: tipo não aceito"); continue
        try:
            data=await f.read()
            if len(data)>30*1024*1024: erros.append(f"{nome}: acima de 30 MB"); continue
            dropbox_rateio.subir_bytes(access,data,f"{a['base']}/{nome}",overwrite=False)  # autorename evita clobber
            salvos.append(nome)
        except Exception as e: erros.append(f"{nome}: {str(e)[:80]}")
    log_frotahub(u["id"],p.get("papel"),a["perm"],"UPAR",f"{len(salvos)} arquivo(s) em {area}")
    return {"salvos":salvos,"erros":erros}

@app.get("/arq/ver")
def arq_ver(request: Request, area: str="", nome: str=""):
    from fastapi import HTTPException
    access=dropbox_rateio.obter_token(); a=_arq_area(access,area)
    if not a: raise HTTPException(400,"área inválida")
    exige(request,a["perm"])
    if not _seg_ok(nome): raise HTTPException(400,"nome inválido")
    data=dropbox_rateio.baixar(access,f"{a['base']}/{nome}")
    if data is None: raise HTTPException(404,"arquivo não encontrado")
    return Response(content=data,media_type=_ctype(nome),
                    headers={"Content-Disposition":f'inline; filename="{urllib.parse.quote(nome)}"'})

@app.post("/arq/excluir")
async def arq_excluir(request: Request):
    from fastapi import HTTPException
    access=dropbox_rateio.obter_token(); b=await request.json(); area=b.get("area")
    a=_arq_area(access,area)
    if not a: raise HTTPException(400,"área inválida")
    u,p=exige(request,a["perm"]); nome=b.get("nome")
    if not _seg_ok(nome): raise HTTPException(400,"nome inválido")
    try: dropbox_rateio.apagar(access,f"{a['base']}/{nome}")   # lixeira do Dropbox (recuperável ~30d)
    except Exception as e: raise HTTPException(500,f"excluir: {str(e)[:120]}")
    log_frotahub(u["id"],p.get("papel"),a["perm"],"EXCLUIR",f"{nome} ({area})")
    return {"ok":True}

# ---- Árvore do Dropbox (pastas 1,2,4,5,9 da manutenção) ----
def _arvore_roots(access, ob):
    out={}
    for n in _ARQ_TREE_NUMS:
        pth=_pasta_manut(access,n,ob)
        if pth: out[os.path.basename(pth)]=pth
    return out
def _arvore_valida(path, roots):
    path=(path or "").strip().strip("/")
    if not path: return ""
    if ".." in path or "\\" in path: return None
    if path.split("/")[0] not in roots: return None
    return path

@app.get("/arq/arvore")
def arq_arvore(request: Request, path: str=""):
    from fastapi import HTTPException
    exige(request,"GERAR_ORCAMENTOS")
    access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
    roots=_arvore_roots(access,ob); path=_arvore_valida(path,roots)
    if path is None: raise HTTPException(403,"caminho fora do escopo")
    if path=="":
        return {"path":"","entradas":[{"nome":k,"dir":True} for k in sorted(roots)]}
    raw=dropbox_rateio.listar_entradas(access,f"{ob}/{path}")
    ents=[{"nome":e["name"],"dir":e["dir"],"tamanho":e.get("size")} for e in raw]
    ents.sort(key=lambda x:(not x["dir"],x["nome"].lower()))
    return {"path":path,"entradas":ents}

@app.get("/arq/arvore_ver")
def arq_arvore_ver(request: Request, path: str=""):
    from fastapi import HTTPException
    exige(request,"GERAR_ORCAMENTOS")
    access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
    roots=_arvore_roots(access,ob); path=_arvore_valida(path,roots)
    if not path: raise HTTPException(403,"caminho inválido")
    data=dropbox_rateio.baixar(access,f"{ob}/{path}")
    if data is None: raise HTTPException(404,"arquivo não encontrado")
    nome=os.path.basename(path)
    return Response(content=data,media_type=_ctype(nome),
                    headers={"Content-Disposition":f'inline; filename="{urllib.parse.quote(nome)}"'})

@app.post("/arq/arvore_excluir")
async def arq_arvore_excluir(request: Request):
    from fastapi import HTTPException
    u,p=exige(request,"GERAR_ORCAMENTOS")
    access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
    roots=_arvore_roots(access,ob); b=await request.json(); path=_arvore_valida(b.get("path"),roots)
    if not path or "/" not in path: raise HTTPException(403,"só é possível excluir itens DENTRO das pastas")
    try: dropbox_rateio.apagar(access,f"{ob}/{path}")
    except Exception as e: raise HTTPException(500,f"excluir: {str(e)[:120]}")
    log_frotahub(u["id"],p.get("papel"),"GERAR_ORCAMENTOS","EXCLUIR_ARVORE",path)
    return {"ok":True}

def _resolve_arq(access, ob, rel):
    """Converte 'NN/sub/arquivo' (guardado no banco) no caminho real do Dropbox."""
    if not rel or "/" not in rel: return None
    num,_,sub=rel.partition("/")
    try: base=_pasta_manut(access, int(num), ob)
    except Exception: base=None
    return f"{base}/{sub}" if base else None

@app.post("/orc/planilha_remover")
async def orc_planilha_remover(request: Request):
    """Remove orçamentos gerados (builder/gerente + PIN): marca como 'removido' (sai da planilha,
       das estatísticas e libera a nota para ser refeita) e manda os arquivos para a lixeira do Dropbox."""
    from fastapi import HTTPException
    u,p=_exige_gestor(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
    ids=[str(x) for x in (b.get("ids") or []) if x]
    if not ids: raise HTTPException(400,"nenhum orçamento selecionado")
    inlist=",".join(urllib.parse.quote(i) for i in ids[:500])
    recs=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?id=in.({inlist})&select=id,ticket,nota_numero,arquivo_pdf,arquivo_doc,rateio,status",SB_KEY) or []
    removidos=0
    try: access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
    except Exception: access=None; ob=None
    for r in recs:
        if r.get("status")!="gerado": continue
        # lixeira: montado (10/11) + não lançado (1/4) pelo nome-base
        if access:
            for rel in (r.get("arquivo_pdf"), r.get("arquivo_doc")):
                full=_resolve_arq(access, ob, rel)
                if full:
                    try: dropbox_rateio.apagar(access, full)
                    except Exception: pass
            base=os.path.basename(r.get("arquivo_pdf") or r.get("arquivo_doc") or "")
            if base:
                pnl=_pasta_manut(access, 4 if r.get("rateio") else 1, ob)
                if pnl:
                    for ext in (".pdf",".docx"):
                        try: dropbox_rateio.apagar(access, f"{pnl}/{os.path.splitext(base)[0]}{ext}")
                        except Exception: pass
        try: _sb_write(f"notas_orcamento?id=eq.{r['id']}", {"status":"removido"}, "PATCH"); removidos+=1
        except Exception: pass
    log_frotahub(u["id"],p.get("papel"),"ORCAMENTOS_GERADOS","REMOVEU",f"{removidos} orçamento(s)")
    return {"ok":True,"removidos":removidos}

@app.get("/orc/planilha_orcamentos_xlsx")
def orc_planilha_xlsx(request: Request, desde: str="", ate: str="", faixa: str=""):
    from fastapi import HTTPException
    u,p=exige(request,"ORCAMENTOS_GERADOS")
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.drawing.image import Image as XLImage
    rows=_planilha_orc_rows(desde,ate,faixa); total=round(sum(r["valor"] for r in rows),2)
    h=_planilha_header(desde,ate,faixa,(p.get("nome_completo") or p.get("nome")),total,len(rows))
    wb=openpyxl.Workbook(); ws=wb.active; ws.title="Orçamentos"
    larg=[8,7.5,37.75,13.13,14,13,42.25,23.45]
    for i,w in enumerate(larg,1): ws.column_dimensions[get_column_letter(i)].width=w
    # ---- cabeçalho (logo + textos) ----
    try:
        tmp=tempfile.NamedTemporaryFile(suffix=".png",delete=False); tmp.write(_b64.b64decode(_LOGO_ORC_B64)); tmp.close()
        img=XLImage(tmp.name); img.width=150; img.height=64; ws.add_image(img,"A1")
    except Exception: pass
    ws.merge_cells("C1:H1"); ws["C1"]=h["titulo"]; ws["C1"].font=Font(bold=True,size=14,color=_MAROM); ws["C1"].alignment=Alignment(horizontal="left",vertical="center")
    ws.merge_cells("C2:H2"); ws["C2"]=f"Período: {h['periodo']}  ·  {h['faixa']}"; ws["C2"].font=Font(size=10,color="555555")
    ws.merge_cells("C3:H3"); ws["C3"]=f"Gerado em {h['gerado']} por {h['responsavel']}  ·  {h['qtd']} orçamento(s)  ·  Total {h['total']}"; ws["C3"].font=Font(size=10,color="555555")
    ws.row_dimensions[1].height=22; ws.row_dimensions[2].height=16; ws.row_dimensions[3].height=16
    # ---- tabela ----
    H0=6   # linha do cabeçalho da tabela
    heads=["Nº","TICKET","LOJA","VALOR","DATA","ORÇAMENTO","CONTA","PCO"]
    borda=Border(*(Side(style="thin",color=_MAROM),)*4)
    for c,txt in enumerate(heads,1):
        cell=ws.cell(H0,c,txt); cell.font=Font(bold=True,size=10,color="FFFFFF")
        cell.fill=PatternFill("solid",fgColor=_MAROM); cell.alignment=Alignment(horizontal="center",vertical="center"); cell.border=borda
    aligns={1:"center",2:"center",3:"left",4:"right",5:"center",6:"center",7:"left",8:"center"}
    fmtD='"R$"\\ #,##0.00'; fmtF='"R$"\\ #,##0.00_);[Red]\\("R$"\\ #,##0.00\\)'
    for k,r in enumerate(rows):
        rr=H0+1+k
        vals=[r["n"],r["ticket"],r["loja"],r["valor"],None,None,r["conta"],None]
        for c in range(1,9):
            cell=ws.cell(rr,c,vals[c-1]); cell.font=Font(size=10); cell.alignment=Alignment(horizontal=aligns[c]); cell.border=borda
        # DATA (E) como data real
        try:
            y,mo,d=r["data"].split("-"); ws.cell(rr,5).value=datetime.date(int(y),int(mo),int(d)); ws.cell(rr,5).number_format="dd/mm/yyyy"
        except Exception: ws.cell(rr,5).value=r["data"]
        ws.cell(rr,4).number_format=fmtD
        ws.cell(rr,6).value=f"=D{rr}"; ws.cell(rr,6).number_format=fmtF   # ORÇAMENTO = VALOR (igual ao modelo)
    ws.freeze_panes=f"A{H0+1}"; ws.auto_filter.ref=f"A{H0}:H{H0+len(rows)}"
    buf=io.BytesIO(); wb.save(buf)
    return Response(content=buf.getvalue(), media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition":'attachment; filename="orcamentos_materiais_frota.xlsx"'})

@app.get("/orc/planilha_orcamentos_pdf")
def orc_planilha_pdf(request: Request, desde: str="", ate: str="", faixa: str=""):
    from fastapi import HTTPException
    u,p=exige(request,"ORCAMENTOS_GERADOS")
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib import colors
    from reportlab.lib.units import cm, mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    rows=_planilha_orc_rows(desde,ate,faixa); total=round(sum(r["valor"] for r in rows),2)
    h=_planilha_header(desde,ate,faixa,(p.get("nome_completo") or p.get("nome")),total,len(rows))
    buf=io.BytesIO()
    doc=SimpleDocTemplate(buf,pagesize=landscape(A4),leftMargin=1.0*cm,rightMargin=1.0*cm,topMargin=0.9*cm,bottomMargin=0.9*cm,title=h["titulo"])
    ss=getSampleStyleSheet()
    tit=ParagraphStyle('t',parent=ss['Title'],textColor=colors.HexColor('#'+_MAROM),fontSize=15,alignment=0,spaceAfter=1)
    sub=ParagraphStyle('s',parent=ss['Normal'],textColor=colors.HexColor('#555555'),fontSize=9,leading=12)
    el=[]
    try:
        tmp=tempfile.NamedTemporaryFile(suffix=".png",delete=False); tmp.write(_b64.b64decode(_LOGO_ORC_B64)); tmp.close()
        logo=RLImage(tmp.name,width=26*mm,height=11*mm)
        cab=Table([[logo, Paragraph(h["titulo"],tit)]],colWidths=[30*mm,None])
        cab.setStyle(TableStyle([('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LEFTPADDING',(0,0),(-1,-1),0)]))
        el.append(cab)
    except Exception: el.append(Paragraph(h["titulo"],tit))
    el.append(Paragraph(f"Período: {h['periodo']} &nbsp;·&nbsp; {h['faixa']}",sub))
    el.append(Paragraph(f"Gerado em {h['gerado']} por {h['responsavel']} &nbsp;·&nbsp; {h['qtd']} orçamento(s) &nbsp;·&nbsp; Total {h['total']}",sub))
    el.append(Spacer(1,7))
    heads=["Nº","TICKET","LOJA","VALOR","DATA","ORÇAMENTO","CONTA","PCO"]
    data=[heads]+[[str(r["n"]),r["ticket"],r["loja"][:44],_reais(r["valor"]),_data_br(r["data"]),_reais(r["valor"]),r["conta"][:40],""] for r in rows]
    t=Table(data,repeatRows=1,colWidths=[1.0*cm,1.7*cm,6.2*cm,2.2*cm,2.0*cm,2.3*cm,6.6*cm,2.2*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,0),colors.HexColor('#'+_MAROM)),('TEXTCOLOR',(0,0),(-1,0),colors.white),
        ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),('FONTSIZE',(0,0),(-1,-1),8),
        ('GRID',(0,0),(-1,-1),0.4,colors.HexColor('#'+_MAROM)),('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ('ALIGN',(0,1),(1,-1),'CENTER'),('ALIGN',(3,1),(5,-1),'RIGHT'),('ALIGN',(4,1),(4,-1),'CENTER'),
        ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,colors.HexColor('#f6f2f2')]),
        ('TOPPADDING',(0,0),(-1,-1),3),('BOTTOMPADDING',(0,0),(-1,-1),3)]))
    el.append(t)
    doc.build(el)
    return Response(content=buf.getvalue(), media_type="application/pdf",
        headers={"Content-Disposition":'attachment; filename="orcamentos_materiais_frota.pdf"'})

ORC_LOTE     = int(os.environ.get("ORC_LOTE", "10"))     # lê N notas
ORC_DESCANSO = int(os.environ.get("ORC_DESCANSO", "60"))  # descansa M segundos entre lotes
_ORC_JOBS = {}   # job_id -> estado/progresso
_ORC_JOB_SEQ = [0]

def _orc_job_run(job, previa, uid, papel):
    import time as _t
    st=_ORC_JOBS[job]
    try:
        access=dropbox_rateio.obter_token()
        _mb=_manut_base(access); _ob=_orc_base(access,_mb)
        ORC_NOTAS=_pasta_manut(access,0,_ob) or (_ob + "/0 - NOTAS PARA ORCAMENTO (COLOCAR AQUI)")
        P6=_pasta_manut(access,6,_ob); P7=_pasta_manut(access,7,_ob); P8=_pasta_manut(access,8,_ob)
        P1=_pasta_manut(access,1,_ob); P9=_pasta_manut(access,9,_ob); P10=_pasta_manut(access,10,_ob)
        arqs=[a for a in dropbox_rateio.listar(access,ORC_NOTAS) if a.lower().endswith((".pdf",".jpg",".jpeg",".png"))]
        def _mime(nm):
            nl=nm.lower(); return "application/pdf" if nl.endswith(".pdf") else ("image/png" if nl.endswith(".png") else "image/jpeg")
        res=st["res"]; _lista=sorted(arqs); st["total"]=len(_lista)
        st["itens"]=[{"nome":n,"idx":k+1,"pct":0,"status":"aguardando","reader":None,"motivo":None} for k,n in enumerate(_lista)]
        if not _lista:                     # pasta vazia -> não é erro, só não processa nada
            st["estado"]="pronto"; st["gerados"]=0; st["vazio"]=True; return
        for _idx,nome in enumerate(_lista):
            fst=st["itens"][_idx]
            if st.get("cancelar"): break          # interrompido pelo usuário
            # lotes: a cada ORC_LOTE notas lidas, descansa ORC_DESCANSO segundos (respeita o limite/min)
            if _idx and _idx % ORC_LOTE == 0:
                st["pausa"]=True; st["retoma_em"]=ORC_DESCANSO
                for _s in range(ORC_DESCANSO):
                    if st.get("cancelar"): break
                    _t.sleep(1); st["retoma_em"]=ORC_DESCANSO-_s-1
                st["pausa"]=False
                if st.get("cancelar"): break
            ext=os.path.splitext(nome)[1] or ".pdf"; is_pdf=nome.lower().endswith(".pdf")
            fst["pct"]=8; fst["status"]="lendo"
            fb=dropbox_rateio.baixar(access,f"{ORC_NOTAS}/{nome}")
            if not fb:
                fst["pct"]=100; fst["status"]="erro"; fst["motivo"]="não baixou"
                res.append({"arquivo":nome,"status":"erro","motivo":"não baixou"}); st["feitas"]=_idx+1; continue
            try:
                notas,reader=_ler_notas_rota(fb,_mime(nome),nome,fst,st)
                fst["reader"]=reader
            except _CotaExcedida:
                fst["pct"]=100; fst["status"]="Não executado por cota diária excedida"; fst["cota"]=True
                res.append({"arquivo":nome,"status":"cota","motivo":"cota diária de leitura (Gemini) excedida — rode de novo mais tarde"})
                st["feitas"]=_idx+1; continue
            except Exception as e:
                fst["pct"]=100; fst["status"]="erro"; fst["motivo"]=str(e)[:120]
                res.append({"arquivo":nome,"status":"erro","motivo":f"leitor: {str(e)[:120]}"}); st["feitas"]=_idx+1; continue
            if not notas:
                fst["pct"]=100; fst["status"]="pendente"; fst["motivo"]="não identifiquei nota"
                info={"arquivo":nome,"status":"pendente","motivo":"não identifiquei nota","destino":"6"}
                if not previa and P6: dropbox_rateio.mover(access,f"{ORC_NOTAS}/{nome}",f"{P6}/{nome}")
                res.append(info); st["feitas"]=_idx+1; continue
            fst["pct"]=65; fst["status"]="Confeccionando orçamento"
            pages=_split_pdf(fb) if (is_pdf and len(notas)>1) else None
            usou_source=False
            for i,nt in enumerate(notas):
                page=pages[i] if (pages and i<len(pages)) else fb
                ticket=_num_limpo(nt.get("ticket")); nota_num=_num_limpo(nt.get("nota_numero")) or "SN"
                itens=nt.get("itens") or []
                valor_nota=0.0
                for it in itens:
                    try: valor_nota+=float(it.get("valor_unit") or 0)*float(it.get("quant") or 0)
                    except Exception: pass
                valor_nota=round(valor_nota,2); valor_orc=round(valor_nota*1.20,2)
                info={"arquivo":nome,"ticket":ticket or None,"nota":nota_num,"itens":len(itens),"valor_nota":valor_nota,"valor_orcamento":valor_orc}
                def _mv_nota(destino_folder, destino_nome):
                    nonlocal usou_source
                    if previa or not destino_folder: return
                    if pages is None:
                        dropbox_rateio.mover(access,f"{ORC_NOTAS}/{nome}",f"{destino_folder}/{destino_nome}"); usou_source=True
                    else:
                        dropbox_rateio.subir_bytes(access,page,f"{destino_folder}/{destino_nome}",overwrite=True)
                if not ticket:
                    info.update(status="pendente",motivo="sem ticket",destino="6")
                    dest6 = nome if pages is None else f"{os.path.splitext(nome)[0]}_p{i+1}{ext}"
                    _mv_nota(P6, dest6)
                    if not previa: _orc_registra({"nota_numero":nota_num if nota_num!="SN" else None,"ticket":None,"status":"sem_ticket","valor_nota":valor_nota,"itens":itens,"arquivo_nota":f"6/{dest6}","criado_por":uid})
                    res.append(info); continue
                lj=_loja_do_ticket(ticket)
                if not lj:
                    info.update(status="pendente",motivo="ticket não encontrado nos chamados",destino="7")
                    dest7 = nome if pages is None else f"TICKET_{ticket}_NOTA_{nota_num}{ext}"
                    _mv_nota(P7, dest7)
                    if not previa: _orc_registra({"nota_numero":nota_num if nota_num!="SN" else None,"ticket":ticket,"status":"ticket_nao_associado","valor_nota":valor_nota,"itens":itens,"arquivo_nota":f"7/{dest7}","criado_por":uid})
                    res.append(info); continue
                loja=lj.get("loja") or {}
                loja_nome=(loja.get("nome") if loja else None) or re.sub(r"^LOJA\s*\d*\s*-?\s*","",lj.get("unidade") or "",flags=re.I).strip() or "—"
                extrap=valor_nota>ORC_EXTRAPOLA
                slug=_slug_loja(loja,lj.get("unidade"))
                mes=_mes_atual()   # roteia pelo mês em que o orçamento é GERADO (não pela data da nota/DAV)
                info.update(status="ok",loja=loja_nome,loja_numero=(loja.get("numero") if loja else None),extrapolado=extrap,mes=mes)
                if _nota_ja_gerada(ticket, nota_num):
                    # nota já processada antes -> orçamento já existe: EXCLUI a nota da pasta 0
                    # (arquivo único vai para a lixeira do Dropbox; multipágina é apagado no fim do laço)
                    info.update(status="duplicada",motivo="nota já processada — excluída da pasta 0",destino="lixeira")
                    if not previa and pages is None:
                        try: dropbox_rateio.apagar(access,f"{ORC_NOTAS}/{nome}"); usou_source=True
                        except Exception as e: info.update(motivo=f"duplicada — não consegui excluir: {str(e)[:80]}")
                    res.append(info); continue
                if not previa:
                    # dados do orçamento
                    hoje=_hoje().strftime("%d/%m/%Y")
                    itens_orc=[{"descricao":it.get("descricao"),"quant":float(it.get("quant") or 0),"unid":it.get("unid") or "UN","valor_unit":float(it.get("valor_unit") or 0)} for it in itens]
                    dados={"num":ticket,"revisao":1,"data":hoje,"loja_nome":loja_nome,
                        "prestador":{"nome":"Frota Macedo Engenharia LTDA","cnpj":"27.363.223/0001-70","forma":"Transferência Bancária 30 dias"},
                        "tomador":{"nome":f"Mercadinhos São Luiz — {loja_nome.title()}","cnpj":(loja.get("cnpj") if loja else None),
                                   "endereco":(loja.get("endereco") if loja else None),
                                   "cidade":((loja.get("cidade") if loja else None) or "")+(" - CE" if (loja.get("cidade") if loja else None) else "")},
                        "itens":itens_orc}
                    base_nome=f"{slug}_{ticket}_NOTA_{nota_num}"
                    doc_bytes=gera_orcamento_docx(dados)
                    arq_pdf=arq_doc=None
                    try:
                        if extrap:
                            if P9: dropbox_rateio.subir_bytes(access,doc_bytes,f"{P9}/{base_nome}.docx",overwrite=True); arq_doc=f"9/{base_nome}.docx"
                        else:
                            pdf_bytes=gera_orcamento_pdf(dados)
                            if P1: dropbox_rateio.subir_bytes(access,pdf_bytes,f"{P1}/{base_nome}.pdf",overwrite=True)
                            if P10:
                                dropbox_rateio.criar_pasta(access,f"{P10}/{mes}"); dropbox_rateio.criar_pasta(access,f"{P10}/{mes}/{slug}")
                                dropbox_rateio.subir_bytes(access,pdf_bytes,f"{P10}/{mes}/{slug}/{base_nome}.pdf",overwrite=True)
                                dropbox_rateio.subir_bytes(access,doc_bytes,f"{P10}/{mes}/{slug}/{base_nome}.docx",overwrite=True)
                                arq_pdf=f"10/{mes}/{slug}/{base_nome}.pdf"; arq_doc=f"10/{mes}/{slug}/{base_nome}.docx"
                    except Exception as e: info.update(motivo=f"salvar: {str(e)[:90]}")
                    # nota renomeada -> pasta 8 / <aba>
                    sub="INSTALACOES" if (lj.get("aba") or "").upper().startswith("INST") else ("CIVIL" if (lj.get("aba") or "").upper().startswith("CIV") else "SEM CLASSIFICACAO")
                    arq_nota=None
                    if P8:
                        dropbox_rateio.criar_pasta(access,f"{P8}/{sub}")
                        _mv_nota(f"{P8}/{sub}", f"TICKET_{ticket}_NOTA_{nota_num}{ext}")
                        arq_nota=f"8/{sub}/TICKET_{ticket}_NOTA_{nota_num}{ext}"
                    _orc_registra({"nota_numero":nota_num if nota_num!="SN" else None,"ticket":ticket,
                        "loja_numero":(loja.get("numero") if loja else None),"loja_nome":loja_nome,"aba":lj.get("aba"),
                        "valor_nota":valor_nota,"valor_orcamento":valor_orc,"status":"gerado","extrapolado":extrap,
                        "itens":itens_orc,"data_nota":_data_iso(nt.get("data_nota")),"mes_ref":mes,
                        "arquivo_nota":arq_nota,"arquivo_pdf":arq_pdf,"arquivo_doc":arq_doc,"criado_por":uid})
                res.append(info)
            # se dividiu em páginas, remove o arquivo-fonte da pasta 0 (páginas já foram espalhadas)
            if not previa and pages is not None and not usou_source:
                try: dropbox_rateio.apagar(access,f"{ORC_NOTAS}/{nome}")
                except Exception: pass
            fst["pct"]=100; fst["status"]="Pronto"
            st["feitas"]=_idx+1
        ok=sum(1 for r in res if r.get("status")=="ok")
        if st.get("cancelar"):
            for f in st["itens"]:                 # marca as que não chegaram a rodar
                if f.get("status") in (None,"aguardando"): f["status"]="interrompido"; f["pct"]=100
            if not previa: log_frotahub(uid,papel,"GERAR_ORCAMENTOS","INTERROMPEU",f"{ok}/{st['feitas']} de {len(_lista)}")
            st["gerados"]=ok; st["estado"]="cancelado"
        else:
            if not previa: log_frotahub(uid,papel,"GERAR_ORCAMENTOS","GEROU",f"{ok}/{len(res)}")
            st["gerados"]=ok; st["estado"]="pronto"
    except Exception as e:
        st["estado"]="erro"; st["erro"]=str(e)[:300]

@app.post("/orc/gerar")
async def orc_gerar(request: Request):
    from fastapi import HTTPException
    import threading
    u,p=exige(request,"GERAR_ORCAMENTOS")
    body={}
    try: body=await request.json()
    except Exception: pass
    previa=bool(body.get("previa"))
    if not (GEMINI_API_KEY or GROQ_KEY): raise HTTPException(500,"Configure GEMINI_API_KEY (imagens) e/ou GROQ_API_KEY (PDFs) no Render")
    # limpa jobs antigos (evita crescer sem fim)
    if len(_ORC_JOBS)>20:
        for k in list(_ORC_JOBS)[:-10]: _ORC_JOBS.pop(k,None)
    _ORC_JOB_SEQ[0]+=1; job=str(_ORC_JOB_SEQ[0])
    _ORC_JOBS[job]={"estado":"rodando","previa":previa,"total":0,"feitas":0,"gerados":0,
                    "res":[],"itens":[],"pausa":False,"retoma_em":0,"lote":ORC_LOTE,"descanso":ORC_DESCANSO}
    t=threading.Thread(target=_orc_job_run,args=(job,previa,u["id"],p["papel"]),daemon=True); t.start()
    return {"job":job,"lote":ORC_LOTE,"descanso":ORC_DESCANSO}

@app.post("/orc/gerar_cancelar")
async def orc_gerar_cancelar(request: Request):
    """Interrompe manualmente um job em andamento. Ele para após terminar a nota atual."""
    from fastapi import HTTPException
    exige(request,"GERAR_ORCAMENTOS")
    b={}
    try: b=await request.json()
    except Exception: pass
    job=str(b.get("job") or "")
    j=_ORC_JOBS.get(job)
    if not j: raise HTTPException(404,"job não encontrado")
    j["cancelar"]=True
    return {"ok":True,"estado":j.get("estado")}

@app.get("/orc/gerar_status")
def orc_gerar_status(request: Request, job: str=""):
    from fastapi import HTTPException
    exige(request,"GERAR_ORCAMENTOS")
    j=_ORC_JOBS.get(job)
    if not j: raise HTTPException(404,"job não encontrado (o motor pode ter reiniciado)")
    return {"estado":j["estado"],"previa":j["previa"],"total":j["total"],"feitas":j["feitas"],
            "gerados":j["gerados"],"pausa":j["pausa"],"retoma_em":j["retoma_em"],
            "lote":j["lote"],"descanso":j["descanso"],"erro":j.get("erro"),
            "itens":j.get("itens",[]),"groq_erro":j.get("groq_erro"),
            "resultados":j["res"] if j["estado"]!="rodando" else j["res"][-1:]}

@app.get("/orc/gemini_cota")
def orc_gemini_cota(request: Request):
    exige(request,"GERAR_ORCAMENTOS")
    return _gemini_cota()

@app.get("/orc/groq_teste")
def orc_groq_teste(request: Request):
    """Diagnóstico da chamada ao Groq: confirma a chave e testa os modelos de texto e visão."""
    exige(request,"GERAR_ORCAMENTOS")
    out={"groq_key_set": bool(GROQ_KEY), "modelo_texto":GROQ_TEXT_MODEL, "modelo_visao":GROQ_VIS_MODEL}
    if not GROQ_KEY:
        out["texto"]={"ok":False,"erro":"GROQ_API_KEY não configurada no Render"}; return out
    # 1) teste de TEXTO
    try:
        c=_groq_chat([{"role":"user","content":'Responda em JSON: {"notas":[{"ticket":"1","itens":[{"descricao":"x","quant":1,"unid":"UN","valor_unit":1}]}]}'}], GROQ_TEXT_FALLBACK)
        out["texto"]={"ok":True,"amostra":(c or "")[:200]}
    except Exception as e:
        out["texto"]={"ok":False,"erro":str(e)[:400]}
    # 2) teste de VISÃO (imagem 1x1 png)
    try:
        px=("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+M9QDwADhgGAWjR9awAAAABJRU5ErkJggg==")
        msgs=[{"role":"user","content":[{"type":"text","text":'Responda em JSON: {"notas":[]}'},
              {"type":"image_url","image_url":{"url":"data:image/png;base64,"+px}}]}]
        c=_groq_chat(msgs, GROQ_VIS_FALLBACK)
        out["visao"]={"ok":True,"amostra":(c or "")[:200]}
    except Exception as e:
        out["visao"]={"ok":False,"erro":str(e)[:400]}
    # 3) modelos disponíveis na conta (destaca os candidatos a visão)
    try:
        ids=sorted(_groq_models())
        vis=[m for m in ids if any(k in m.lower() for k in ("scout","maverick","vision","llama-4"))]
        out["modelos_disponiveis"]=ids
        out["modelos_visao_candidatos"]=vis
        out["modo_escaneado"]=("Groq visão" if (GROQ_VIS_MODEL and vis) else "OCR local (tesseract) + Groq texto")
    except Exception as e:
        out["modelos_disponiveis"]={"erro":str(e)[:300]}
    return out

@app.get("/orc/notas_pasta")
def orc_notas_pasta(request: Request):
    """Lista os arquivos da pasta 0 (nome + contagem) para a tela de gerar orçamentos."""
    from fastapi import HTTPException
    exige(request,"GERAR_ORCAMENTOS")
    try:
        access=dropbox_rateio.obter_token()
        _mb=_manut_base(access); _ob=_orc_base(access,_mb)
        ORC_NOTAS=_pasta_manut(access,0,_ob) or (_ob + "/0 - NOTAS PARA ORCAMENTO (COLOCAR AQUI)")
        arqs=[a for a in dropbox_rateio.listar(access,ORC_NOTAS) if a.lower().endswith((".pdf",".jpg",".jpeg",".png"))]
        arqs=sorted(arqs)
        return {"total":len(arqs),"arquivos":arqs}
    except Exception as e:
        raise HTTPException(500,f"lista: {str(e)[:160]}")

# ==================================================================
#  MEMÓRIA DE FORMATOS POR FORNECEDOR (aprende o layout por CNPJ)
# ==================================================================
_FROTA_CNPJ = re.sub(r"\D","",FAT_CNPJ)
def _cnpjs(texto):
    return [re.sub(r"\D","",m) for m in re.findall(r"\d{2}\.?\d{3}\.?\d{3}/?\d{4}-?\d{2}", texto or "")]
def _cnpj_emitente(texto):
    """CNPJ de quem EMITIU a nota (ignora o CNPJ da Frota, que é o destinatário)."""
    for c in _cnpjs(texto):
        if len(c)==14 and c!=_FROTA_CNPJ: return c
    return None
def _nome_emitente(texto):
    m=(re.search(r"RECEBEMOS DE\s+(.+?)\s+OS PRODUTOS", texto or "", re.I)
       or re.search(r"Raz[ãa]o Social\s*:?\s*(.+)", texto or "", re.I)
       or re.search(r"EMITENTE\s*:?\s*(.+)", texto or "", re.I))
    return (re.sub(r"\s{2,}"," ",m.group(1)).strip()[:80] if m else None)
def _qual_formato(texto):
    up=(texto or "").upper()
    if "DADOS DO PRODUTO" in up or "DOCUMENTO AUXILIAR DA" in up: return "danfe"
    return "dav"
def _formato_registra(texto, formato, nota):
    """Grava/incrementa o formato aprendido para o CNPJ do emitente. Best-effort."""
    try:
        cnpj=_cnpj_emitente(texto)
        if not cnpj: return
        _sb_json(f"{SB_URL}/rest/v1/rpc/formato_inc", SB_KEY, data={
            "p_cnpj":cnpj,"p_nome":_nome_emitente(texto),"p_formato":formato,"p_nota":str(nota or "")}, method="POST")
    except Exception: pass

@app.get("/config/formatos")
def config_formatos(request: Request):
    """Lista os formatos já aprendidos por fornecedor (memória que cresce com o uso)."""
    exige(request,"GERAR_ORCAMENTOS")
    try:
        rows=_sb_json(f"{SB_URL}/rest/v1/formatos_fornecedor?select=cnpj,nome,formato,acertos,ultima_em&order=acertos.desc&limit=500",SB_KEY) or []
    except Exception: rows=[]
    return {"formatos":rows}

# ==================================================================
#  RATEIO DE NOTAS (1 nota dividida entre vários chamados)
# ==================================================================
def _mime_de(nome):
    nl=(nome or "").lower()
    if nl.endswith(".pdf"): return "application/pdf"
    if nl.endswith((".xlsx",".xls")): return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    return "image/png" if nl.endswith(".png") else "image/jpeg"

def _rateio_base(access):
    mb=_manut_base(access); ob=_orc_base(access,mb); return ob

def _rateio_pasta3(access, ob=None):
    ob=ob or _rateio_base(access)
    return _pasta_manut(access,3,ob) or (ob + "/3 - NOTAS PARA RATEIO")

def _rateio_leitura(fb, mime, nome):
    """Lê a nota do rateio. Excel -> parser de planilha; senão a cascata (texto/OCR -> local -> Groq).
       Devolve (notas, texto)."""
    nl=(nome or "").lower()
    if nl.endswith(".xlsx") or nl.endswith(".xls"):
        notas=_parse_excel_nota(fb)
        if notas and _notas_seguras(notas):
            try: _formato_registra("", "excel", notas[0].get("nota_numero"))
            except Exception: pass
        return (notas or []), ""
    is_pdf=(nome or "").lower().endswith(".pdf") or mime=="application/pdf"
    texto=""
    if is_pdf: texto=_pdf_texto(fb)
    if len(texto)<40:
        try: texto=_ocr_texto(fb, mime, nome)
        except Exception: texto=""
    notas=_parse_nota_local(texto)
    if notas and _leitura_ok(notas, texto):
        _formato_registra(texto,_qual_formato(texto),notas[0].get("nota_numero")); return notas, texto
    if GROQ_KEY and texto:
        try:
            g=_groq_notas_texto(texto)
            if g and _leitura_ok(g, texto):
                _formato_registra(texto,"groq",g[0].get("nota_numero")); return g, texto
        except Exception: pass
    return (notas or []), texto

_RAT_INFO={}   # cache de leitura por arquivo
def _rateio_info(access, pasta3, arquivo):
    if arquivo in _RAT_INFO: return _RAT_INFO[arquivo]
    fb=dropbox_rateio.baixar(access,f"{pasta3}/{arquivo}")
    if not fb: return {"ok":False,"erro":"não baixou o arquivo"}
    notas,texto=_rateio_leitura(fb,_mime_de(arquivo),arquivo)
    if not notas:
        info={"ok":False,"erro":"não consegui ler a nota (formato novo?)"}
    else:
        nt=notas[0]
        info={"ok":True,"arquivo":arquivo,"fornecedor":nt.get("fornecedor") or _nome_emitente(texto) or "—",
              "data":nt.get("data_nota"),"valor":_soma_itens(notas),
              "nota_numero":nt.get("nota_numero"),"cnpj":nt.get("cnpj") or _cnpj_emitente(texto),
              "itens":nt.get("itens") or []}
    _RAT_INFO[arquivo]=info
    if len(_RAT_INFO)>60:
        for k in list(_RAT_INFO)[:30]: _RAT_INFO.pop(k,None)
    return info

def _linha_valor(a): return _num_br(a.get("quant"))*_num_br(a.get("valor_unit"))

def _groq_casa_materiais(materiais, rec):
    """Para cada MATERIAL, o Groq lista os chamados cujo PROBLEMA realmente usa aquele material.
       Devolve {idx_material: [numeros_ticket]}."""
    it_txt="\n".join(f"{i}: {m.get('descricao')}" for i,m in enumerate(materiais))
    tk_txt="\n".join(f"{t['numero']}: {t.get('descricao') or '(sem descrição)'}" for t in rec)
    prompt=("Você relaciona MATERIAIS de uma nota com CHAMADOS de manutenção. "
      "Para CADA material, liste os números dos chamados cujo PROBLEMA realmente USA aquele material "
      "(ex.: 'lâmpada/lamp' casa com chamado de TROCA DE LÂMPADA; 'refletor' NÃO casa com troca de lâmpada; "
      "'painel'/'tampão' que não aparecem em nenhuma descrição -> lista vazia). "
      "Seja rigoroso: só inclua o chamado se o material faz sentido para aquele problema. "
      'Responda só JSON: {"relacoes":[{"material":0,"tickets":["126460","126465"]}]}.'
      "\n\nMATERIAIS:\n"+it_txt+"\n\nCHAMADOS:\n"+tk_txt)
    c=_groq_chat([{"role":"user","content":prompt}], GROQ_TEXT_FALLBACK)
    d=_parse_json(c); out={}
    for a in (d.get("relacoes") or []):
        try: out[int(a.get("material"))]=[str(x) for x in (a.get("tickets") or [])]
        except Exception: pass
    return out

def _distribui(total, tickets):
    """Distribui uma quantidade entre tickets. Inteiro -> unidades inteiras (resto nos primeiros)."""
    n=len(tickets)
    if n==0: return {}
    if abs(total-round(total))<1e-9:
        tot=int(round(total)); base=tot//n; resto=tot-base*n
        return {t:(base+(1 if i<resto else 0)) for i,t in enumerate(tickets)}
    q=round(total/n,4); d={t:q for t in tickets}
    d[tickets[-1]]=round(total-q*(n-1),4)
    return d

def _ratear_matriz(itens, rec, nao):
    """Monta a matriz material × ticket. Casa cada material com os chamados que o usam (Groq),
       divide entre eles; material que não casa com ninguém é dividido entre TODOS.
       Garante >=1 item por ticket. Devolve (materiais, todos_tickets)."""
    rec_nums=[t["numero"] for t in rec]; todos=rec_nums+[t for t in nao if t not in rec_nums]
    mapa={}
    if rec and itens:
        try: mapa=_groq_casa_materiais(itens, rec)
        except Exception: mapa={}
    materiais=[]
    for i,it in enumerate(itens):
        q=_num_br(it.get("quant")); vu=_num_br(it.get("valor_unit"))
        matched=[t for t in (mapa.get(i) or []) if t in rec_nums]
        destinos = matched if matched else todos
        alloc=_distribui(q, destinos)
        full={t:round(alloc.get(t,0),4) for t in todos}
        materiais.append({"descricao":it.get("descricao"),"unid":it.get("unid") or "UN",
                          "valor_unit":vu,"quant_total":q,"alloc":full})
    # garante >=1 item por ticket (transfere metade da maior alocação de algum material)
    def _tem(t): return any(m["alloc"].get(t,0)>0 for m in materiais)
    for t in todos:
        if _tem(t) or not materiais: continue
        best=max(materiais, key=lambda m: m["quant_total"]*m["valor_unit"])
        doador=max(best["alloc"], key=lambda k: best["alloc"][k])
        mov=round(best["alloc"][doador]/2,4) or best["alloc"][doador]
        best["alloc"][doador]=round(best["alloc"][doador]-mov,4); best["alloc"][t]=round(best["alloc"].get(t,0)+mov,4)
    return materiais, todos

@app.get("/rateio/notas")
def rateio_notas(request: Request):
    from fastapi import HTTPException
    exige(request,"RATEIO_NOTAS")
    try:
        access=dropbox_rateio.obter_token(); p3=_rateio_pasta3(access)
        arqs=sorted([a for a in dropbox_rateio.listar(access,p3) if a.lower().endswith((".pdf",".jpg",".jpeg",".png",".xlsx",".xls"))])
        return {"arquivos":arqs,"total":len(arqs)}
    except Exception as e: raise HTTPException(500,f"lista: {str(e)[:160]}")

@app.get("/rateio/nota_info")
def rateio_nota_info(request: Request, arquivo: str=""):
    from fastapi import HTTPException
    exige(request,"RATEIO_NOTAS")
    try:
        access=dropbox_rateio.obter_token(); p3=_rateio_pasta3(access)
        return _rateio_info(access,p3,arquivo)
    except Exception as e: raise HTTPException(500,f"info: {str(e)[:160]}")

@app.get("/rateio/pdf")
def rateio_pdf(request: Request, arquivo: str=""):
    from fastapi import HTTPException
    exige(request,"RATEIO_NOTAS")
    try:
        access=dropbox_rateio.obter_token(); p3=_rateio_pasta3(access)
        fb=dropbox_rateio.baixar(access,f"{p3}/{arquivo}")
        if not fb: raise HTTPException(404,"arquivo não encontrado")
        return Response(content=fb, media_type=_mime_de(arquivo))
    except HTTPException: raise
    except Exception as e: raise HTTPException(500,f"pdf: {str(e)[:160]}")

def _ticket_info(tk):
    """Dados do chamado p/ o rateio: descrição, aba e a loja (do cadastro)."""
    ch=_sb_json(f"{SB_URL}/rest/v1/chamados?numero=eq.{urllib.parse.quote(tk)}&select=loja,aba,descricao&limit=1",SB_KEY) or []
    if not ch: return {"numero":tk,"reconhecido":False,"descricao":None,"aba":None,"loja":None}
    lj=_loja_do_ticket(tk) or {}
    return {"numero":tk,"reconhecido":True,"descricao":ch[0].get("descricao"),
            "aba":ch[0].get("aba"),"loja":lj.get("loja")}

@app.post("/rateio/preparar")
async def rateio_preparar(request: Request):
    from fastapi import HTTPException
    exige(request,"RATEIO_NOTAS")
    b=await request.json()
    arquivo=b.get("arquivo"); tickets=[re.sub(r"\D","",str(t))[:6] for t in (b.get("tickets") or []) if re.sub(r"\D","",str(t))]
    tickets=[t for i,t in enumerate(tickets) if t and t not in tickets[:i]]   # únicos, ordem preservada
    if len(tickets)<2: raise HTTPException(400,"informe pelo menos 2 chamados")
    access=dropbox_rateio.obter_token(); p3=_rateio_pasta3(access)
    info=_rateio_info(access,p3,arquivo)
    if not info.get("ok"): raise HTTPException(400, info.get("erro") or "não consegui ler a nota")
    itens=info["itens"]
    rec=[]; nao=[]; loja=None; aba_pad=None; lojas=set()
    for tk in tickets:
        ti=_ticket_info(tk)
        if ti["reconhecido"]:
            rec.append(ti)
            if ti.get("loja"): lojas.add(ti["loja"].get("numero"))
            if loja is None and ti.get("loja"): loja=ti["loja"]; aba_pad=ti.get("aba")
        else: nao.append(tk)
    if loja is None: raise HTTPException(400,"nenhum chamado reconhecido no banco — não sei a loja da nota")
    if len(lojas)>1: raise HTTPException(400,"os chamados são de lojas diferentes (o rateio exige a mesma loja)")
    materiais, todos=_ratear_matriz(itens, rec, nao)
    abas={t["numero"]:t.get("aba") for t in rec}
    tks=[{"numero":t,"reconhecido":t in abas,"aba":abas.get(t) or aba_pad,
          "descricao":next((x.get("descricao") for x in rec if x["numero"]==t),None)} for t in todos]
    return {"ok":True,"arquivo":arquivo,"fornecedor":info.get("fornecedor"),"nota_numero":info.get("nota_numero"),
            "data_nota":info.get("data"),"total_nota":info.get("valor"),
            "loja":{"numero":loja.get("numero"),"nome":loja.get("nome"),"cnpj":loja.get("cnpj"),
                    "endereco":loja.get("endereco"),"cidade":loja.get("cidade")},
            "tickets":tks,"materiais":materiais}

_RAT_JOBS={}; _RAT_SEQ=[0]; _RAT_LOCK=set()   # notas em processamento AGORA (trava duplo-clique / rodar 2x)
def _rat_job_run(job, dados_job, uid, papel):
    st=_RAT_JOBS[job]
    try:
        access=dropbox_rateio.obter_token(); ob=_rateio_base(access)
        P4=_pasta_manut(access,4,ob); P11=_pasta_manut(access,11,ob)
        P8=_pasta_manut(access,8,ob); P9=_pasta_manut(access,9,ob); p3=_rateio_pasta3(access,ob)
        arquivo=dados_job["arquivo"]; nota_num=_num_limpo(dados_job.get("nota_numero")) or "SN"
        loja=dados_job["loja"]; alocacao=dados_job["alocacao"]; ext=os.path.splitext(arquivo)[1] or ".pdf"
        loja_nome=loja.get("nome") or "—"; slug=_slug_loja(loja,"")
        mes=_mes_atual(); hoje=_hoje().strftime("%d/%m/%Y")
        st["total"]=len(alocacao); res=st["res"]
        for _i,al in enumerate(alocacao):
            tk=str(al.get("ticket")); aba=al.get("aba")
            itens_orc=[{"descricao":x.get("descricao"),"quant":_num_br(x.get("quant")),
                        "unid":x.get("unid") or "UN","valor_unit":_num_br(x.get("valor_unit"))} for x in (al.get("itens") or [])]
            valor_nota=round(sum(i["valor_unit"]*i["quant"] for i in itens_orc),2)
            valor_orc=round(valor_nota*1.20,2)
            r={"ticket":tk,"itens":len(itens_orc),"valor_nota":valor_nota,"valor_orcamento":valor_orc}
            base_nome=f"{slug}_{tk}_NOTA_{nota_num}_RATEIO"; arq_pdf=arq_doc=None
            if not itens_orc: r["status"]="pulado"; r["motivo"]="sem itens"; res.append(r); st["feitas"]=_i+1; continue
            # TRAVA DE DUPLICIDADE (cobre nota SEM número pelo nome-base do arquivo)
            if _nota_ja_gerada(tk, nota_num, base_nome):
                r["status"]="duplicada"; r["motivo"]="orçamento já existe para este ticket/nota — não gerado de novo"; res.append(r); st["feitas"]=_i+1; continue
            # TETO: nenhum orçamento de rateio (JÁ com +20%) pode passar de RATEIO_TETO
            if valor_orc>RATEIO_TETO:
                r["status"]="acima_limite"; r["motivo"]=f"orçamento {_rs(valor_orc)} acima do teto de {_rs(RATEIO_TETO)} — rateie em mais tickets ou faça manual"
                res.append(r); st["feitas"]=_i+1; continue
            extrap=valor_nota>ORC_EXTRAPOLA
            dados={"num":tk,"revisao":1,"data":hoje,"loja_nome":loja_nome,
                "prestador":{"nome":"Frota Macedo Engenharia LTDA","cnpj":"27.363.223/0001-70","forma":"Transferência Bancária 30 dias"},
                "tomador":{"nome":f"Mercadinhos São Luiz — {loja_nome.title()}","cnpj":loja.get("cnpj"),
                           "endereco":loja.get("endereco"),
                           "cidade":((loja.get("cidade") or ""))+(" - CE" if loja.get("cidade") else "")},
                "itens":itens_orc}
            try:
                doc_bytes=gera_orcamento_docx(dados)
                if extrap:
                    if P9: dropbox_rateio.subir_bytes(access,doc_bytes,f"{P9}/{base_nome}.docx",overwrite=True); arq_doc=f"9/{base_nome}.docx"
                else:
                    pdf_bytes=gera_orcamento_pdf(dados)
                    if P4: dropbox_rateio.subir_bytes(access,pdf_bytes,f"{P4}/{base_nome}.pdf",overwrite=True); arq_pdf=f"4/{base_nome}.pdf"
                    if P11:
                        dropbox_rateio.criar_pasta(access,f"{P11}/{mes}"); dropbox_rateio.criar_pasta(access,f"{P11}/{mes}/{slug}")
                        dropbox_rateio.subir_bytes(access,pdf_bytes,f"{P11}/{mes}/{slug}/{base_nome}.pdf",overwrite=True)
                        dropbox_rateio.subir_bytes(access,doc_bytes,f"{P11}/{mes}/{slug}/{base_nome}.docx",overwrite=True)
                        arq_pdf=f"11/{mes}/{slug}/{base_nome}.pdf"; arq_doc=f"11/{mes}/{slug}/{base_nome}.docx"
            except Exception as e: r["motivo"]=f"salvar: {str(e)[:90]}"
            _orc_registra({"nota_numero":nota_num if nota_num!="SN" else None,"ticket":tk,
                "loja_numero":loja.get("numero"),"loja_nome":loja_nome,"aba":aba,
                "valor_nota":valor_nota,"valor_orcamento":valor_orc,"status":"gerado","extrapolado":extrap,
                "rateio":True,"itens":itens_orc,"data_nota":_data_iso(dados_job.get("data_nota")),"mes_ref":mes,
                "arquivo_pdf":arq_pdf,"arquivo_doc":arq_doc,"criado_por":uid})
            r["status"]="ok"; r["loja"]=loja_nome; res.append(r); st["feitas"]=_i+1
        # move a nota para a pasta 8 (uma vez, no fim)
        sub="INSTALACOES" if (alocacao and (alocacao[0].get("aba") or "").upper().startswith("INST")) else ("CIVIL" if (alocacao and (alocacao[0].get("aba") or "").upper().startswith("CIV")) else "SEM CLASSIFICACAO")
        if P8:
            try:
                dropbox_rateio.criar_pasta(access,f"{P8}/{sub}")
                dropbox_rateio.mover(access,f"{p3}/{arquivo}",f"{P8}/{sub}/RATEIO_{nota_num}_{os.path.splitext(arquivo)[0]}{ext}")
            except Exception: pass
        _RAT_INFO.pop(arquivo,None)
        ok=sum(1 for x in res if x.get("status")=="ok")
        log_frotahub(uid,papel,"RATEIO_NOTAS","RATEOU",f"{ok}/{len(alocacao)} · nota {nota_num}")
        st["gerados"]=ok; st["estado"]="pronto"
    except Exception as e:
        st["estado"]="erro"; st["erro"]=str(e)[:300]
    finally:
        _RAT_LOCK.discard(dados_job.get("arquivo"))

@app.post("/rateio/gerar")
async def rateio_gerar(request: Request):
    from fastapi import HTTPException
    import threading
    u,p=exige(request,"RATEIO_NOTAS")
    b=await request.json()
    arquivo=b.get("arquivo"); alocacao=b.get("alocacao") or []; loja_numero=b.get("loja_numero")
    total_nota=_num_br(b.get("total_nota"))
    if not arquivo or not alocacao: raise HTTPException(400,"faltam dados do rateio")
    # confere conservação (a soma dos rateios tem que fechar o valor da nota)
    soma=round(sum(_num_br(x.get("valor_unit"))*_num_br(x.get("quant")) for al in alocacao for x in (al.get("itens") or [])),2)
    if total_nota>0 and abs(soma-total_nota)>max(0.05,0.01*total_nota):
        raise HTTPException(400,f"a soma dos rateios (R$ {soma}) não fecha o total da nota (R$ {total_nota})")
    if any(not (al.get("itens")) for al in alocacao):
        raise HTTPException(400,"todo chamado precisa de pelo menos 1 item")
    # loja oficial (do cadastro), nunca do cliente
    lj=(_sb_json(f"{SB_URL}/rest/v1/lojas?numero=eq.{int(loja_numero)}&limit=1",SB_KEY) or [None])[0] if loja_numero is not None else None
    if not lj: raise HTTPException(400,"loja não encontrada no cadastro")
    # TRAVA de concorrência: mesma nota não pode ser rateada 2x ao mesmo tempo (duplo clique / 2 abas)
    if arquivo in _RAT_LOCK: raise HTTPException(409,"esta nota já está sendo rateada agora — aguarde terminar")
    _RAT_LOCK.add(arquivo)
    _RAT_SEQ[0]+=1; job=str(_RAT_SEQ[0])
    if len(_RAT_JOBS)>20:
        for k in list(_RAT_JOBS)[:-10]: _RAT_JOBS.pop(k,None)
    _RAT_JOBS[job]={"estado":"rodando","total":0,"feitas":0,"gerados":0,"res":[]}
    dados_job={"arquivo":arquivo,"nota_numero":b.get("nota_numero"),"data_nota":b.get("data_nota"),
               "loja":lj,"alocacao":alocacao}
    threading.Thread(target=_rat_job_run,args=(job,dados_job,u["id"],p["papel"]),daemon=True).start()
    return {"job":job}

@app.get("/rateio/gerar_status")
def rateio_gerar_status(request: Request, job: str=""):
    from fastapi import HTTPException
    exige(request,"RATEIO_NOTAS")
    j=_RAT_JOBS.get(job)
    if not j: raise HTTPException(404,"job não encontrado")
    return {"estado":j["estado"],"total":j["total"],"feitas":j["feitas"],"gerados":j["gerados"],
            "erro":j.get("erro"),"resultados":j["res"]}

# ==================================================================
#  CORRIGIR NOTAS PENDENTES + GERAR ORÇAMENTOS CORRIGIDOS
# ==================================================================
@app.get("/orc/corrigir_listar")
def orc_corrigir_listar(request: Request):
    """Notas pendentes no BD (sem ticket / ticket não associado) + as já corrigidas."""
    exige(request,"CORRIGIR_NOTAS")
    q=("select=id,nota_numero,ticket,valor_nota,itens,status,arquivo_nota,loja_nome,criado_em"
       "&status=in.(sem_ticket,ticket_nao_associado,corrigido)&order=criado_em.desc&limit=1000")
    try: rows=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?{q}",SB_KEY) or []
    except Exception: rows=[]
    return {"itens":rows}

@app.post("/orc/corrigir_salvar")
async def orc_corrigir_salvar(request: Request):
    from fastapi import HTTPException
    u,p=exige(request,"CORRIGIR_NOTAS")
    b=await request.json()
    _id=b.get("id"); ticket=re.sub(r"\D","",str(b.get("ticket") or ""))[:6]; nota_num=_num_limpo(b.get("nota_numero"))
    if not _id: raise HTTPException(400,"id ausente")
    if len(ticket)<4: raise HTTPException(400,"informe um ticket válido (6 dígitos)")
    lj=_loja_do_ticket(ticket)
    if not lj: raise HTTPException(400,"ticket não encontrado nos chamados — atualize a Lista do Trílogo e tente de novo")
    loja=lj.get("loja") or {}
    ch=_sb_json(f"{SB_URL}/rest/v1/chamados?numero=eq.{urllib.parse.quote(ticket)}&select=descricao&limit=1",SB_KEY) or []
    desc=(ch[0].get("descricao") if ch else None)
    patch={"ticket":ticket,"nota_numero":(nota_num or None),"status":"corrigido",
           "loja_numero":loja.get("numero"),"loja_nome":loja.get("nome"),"aba":lj.get("aba")}
    try: _sb_write(f"notas_orcamento?id=eq.{_id}", patch, "PATCH")
    except Exception as e: raise HTTPException(400,f"salvar: {str(e)[:120]}")
    log_frotahub(u["id"],p.get("papel"),"CORRIGIR_NOTAS","CORRIGIU",f"{ticket}/NF {nota_num or '—'}")
    return {"ok":True,"loja_numero":loja.get("numero"),"loja_nome":loja.get("nome"),"aba":lj.get("aba"),"descricao":desc}

_CORR_JOBS={}; _CORR_SEQ=[0]
def _corr_job_run(job, uid, papel):
    st=_CORR_JOBS[job]
    try:
        access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
        P1=_pasta_manut(access,1,ob); P8=_pasta_manut(access,8,ob); P9=_pasta_manut(access,9,ob)
        P10=_pasta_manut(access,10,ob); P6=_pasta_manut(access,6,ob); P7=_pasta_manut(access,7,ob)
        recs=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?select=id,nota_numero,ticket,itens,aba,loja_numero,loja_nome,arquivo_nota&status=eq.corrigido&limit=1000",SB_KEY) or []
        st["total"]=len(recs); res=st["res"]; mes=_mes_atual(); hoje=_hoje().strftime("%d/%m/%Y")
        for _i,rec in enumerate(recs):
            tk=str(rec.get("ticket") or ""); nota_num=_num_limpo(rec.get("nota_numero")) or "SN"
            itens=rec.get("itens") or []
            r={"ticket":tk,"nota":nota_num,"itens":len(itens)}
            lj=_loja_do_ticket(tk); loja=(lj or {}).get("loja") or {}
            if not lj: r.update(status="erro",motivo="ticket não encontrado"); res.append(r); st["feitas"]=_i+1; continue
            if _nota_ja_gerada(tk, nota_num):
                _sb_write(f"notas_orcamento?id=eq.{rec['id']}", {"status":"duplicada"}, "PATCH")
                r.update(status="duplicada",motivo="orçamento já existe"); res.append(r); st["feitas"]=_i+1; continue
            loja_nome=loja.get("nome") or rec.get("loja_nome") or "—"
            valor_nota=round(sum(_num_br(x.get("valor_unit"))*_num_br(x.get("quant")) for x in itens),2)
            valor_orc=round(valor_nota*1.20,2); extrap=valor_nota>ORC_EXTRAPOLA
            slug=_slug_loja(loja,(lj or {}).get("unidade"))
            itens_orc=[{"descricao":x.get("descricao"),"quant":_num_br(x.get("quant")),"unid":x.get("unid") or "UN","valor_unit":_num_br(x.get("valor_unit"))} for x in itens]
            dados={"num":tk,"revisao":1,"data":hoje,"loja_nome":loja_nome,
                "prestador":{"nome":"Frota Macedo Engenharia LTDA","cnpj":"27.363.223/0001-70","forma":"Transferência Bancária 30 dias"},
                "tomador":{"nome":f"Mercadinhos São Luiz — {loja_nome.title()}","cnpj":loja.get("cnpj"),
                           "endereco":loja.get("endereco"),"cidade":((loja.get("cidade") or ""))+(" - CE" if loja.get("cidade") else "")},
                "itens":itens_orc}
            base_nome=f"{slug}_{tk}_NOTA_{nota_num}"; arq_pdf=arq_doc=None
            try:
                doc_bytes=gera_orcamento_docx(dados)
                if extrap:
                    if P9: dropbox_rateio.subir_bytes(access,doc_bytes,f"{P9}/{base_nome}.docx",overwrite=True); arq_doc=f"9/{base_nome}.docx"
                else:
                    pdf_bytes=gera_orcamento_pdf(dados)
                    if P1: dropbox_rateio.subir_bytes(access,pdf_bytes,f"{P1}/{base_nome}.pdf",overwrite=True)
                    if P10:
                        dropbox_rateio.criar_pasta(access,f"{P10}/{mes}"); dropbox_rateio.criar_pasta(access,f"{P10}/{mes}/{slug}")
                        dropbox_rateio.subir_bytes(access,pdf_bytes,f"{P10}/{mes}/{slug}/{base_nome}.pdf",overwrite=True)
                        dropbox_rateio.subir_bytes(access,doc_bytes,f"{P10}/{mes}/{slug}/{base_nome}.docx",overwrite=True)
                        arq_pdf=f"10/{mes}/{slug}/{base_nome}.pdf"; arq_doc=f"10/{mes}/{slug}/{base_nome}.docx"
            except Exception as e: r["motivo"]=f"salvar: {str(e)[:80]}"
            # move a nota da pasta 6/7 -> 8/<aba>
            aba=(lj or {}).get("aba") or rec.get("aba") or ""
            sub="INSTALACOES" if aba.upper().startswith("INST") else ("CIVIL" if aba.upper().startswith("CIV") else "SEM CLASSIFICACAO")
            arq_nota=rec.get("arquivo_nota"); dest_nota=None
            if P8 and arq_nota and "/" in arq_nota:
                npasta,_,nfile=arq_nota.partition("/")
                origem = P6 if npasta=="6" else (P7 if npasta=="7" else None)
                if origem:
                    try:
                        dropbox_rateio.criar_pasta(access,f"{P8}/{sub}")
                        ext=os.path.splitext(nfile)[1] or ".pdf"
                        dest_nota=f"8/{sub}/TICKET_{tk}_NOTA_{nota_num}{ext}"
                        dropbox_rateio.mover(access,f"{origem}/{nfile}",f"{P8}/{sub}/TICKET_{tk}_NOTA_{nota_num}{ext}")
                    except Exception: pass
            _sb_write(f"notas_orcamento?id=eq.{rec['id']}",
                {"status":"gerado","valor_nota":valor_nota,"valor_orcamento":valor_orc,"extrapolado":extrap,
                 "itens":itens_orc,"mes_ref":mes,"arquivo_pdf":arq_pdf,"arquivo_doc":arq_doc,"arquivo_nota":dest_nota or arq_nota}, "PATCH")
            r.update(status="ok",loja=loja_nome,valor_orcamento=valor_orc); res.append(r); st["feitas"]=_i+1
        ok=sum(1 for x in res if x.get("status")=="ok")
        log_frotahub(uid,papel,"GERAR_ORCAMENTOS_CORRIGIDOS","GEROU",f"{ok}/{len(recs)}")
        st["gerados"]=ok; st["estado"]="pronto"
    except Exception as e:
        st["estado"]="erro"; st["erro"]=str(e)[:300]

@app.post("/orc/corrigir_gerar")
async def orc_corrigir_gerar(request: Request):
    from fastapi import HTTPException
    import threading
    u,p=exige(request,"GERAR_ORCAMENTOS_CORRIGIDOS")
    _CORR_SEQ[0]+=1; job=str(_CORR_SEQ[0])
    if len(_CORR_JOBS)>20:
        for k in list(_CORR_JOBS)[:-10]: _CORR_JOBS.pop(k,None)
    _CORR_JOBS[job]={"estado":"rodando","total":0,"feitas":0,"gerados":0,"res":[]}
    threading.Thread(target=_corr_job_run,args=(job,u["id"],p["papel"]),daemon=True).start()
    return {"job":job}

@app.get("/orc/corrigir_gerar_status")
def orc_corrigir_gerar_status(request: Request, job: str=""):
    from fastapi import HTTPException
    exige(request,"GERAR_ORCAMENTOS_CORRIGIDOS")
    j=_CORR_JOBS.get(job)
    if not j: raise HTTPException(404,"job não encontrado")
    return {"estado":j["estado"],"total":j["total"],"feitas":j["feitas"],"gerados":j["gerados"],
            "erro":j.get("erro"),"resultados":j["res"]}

# ==================================================================
#  ORÇAR NOTAS A PARTIR DE TXT (plano C — quando a leitura automática falha)
#  Requer builder + PIN. NÃO roteia a nota física, só gera/roteia os orçamentos.
# ==================================================================
def _exige_builder(request):
    from fastapi import HTTPException
    u=auth_user(_bearer(request))
    if not u or not u.get("id"): raise HTTPException(401,"não autenticado")
    p=perfil_de(u["id"])
    if not p or p.get("nivel")!="builder": raise HTTPException(403,"apenas o builder")
    return u,p

def _parse_txt_notas(texto):
    """Lê o TXT gerado externamente. Blocos separados por '---' (ou por cada 'TICKET').
       Cada bloco: TICKET:, NOTA:, DATA:, e linhas ITEM: descrição | unidade | quant | preço_unit [| preço_total]."""
    T=texto or ""
    blocos=re.split(r"\n\s*-{3,}\s*\n|\n\s*={3,}\s*\n", T)
    if len(blocos)<=1:
        parts=re.split(r"(?=^\s*TICKET\b)", T, flags=re.I|re.M); blocos=[p for p in parts if p.strip()]
    notas=[]
    for bl in blocos:
        if not bl.strip(): continue
        mt=re.search(r"TICKET\s*[:\-]?\s*(\d{4,})", bl, re.I) or re.search(r"#\s*(\d{4,})", bl)
        mn=re.search(r"(?:NOTA|N[ºo°]?\s*DOC\w*|DOCUMENTO)\s*[:\-]?\s*(\d+)", bl, re.I)
        md=re.search(r"DATA\s*[:\-]?\s*(\d{2}/\d{2}/\d{4})", bl, re.I) or re.search(r"(\d{2}/\d{2}/\d{4})", bl)
        itens=[]
        for ln in bl.splitlines():
            m=re.match(r"\s*ITEM\s*[:\-]?\s*(.+)", ln, re.I)
            if not m: continue
            c=[x.strip() for x in m.group(1).split("|")]
            if len(c)<3: continue
            if len(c)>=4:
                desc=c[0]; unid=c[1] or "UN"; q=_num_br(c[2]); vu=_num_br(c[3]); vt=_num_br(c[4]) if len(c)>=5 else 0
            else:
                desc=c[0]; unid="UN"; q=_num_br(c[1]); vu=_num_br(c[2]); vt=0
            if vu<=0 and vt>0 and q>0: vu=round(vt/q,4)
            if q>0 and vu>0 and desc:
                itens.append({"descricao":desc[:120],"unid":(unid[:6] or "UN"),"quant":q,"valor_unit":_reconcilia(q,vu,vt)})
        if not itens: continue
        notas.append({"ticket":(mt.group(1) if mt else None),"nota_numero":(mn.group(1) if mn else None),
                      "data_nota":(md.group(1) if md else None),"itens":itens})
    return notas

@app.post("/orc/txt_parse")
async def orc_txt_parse(request: Request):
    from fastapi import HTTPException
    u,p=_exige_builder(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN do builder incorreto")
    notas=_parse_txt_notas(b.get("texto") or "")
    if not notas: raise HTTPException(400,"não encontrei nenhuma nota no TXT (confira o formato)")
    out=[]
    for i,nt in enumerate(notas):
        tk=_num_limpo(nt.get("ticket")); nota_num=_num_limpo(nt.get("nota_numero")) or "SN"
        itens=nt.get("itens") or []
        valor=round(sum(_num_br(x.get("valor_unit"))*_num_br(x.get("quant")) for x in itens),2)
        info={"i":i,"ticket":tk or None,"nota_numero":nt.get("nota_numero"),"data_nota":nt.get("data_nota"),
              "itens":itens,"valor":valor,"reconhecido":False,"loja_nome":None,"duplicada":False,"motivo":None}
        if not tk: info["motivo"]="sem ticket — irá para correção"
        else:
            lj=_loja_do_ticket(tk); loja=(lj or {}).get("loja") or {}
            if not lj: info["motivo"]="ticket não associado — irá para correção"
            else:
                info["reconhecido"]=True; info["loja_nome"]=loja.get("nome"); info["aba"]=lj.get("aba")
                if _nota_ja_gerada(tk, nota_num): info["duplicada"]=True; info["motivo"]="orçamento já existe para este ticket/nota"
        out.append(info)
    return {"notas":out}

def _nota_ja_registrada(ticket, nota_num):
    """True se já existe QUALQUER registro (pendente ou gerado) para essa nota — evita duplicar."""
    if not nota_num or str(nota_num)=="SN": return False
    parts=[f"nota_numero=eq.{nota_num}","select=ticket","limit=1"]
    parts.insert(0, f"ticket=eq.{ticket}" if ticket else "ticket=is.null")
    try: return bool(_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?{'&'.join(parts)}",SB_KEY))
    except Exception: return False

_TXT_JOBS={}; _TXT_SEQ=[0]
def _txt_job_run(job, notas, uid, papel):
    st=_TXT_JOBS[job]
    try:
        access=dropbox_rateio.obter_token(); ob=_orc_base(access,_manut_base(access))
        P1=_pasta_manut(access,1,ob); P9=_pasta_manut(access,9,ob); P10=_pasta_manut(access,10,ob)
        st["total"]=len(notas); res=st["res"]; mes=_mes_atual(); hoje=_hoje().strftime("%d/%m/%Y")
        for _i,nt in enumerate(notas):
            tk=_num_limpo(nt.get("ticket")); nota_num=_num_limpo(nt.get("nota_numero")) or "SN"
            itens=nt.get("itens") or []
            r={"ticket":tk,"nota":nota_num,"itens":len(itens)}
            valor_pend=round(sum(_num_br(x.get("valor_unit"))*_num_br(x.get("quant")) for x in itens),2)
            # SEM TICKET -> vai para o BD como pendente (aparece em "Corrigir notas"); nada de nota física
            if not tk:
                if _nota_ja_registrada(None, nota_num):
                    r.update(status="duplicada",motivo="nota já registrada"); res.append(r); st["feitas"]=_i+1; continue
                _orc_registra({"nota_numero":nota_num if nota_num!="SN" else None,"ticket":None,"status":"sem_ticket",
                               "valor_nota":valor_pend,"itens":itens,"data_nota":_data_iso(nt.get("data_nota")),"criado_por":uid})
                r.update(status="sem_ticket",motivo="enviada para correção"); res.append(r); st["feitas"]=_i+1; continue
            lj=_loja_do_ticket(tk); loja=(lj or {}).get("loja") or {}
            # TICKET NÃO ASSOCIADO -> BD como pendente
            if not lj:
                if _nota_ja_registrada(tk, nota_num):
                    r.update(status="duplicada",motivo="nota já registrada"); res.append(r); st["feitas"]=_i+1; continue
                _orc_registra({"nota_numero":nota_num if nota_num!="SN" else None,"ticket":tk,"status":"ticket_nao_associado",
                               "valor_nota":valor_pend,"itens":itens,"data_nota":_data_iso(nt.get("data_nota")),"criado_por":uid})
                r.update(status="ticket_nao_associado",motivo="enviada para correção"); res.append(r); st["feitas"]=_i+1; continue
            if _nota_ja_gerada(tk, nota_num):   # NUNCA processa a mesma nota 2x
                r.update(status="duplicada",motivo="orçamento já existe"); res.append(r); st["feitas"]=_i+1; continue
            loja_nome=loja.get("nome") or "—"
            valor_nota=round(sum(_num_br(x.get("valor_unit"))*_num_br(x.get("quant")) for x in itens),2)
            valor_orc=round(valor_nota*1.20,2); extrap=valor_nota>ORC_EXTRAPOLA
            slug=_slug_loja(loja,(lj or {}).get("unidade"))
            itens_orc=[{"descricao":x.get("descricao"),"quant":_num_br(x.get("quant")),"unid":x.get("unid") or "UN","valor_unit":_num_br(x.get("valor_unit"))} for x in itens]
            dados={"num":tk,"revisao":1,"data":hoje,"loja_nome":loja_nome,
                "prestador":{"nome":"Frota Macedo Engenharia LTDA","cnpj":"27.363.223/0001-70","forma":"Transferência Bancária 30 dias"},
                "tomador":{"nome":f"Mercadinhos São Luiz — {loja_nome.title()}","cnpj":loja.get("cnpj"),
                           "endereco":loja.get("endereco"),"cidade":((loja.get("cidade") or ""))+(" - CE" if loja.get("cidade") else "")},
                "itens":itens_orc}
            base_nome=f"{slug}_{tk}_NOTA_{nota_num}"; arq_pdf=arq_doc=None
            try:
                doc_bytes=gera_orcamento_docx(dados)
                if extrap:
                    if P9: dropbox_rateio.subir_bytes(access,doc_bytes,f"{P9}/{base_nome}.docx",overwrite=True); arq_doc=f"9/{base_nome}.docx"
                else:
                    pdf_bytes=gera_orcamento_pdf(dados)
                    if P1: dropbox_rateio.subir_bytes(access,pdf_bytes,f"{P1}/{base_nome}.pdf",overwrite=True)
                    if P10:
                        dropbox_rateio.criar_pasta(access,f"{P10}/{mes}"); dropbox_rateio.criar_pasta(access,f"{P10}/{mes}/{slug}")
                        dropbox_rateio.subir_bytes(access,pdf_bytes,f"{P10}/{mes}/{slug}/{base_nome}.pdf",overwrite=True)
                        dropbox_rateio.subir_bytes(access,doc_bytes,f"{P10}/{mes}/{slug}/{base_nome}.docx",overwrite=True)
                        arq_pdf=f"10/{mes}/{slug}/{base_nome}.pdf"; arq_doc=f"10/{mes}/{slug}/{base_nome}.docx"
            except Exception as e: r["motivo"]=f"salvar: {str(e)[:80]}"
            _orc_registra({"nota_numero":nota_num if nota_num!="SN" else None,"ticket":tk,
                "loja_numero":loja.get("numero"),"loja_nome":loja_nome,"aba":lj.get("aba"),
                "valor_nota":valor_nota,"valor_orcamento":valor_orc,"status":"gerado","extrapolado":extrap,
                "itens":itens_orc,"data_nota":_data_iso(nt.get("data_nota")),"mes_ref":mes,
                "arquivo_pdf":arq_pdf,"arquivo_doc":arq_doc,"criado_por":uid})
            r.update(status="ok",loja=loja_nome,valor_orcamento=valor_orc); res.append(r); st["feitas"]=_i+1
        ok=sum(1 for x in res if x.get("status")=="ok")
        log_frotahub(uid,papel,"ORCAR_TXT","GEROU",f"{ok}/{len(notas)}")
        st["gerados"]=ok; st["estado"]="pronto"
    except Exception as e:
        st["estado"]="erro"; st["erro"]=str(e)[:300]

@app.post("/orc/txt_gerar")
async def orc_txt_gerar(request: Request):
    from fastapi import HTTPException
    import threading
    u,p=_exige_builder(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN do builder incorreto")
    notas=b.get("notas") or []
    if not notas: raise HTTPException(400,"nenhuma nota para gerar")
    _TXT_SEQ[0]+=1; job=str(_TXT_SEQ[0])
    if len(_TXT_JOBS)>20:
        for k in list(_TXT_JOBS)[:-10]: _TXT_JOBS.pop(k,None)
    _TXT_JOBS[job]={"estado":"rodando","total":0,"feitas":0,"gerados":0,"res":[]}
    threading.Thread(target=_txt_job_run,args=(job,notas,u["id"],p["papel"]),daemon=True).start()
    return {"job":job}

@app.get("/orc/txt_gerar_status")
def orc_txt_gerar_status(request: Request, job: str=""):
    from fastapi import HTTPException
    _exige_builder(request)
    j=_TXT_JOBS.get(job)
    if not j: raise HTTPException(404,"job não encontrado")
    return {"estado":j["estado"],"total":j["total"],"feitas":j["feitas"],"gerados":j["gerados"],
            "erro":j.get("erro"),"resultados":j["res"]}

# ==================================================================
#  DASHBOARD (métricas gerais desde 01/07/2026)
# ==================================================================
DASH_INICIO = os.environ.get("DASH_INICIO","2026-07-01")
META_MATERIAL_MES = float(os.environ.get("META_MATERIAL_MES","80000"))  # 40% do contrato = R$ 80.000/mês
def _seg_semana(iso):
    """Devolve a segunda-feira (YYYY-MM-DD) da semana da data iso YYYY-MM-DD."""
    try:
        d=datetime.date.fromisoformat(iso[:10]); d=d-datetime.timedelta(days=d.weekday()); return d.isoformat()
    except Exception: return None

@app.get("/orc/dashboard")
def orc_dashboard(request: Request):
    from fastapi import HTTPException
    exige(request,"DASHBOARD")
    try:
        # ---- atendidos (marco vistoriado) desde o início ----
        ap=["select=numero,aba,loja,tipo_predial,responsavel,atendido_em","atendido=is.true",
            f"atendido_em=gte.{DASH_INICIO}","limit=50000"]
        at=_sb_json(f"{SB_URL}/rest/v1/chamados?"+"&".join(ap),SB_KEY) or []
        # ---- orçamentos gerados (custo de material) ----
        op=["select=ticket,loja_nome,aba,valor_nota,valor_orcamento,criado_em","status=eq.gerado","limit=50000"]
        orc=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?"+"&".join(op),SB_KEY) or []
    except Exception as e:
        raise HTTPException(500,f"dashboard: {str(e)[:160]}")

    por_dia={}; por_sem={}; por_mes={}; por_tipo={}; por_loja={}; por_resp={}; hab={}
    atendidos_set=set()
    for r in at:
        num=str(r.get("numero") or ""); atendidos_set.add(num)
        d=(r.get("atendido_em") or "")[:10]
        if d: por_dia[d]=por_dia.get(d,0)+1
        sw=_seg_semana(d)
        if sw: por_sem[sw]=por_sem.get(sw,0)+1
        mm=d[:7]
        if mm: por_mes[mm]=por_mes.get(mm,0)+1
        tp=(r.get("tipo_predial") or "—").strip() or "—"; por_tipo[tp]=por_tipo.get(tp,0)+1
        lj=_loja_padrao(r.get("loja")); por_loja[lj]=por_loja.get(lj,0)+1
        rp=(r.get("responsavel") or "—").strip() or "—"; por_resp[rp]=por_resp.get(rp,0)+1
        h=hab.setdefault(rp,{"total":0,"por_tipo":{}}); h["total"]+=1
        h["por_tipo"][tp]=h["por_tipo"].get(tp,0)+1

    # ---- custo de material por mês + tickets com custo ----
    custo_mes={}; custo_set=set(); v_orc_total=0.0; v_nota_total=0.0
    for o in orc:
        tk=str(o.get("ticket") or "")
        if tk: custo_set.add(tk)
        vo=_numf(o.get("valor_orcamento")); vn=_numf(o.get("valor_nota"))
        v_orc_total+=vo; v_nota_total+=vn
        mm=(o.get("criado_em") or "")[:7]
        if mm:
            c=custo_mes.setdefault(mm,{"orc":0.0,"nota":0.0,"n":0}); c["orc"]+=vo; c["nota"]+=vn; c["n"]+=1

    com_custo=len(atendidos_set & custo_set)
    total_at=len(atendidos_set); sem_custo=max(0,total_at-com_custo)

    def _serie(d, label_fmt=None):
        return [{"x":k,"y":v} for k,v in sorted(d.items())]
    def _rank(d, n=None):
        it=sorted(({"nome":k,"total":v} for k,v in d.items()), key=lambda x:-x["total"])
        return it[:n] if n else it
    # meta por mês
    meta_serie=[]
    for mm in sorted(custo_mes):
        c=custo_mes[mm]["orc"]
        meta_serie.append({"mes":mm,"custo":round(c,2),"meta":META_MATERIAL_MES,
                           "pct":round(c*100.0/META_MATERIAL_MES,1) if META_MATERIAL_MES else None,
                           "saldo":round(META_MATERIAL_MES-c,2)})
    # habilidade -> top tipo por responsável
    hab_out=[]
    for rp,h in sorted(hab.items(), key=lambda x:-x[1]["total"]):
        tipos=sorted(({"tipo":k,"n":v} for k,v in h["por_tipo"].items()), key=lambda x:-x["n"])
        forte=tipos[0]["tipo"] if tipos else "—"
        hab_out.append({"responsavel":rp,"total":h["total"],"forte":forte,"tipos":tipos[:6]})

    return {
        "inicio":DASH_INICIO,"meta_mes":META_MATERIAL_MES,
        "total_atendidos":total_at,"com_custo":com_custo,"sem_custo":sem_custo,
        "pct_com_custo":(round(com_custo*100.0/total_at,1) if total_at else 0),
        "valor_orcamentos":round(v_orc_total,2),"valor_notas":round(v_nota_total,2),
        "atendidos_dia":_serie(por_dia),"atendidos_semana":_serie(por_sem),"atendidos_mes":_serie(por_mes),
        "por_tipo":_rank(por_tipo),"por_loja":_rank(por_loja),"por_responsavel":_rank(por_resp),
        "custo_mes":meta_serie,"habilidade":hab_out,
    }

# ==================================================================
#  ESTATÍSTICA CONSOLIDADA (3 escopos + por loja + tempos)
# ==================================================================
def _pd(s):
    try: return datetime.date.fromisoformat(str(s)[:10])
    except Exception: return None
def _pdt(s):
    """Data OU timestamp -> datetime naive (sem tz)."""
    if not s: return None
    s=str(s)
    try:
        d=datetime.datetime.fromisoformat(s.replace("Z","+00:00"))
        return d.replace(tzinfo=None)
    except Exception:
        d=_pd(s); return datetime.datetime.combine(d,datetime.time()) if d else None
def _tendencia(serie):
    """serie: lista de {x,y} ordenada. Compara 1ª metade x 2ª metade."""
    ys=[float(p.get("y") or 0) for p in serie]; n=len(ys)
    if n<2: return {"direcao":"estavel","variacao_pct":0}
    h=n//2; a=sum(ys[:h])/max(1,h); b=sum(ys[h:])/max(1,n-h)
    if a==0: return {"direcao":("alta" if b>0 else "estavel"),"variacao_pct":(100.0 if b>0 else 0)}
    var=(b-a)*100.0/a
    return {"direcao":("alta" if var>5 else ("baixa" if var<-5 else "estavel")),"variacao_pct":round(var,1)}
def _serie_periodo(datas_val, gran):
    """datas_val: lista de (date, valor). gran: 'dia'|'semana'|'mes'. Soma por bucket."""
    ac={}
    for d,v in datas_val:
        if not d: continue
        if gran=="mes": k=d.isoformat()[:7]
        elif gran=="semana": k=_seg_semana(d.isoformat())
        else: k=d.isoformat()
        if k: ac[k]=ac.get(k,0)+v
    return [{"x":k,"y":round(av,2)} for k,av in sorted(ac.items())]

@app.get("/orc/estat")
def orc_estat(request: Request, modo: str="inicio", de: str="", ate: str="", mes: str="", loja: str=""):
    from fastapi import HTTPException
    exige(request,"DASHBOARD")
    hoje=_hoje()
    ini=_pd(DASH_INICIO) or datetime.date(2026,7,1)
    if modo=="periodo":
        d0=_pd(de) or ini; d1=_pd(ate) or hoje
    elif modo=="mensal":
        mm=(mes or hoje.isoformat())[:7]
        try:
            y,m=int(mm[:4]),int(mm[5:7]); d0=datetime.date(y,m,1)
            d1=datetime.date(y+(1 if m==12 else 0),(1 if m==12 else m+1),1)-datetime.timedelta(days=1)
        except Exception: d0,d1=ini,hoje; modo="inicio"
    else:
        modo="inicio"; d0,d1=ini,hoje
    if d1<d0: d0,d1=d1,d0
    alvo=_loja_padrao(loja) if loja else None
    try:
        ap=["select=numero,aba,loja,tipo_predial,atendido_em","atendido=is.true",
            f"atendido_em=gte.{DASH_INICIO}","limit=50000"]
        at=_sb_json(f"{SB_URL}/rest/v1/chamados?"+"&".join(ap),SB_KEY) or []
        op=["select=ticket,loja_nome,aba,valor_nota,valor_orcamento,criado_em","status=eq.gerado","limit=50000"]
        orc=_sb_json(f"{SB_URL}/rest/v1/notas_orcamento?"+"&".join(op),SB_KEY) or []
        mu=_sb_json(f"{SB_URL}/rest/v1/chamado_mau_uso?select=numero,categoria,tem_mau_uso,loja,data_criacao&limit=50000",SB_KEY) or []
    except Exception as e:
        raise HTTPException(500,f"estat: {str(e)[:160]}")
    # granularidade da tendência conforme o span
    span=(d1-d0).days
    gran = "dia" if span<=92 else ("semana" if span<=400 else "mes")
    # ---- chamados atendidos no escopo ----
    at_set=set(); por_tipo={}; por_loja={}; at_datas=[]
    for r in at:
        d=_pd(r.get("atendido_em"))
        if not d or not (d0<=d<=d1): continue
        if alvo and _loja_padrao(r.get("loja"))!=alvo: continue
        num=str(r.get("numero") or ""); at_set.add(num); at_datas.append((d,1))
        tp=(r.get("tipo_predial") or "—").strip() or "—"; por_tipo[tp]=por_tipo.get(tp,0)+1
        lj=_loja_padrao(r.get("loja")); por_loja[lj]=por_loja.get(lj,0)+1
    # ---- custo de material no escopo ----
    custo_set=set(); custo_datas=[]; v_orc=0.0; v_nota=0.0
    for o in orc:
        d=_pd(o.get("criado_em"))
        if not d or not (d0<=d<=d1): continue
        if alvo and _loja_padrao(o.get("loja_nome"))!=alvo: continue
        tk=str(o.get("ticket") or "")
        if tk: custo_set.add(tk)
        vo=_numf(o.get("valor_orcamento")); vn=_numf(o.get("valor_nota"))
        v_orc+=vo; v_nota+=vn; custo_datas.append((d,vo))
    total_at=len(at_set); com=len(at_set & custo_set); sem=max(0,total_at-com)
    # meta do período (proporcional aos dias; mensal = 1 mês cheio)
    meses = 1.0 if modo=="mensal" else max(0.03,(span+1)/30.44)
    meta_periodo = META_MATERIAL_MES*meses
    serie_ch=_serie_periodo(at_datas,gran)
    serie_cst=_serie_periodo(custo_datas,gran)
    # ---- mau uso no escopo ----
    mu_com=0; mu_tot=0; mu_cat={}; mu_loja={}
    for m in mu:
        d=_pd(m.get("data_criacao"))
        if d and not (d0<=d<=d1): continue
        if alvo and _loja_padrao(m.get("loja"))!=alvo: continue
        mu_tot+=1
        cat=m.get("categoria") or "INDETERMINADO"
        if cat in MAU_USO_SEM: continue
        mu_com+=1; mu_cat[cat]=mu_cat.get(cat,0)+1
        lj=_loja_padrao(m.get("loja")); mu_loja[lj]=mu_loja.get(lj,0)+1
    def _rank(d,n=12):
        return sorted(({"nome":k,"total":v} for k,v in d.items()),key=lambda x:-x["total"])[:n]
    mu_cats=sorted(({"categoria":k,"nome":MAU_USO_MAP.get(k,k),"total":v} for k,v in mu_cat.items()),key=lambda x:-x["total"])
    return {
        "modo":modo,"de":d0.isoformat(),"ate":d1.isoformat(),"loja":(alvo or ""),
        "meta_mes":META_MATERIAL_MES,"meta_periodo":round(meta_periodo,2),
        "custo_material":round(v_orc,2),"valor_notas":round(v_nota,2),
        "meta_pct":(round(v_orc*100.0/meta_periodo,1) if meta_periodo else None),
        "meta_saldo":round(meta_periodo-v_orc,2),
        "total_atendidos":total_at,"com_custo":com,"sem_custo":sem,
        "pct_com_custo":(round(com*100.0/total_at,1) if total_at else 0),
        "serie_chamados":serie_ch,"tend_chamados":_tendencia(serie_ch),
        "serie_custo":serie_cst,"tend_custo":_tendencia(serie_cst),"gran":gran,
        "por_tipo":_rank(por_tipo),"por_loja":_rank(por_loja),
        "mau_uso":{"total":mu_tot,"com_mau_uso":mu_com,
                   "pct":(round(mu_com*100.0/mu_tot,1) if mu_tot else 0),
                   "por_categoria":mu_cats[:12],"por_loja":_rank(mu_loja)},
        "lojas": sorted({_loja_padrao(r.get("loja")) for r in at if r.get("loja")}),
    }

@app.get("/orc/estat_tempos")
def orc_estat_tempos(request: Request, lojas: str=""):
    from fastapi import HTTPException
    exige(request,"DASHBOARD")
    # tolerante: se as colunas de transição ainda não existem (SQL não rodado), carrega sem elas
    ch=None
    for cols in ("numero,aba,loja,data_criacao,em_execucao_em,executado_em,atendido_em",
                 "numero,aba,loja,data_criacao,atendido_em"):
        try:
            ch=_sb_json(f"{SB_URL}/rest/v1/chamados?select={cols}&data_criacao=gte.{DASH_INICIO}&limit=50000",SB_KEY) or []
            break
        except Exception as e:
            last=str(e)[:160]; ch=None
    if ch is None:
        raise HTTPException(500,f"tempos: {last}")
    disp=sorted({_loja_padrao(c.get("loja")) for c in ch if c.get("loja")})
    sel=[_loja_padrao(x) for x in (lojas or "").split("||") if x.strip()]
    base=[c for c in ch if (not sel or _loja_padrao(c.get("loja")) in sel)]
    esperas=[]; reparos=[]; totais=[]; aberturas=[]
    for c in base:
        dc=_pdt(c.get("data_criacao")); ex=_pdt(c.get("em_execucao_em"))
        fx=_pdt(c.get("executado_em")) or _pdt(c.get("atendido_em"))
        if dc: aberturas.append(dc)
        if dc and ex and ex>=dc: esperas.append((ex-dc).total_seconds())
        if ex and fx and fx>=ex: reparos.append((fx-ex).total_seconds())
        if dc and fx and fx>=dc: totais.append((fx-dc).total_seconds())
    aberturas.sort()
    gaps=[(aberturas[i]-aberturas[i-1]).total_seconds() for i in range(1,len(aberturas))]
    def med(lst): return (round(sum(lst)/len(lst)/86400.0,1) if lst else None)
    return {
        "lojas_disponiveis":disp,"selecionadas":sel,"n_chamados":len(base),
        "espera":{"dias":med(esperas),"n":len(esperas)},
        "reparo":{"dias":med(reparos),"n":len(reparos)},
        "total":{"dias":med(totais),"n":len(totais)},
        "entre":{"dias":med(gaps),"n":len(gaps)},
    }

# ==================================================================
#  ANÁLISE DE MAU USO (más condutas prováveis por trás do chamado)
# ==================================================================
# Taxonomia com distinção MAU USO (comportamento/operação) x ESTRUTURAL/DESGASTE.
# Prompt em modo LIBERAL: o Groq aponta a provável má conduta com liberdade (ver _groq_mau_uso).
MAU_USO_CATS = [
 # --- MAU USO real ---
 ("ENTUPIMENTO_DESCARTE","Entupimento por descarte (gordura/resíduo em ralo/caixa de gordura)"),
 ("IMPACTO_OPERACAO","Impacto de carrinho/empilhadeira/paleteira"),
 ("FERRAGEM_FORCA","Ferragem/fechadura/maçaneta quebrada por força/uso"),
 ("SOBRECARGA_ELETRICA","Sobrecarga elétrica / uso indevido de tomada"),
 ("AGUA_AREA_IMPROPRIA","Água/umidade indevida por lavagem/descuido (não telhado)"),
 ("OPERACAO_EQUIPAMENTO","Operação incorreta de equipamento"),
 ("FALTA_LIMPEZA","Falta de limpeza/conservação que gerou dano"),
 # --- NÃO é mau uso (estrutural/desgaste) ---
 ("INFILTRACAO_ESTRUTURAL","Infiltração/goteira de teto/telhado/laje (estrutural)"),
 ("DESGASTE_ACABAMENTO","Desgaste de piso/cerâmica/pintura/forro"),
 ("HIDRAULICA_DESGASTE","Vazamento de tubulação/registro/louça por desgaste"),
 ("OUTRO_ESTRUTURAL","Outro reparo predial sem indício de mau uso"),
 ("INDETERMINADO","Indeterminado / descrição insuficiente"),
]
MAU_USO_MAP=dict(MAU_USO_CATS)
MAU_USO_SEM={"INFILTRACAO_ESTRUTURAL","DESGASTE_ACABAMENTO","HIDRAULICA_DESGASTE","OUTRO_ESTRUTURAL","INDETERMINADO"}

def _groq_mau_uso(descricoes):
    """Classifica um LOTE de descrições. Devolve {idx: (categoria, explicacao)}."""
    cats=" | ".join(f"{c}={n}" for c,n in MAU_USO_CATS)
    corpo="\n".join(f"{i}: {d}" for i,d in enumerate(descricoes))
    prompt=("Você é engenheiro de manutenção predial de uma rede de supermercados. Para CADA chamado, "
      "identifique a PROVÁVEL MÁ CONDUTA / uso inadequado do lojista que pode ter contribuído para o problema. "
      "Seja criterioso, mas NÃO seja tímido: sempre que houver um indício plausível de que operação, descuido, "
      "descarte errado, sobrecarga ou impacto contribuíram, aponte a categoria de MAU USO correspondente. "
      "Use as categorias estruturais/desgaste apenas quando NÃO houver qualquer contribuição plausível de comportamento.\n"
      "Guia (aponte mau uso quando fizer sentido):\n"
      "- entupimento de ralo/caixa de gordura/esgoto em food service, padaria, açougue, manipulação -> ENTUPIMENTO_DESCARTE\n"
      "- porta/gôndola/parede/quina batida, amassada ou danificada por carrinho/empilhadeira/paleteira/manobra -> IMPACTO_OPERACAO\n"
      "- fechadura/maçaneta/trinco/dobradiça/puxador quebrado, forçado ou 'quebrou de novo' -> FERRAGEM_FORCA\n"
      "- tomada/disjuntor que cai ou queima com equipamento, sobrecarga, gambiarra, extensão -> SOBRECARGA_ELETRICA\n"
      "- água/umidade por lavagem, mangueira, descuido em área imprópria (NÃO telhado) -> AGUA_AREA_IMPROPRIA\n"
      "- equipamento operado errado (câmara, fritadeira, balança, motor) -> OPERACAO_EQUIPAMENTO\n"
      "- sujeira/gordura/falta de limpeza que causou o dano ou entupimento -> FALTA_LIMPEZA\n"
      "Estrutural/desgaste (só quando o uso claramente não contribuiu):\n"
      "- goteira/infiltração de teto/telhado/laje na chuva -> INFILTRACAO_ESTRUTURAL\n"
      "- cerâmica/rejunte/pintura/piso/forro por tempo e tráfego -> DESGASTE_ACABAMENTO\n"
      "- vazamento de tubulação/registro/louça por desgaste natural -> HIDRAULICA_DESGASTE\n"
      "- reparo predial comum sem indício -> OUTRO_ESTRUTURAL ; descrição vaga demais -> INDETERMINADO\n"
      "Obs.: água que atinge tomada/quadro POR infiltração de telhado é INFILTRACAO_ESTRUTURAL.\n"
      "CATEGORIAS (use o CÓDIGO): "+cats+"\n\n"
      "Para cada item responda: i, categoria (código), explicacao (frase curta em pt-BR até 120 caracteres com a "
      "provável má conduta/causa). Responda só JSON: "
      '{"itens":[{"i":0,"categoria":"ENTUPIMENTO_DESCARTE","explicacao":"..."}]}\n\nCHAMADOS:\n'+corpo)
    c=_groq_chat([{"role":"user","content":prompt}], GROQ_TEXT_FALLBACK)
    d=_parse_json(c); out={}
    for a in (d.get("itens") or []):
        try:
            i=int(a.get("i")); cat=str(a.get("categoria") or "").strip().upper()
            if cat not in MAU_USO_MAP: cat="INDETERMINADO"
            out[i]=(cat, str(a.get("explicacao") or "")[:160])
        except Exception: pass
    return out

_MU_JOBS={}; _MU_SEQ=[0]
def _mau_uso_job(job, uid, papel, reset=False):
    st=_MU_JOBS[job]
    try:
        if reset:   # taxonomia mudou -> reanálise total: limpa SÓ o que a IA classificou (preserva manual)
            try:
                rq=urllib.request.Request(f"{SB_URL}/rest/v1/chamado_mau_uso?manual=is.false",method="DELETE",
                    headers={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}","prefer":"return=minimal"})
                urllib.request.urlopen(rq,timeout=60)
            except Exception as e: st["erro"]=f"reset: {str(e)[:120]}"
        # chamados desde o início, com descrição
        ap=["select=numero,aba,loja,tipo_predial,descricao,data_criacao",f"data_criacao=gte.{DASH_INICIO}","limit=50000"]
        ch=_sb_json(f"{SB_URL}/rest/v1/chamados?"+"&".join(ap),SB_KEY) or []
        # não reanalisa quem já foi classificado; no reset, mantém as classificações MANUAIS
        _q="chamado_mau_uso?select=numero,aba&"+("manual=is.true&" if reset else "")+"limit=50000"
        ja=_sb_json(f"{SB_URL}/rest/v1/{_q}",SB_KEY) or []
        feitos={(str(x.get("numero")),str(x.get("aba") or "")) for x in ja}
        pend=[c for c in ch if (str(c.get("numero")),str(c.get("aba") or "")) not in feitos and (c.get("descricao") or "").strip()]
        st["total"]=len(pend); st["feitas"]=0
        if not pend: st["estado"]="pronto"; return
        LOTE=15
        for k in range(0,len(pend),LOTE):
            grupo=pend[k:k+LOTE]
            descs=[f"{(c.get('tipo_predial') or '')}: {c.get('descricao')}" for c in grupo]
            try: res=_groq_mau_uso(descs)
            except Exception as e:
                st["erro"]=f"Groq: {str(e)[:150]}"; res={}
            linhas=[]
            for i,c in enumerate(grupo):
                cat,exp = res.get(i,("INDETERMINADO",""))
                linhas.append({"numero":str(c.get("numero")),"aba":str(c.get("aba") or ""),
                    "categoria":cat,"tem_mau_uso":(cat not in MAU_USO_SEM),"explicacao":exp,
                    "loja":c.get("loja"),"tipo_predial":c.get("tipo_predial"),"data_criacao":c.get("data_criacao")})
            try:
                body=json.dumps(linhas,ensure_ascii=False).encode()
                rq=urllib.request.Request(f"{SB_URL}/rest/v1/chamado_mau_uso?on_conflict=numero,aba",data=body,method="POST",
                    headers={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}","content-type":"application/json","prefer":"resolution=merge-duplicates,return=minimal"})
                urllib.request.urlopen(rq,timeout=40)
            except Exception as e: st["erro"]=f"gravar: {str(e)[:150]}"
            st["feitas"]=min(len(pend),k+LOTE); st["novos"]=st.get("novos",0)+len(linhas)
        log_frotahub(uid,papel,"MAU_USO","ANALISOU",f"{st['feitas']} chamados")
        st["estado"]="pronto"
    except Exception as e:
        st["estado"]="erro"; st["erro"]=str(e)[:300]

@app.post("/orc/mau_uso_analisar")
async def orc_mau_uso_analisar(request: Request, reset: int=0):
    from fastapi import HTTPException
    import threading
    u,p=exige(request,"MAU_USO")
    if not GROQ_KEY: raise HTTPException(400,"configure a GROQ_API_KEY no Render para a análise")
    _MU_SEQ[0]+=1; job=str(_MU_SEQ[0])
    if len(_MU_JOBS)>10:
        for k in list(_MU_JOBS)[:-5]: _MU_JOBS.pop(k,None)
    _MU_JOBS[job]={"estado":"rodando","total":0,"feitas":0,"novos":0}
    threading.Thread(target=_mau_uso_job,args=(job,u["id"],p["papel"],bool(reset)),daemon=True).start()
    return {"job":job}

@app.get("/orc/mau_uso_status")
def orc_mau_uso_status(request: Request, job: str=""):
    from fastapi import HTTPException
    exige(request,"MAU_USO")
    j=_MU_JOBS.get(job)
    if not j: raise HTTPException(404,"job não encontrado")
    return {"estado":j["estado"],"total":j["total"],"feitas":j["feitas"],"novos":j.get("novos",0),"erro":j.get("erro")}

@app.get("/orc/mau_uso_chamado")
def orc_mau_uso_chamado(request: Request, numero: str=""):
    exige(request,"MAU_USO")
    numero=re.sub(r"\D","",numero or "")
    if not numero: return {"ok":False,"erro":"informe o número do chamado"}
    r=_sb_json(f"{SB_URL}/rest/v1/chamado_mau_uso?numero=eq.{numero}&select=numero,aba,categoria,tem_mau_uso,explicacao,loja,tipo_predial&limit=4",SB_KEY) or []
    ch=_sb_json(f"{SB_URL}/rest/v1/chamados?numero=eq.{numero}&select=loja,descricao,tipo_predial&limit=1",SB_KEY) or []
    descricao=(ch[0].get("descricao") if ch else None)
    itens=[{**x,"categoria_nome":MAU_USO_MAP.get(x.get("categoria"),x.get("categoria")),"loja_padrao":_loja_padrao(x.get("loja"))} for x in r]
    return {"ok":bool(itens),"numero":numero,"descricao":descricao,"itens":itens,
            "erro":(None if itens else "chamado ainda não analisado (rode a análise)")}

@app.get("/orc/mau_uso_stats")
def orc_mau_uso_stats(request: Request):
    exige(request,"MAU_USO")
    rows=_sb_json(f"{SB_URL}/rest/v1/chamado_mau_uso?select=categoria,tem_mau_uso,loja&limit=50000",SB_KEY) or []
    total=len(rows); com=sum(1 for r in rows if r.get("tem_mau_uso"))
    por_cat={}; por_loja={}; rank_cat={}
    for r in rows:
        cat=r.get("categoria") or "INDETERMINADO"
        if cat in MAU_USO_SEM: continue          # só conta mau uso real
        lj=_loja_padrao(r.get("loja"))
        por_cat[cat]=por_cat.get(cat,0)+1
        por_loja[lj]=por_loja.get(lj,0)+1
        rank_cat.setdefault(cat,{}); rank_cat[cat][lj]=rank_cat[cat].get(lj,0)+1
    cats=sorted(({"categoria":k,"nome":MAU_USO_MAP.get(k,k),"total":v} for k,v in por_cat.items()),key=lambda x:-x["total"])
    lojas=sorted(({"loja":k,"total":v} for k,v in por_loja.items()),key=lambda x:-x["total"])
    ranking=[]
    for c in cats:
        tops=sorted(({"loja":k,"total":v} for k,v in rank_cat.get(c["categoria"],{}).items()),key=lambda x:-x["total"])[:5]
        ranking.append({"categoria":c["categoria"],"nome":c["nome"],"lojas":tops})
    return {"total":total,"com_mau_uso":com,"sem_mau_uso":total-com,
            "pct":(round(com*100.0/total,1) if total else 0),
            "por_categoria":cats,"por_loja":lojas,"ranking":ranking}

@app.get("/orc/mau_uso_pendentes")
def orc_mau_uso_pendentes(request: Request, filtro: str="", limite: int=100):
    """Lista os chamados que a IA NÃO classificou (não analisados ou marcados INDETERMINADO),
    para classificação manual. Devolve a DESCRIÇÃO de cada chamado + as categorias disponíveis."""
    exige(request,"MAU_USO")
    ap=["select=numero,aba,loja,tipo_predial,descricao,data_criacao",f"data_criacao=gte.{DASH_INICIO}","limit=50000"]
    ch=_sb_json(f"{SB_URL}/rest/v1/chamados?"+"&".join(ap),SB_KEY) or []
    cls=_sb_json(f"{SB_URL}/rest/v1/chamado_mau_uso?select=numero,aba,categoria,manual&limit=50000",SB_KEY) or []
    mp={(str(x.get("numero")),str(x.get("aba") or "")):x for x in cls}
    f=(filtro or "").strip().lower()
    pend=[]
    for c in ch:
        d=(c.get("descricao") or "").strip()
        if not d: continue
        cur=mp.get((str(c.get("numero")),str(c.get("aba") or "")))
        # pendente = sem classificação OU classificado como INDETERMINADO pela IA (e não travado por humano)
        if cur and (cur.get("manual") or cur.get("categoria")!="INDETERMINADO"): continue
        if f and (f not in d.lower()) and (f not in str(c.get("numero") or "")): continue
        pend.append({"numero":str(c.get("numero")),"aba":str(c.get("aba") or ""),
            "loja":_loja_padrao(c.get("loja")),"loja_raw":c.get("loja"),
            "tipo_predial":c.get("tipo_predial"),"descricao":d,"data_criacao":c.get("data_criacao"),
            "categoria_atual":(cur.get("categoria") if cur else None)})
    total=len(pend)
    try: lim=max(1,min(int(limite or 100),500))
    except Exception: lim=100
    cats=[{"codigo":c,"nome":n,"mau_uso":(c not in MAU_USO_SEM)} for c,n in MAU_USO_CATS]
    return {"total":total,"mostrando":min(total,lim),"itens":pend[:lim],"categorias":cats}

@app.post("/orc/mau_uso_manual")
async def orc_mau_uso_manual(request: Request):
    """Salva classificações MANUAIS (uma ou várias). Ficam com manual=true e o
    'Reanalisar tudo' não as sobrescreve."""
    from fastapi import HTTPException
    u,p=exige(request,"MAU_USO")
    b=await request.json()
    itens=b.get("itens") if isinstance(b.get("itens"),list) else ([b] if b.get("numero") else [])
    linhas=[]
    for it in itens:
        num=re.sub(r"\D","",str(it.get("numero") or ""))
        cat=str(it.get("categoria") or "").strip().upper()
        if not num or cat not in MAU_USO_MAP: continue
        linhas.append({"numero":num,"aba":str(it.get("aba") or ""),
            "categoria":cat,"tem_mau_uso":(cat not in MAU_USO_SEM),
            "explicacao":"classificação manual","manual":True,
            "loja":it.get("loja") or it.get("loja_raw"),"tipo_predial":it.get("tipo_predial"),
            "data_criacao":it.get("data_criacao")})
    if not linhas: raise HTTPException(400,"nada para salvar (informe número e categoria válida)")
    try:
        body=json.dumps(linhas,ensure_ascii=False).encode()
        rq=urllib.request.Request(f"{SB_URL}/rest/v1/chamado_mau_uso?on_conflict=numero,aba",data=body,method="POST",
            headers={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}","content-type":"application/json","prefer":"resolution=merge-duplicates,return=minimal"})
        urllib.request.urlopen(rq,timeout=40)
    except Exception as e: raise HTTPException(500,f"falha ao salvar: {str(e)[:150]}")
    log_frotahub(u["id"],p.get("papel"),"MAU_USO","MANUAL",f"{len(linhas)} chamados")
    return {"ok":True,"salvos":len(linhas)}

# ================= USUÁRIOS / LOGINS / CATEGORIAS / PIN =================
import hashlib
PIN_PEPPER  = os.environ.get("PIN_PEPPER", "frotahub-pin-v1")
USUARIO_DOM = os.environ.get("USUARIO_DOMINIO", "frotahub.local")

def _pin_hash(uid, pin): return hashlib.sha256(f"{PIN_PEPPER}:{uid}:{pin}".encode()).hexdigest()
def _slug(s):
    s=unicodedata.normalize("NFKD",s or "").encode("ascii","ignore").decode().lower()
    s=re.sub(r"[^a-z0-9]+","_",s).strip("_"); return s or "cat"
def _usuario_email(usuario):
    u=(usuario or "").strip()
    return u if "@" in u else (re.sub(r"\s+","",u.lower())+"@"+USUARIO_DOM)

def _sb_write(path, data, method="POST", prefer="return=minimal"):
    url=f"{SB_URL}/rest/v1/{path}"
    hdrs={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}","content-type":"application/json","prefer":prefer}
    req=urllib.request.Request(url,data=json.dumps(data).encode(),headers=hdrs,method=method)
    with urllib.request.urlopen(req,timeout=30) as r:
        b=r.read().decode(); return json.loads(b) if b else None

def _sb_admin(path, data=None, method="GET"):
    url=f"{SB_URL}/auth/v1{path}"
    hdrs={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}"}
    if data is not None: hdrs["content-type"]="application/json"
    req=urllib.request.Request(url,data=(json.dumps(data).encode() if data is not None else None),headers=hdrs,method=method)
    with urllib.request.urlopen(req,timeout=30) as r:
        b=r.read().decode(); return json.loads(b) if b else None

def _verifica_senha(email, senha):
    if not email or not senha: return False
    try:
        _sb_json(f"{SB_URL}/auth/v1/token?grant_type=password", SB_ANON,
                 data={"email":email,"password":senha}, method="POST"); return True
    except Exception: return False

def _pin_do(uid):
    d=_sb_json(f"{SB_URL}/rest/v1/perfis?id=eq.{uid}&select=pin_hash&limit=1",SB_KEY) or []
    return d[0].get("pin_hash") if d else None
def _verifica_pin(uid, nivel, pin):
    pin=str(pin or "").strip()
    if not pin: return False
    ph=_pin_do(uid)
    if not ph: return nivel=="builder" and pin=="1234"   # PIN inicial do builder
    return _pin_hash(uid,pin)==ph

def _exige_gestor(request):
    from fastapi import HTTPException
    u=auth_user(_bearer(request))
    if not u or not u.get("id"): raise HTTPException(401,"não autenticado")
    p=perfil_de(u["id"])
    if not p or p.get("ativo") is False: raise HTTPException(403,"sem perfil ativo")
    if p.get("nivel") not in ("builder","gerente"): raise HTTPException(403,"apenas builder e gerente")
    return u,p

def _nivel_categoria(cat_id):
    c=_sb_json(f"{SB_URL}/rest/v1/categorias?id=eq.{urllib.parse.quote(str(cat_id))}&select=nivel,protegida&limit=1",SB_KEY) or []
    return (c[0] if c else {"nivel":"comum","protegida":False})

@app.get("/usuarios/listar")
def usuarios_listar(request: Request):
    _exige_gestor(request)
    rows=_sb_json(f"{SB_URL}/rest/v1/perfis?select=id,usuario,nome,nome_completo,cpf,ativo,categoria_id,primeiro_acesso,must_change_pw,categorias(nivel,nome)&order=usuario",SB_KEY) or []
    for r in rows:
        c=r.get("categorias") or {}; r["nivel"]=c.get("nivel"); r["categoria_nome"]=c.get("nome"); r.pop("categorias",None)
    return {"itens":rows}

@app.post("/usuarios/criar")
async def usuarios_criar(request: Request):
    from fastapi import HTTPException
    u,p=_exige_gestor(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
    usuario=(b.get("usuario") or "").strip(); senha=b.get("senha") or ""; cat=(b.get("categoria_id") or "").strip()
    nome=(b.get("nome_completo") or "").strip()
    if not (usuario and senha and cat): raise HTTPException(400,"usuário, senha e categoria são obrigatórios")
    if len(senha)<6: raise HTTPException(400,"senha muito curta (mín. 6)")
    nc=_nivel_categoria(cat); niv=nc.get("nivel")
    if niv in ("builder","gerente") and p.get("nivel")!="builder":
        raise HTTPException(403,"apenas o builder cria/edita logins builder ou gerente")
    papel={"builder":"builder","gerente":"gerente"}.get(niv,"administrativo")
    try:
        novo=_sb_admin("/admin/users", {"email":_usuario_email(usuario),"password":senha,"email_confirm":True}, "POST")
    except urllib.error.HTTPError as e:
        raise HTTPException(400,f"criar usuário: {e.read().decode()[:180]}")
    uid=novo.get("id")
    try:
        _sb_write("perfis", {"id":uid,"usuario":usuario,"nome":nome or usuario,"nome_completo":nome or None,
            "papel":papel,"categoria_id":cat,"ativo":True,"must_change_pw":True,"primeiro_acesso":True})
    except Exception as e:
        try: _sb_admin(f"/admin/users/{uid}", method="DELETE")   # desfaz se o perfil falhar
        except Exception: pass
        raise HTTPException(400,f"criar perfil: {str(e)[:180]}")
    log_frotahub(u["id"],p.get("papel"),"CONFIG_USUARIOS","CRIOU_LOGIN",usuario)
    return {"ok":True,"id":uid,"login":_usuario_email(usuario)}

@app.post("/usuarios/editar")
async def usuarios_editar(request: Request):
    from fastapi import HTTPException
    u,p=_exige_gestor(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
    uid=b.get("id");  alvo=_sb_json(f"{SB_URL}/rest/v1/perfis?id=eq.{uid}&select=categoria_id,categorias(nivel)&limit=1",SB_KEY) or []
    if not alvo: raise HTTPException(404,"login não encontrado")
    niv_alvo=(alvo[0].get("categorias") or {}).get("nivel")
    if niv_alvo in ("builder","gerente") and p.get("nivel")!="builder": raise HTTPException(403,"apenas o builder edita builder/gerente")
    patch={}
    for k in ("usuario","nome_completo","cpf","categoria_id","ativo"):
        if k in b and b[k] is not None: patch[k]=b[k]
    if "nome_completo" in patch: patch["nome"]=patch["nome_completo"]
    if patch.get("categoria_id"):
        nc=_nivel_categoria(patch["categoria_id"])
        if nc.get("nivel") in ("builder","gerente") and p.get("nivel")!="builder": raise HTTPException(403,"só o builder promove a builder/gerente")
        patch["papel"]={"builder":"builder","gerente":"gerente"}.get(nc.get("nivel"),"administrativo")
    if patch: _sb_write(f"perfis?id=eq.{uid}", patch, "PATCH")
    log_frotahub(u["id"],p.get("papel"),"CONFIG_USUARIOS","EDITOU_LOGIN",str(uid))
    return {"ok":True}

@app.post("/usuarios/excluir")
async def usuarios_excluir(request: Request):
    from fastapi import HTTPException
    u,p=_exige_gestor(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
    uid=b.get("id")
    alvo=_sb_json(f"{SB_URL}/rest/v1/perfis?id=eq.{uid}&select=usuario,categorias(nivel)&limit=1",SB_KEY) or []
    if not alvo: raise HTTPException(404,"login não encontrado")
    niv=(alvo[0].get("categorias") or {}).get("nivel")
    if niv=="builder": raise HTTPException(403,"o login builder não pode ser excluído")
    if niv=="gerente" and p.get("nivel")!="builder": raise HTTPException(403,"apenas o builder exclui gerentes")
    # desvincula os logs do usuário (o FK log_atividades.user_id -> auth.users bloqueia o delete).
    # preserva o histórico com user_id nulo em vez de apagá-lo.
    try: _sb_write(f"log_atividades?user_id=eq.{uid}", {"user_id":None}, "PATCH")
    except Exception: pass
    try: _sb_admin(f"/admin/users/{uid}", method="DELETE")   # perfil cai por cascade (FK on delete cascade)
    except urllib.error.HTTPError as e: raise HTTPException(400,f"excluir: {e.read().decode()[:160]}")
    log_frotahub(u["id"],p.get("papel"),"CONFIG_USUARIOS","EXCLUIU_LOGIN",alvo[0].get("usuario") or str(uid))
    return {"ok":True}

@app.post("/usuarios/reset_senha")
async def usuarios_reset_senha(request: Request):
    from fastapi import HTTPException
    u,p=_exige_gestor(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
    uid=b.get("id"); nova=b.get("nova_senha") or ""
    if len(nova)<6: raise HTTPException(400,"senha muito curta (mín. 6)")
    alvo=_sb_json(f"{SB_URL}/rest/v1/perfis?id=eq.{uid}&select=categorias(nivel)&limit=1",SB_KEY) or []
    if not alvo: raise HTTPException(404,"login não encontrado")
    if (alvo[0].get("categorias") or {}).get("nivel")=="gerente" and p.get("nivel")!="builder":
        raise HTTPException(403,"apenas o builder reseta senha de gerente")
    try: _sb_admin(f"/admin/users/{uid}", {"password":nova}, "PUT")
    except urllib.error.HTTPError as e: raise HTTPException(400,f"reset: {e.read().decode()[:160]}")
    _sb_write(f"perfis?id=eq.{uid}", {"must_change_pw":True}, "PATCH")
    log_frotahub(u["id"],p.get("papel"),"CONFIG_USUARIOS","RESETOU_SENHA",str(uid))
    return {"ok":True}

@app.post("/usuarios/primeiro_acesso")
async def usuarios_primeiro_acesso(request: Request):
    from fastapi import HTTPException
    u=auth_user(_bearer(request))
    if not u or not u.get("id"): raise HTTPException(401,"não autenticado")
    b=await request.json()
    pin=str(b.get("pin") or "").strip()
    if len(pin)!=4 or not pin.isdigit(): raise HTTPException(400,"o PIN deve ter 4 dígitos")
    nome=(b.get("nome_completo") or "").strip(); cpf=(b.get("cpf") or "").strip()
    if not nome: raise HTTPException(400,"informe o nome completo")
    _sb_write(f"perfis?id=eq.{u['id']}", {"nome_completo":nome,"nome":nome,"cpf":cpf or None,
        "pin_hash":_pin_hash(u["id"],pin),"primeiro_acesso":False,"must_change_pw":False}, "PATCH")
    return {"ok":True}

@app.post("/usuarios/senha_trocada")
async def usuarios_senha_trocada(request: Request):
    """Chamado após o usuário trocar a senha no 1º acesso/reset (limpa a flag)."""
    from fastapi import HTTPException
    u=auth_user(_bearer(request))
    if not u or not u.get("id"): raise HTTPException(401,"não autenticado")
    _sb_write(f"perfis?id=eq.{u['id']}", {"must_change_pw":False}, "PATCH")
    return {"ok":True}

@app.post("/usuarios/mudar_pin")
async def usuarios_mudar_pin(request: Request):
    from fastapi import HTTPException
    u=auth_user(_bearer(request));  p=perfil_de(u["id"]) if u else None
    if not u or not p: raise HTTPException(401,"não autenticado")
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin_atual")): raise HTTPException(403,"PIN atual incorreto")
    novo=str(b.get("pin_novo") or "").strip()
    if len(novo)!=4 or not novo.isdigit(): raise HTTPException(400,"o novo PIN deve ter 4 dígitos")
    _sb_write(f"perfis?id=eq.{u['id']}", {"pin_hash":_pin_hash(u["id"],novo)}, "PATCH")
    return {"ok":True}

@app.post("/usuarios/verifica_pin")
async def usuarios_verifica_pin(request: Request):
    from fastapi import HTTPException
    u=auth_user(_bearer(request));  p=perfil_de(u["id"]) if u else None
    if not u or not p: raise HTTPException(401,"não autenticado")
    b=await request.json()
    return {"ok": _verifica_pin(u["id"], p.get("nivel"), b.get("pin"))}

# ---- categorias e matriz de permissões ----
@app.get("/config/categorias")
def config_categorias(request: Request):
    _exige_gestor(request)
    cats=_sb_json(f"{SB_URL}/rest/v1/categorias?select=id,nome,nivel,protegida&order=nome",SB_KEY) or []
    perms=_sb_json(f"{SB_URL}/rest/v1/categoria_permissoes?select=categoria_id,rotina,pode",SB_KEY) or []
    return {"categorias":cats,"permissoes":perms}

@app.post("/config/categoria_criar")
async def config_categoria_criar(request: Request):
    from fastapi import HTTPException
    u,p=_exige_gestor(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
    nome=(b.get("nome") or "").strip()
    if not nome: raise HTTPException(400,"informe o nome da categoria")
    cid=_slug(nome)
    try: _sb_write("categorias", {"id":cid,"nome":nome,"nivel":"comum","protegida":False})
    except urllib.error.HTTPError as e: raise HTTPException(400,f"criar categoria: {e.read().decode()[:160]}")
    log_frotahub(u["id"],p.get("papel"),"CONFIG_CATEGORIAS","CRIOU_CATEGORIA",cid)
    return {"ok":True,"id":cid}

@app.post("/config/categoria_excluir")
async def config_categoria_excluir(request: Request):
    from fastapi import HTTPException
    u,p=_exige_gestor(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
    cid=b.get("id"); nc=_nivel_categoria(cid)
    if nc.get("protegida"): raise HTTPException(403,"categoria protegida não pode ser excluída")
    emuso=_sb_json(f"{SB_URL}/rest/v1/perfis?categoria_id=eq.{urllib.parse.quote(str(cid))}&select=id&limit=1",SB_KEY) or []
    if emuso: raise HTTPException(400,"há logins usando esta categoria")
    try:
        req=urllib.request.Request(f"{SB_URL}/rest/v1/categorias?id=eq.{urllib.parse.quote(str(cid))}",method="DELETE",
            headers={"apikey":SB_KEY,"authorization":f"Bearer {SB_KEY}","prefer":"return=minimal"})
        urllib.request.urlopen(req,timeout=20)
    except Exception as e: raise HTTPException(400,f"excluir: {str(e)[:120]}")
    return {"ok":True}

@app.post("/config/permissoes_set")
async def config_permissoes_set(request: Request):
    from fastapi import HTTPException
    u,p=_exige_gestor(request)
    b=await request.json()
    if not _verifica_pin(u["id"], p.get("nivel"), b.get("pin")): raise HTTPException(403,"PIN incorreto")
    cid=b.get("categoria_id"); nc=_nivel_categoria(cid)
    if nc.get("protegida") and p.get("nivel")!="builder": raise HTTPException(403,"categoria protegida — só o builder altera")
    rows=[{"categoria_id":cid,"rotina":it.get("rotina"),"pode":bool(it.get("pode"))} for it in (b.get("permissoes") or []) if it.get("rotina")]
    if rows: _sb_write("categoria_permissoes", rows, "POST", "resolution=merge-duplicates,return=minimal")
    log_frotahub(u["id"],p.get("papel"),"CONFIG_CATEGORIAS","AJUSTOU_PERMISSOES",str(cid))
    return {"ok":True,"n":len(rows)}

# ---- política de PIN (builder-only, exige senha do builder) ----
@app.get("/config/pin_policy")
def config_pin_policy(request: Request):
    _exige_gestor(request)
    rows=_sb_json(f"{SB_URL}/rest/v1/pin_policy?select=rotina,exige_pin",SB_KEY) or []
    return {"itens":rows}

@app.post("/config/pin_policy_set")
async def config_pin_policy_set(request: Request):
    from fastapi import HTTPException
    u,p=_exige_gestor(request)
    if p.get("nivel")!="builder": raise HTTPException(403,"apenas o builder altera a política de PIN")
    b=await request.json()
    if not _verifica_senha(u.get("email"), b.get("senha")): raise HTTPException(403,"senha do builder incorreta")
    rows=[{"rotina":it.get("rotina"),"exige_pin":bool(it.get("exige_pin"))} for it in (b.get("itens") or []) if it.get("rotina")]
    if rows: _sb_write("pin_policy", rows, "POST", "resolution=merge-duplicates,return=minimal")
    log_frotahub(u["id"],p.get("papel"),"CONFIG_PIN","AJUSTOU_POLITICA_PIN",f"{len(rows)} itens")
    return {"ok":True}

# ---- Camada de arquivos hot/cold + Backup (Supabase Storage <-> Dropbox) ----
try:
    import supabase_storage as sst
    import backup_endpoints
    backup_endpoints.montar(app, dict(
        exige=exige, verifica_pin=_verifica_pin, dropbox=dropbox_rateio,
        sb_url=SB_URL, sb_key=SB_KEY, bucket=os.environ.get("SUPABASE_BUCKET", "frotahub"),
        agora=_agora, hoje=_hoje, cfg_get=_cfg_get, cfg_set=_cfg_set, log=log_frotahub, sst=sst,
    ))
    print("backup_endpoints montado (rev 114)")
except Exception as _e:
    print("AVISO: backup_endpoints não montado:", _e)

# ---- Servir o FrotaHub (front estático) pelo próprio motor, em /app ----
# Alternativa ao Netlify: coloque os arquivos do site na pasta "frontend/" ao lado
# do app.py (index.html, manifest.webmanifest, sw.js, share.html, favicon.ico, icons/).
try:
    from fastapi.staticfiles import StaticFiles
    _FRONT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "frontend")
    if os.path.isdir(_FRONT):
        app.mount("/app", StaticFiles(directory=_FRONT, html=True), name="frontend")
        print("front estático montado em /app")
except Exception as _e:
    print("AVISO: front estático não montado:", _e)


def _basic_auth(app, user, pw):
    """Protege TODAS as rotas com usuário/senha (HTTP Basic), na camada ASGI."""
    import base64, secrets
    async def wrapped(scope, receive, send):
        # rotas do FrotaHub têm autenticação própria (token Supabase) — não pedir Basic
        _path = scope.get("path", "") or ""
        if scope["type"] == "http" and (_path.startswith("/pco") or _path.startswith("/api")
                                        or _path.startswith("/notas") or _path.startswith("/orc") or _path.startswith("/migrar")
                                        or _path.startswith("/desfazer") or _path.startswith("/usuarios") or _path.startswith("/config")
                                        or _path.startswith("/rateio") or _path.startswith("/arq") or _path.startswith("/robot")
                                        or _path.startswith("/backup") or _path.startswith("/app")
                                        or scope.get("method") == "OPTIONS"):
            await app(scope, receive, send); return
        if scope["type"] in ("http", "websocket"):
            hdrs = dict(scope.get("headers") or [])
            auth = hdrs.get(b"authorization", b"").decode()
            ok = False
            if auth.startswith("Basic "):
                try:
                    usr, _, pwd = base64.b64decode(auth[6:]).decode("utf-8", "ignore").partition(":")
                    ok = secrets.compare_digest(usr, user) and secrets.compare_digest(pwd, pw)
                except Exception:
                    ok = False
            if not ok:
                if scope["type"] == "http":
                    await send({"type": "http.response.start", "status": 401,
                                "headers": [(b"www-authenticate", b'Basic realm="Motor"'),
                                            (b"content-type", b"text/plain; charset=utf-8")]})
                    await send({"type": "http.response.body", "body": "Autenticação necessária".encode()})
                else:
                    await send({"type": "websocket.close", "code": 1008})
                return
        await app(scope, receive, send)
    return wrapped

if __name__ == "__main__":
    port = int(os.environ.get("PORT", "7860"))    # Render/Cloud injeta PORT
    u, p = os.environ.get("APP_USER", ""), os.environ.get("APP_PASS", "")
    import uvicorn
    asgi = _basic_auth(app, u, p) if (u and p) else app
    uvicorn.run(asgi, host="0.0.0.0", port=port)
